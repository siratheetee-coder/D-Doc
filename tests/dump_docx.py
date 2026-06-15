"""อ่านโครงสร้างไฟล์ .docx ตัวอย่าง เพื่อศึกษารูปแบบเอกสารจริง"""
import sys
from docx import Document

path = sys.argv[1]
doc = Document(path)

print(f"===== ไฟล์: {path} =====")
print(f"จำนวนย่อหน้า: {len(doc.paragraphs)} | จำนวนตาราง: {len(doc.tables)}")
print(f"จำนวนรูปภาพ (โดยประมาณ): {len(doc.inline_shapes)}")
print("=" * 60)

# เดินไล่ทั้งย่อหน้าและตารางตามลำดับในเอกสาร
from docx.oxml.ns import qn

body = doc.element.body
para_idx = 0
tbl_idx = 0
block_no = 0
for child in body.iterchildren():
    if child.tag == qn('w:p'):
        # หา paragraph object ตรงกับ element นี้
        text = "".join(node.text or "" for node in child.iter(qn('w:t')))
        if text.strip():
            block_no += 1
            print(f"[P{block_no}] {text.strip()}")
    elif child.tag == qn('w:tbl'):
        tbl_idx += 1
        print(f"\n--- ตารางที่ {tbl_idx} ---")
        # หา table object
        for t in doc.tables:
            if t._tbl is child:
                for r_i, row in enumerate(t.rows):
                    cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                    print(f"   แถว{r_i+1}: " + " | ".join(cells))
                break
        print("--- จบตาราง ---\n")
