# -*- coding: utf-8 -*-
"""ตรวจ: ตารางแนบท้าย 7 คอลัมน์ + เอกสารใบเบิกจ่ายใหม่ + ช่องลงนามตาราง"""
import tempfile, os
from types import SimpleNamespace as NS
from docxtpl import DocxTemplate
from docx import Document
from app.services.render import build_context, TEMPLATES_DIR, TEMPLATE_FILES

class Sch:
    name="โรงเรียนเพียงหลวง ๑๘"; address="ต.ปากตม อ.เชียงคาน จ.เลย"; director_name="นายสุรชาติ จันปัดถา"
    officer_name="นางสาวศิริลักษณ์ ปองดี"; head_officer_name="นายสุรชาติ จันปัดถา"; doc_prefix="ศธ"
    director_position="ผู้อำนวยการโรงเรียน"; finance_officer_name="นางกองแก้ว กันยะหา"

def mkproc(vat, overdue=0):
    items=[NS(name="ปากกา",quantity=2,unit="กล่อง",unit_price=100,amount=200),
           NS(name="กระดาษ A4",quantity=25,unit="รีม",unit_price=111,amount=2775)]
    return NS(items=items, committees=[], memo_no="11/2569", spec_memo_no="", inspect_memo_no="13/2569",
              result_memo_no="12/2569", request_date=None, proc_type="ซื้อ", subject="วัสดุสำนักงาน",
              project_name="โครงการ X", department="วิชาการ", purpose="ใช้สอน", method="เฉพาะเจาะจง",
              budget_source="อุดหนุน", delivery_days=3, total_amount=2975, inspection_mode="committee",
              vendor=NS(name="หจก.จัสมิน",address="เลย",tax_id="0103",bank_account="-",phone="-"),
              order_no="1/2569", command_no="128/2569", penalty_rate=0.2,
              vat_mode=("include" if vat else "none"), order_signer="director", overdue_days=overdue,
              delivery_note_no="5", delivery_note_book="2", order_date=None, delivery_due_date=None,
              inspect_date=None)

# 1) ตารางแนบท้าย 7 คอลัมน์ ในรายงานขอซื้อ
tpl=DocxTemplate(str(TEMPLATES_DIR/TEMPLATE_FILES["รายงานขอซื้อ"]))
tpl.render(build_context(mkproc(True), Sch()))
f=os.path.join(tempfile.gettempdir(),"pr.docx"); tpl.save(f)
d=Document(f)
big=[t for t in d.tables if len(t.columns)==7]
assert big, "ไม่พบตารางแนบท้าย 7 คอลัมน์"
t=big[0]
print("=== ตารางแนบท้าย (7 คอลัมน์) ===")
for r in t.rows[:6]:
    print(" | ".join(c.text.strip() for c in r.cells))
print("... แถวสรุปท้าย ...")
for r in t.rows[-4:]:
    print(" | ".join(c.text.strip() for c in r.cells))

# 2) เอกสารใบเบิกจ่ายใหม่
tpl2=DocxTemplate(str(TEMPLATES_DIR/TEMPLATE_FILES["รายงานผลตรวจรับและเบิกจ่าย"]))
tpl2.render(build_context(mkproc(True, overdue=2), Sch()))
f2=os.path.join(tempfile.gettempdir(),"db.docx"); tpl2.save(f2)
d2=Document(f2)
txt="\n".join(p.text for p in d2.paragraphs)
txt += "\n" + "\n".join(c.text for t in d2.tables for r in t.rows for c in r.cells)
print("\n=== ใบเบิกจ่าย: บรรทัดการเงิน ===")
for line in txt.splitlines():
    if any(k in line for k in ("มูลค่าสินค้า","ภาษีมูลค่าเพิ่ม","ขอเบิกทั้งสิ้น","ภาษีเงินได้","ค่าปรับ","คงเหลือ")):
        print(line.strip())
assert "คงเหลือจ่ายจริง" in txt and "นางกองแก้ว" in txt
print("\nOK: ใบเบิกจ่าย + ตาราง 7 คอลัมน์ ทำงานถูกต้อง")
