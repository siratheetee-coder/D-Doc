"""วินิจฉัยว่า construct ไหนของ docxtpl พัง: table-row loop vs paragraph if/for"""
import tempfile, os
from docx import Document
from docxtpl import DocxTemplate


def build_table_single_row(path):
    d = Document()
    t = d.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "ลำดับ"; t.rows[0].cells[1].text = "ชื่อ"
    t.rows[1].cells[0].text = "{%tr for it in items %}{{ loop.index }}"
    t.rows[1].cells[1].text = "{{ it.name }}{%tr endfor %}"
    d.save(path)


def build_table_three_row(path):
    d = Document()
    t = d.add_table(rows=4, cols=2)
    t.rows[0].cells[0].text = "ลำดับ"; t.rows[0].cells[1].text = "ชื่อ"
    t.rows[1].cells[0].text = "{%tr for it in items %}"
    t.rows[2].cells[0].text = "{{ loop.index }}"; t.rows[2].cells[1].text = "{{ it.name }}"
    t.rows[3].cells[0].text = "{%tr endfor %}"
    d.save(path)


def build_par_if_for(path):
    d = Document()
    d.add_paragraph("{%p if mode == 'single' %}")
    d.add_paragraph("คนเดียว {{ who }}")
    d.add_paragraph("{%p else %}")
    d.add_paragraph("{%p for m in members %}")
    d.add_paragraph("- {{ m }}")
    d.add_paragraph("{%p endfor %}")
    d.add_paragraph("{%p endif %}")
    d.save(path)


def try_render(name, builder, ctx):
    path = os.path.join(tempfile.gettempdir(), name)
    builder(path)
    try:
        tpl = DocxTemplate(path)
        tpl.render(ctx)
        print(f"[OK] {name}")
    except Exception as e:
        print(f"[FAIL] {name}: {type(e).__name__}: {e}")


if __name__ == "__main__":
    try_render("t_single.docx", build_table_single_row, {"items": [{"name": "ก"}, {"name": "ข"}]})
    try_render("t_three.docx", build_table_three_row, {"items": [{"name": "ก"}, {"name": "ข"}]})
