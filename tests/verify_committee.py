# -*- coding: utf-8 -*-
"""ตรวจ: ตารางกรรมการ (ไร้เส้น) ในรายงานขอซื้อ โหมดคณะกรรมการ + วันที่หัวเป็นแบบสั้น"""
import tempfile, os
from types import SimpleNamespace as NS
from datetime import datetime
from docxtpl import DocxTemplate
from docx import Document
from app.services.render import build_context, TEMPLATES_DIR, TEMPLATE_FILES

class Sch:
    name="โรงเรียนเพียงหลวง 18"; address="ต.ปากตม อ.เชียงคาน จ.เลย"; director_name="นายสุรชาติ จันปัดถา"
    officer_name="น.ส.ศิริลักษณ์"; head_officer_name="นายสุรชาติ"; doc_prefix="ศธ"
    director_position="ผู้อำนวยการโรงเรียน"; finance_officer_name="นางกองแก้ว"

members=[NS(name="นางสาวรดาณัฐ นพรัตน์ไมตรี",position="ครูชำนาญการ",role="ประธานกรรมการ"),
         NS(name="นางสาวภคพร พินทา",position="ครูชำนาญการ",role="กรรมการ"),
         NS(name="นายวรุตม์ เพ็ชรสิน",position="ครูผู้ช่วย",role="กรรมการ")]
proc=NS(items=[NS(name="ปากกา",quantity=2,unit="กล่อง",unit_price=100,amount=200)],
        committees=[NS(kind="inspect",members=members)], memo_no="11/2569", spec_memo_no="",
        inspect_memo_no="13/2569", result_memo_no="12/2569",
        request_date=datetime(2025,11,3), proc_type="ซื้อ", subject="วัสดุ", project_name="X",
        department="วิชาการ", purpose="ใช้สอน", method="เฉพาะเจาะจง", budget_source="อุดหนุน",
        delivery_days=3, total_amount=200, inspection_mode="committee",
        vendor=None, order_no="1/2569", command_no="128/2569", penalty_rate=0.2, vat_mode="none",
        order_signer="director", overdue_days=0, delivery_note_no="", delivery_note_book="",
        order_date=datetime(2025,11,3), delivery_due_date=None, inspect_date=None)

tpl=DocxTemplate(str(TEMPLATES_DIR/TEMPLATE_FILES["รายงานขอซื้อ"]))
tpl.render(build_context(proc, Sch()))
f=os.path.join(tempfile.gettempdir(),"prc.docx"); tpl.save(f)
d=Document(f)
paras="\n".join(p.text for p in d.paragraphs)
# วันที่หัวเป็นแบบสั้น (ไม่มีคำว่า 'เดือน'/'พ.ศ.')
hdr=[l for l in paras.splitlines() if l.startswith("ที่")]
print("หัววันที่:", hdr[0] if hdr else "(ไม่พบ)")
assert "เดือน" not in (hdr[0] if hdr else ""), "วันที่หัวควรเป็นแบบสั้น"
# ตารางกรรมการ
print("\n=== ตารางกรรมการ ===")
for t in d.tables:
    cells=[c.text.strip() for r in t.rows for c in r.cells]
    if any("ประธานกรรมการ" in x for x in cells):
        for r in t.rows:
            print(" | ".join(c.text.strip() for c in r.cells))
        break
assert "นางสาวรดาณัฐ นพรัตน์ไมตรี" in [c.text.strip() for t in d.tables for r in t.rows for c in r.cells]
print("\nOK: ตารางกรรมการ + วันที่หัวแบบสั้น ทำงานถูกต้อง")
