# -*- coding: utf-8 -*-
"""ตรวจฟอร์แมตละเอียดของ docx: ฟอนต์ ขนาด ระยะขอบ เส้นตาราง การจัดวาง"""
import sys
from docx import Document
from docx.shared import Pt

path = sys.argv[1]
d = Document(path)

# section margins / page size
for i, s in enumerate(d.sections):
    print(f"[SECTION {i}] page {s.page_width and round(s.page_width.cm,2)} x {s.page_height and round(s.page_height.cm,2)} cm | "
          f"margins L{round(s.left_margin.cm,2)} R{round(s.right_margin.cm,2)} T{round(s.top_margin.cm,2)} B{round(s.bottom_margin.cm,2)} cm")

# default style
ns = d.styles['Normal']
print(f"[Normal style] font={ns.font.name} size={ns.font.size}")

def run_info(r):
    sz = r.font.size.pt if r.font.size else None
    # szCs (complex script size)
    rpr = r._element.rPr
    cs = None; rfonts = None; lang=None
    if rpr is not None:
        szcs = rpr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}szCs')
        if szcs is not None: cs = int(szcs.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'))/2
        rf = rpr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
        if rf is not None: rfonts = rf.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs') or rf.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii')
    return f"font={rfonts} sz={sz} szCs={cs} bold={r.font.bold}"

print("\n--- first 25 paragraphs: alignment + first-run font ---")
for i, p in enumerate(d.paragraphs[:25]):
    al = p.alignment
    pf = p.paragraph_format
    ls = pf.line_spacing
    txt = (p.text[:40] + '...') if len(p.text) > 40 else p.text
    ri = run_info(p.runs[0]) if p.runs else "(no run)"
    print(f"[P{i}] align={al} ls={ls} | {ri} | {txt}")

# table borders of table 2 (item table)
print("\n--- table cell/border sample (table index 1) ---")
if len(d.tables) > 1:
    t = d.tables[1]
    tblPr = t._element.tblPr
    from lxml import etree
    print(etree.tostring(tblPr, pretty_print=True).decode()[:1500])
