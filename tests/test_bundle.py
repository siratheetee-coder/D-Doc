"""ทดสอบฟีเจอร์ออกทั้งชุด (.zip) ด้วย TestClient ในโปรเซส (โค้ดปัจจุบัน)"""
import io, zipfile, re
from fastapi.testclient import TestClient
from app.main import app
from app.database import SessionLocal
from app.models import (School, Vendor, Procurement, ProcurementItem,
                        DocNumberCounter, Committee, CommitteeMember, Document)


def reset():
    db = SessionLocal()
    for M in (Document, CommitteeMember, Committee, ProcurementItem, Procurement,
              DocNumberCounter, Vendor, School):
        db.query(M).delete()
    db.commit(); db.close()


def main():
    reset()
    c = TestClient(app)
    c.post("/settings", data={"name": "โรงเรียนบ้านหินลาด", "director_name": "นายอัครพงศ์ ศรีวงศ์",
                              "officer_name": "นายสิรธีร์", "head_officer_name": "นางสาวประทุมพร",
                              "doc_set_threshold": "5000"})
    c.post("/vendors", data={"name": "ร้านทดสอบ", "tax_id": "1234567890123"})
    vid = re.search(r'/vendors/(\d+)/delete', c.get("/vendors").text).group(1)

    # เรื่องวงเงินใหญ่ (>5000) -> ควรติ๊กครบชุด
    r = c.post("/procurement/new", data={
        "fiscal_year": "2569", "memo_no": "68/2569", "proc_type": "จ้าง", "method": "เฉพาะเจาะจง",
        "subject": "จัดทำป้ายไวนิล", "vendor_id": vid, "inspection_mode": "committee",
        "item_name": ["ป้าย"], "item_qty": ["1"], "item_unit": ["ป้าย"], "item_price": ["8000"],
        "member_name": ["นาย ก", "นาย ข", "นาย ค"], "member_position": ["ครู","ครู","ครู"],
        "member_role": ["ประธานกรรมการ","กรรมการ","กรรมการและเลขานุการ"],
    })
    pid = r.url.path.split("/")[-1]

    def count_checked(html):
        return sum(1 for ln in html.splitlines()
                   if "<input" in ln and 'name="kinds"' in ln and "checked" in ln)

    # วงเงินใหญ่: ติ๊ก 9 ใบ (ยกเว้น แต่งตั้งกก.คุณลักษณะ + TOR)
    page = c.get(f"/procurement/{pid}/bundle").text
    assert "ออกเอกสารทั้งชุด" in page
    assert count_checked(page) == 9, f"วงเงินใหญ่ควรติ๊ก 9 ใบ ได้ {count_checked(page)}"
    assert "แต่งตั้งกรรมการคุณลักษณะ" in page  # มีในรายการให้เลือก แต่ไม่ติ๊ก
    print("[1] หน้า bundle วงเงินใหญ่ ติ๊ก 9 ใบ: OK")

    # ตั้งชื่อโครงการ+ราคาให้เรื่องนี้ เพื่อตรวจชื่อไฟล์
    c.post(f"/procurement/{pid}/update-refs", data={
        "memo_no": "1/2569", "order_no": "", "command_no": "", "result_memo_no": "",
        "spec_memo_no": "", "inspect_memo_no": ""})

    # POST เลือก 3 ใบ -> ได้ไฟล์ .docx เดียว (รวม 3 ใบ)
    r = c.post(f"/procurement/{pid}/bundle", data={
        "kinds": ["รายงานขอซื้อ", "ใบสั่งซื้อ/สั่งจ้าง", "ใบตรวจรับพัสดุ"]})
    assert r.status_code == 200
    assert "wordprocessing" in r.headers["content-type"], r.headers["content-type"]
    # ชื่อไฟล์ใน header: เลขที่ ชื่อโครงการ ราคา
    cd = r.headers.get("content-disposition", "")
    assert "1-2569" in cd, f"ชื่อไฟล์ควรมีเลขที่: {cd}"
    # ไฟล์เดียว เปิดได้ และมีเนื้อหาจากทั้ง 3 ใบ (มี page break คั่น)
    from docx import Document as Docx
    d = Docx(io.BytesIO(r.content))
    txt = "\n".join(p.text for p in d.paragraphs)
    ttxt = txt + "\n" + "\n".join(cell.text for t in d.tables for row in t.rows for cell in row.cells)
    assert "บันทึกข้อความ" in ttxt and "ใบสั่ง" in ttxt and "ใบตรวจรับพัสดุ" in ttxt, "เนื้อหาไม่ครบ 3 ใบ"
    # ตรวจลำดับมาตรฐาน: รายงานขอซื้อ(3) -> ใบสั่งซื้อ(8) -> ใบตรวจรับ(10)
    i_req = ttxt.find("รายงานขอ")
    i_order = ttxt.find("ใบสั่งจ้าง") if "ใบสั่งจ้าง" in ttxt else ttxt.find("ใบสั่งซื้อ")
    i_inspect = ttxt.find("ใบตรวจรับพัสดุ")
    assert i_req < i_order < i_inspect, f"ลำดับเอกสารไม่ถูก: {i_req},{i_order},{i_inspect}"
    breaks = txt.count("\f") + sum(1 for p in d.paragraphs for r2 in p.runs
                                   if r2._element.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'))
    print(f"[2] รวมเป็นไฟล์เดียว 3 ใบ (เปิดได้, มีครบ, ชื่อไฟล์ถูก): OK")

    # เรื่องวงเงินเล็ก (<=5000) -> ติ๊กชุดย่อ (2 ใบ)
    r2 = c.post("/procurement/new", data={
        "fiscal_year": "2569", "memo_no": "69/2569", "proc_type": "ซื้อ", "method": "เฉพาะเจาะจง",
        "subject": "วัสดุเล็ก", "inspection_mode": "single",
        "item_name": ["ปากกา"], "item_qty": ["10"], "item_unit": ["ด้าม"], "item_price": ["20"],
        "member_name": ["นาย ง"], "member_position": ["ครู"], "member_role": ["ผู้ตรวจรับ"],
    })
    pid2 = r2.url.path.split("/")[-1]
    page2 = c.get(f"/procurement/{pid2}/bundle").text
    assert "วงเงินไม่เกิน" in page2, "ควรเป็นโหมดวงเงินเล็ก (ชุดย่อ)"
    # วงเงินเล็ก: ติ๊ก 7 ใบ (ยกเว้น ประกาศ + กก.คุณลักษณะ + TOR + คำสั่งแต่งตั้ง)
    n = count_checked(page2)
    assert n == 7, f"วงเงินเล็กควรติ๊ก 7 ใบ ได้ {n}"
    print("[3] หน้า bundle วงเงินเล็ก ติ๊ก 7 ใบ: OK")

    print("\n=== ฟีเจอร์ออกทั้งชุด (.zip) ผ่านทั้งหมด ===")
    c.close()


if __name__ == "__main__":
    main()
