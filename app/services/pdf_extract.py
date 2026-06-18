# -*- coding: utf-8 -*-
"""
pdf_extract.py
--------------
อ่านข้อมูลจากไฟล์ PDF หนังสือราชการ (text-based เช่นที่ออกจาก AMSS)
แล้วแยกฟิลด์สำคัญด้วย heuristic ตามรูปแบบหนังสือราชการ — ไม่ต้องใช้ OCR

คืน dict: {letter_no, letter_date(datetime|None), subject, addressee, from_org, ok, raw_text}
ฟิลด์ที่อ่านไม่ได้จะเป็นค่าว่าง ให้ผู้ใช้กรอก/แก้เองในขั้นตอนตรวจก่อนบันทึก
"""
import re
from datetime import datetime

from app.thai_utils import _THAI_MONTHS

# ตารางแปลงเลขไทย -> อารบิก
_THAI_ARABIC = str.maketrans("๐๑๒๓๔๕๖๗๘๙", "0123456789")


def _arabic(s: str) -> str:
    """แปลงเลขไทยในข้อความเป็นเลขอารบิก เช่น 'ศธ ๐๔๑๑๒/๒๑๓๘' -> 'ศธ 04112/2138'"""
    return (s or "").translate(_THAI_ARABIC)

# เดือนไทย (ข้าม index 0 ที่ว่าง) -> ใช้สร้าง regex และ map ชื่อ->เลขเดือน
_MONTHS = [m for m in _THAI_MONTHS if m]
_MONTH_TO_NUM = {m: i for i, m in enumerate(_THAI_MONTHS) if m}
_MONTH_ALT = "|".join(_MONTHS)


def extract_text(pdf_path: str) -> str:
    """ดึงข้อความทุกหน้าจาก PDF (คืน '' ถ้าอ่านไม่ได้/เป็นไฟล์สแกน)"""
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_text_docx(path: str) -> str:
    """ดึงข้อความจากไฟล์ Word (.docx): ย่อหน้า + เซลล์ตาราง"""
    try:
        from docx import Document as _Docx
        doc = _Docx(path)
        parts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""


def extract_text_any(path: str) -> str:
    """ดึงข้อความตามนามสกุลไฟล์: .docx ใช้ python-docx, อื่น ๆ ใช้ pdfplumber"""
    if path.lower().endswith(".docx"):
        return _extract_text_docx(path)
    return extract_text(path)


def _parse_thai_date(text: str):
    """หาวันที่ไทยในข้อความ เช่น '9 มิถุนายน 2569' -> datetime (พ.ศ.->ค.ศ.)"""
    m = re.search(r"(\d{1,2})\s*(?:เดือน\s*)?(" + _MONTH_ALT + r")\s*(?:พ\.?ศ\.?\s*)?(\d{4})", text)
    if not m:
        return None
    day = int(m.group(1))
    mon = _MONTH_TO_NUM.get(m.group(2))
    year = int(m.group(3))
    if year > 2400:          # พ.ศ. -> ค.ศ.
        year -= 543
    try:
        return datetime(year, mon, day)
    except (ValueError, TypeError):
        return None


def _first(pattern: str, text: str, group: int = 1) -> str:
    m = re.search(pattern, text)
    return m.group(group).strip() if m else ""


def extract_letter_fields(pdf_path: str) -> dict:
    """อ่านฟิลด์หนังสือราชการจากไฟล์ (PDF หรือ Word .docx)"""
    text = extract_text_any(pdf_path)
    result = {"letter_no": "", "letter_date": None, "subject": "",
              "addressee": "", "from_org": "", "ok": bool(text.strip()), "raw_text": text}
    if not text.strip():
        return result   # อ่านข้อความไม่ได้ (อาจเป็นไฟล์สแกน) -> ให้กรอกเอง

    # เลขที่หนังสือ: จับที่ตัว "ศธ" โดยตรง (รหัสกระทรวงศึกษาธิการ) รองรับเลขไทย/อารบิก + เว้นวรรค
    # เช่น "ศธ ๐๔๑๑๒ /2138", "ศธ 04123/ว 456"  — ไม่ต้องมีคำว่า "ที่" นำหน้า
    no = _first(r"(ศธ\s*[0-9๐-๙]{2,8}\s*/\s*(?:ว\s*)?[0-9๐-๙]+)", text)
    if not no:
        # สำรอง: เลขที่ขึ้นต้นด้วยรหัส 2 ตัวอักษรไทยอื่น ๆ (นร/กค/มท ฯลฯ) แล้วตามด้วย /เลข
        no = _first(r"([ก-ฮ]{2}\s*[0-9๐-๙]{2,8}\s*/\s*(?:ว\s*)?[0-9๐-๙]+)", text)
    if not no:
        # สำรองสุดท้าย: หลังคำว่า "ที่"
        no = _first(r"(?:^|\n)\s*ที่\s+([^\n]{2,40})", text)
    # แปลงเลขไทย -> อารบิก + ยุบช่องว่างซ้ำ
    result["letter_no"] = _arabic(re.sub(r"\s+", " ", no).strip(" ."))

    # ลงวันที่
    result["letter_date"] = _parse_thai_date(text)

    # เรื่อง (ป้ายกำกับอยู่ต้นบรรทัด — anchor กัน match คำอื่น)
    result["subject"] = _first(r"(?:^|\n)\s*เรื่อง\s*[:：]?\s*(.+)", text)

    # เรียน (ผู้รับ) — anchor ต้นบรรทัด กันไปจับ "เรียน" ในคำว่า "โรงเรียน"
    result["addressee"] = _first(r"(?:^|\n)\s*เรียน\s*[:：]?\s*(.+)", text)

    # หน่วยงานต้นทาง: เดาจากบรรทัดที่มีคำนำหน้าหน่วยงาน (เอาบรรทัดแรกที่เจอช่วงหัวกระดาษ)
    for line in text.split("\n")[:15]:
        s = line.strip()
        if re.match(r"^(สำนักงาน|โรงเรียน|ที่ทำการ|กรม|กระทรวง|องค์การ|เทศบาล|มหาวิทยาลัย)", s):
            result["from_org"] = s
            break

    return result


def _tables_pdf(path: str):
    try:
        import pdfplumber
        out = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                out.extend(page.extract_tables() or [])
        return out
    except Exception:
        return []


def _tables_docx(path: str):
    try:
        from docx import Document as _D
        d = _D(path)
        return [[[c.text for c in row.cells] for row in t.rows] for t in d.tables]
    except Exception:
        return []


def _extract_tables_any(path: str):
    return _tables_docx(path) if path.lower().endswith(".docx") else _tables_pdf(path)


def _num(s) -> str:
    return re.sub(r"[^\d.]", "", _arabic(s or ""))


def _parse_items_from_tables(tables) -> list:
    """หาตารางที่มีหัวคอลัมน์ รายการ/จำนวน/ราคา แล้วดึงรายการพัสดุ (best-effort)"""
    for tb in tables or []:
        if not tb or len(tb) < 2:
            continue
        header, hidx = None, 0
        for i, row in enumerate(tb[:3]):
            joined = " ".join((c or "") for c in row)
            if ("รายการ" in joined or "รายละเอียด" in joined) and ("จำนวน" in joined or "ราคา" in joined):
                header, hidx = row, i
                break
        if header is None:
            continue

        def col(*keys):
            for j, c in enumerate(header):
                if any(k in (c or "") for k in keys):
                    return j
            return None

        ci_name = col("รายการ", "รายละเอียด")
        ci_qty, ci_unit = col("จำนวน"), col("หน่วย")
        ci_price = col("ราคาต่อหน่วย", "ราคา/หน่วย", "หน่วยละ", "ราคา")
        if ci_name is None:
            continue
        # คำที่บอกว่าเป็นแถวหัวตาราง/ยอดรวม/คำอ่านบาท ไม่ใช่รายการพัสดุจริง
        _JUNK = ("รวม", "ราคากลาง", "บาทถ้วน", "รายละเอียดพัสดุ", "รายการ",
                 "ลงชื่อ", "หมายเหตุ", "จำนวนหน่วย")
        items = []
        for row in tb[hidx + 1:]:
            if not row or ci_name >= len(row):
                continue
            name = (row[ci_name] or "").strip().replace("\n", " ")
            row_text = " ".join((c or "") for c in row)
            # ข้าม: ว่าง / ขึ้นต้นวงเล็บ (คำอ่านบาท) / มีคำว่ายอดรวม/บาทถ้วน / เป็นหัวตารางซ้ำ
            if (not name or name.startswith("(") or "บาทถ้วน" in row_text
                    or any(k in name for k in _JUNK)):
                continue
            qty = _num(row[ci_qty]) if (ci_qty is not None and ci_qty < len(row)) else ""
            unit = (row[ci_unit].strip() if (ci_unit is not None and ci_unit < len(row) and row[ci_unit]) else "")
            # หน่วยที่เป็นคำอ่านบาท (ขึ้นต้นวงเล็บ) = แถวยอดรวม ข้าม
            if unit.startswith("("):
                continue
            price = _num(row[ci_price]) if (ci_price is not None and ci_price < len(row)) else ""
            items.append({"name": name, "qty": qty or 1, "unit": unit or "ชิ้น", "unit_price": price or 0})
        if items:
            return items
    return []


def _parse_inspectors(text: str) -> list:
    """ดึงชื่อกรรมการตรวจรับจากช่วงข้อความใกล้คำว่า 'ตรวจรับ' (best-effort)"""
    idx = -1
    for kw in ("คณะกรรมการตรวจรับ", "ผู้ตรวจรับ", "ตรวจรับพัสดุ", "ตรวจรับ"):
        idx = text.find(kw)
        if idx != -1:
            break
    if idx == -1:
        return []
    seg = text[idx: idx + 700]
    out, seen = [], set()
    _w = r"[ก-๎]+"   # คำไทย (รวมสระนำหน้า เ แ โ ใ ไ)
    for m in re.finditer(r"(?:นาย|นาง|นางสาว|ว่าที่ร้อยตรี|ว่าที่)\s*" + _w + r"(?:\s+" + _w + r")?", seg):
        nm = re.sub(r"\s+", " ", m.group(0)).strip()
        if nm in seen:
            continue
        seen.add(nm)
        out.append({"name": nm, "position": "ครู", "role": "กรรมการ"})
        if len(out) >= 5:
            break
    return out


def extract_procurement_fields(path: str) -> dict:
    """อ่านฟิลด์เรื่องจัดซื้อ/จัดจ้างจากไฟล์ (heuristic ออฟไลน์) — ดึงเพิ่มจากฟิลด์หนังสือทั่วไป
    รองรับรูปแบบ 'รายงานขอซื้อ/จ้าง' มาตรฐาน คืน dict (ฟิลด์ที่อ่านไม่ได้เป็นค่าว่าง)
    หมายเหตุ: รายการพัสดุ/กรรมการ เป็น best-effort -> ตรวจ/แก้ก่อนบันทึกเสมอ"""
    text = extract_text_any(path)
    base = extract_letter_fields(path)   # ใช้ letter_no/date/subject ที่ได้มาแล้ว
    # เลขที่บันทึก: ดึงเฉพาะรูปแบบ N/ปปปป (กันคว้า "วันที่..." ที่อยู่บรรทัดเดียวกัน)
    memo = _first(r"(?:^|\n)\s*ที่\s+([0-9๐-๙]+\s*/\s*[0-9๐-๙]{3,4})", text) or base.get("letter_no", "")
    r = {
        "proc_type": "", "subject": "", "memo_no": _arabic(re.sub(r"\s+", "", memo)),
        "request_date": base.get("letter_date"), "department": "", "project_name": "",
        "purpose": "", "budget_source": "", "delivery_days": "", "vendor_name": "",
        "items": [], "inspectors": [], "raw_text": text, "ok": bool(text.strip()),
    }
    if not text.strip():
        return r

    # ประเภท ซื้อ/จ้าง จากคำว่า "ขอซื้อ"/"ขอจ้าง"
    mt = re.search(r"ขอ(ซื้อ|จ้าง)", text)
    r["proc_type"] = mt.group(1) if mt else ""

    # เรื่อง: ตัด "รายงานขอซื้อ/จ้าง" นำหน้าออก เหลือชื่อรายการจริง
    subj = base.get("subject", "")
    subj = re.sub(r"^\s*รายงานขอ(ซื้อ|จ้าง)\s*", "", subj).strip()
    r["subject"] = subj

    # ฝ่าย/งาน: "ด้วย <ฝ่าย> มีความประสงค์"
    r["department"] = _first(r"ด้วย\s+(.+?)\s*มีความประสงค์", text)
    # โครงการ: "ตามโครงการ <X> จำนวน" หรือ "โครงการ<X>"
    r["project_name"] = _first(r"ตามโครงการ\s*(.+?)\s*(?:จำนวน|เป็นเงิน|$)", text)
    # เหตุผลความจำเป็น: "...คือ <X>"
    r["purpose"] = _first(r"เหตุผลและความจำเป็น[^\n]*?คือ\s+(.+)", text)
    # แหล่งงบ: "จากเงิน<X> "
    r["budget_source"] = _first(r"จากเงิน\s*([^\s]+)", text)
    # กำหนดส่งมอบ: "ภายใน <N> วัน"
    dd = _first(r"ภายใน\s+([0-9๐-๙]+)\s*วัน", text)
    r["delivery_days"] = _arabic(dd)
    # ผู้ขาย: "กับ <X> ซึ่งมีอาชีพ" (จากรายงานผลพิจารณา) เผื่อมี
    r["vendor_name"] = _first(r"(?:ตกลงราคากับ|จาก)\s+(.+?)\s*(?:ซึ่งมีอาชีพ|เป็นผู้)", text)

    # รายการพัสดุ (จากตาราง) + กรรมการตรวจรับ (จากข้อความ) — best-effort
    try:
        r["items"] = _parse_items_from_tables(_extract_tables_any(path))
    except Exception:
        r["items"] = []
    try:
        r["inspectors"] = _parse_inspectors(text)
    except Exception:
        r["inspectors"] = []

    return r
