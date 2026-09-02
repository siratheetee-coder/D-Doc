# -*- coding: utf-8 -*-
"""
receipt_voucher.py - ใบสำคัญรับเงิน (ฟอร์มกลาง ใช้ได้ทั้งงานพัสดุและอาหารกลางวัน)
อ้างรูปแบบจากแม่แบบราชการ: หัวเรื่อง + ข้อมูลผู้รับเงิน (ที่อยู่) + ตารางรายการ/จำนวนเงิน
+ จำนวนเงินตัวอักษร + ช่องลงชื่อผู้รับเงิน/ผู้จ่ายเงิน
"""
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.services.doc_page import set_a4


def _fixed_grid(tbl, widths_cm):
    """บังคับตารางเป็น fixed layout + กำหนดความกว้างคอลัมน์เป๊ะ (กันตารางล้นขอบ/autofit เพี้ยน)"""
    tblPr = tbl._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout"); tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa"); tblW.set(qn("w:w"), str(int(sum(widths_cm) * 567)))
    old = tbl._tbl.find(qn("w:tblGrid"))
    if old is not None:
        tbl._tbl.remove(old)
    grid = OxmlElement("w:tblGrid")
    for w in widths_cm:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(int(w * 567))); grid.append(gc)
    tbl._tbl.insert(list(tbl._tbl).index(tblPr) + 1, grid)
from app.services.build_templates import (
    _font, _p, _set_cell, _sign_table, THAI_FONT, _csize, _bcs,
)
from app.database import get_data_dir
from app.thai_utils import bahttext, thai_date

_BLANK = "................................"
_DOT = "..........................................."


def _money(v) -> str:
    try:
        return f"{float(v):,.2f}"
    except Exception:
        return "0.00"


def _split_satang(v):
    """แยกบาท/สตางค์ (คืนสตริง) เช่น 1234.50 -> ('1,234', '50')"""
    try:
        total = round(float(v or 0), 2)
    except Exception:
        total = 0.0
    baht = int(total)
    satang = int(round((total - baht) * 100))
    return f"{baht:,}", f"{satang:02d}"


def _safe(name: str) -> str:
    for ch in '\\/:*?"<>|\n\r\t':
        name = name.replace(ch, " ")
    return " ".join(name.split()).strip() or "ใบสำคัญรับเงิน"


def render_receipt_voucher(school, *, payee="", payee_address="", items=None, total=None,
                           date=None, payer="", received_from="", subject="", doc=None) -> str:
    """ออกใบสำคัญรับเงิน 1 ฉบับ
      payee          = ผู้รับเงิน (ข้าพเจ้า)
      payee_address  = ที่อยู่ผู้รับเงิน (พิมพ์ต่อท้าย ถ้ามี)
      items          = [(รายการ, จำนวนเงิน float), ...]
      total          = ยอดรวม (ไม่ส่งมา = ผลรวม items)
      date           = วันที่รับเงิน (datetime)
      payer          = ผู้จ่ายเงิน (เจ้าหน้าที่การเงิน/ผอ.)
      received_from  = ได้รับเงินจาก (ปกติ = ชื่อโรงเรียน)
      subject        = ใช้ตั้งชื่อไฟล์
    """
    own = doc is None
    if own:
        doc = Document(); set_a4(doc); _font(doc)
    elif doc.paragraphs or doc.tables:
        doc.add_page_break()

    sname = (getattr(school, "name", "") or "").strip()
    received_from = (received_from or sname or _BLANK)
    items = items or []
    if total is None:
        total = sum((a or 0) for _, a in items)
    total = round(float(total or 0), 2)

    _p(doc, "ใบสำคัญรับเงิน", align="center", bold=True, size=20, after=2)
    if sname:
        _p(doc, sname, align="center", size=15, after=8)

    # วันที่รับเงิน (ชิดขวา)
    _p(doc, f"วันที่ {thai_date(date) if date else _BLANK}", align="right", after=6)

    # ข้อมูลผู้รับเงิน
    who = (payee or "").strip() or _BLANK
    line = f"ข้าพเจ้า {who}"
    if (payee_address or "").strip():
        line += f"  ที่อยู่ {payee_address.strip()}"
    _p(doc, line, after=0, indent=0.5)
    _p(doc, f"ได้รับเงินจาก {received_from} ดังรายการต่อไปนี้", after=6, indent=0.5)

    # ตารางรายการ / จำนวนเงิน (หัว 2 แถว: รายการ | จำนวนเงิน[บาท|สตางค์])
    from docx.enum.table import WD_ALIGN_VERTICAL
    tbl = doc.add_table(rows=2, cols=3)
    tbl.style = "Table Grid"
    tbl.autofit = False
    _wcm = [10.0, 3.5, 1.5]                  # รวม 15.0 <= พื้นที่พิมพ์ 16.5 (ขอบ 3/1.5) กันล้น
    _fixed_grid(tbl, _wcm)                   # fixed layout กันตารางล้นขอบ
    widths = [Cm(w) for w in _wcm]
    h0 = tbl.rows[0].cells
    h1 = tbl.rows[1].cells
    cell_rai = h0[0].merge(h1[0])              # "รายการ" รวม 2 แถวแนวตั้ง
    _set_cell(cell_rai, "รายการ", bold=True, align="center", size=15)
    cell_rai.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_amt = h0[1].merge(h0[2])              # "จำนวนเงิน" รวม 2 คอลัมน์
    _set_cell(cell_amt, "จำนวนเงิน", bold=True, align="center", size=15)
    _set_cell(h1[1], "บาท", bold=True, align="center", size=14)
    _set_cell(h1[2], "ส.ต.", bold=True, align="center", size=14)
    cell_rai.width = widths[0]
    h1[1].width = widths[1]
    h1[2].width = widths[2]
    body = list(items)
    while len(body) < 5:              # เว้นบรรทัดว่างให้กรอกมือได้
        body.append(("", None))
    for name, amt in body:
        row = tbl.add_row().cells
        baht, sat = _split_satang(amt) if amt not in (None, "") else ("", "")
        _set_cell(row[0], name or "", align="left", size=14)
        _set_cell(row[1], baht, align="right", size=14)
        _set_cell(row[2], sat, align="right", size=14)
        for c, w in zip(row, widths):
            c.width = w
    # แถวรวม
    tb, ts = _split_satang(total)
    tot = tbl.add_row().cells
    _set_cell(tot[0], "รวมทั้งสิ้น", bold=True, align="right", size=14)
    _set_cell(tot[1], tb, bold=True, align="right", size=14)
    _set_cell(tot[2], ts, bold=True, align="right", size=14)
    for c, w in zip(tot, widths):
        c.width = w

    _p(doc, f"จำนวนเงิน (ตัวอักษร)  {bahttext(total)}", bold=True, before=6, after=16, indent=0.5)

    # ช่องลงชื่อ ผู้รับเงิน / ผู้จ่ายเงิน (อยู่ด้านล่าง)
    payer_disp = (payer or "").strip()
    _sign_table(doc, [
        [(f"ลงชื่อ {_DOT} ผู้รับเงิน", "center"),
         (f"( {who} )", "center")],
    ])
    _p(doc, "", after=10)
    _sign_table(doc, [
        [(f"ลงชื่อ {_DOT} ผู้จ่ายเงิน", "center"),
         (f"( {payer_disp or _BLANK} )", "center")],
    ])

    if not own:
        return doc
    out_dir = get_data_dir() / "documents"
    out_dir.mkdir(exist_ok=True)
    fname = _safe(f"ใบสำคัญรับเงิน {subject}".strip())
    out_path = out_dir / (fname + ".docx")
    doc.save(str(out_path))
    return str(out_path)
