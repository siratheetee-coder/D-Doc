# -*- coding: utf-8 -*-
"""
nutrition_report.py - รายงานภาวะโภชนาการนักเรียน (Word)
สรุปน้ำหนักตามเกณฑ์ส่วนสูง (ผอม/สมส่วน/อ้วน ฯลฯ) แยกตามระดับชั้นและเพศ
+ กราฟแท่งแบบข้อความสี (ไม่พึ่ง matplotlib) · ใช้ผลการชั่งเทอมล่าสุดของแต่ละคน
"""
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from app.services.doc_page import set_a4

from app.database import get_data_dir
from app.thai_utils import thai_date

THAI_FONT = "TH Sarabun New"
# สีประจำภาวะ (ใช้กับกราฟแท่ง)
WH_COLOR = {"ผอม": "C0392B", "ค่อนข้างผอม": "E67E22", "สมส่วน": "27AE60",
            "ท้วม": "2980B9", "เริ่มอ้วน": "E67E22", "อ้วน": "C0392B"}


def _safe(text: str) -> str:
    for ch in '<>:"/\\|?*':
        text = text.replace(ch, "_")
    return text.strip()


def _p(doc, text="", *, align="left", bold=False, size=14, after=2, color=None):
    p = doc.add_paragraph()
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = THAI_FONT
    r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def _set_cell(cell, text, *, bold=False, align="center", size=13, fill=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = THAI_FONT
    r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if fill is not None:
        tcpr = cell._tc.get_or_add_tcPr()
        shd = tcpr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): fill})
        tcpr.append(shd)


def _pct(n, total):
    return (100.0 * n / total) if total else 0.0


def render_nutrition_report(school, cats, class_counts, sex_counts, totals, assessed, as_of=None) -> str:
    """cats: รายชื่อภาวะ (WH_LABELS) · class_counts: [{level, counts{cat:n}, total}]
    sex_counts: {'ชาย':{cat:n}, 'หญิง':{cat:n}} · totals: {cat:n} · assessed: int"""
    doc = Document(); set_a4(doc)
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.left_margin = sec.right_margin = Cm(1.5)
    sec.top_margin = sec.bottom_margin = Cm(1.5)
    base = doc.styles["Normal"]; base.font.name = THAI_FONT; base.font.size = Pt(13)
    base._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)

    _p(doc, "รายงานภาวะโภชนาการนักเรียน", align="center", bold=True, size=18, after=0)
    _p(doc, (school.name or ""), align="center", bold=True, size=15, after=0)
    _p(doc, f"ประเมินแล้ว {assessed} คน · ณ {thai_date(as_of) if as_of else ''}", align="center", size=13, after=8)

    # ---- ตารางสรุปรวม (ภาวะ × เพศ) ----
    _p(doc, "สรุปน้ำหนักตามเกณฑ์ส่วนสูง (แยกตามเพศ)", bold=True, size=14, after=2)
    head = ["ภาวะโภชนาการ", "ชาย", "หญิง", "รวม", "ร้อยละ"]
    t = doc.add_table(rows=1, cols=len(head)); t.style = "Table Grid"
    for c, h in enumerate(head):
        _set_cell(t.rows[0].cells[c], h, bold=True, fill="DCFCE7")
    for cat in cats:
        m = sex_counts.get("ชาย", {}).get(cat, 0)
        f = sex_counts.get("หญิง", {}).get(cat, 0)
        tot = totals.get(cat, 0)
        cells = t.add_row().cells
        _set_cell(cells[0], cat, align="left", color=WH_COLOR.get(cat))
        _set_cell(cells[1], str(m)); _set_cell(cells[2], str(f)); _set_cell(cells[3], str(tot), bold=True)
        _set_cell(cells[4], f"{_pct(tot, assessed):.1f}%")
    tr = t.add_row().cells
    _set_cell(tr[0], "รวม", bold=True, align="left", fill="F1F5F9")
    _set_cell(tr[1], str(sum(sex_counts.get("ชาย", {}).values())), bold=True, fill="F1F5F9")
    _set_cell(tr[2], str(sum(sex_counts.get("หญิง", {}).values())), bold=True, fill="F1F5F9")
    _set_cell(tr[3], str(assessed), bold=True, fill="F1F5F9")
    _set_cell(tr[4], "100.0%" if assessed else "-", bold=True, fill="F1F5F9")

    # ---- กราฟแท่ง (ข้อความสี) ----
    _p(doc, "", after=2)
    _p(doc, "กราฟภาวะโภชนาการ (ร้อยละ)", bold=True, size=14, after=2)
    for cat in cats:
        tot = totals.get(cat, 0)
        pct = _pct(tot, assessed)
        blocks = "█" * int(round(pct / 2.5))    # 1 บล็อก ≈ 2.5% (เต็ม 100% = 40 บล็อก)
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
        r0 = p.add_run(f"{cat:<14} "); r0.font.name = THAI_FONT; r0.font.size = Pt(13)
        r0._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
        rb = p.add_run(blocks or "▏"); rb.font.name = THAI_FONT; rb.font.size = Pt(13)
        rb.font.color.rgb = RGBColor.from_string(WH_COLOR.get(cat, "888888"))
        rt = p.add_run(f"  {tot} คน ({pct:.1f}%)"); rt.font.name = THAI_FONT; rt.font.size = Pt(13)
        rt._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)

    # ---- ตารางแยกระดับชั้น ----
    _p(doc, "", after=4)
    _p(doc, "แยกตามระดับชั้น", bold=True, size=14, after=2)
    head2 = ["ระดับชั้น"] + cats + ["ประเมิน"]
    widths = [Cm(2.4)] + [Cm(2.5)] * len(cats) + [Cm(2.2)]
    t2 = doc.add_table(rows=1, cols=len(head2)); t2.style = "Table Grid"
    for c, (h, w) in enumerate(zip(head2, widths)):
        _set_cell(t2.rows[0].cells[c], h, bold=True, fill="DCFCE7", size=12)
        t2.rows[0].cells[c].width = w
    for cr in class_counts:
        cells = t2.add_row().cells
        _set_cell(cells[0], cr["level"], align="left")
        for i, cat in enumerate(cats):
            _set_cell(cells[1 + i], str(cr["counts"].get(cat, 0) or ""))
        _set_cell(cells[1 + len(cats)], str(cr["total"]), bold=True)
        for c, w in enumerate(widths):
            cells[c].width = w
    trc = t2.add_row().cells
    _set_cell(trc[0], "รวม", bold=True, align="left", fill="F1F5F9")
    for i, cat in enumerate(cats):
        _set_cell(trc[1 + i], str(totals.get(cat, 0) or ""), bold=True, fill="F1F5F9")
    _set_cell(trc[1 + len(cats)], str(assessed), bold=True, fill="F1F5F9")
    for c, w in enumerate(widths):
        trc[c].width = w

    _p(doc, "", after=6)
    _p(doc, "เกณฑ์อ้างอิง: กราฟการเจริญเติบโตของกรมอนามัย กระทรวงสาธารณสุข (น้ำหนักตามเกณฑ์ส่วนสูง)",
       size=12, after=10, color="666666")

    officer = (getattr(school, "officer_name", "") or "").strip()
    _p(doc, "ลงชื่อ.............................................ผู้จัดทำรายงาน", align="center", after=0)
    _p(doc, f"( {officer} )", align="center", after=0)
    _p(doc, "เจ้าหน้าที่โครงการอาหารกลางวัน", align="center", after=8)
    _p(doc, "ลงชื่อ.............................................", align="center", after=0)
    _p(doc, f"( {(getattr(school, 'director_name', '') or '').strip()} )", align="center", after=0)
    dpos = ("ผู้อำนวยการ" + school.name) if (school.name or "").startswith("โรงเรียน") else "ผู้อำนวยการโรงเรียน"
    _p(doc, dpos, align="center", after=2)

    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    out = out_dir / (_safe("รายงานภาวะโภชนาการ") + ".docx")
    doc.save(str(out))
    return str(out)
