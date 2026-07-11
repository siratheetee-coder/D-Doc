"""
build_templates.py
------------------
สร้างไฟล์ "แม่แบบ" .docx (มี placeholder ของ docxtpl) ลงในโฟลเดอร์ app/doc_templates/
ยึดข้อความ/กฎหมายตามฟอร์มจริงของโรงเรียน (โฟลเดอร์ จัดซื้อ/)

รันครั้งเดียวเพื่อสร้าง/อัปเดตแม่แบบ:
    .venv\\Scripts\\python.exe -m app.services.build_templates

ผู้ใช้ปลายทางสามารถเปิดไฟล์แม่แบบใน Word มาแก้หน้าตา/ถ้อยคำได้เอง
(ห้ามลบ {{ ... }} และ {% ... %} ที่เป็นช่องเติมข้อมูล)
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

THAI_FONT = "TH Sarabun New"
TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "doc_templates"
ASSETS_DIR = TEMPLATES_DIR / "assets"
# ไฟล์รูปตราครุฑ (วางไฟล์นี้เองได้ ระบบจะฝังให้อัตโนมัติ) รองรับ .png/.jpg
KRUT_CANDIDATES = ["krut.png", "krut.jpg", "ครุฑ.png", "ครุฑ.jpg"]


def _krut_path():
    """หาไฟล์ตราครุฑในโฟลเดอร์ assets/ (คืน None ถ้าไม่มี)"""
    for name in KRUT_CANDIDATES:
        p = ASSETS_DIR / name
        if p.exists():
            return p
    return None


def _krut_and_title(doc):
    """
    หัวบันทึกข้อความตามระเบียบงานสารบรรณ:
    ตราครุฑ 1.5 ซม. ชิดซ้าย และคำว่า "บันทึกข้อความ" (29pt) อยู่กึ่งกลางระดับเดียวกัน
    ใช้ tab แบบ center ที่กึ่งกลางหน้ากระดาษ
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    # tab กึ่งกลาง ~8 ซม. (กึ่งกลางพื้นที่พิมพ์บน A4 ขอบ 2.5 ซม.)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(8), WD_TAB_ALIGNMENT.CENTER)
    krut = _krut_path()
    if krut:
        p.add_run().add_picture(str(krut), height=Cm(1.5))
    else:
        rk = p.add_run("[ครุฑ: วางไฟล์ doc_templates/assets/krut.png]")
        rk.font.size = Pt(9)
        rk.font.name = THAI_FONT
    p.add_run("\t")
    rt = p.add_run("บันทึกข้อความ")
    _csize(rt, 29)
    _bcs(rt, True)
    rt.font.name = THAI_FONT
    rt._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)


def _p_runs(doc, segments, *, indent=None, align=None, tab_cm=None, size=16, after=2):
    """
    ย่อหน้าที่มีหลายช่วงข้อความ (กำหนดตัวหนาแยกได้)
    segments = [(text, bold), ...]  ใช้ "\\t" เพื่อเว้นไป tab stop
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if indent is not None:
        p.paragraph_format.first_line_indent = Cm(indent)
    if tab_cm is not None:
        p.paragraph_format.tab_stops.add_tab_stop(Cm(tab_cm), WD_TAB_ALIGNMENT.LEFT)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, bold in segments:
        r = p.add_run(text)
        _csize(r, size)
        _bcs(r, bold)
        r.font.name = THAI_FONT
        r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    return p


def _hr(doc):
    """เส้นขีดคั่นแนวนอนเต็มความกว้าง (ใช้คั่นใต้ 'เรื่อง')"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    pPr.append(pbdr)

LEGAL_INTRO = (
    "งานพัสดุได้ตรวจสอบแล้วเห็นควรจัด{{ proc_type }}ตามเสนอ และเพื่อให้เป็นไปตาม"
    "พระราชบัญญัติการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560 ข้อ 56 วรรคหนึ่ง (2) (ข) "
    "และระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อ/จัดจ้างและบริหารพัสดุภาครัฐ พ.ศ. 2560 "
    "ข้อ 22 ข้อ 79 ข้อ 25 (5) และกฎกระทรวงวงเงินการจัดซื้อจัดจ้างพัสดุโดยวิธีเจาะจง "
    "วงเงินการจัดซื้อจัดจ้างที่ไม่ทำข้อตกลงเป็นหนังสือ และวงเงินการจัดซื้อจัดจ้างในการ"
    "แต่งตั้งผู้ตรวจรับพัสดุ พ.ศ. 2560 ข้อ 1 และ ข้อ 5 จึงรายงานขอ{{ proc_type }} ดังนี้"
)


def _margins(doc):
    """ตั้งระยะขอบกระดาษตามมาตรฐานหนังสือราชการไทย (ซ้ายกว้างสำหรับเย็บเล่ม)
    ซ้าย 3.0 / ขวา 1.75 / บน 1.5 / ล่าง 0.5 ซม. (อิงแบบฟอร์มที่ใช้กันทั่วไป)"""
    for s in doc.sections:
        s.page_width = Cm(21.0)     # A4 (ไม่ใช่ Letter ที่เป็นค่าปริยาย)
        s.page_height = Cm(29.7)
        s.left_margin = Cm(3.0)
        s.right_margin = Cm(1.75)
        s.top_margin = Cm(1.5)
        s.bottom_margin = Cm(0.5)


def _font(doc):
    _margins(doc)
    style = doc.styles["Normal"]
    style.font.name = THAI_FONT
    style.font.size = Pt(16)   # มาตรฐานงานสารบรรณ: 16 pt
    # ระยะห่างบรรทัด 1.0 (single) และไม่เว้นช่องว่างท้ายย่อหน้าโดยปริยาย
    pf = style.paragraph_format
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for k in ("w:cs", "w:ascii", "w:hAnsi", "w:eastAsia"):
        rfonts.set(qn(k), THAI_FONT)
    # กำกับภาษาไทย (th-TH) เพื่อให้ Word ตัดคำไทยด้วยพจนานุกรม
    # -> บรรทัดเติมเต็ม ไม่เกิดการยืดช่องว่าง/อักษรเวลาจัดชิดขอบ
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "th-TH")
    lang.set(qn("w:bidi"), "th-TH")
    lang.set(qn("w:eastAsia"), "th-TH")
    rpr.append(lang)
    # เปิดการตัดคำแบบ complex script (ภาษาไทย)
    style.font.complex_script = True
    # สำคัญ: ภาษาไทยใช้ขนาดจาก szCs (complex script) ต้องตั้งคู่กับ sz เสมอ
    # ไม่งั้น Word จะ fallback เป็น 11pt
    _set_szcs(rpr, 16)


def _set_szcs(rpr, pt):
    """ตั้งค่า w:szCs (ขนาดฟอนต์ complex script) เป็น pt — แก้ปัญหาไทยกลายเป็น 11pt"""
    szcs = rpr.find(qn("w:szCs"))
    if szcs is None:
        szcs = OxmlElement("w:szCs")
        rpr.append(szcs)
    szcs.set(qn("w:val"), str(int(pt * 2)))   # หน่วยเป็นครึ่งพอยต์


def _csize(run, pt):
    """ตั้งขนาดฟอนต์ของ run ทั้ง sz (ปกติ) และ szCs (ภาษาไทย/complex) ให้ตรงกัน"""
    run.font.size = Pt(pt)
    _set_szcs(run._element.get_or_add_rPr(), pt)


def _bcs(run, bold):
    """ตั้งตัวหนาให้ครบทั้ง w:b (ปกติ) และ w:bCs (ภาษาไทย/complex script)
    สำคัญ: ภาษาไทยจะหนาจริงต่อเมื่อมี w:bCs ด้วย (เหมือน szCs ของขนาดฟอนต์)"""
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    el = rpr.find(qn("w:bCs"))
    if el is None:
        el = OxmlElement("w:bCs")
        rpr.append(el)
    el.set(qn("w:val"), "true" if bold else "false")


def _p(doc, text="", *, align=None, bold=False, size=16, indent=None, before=0, after=2):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent is not None:
        pf.first_line_indent = Cm(indent)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        # ใช้ "จัดกระจายแบบไทย" (thaiDistribute) ไม่ใช่ justify แบบอังกฤษ (both)
        # เพื่อให้กระจายระยะห่างทีละอักษรสม่ำเสมอ ไม่เกิดช่องว่างยักษ์
        p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    run = p.add_run(text)
    _csize(run, size)
    _bcs(run, bold)
    run.font.name = THAI_FONT
    run._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    return p


def _set_cell(cell, text, *, bold=False, align="left", size=14):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    r = p.add_run(text)
    r.font.name = THAI_FONT
    _csize(r, size)
    _bcs(r, bold)
    r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)


def _repeat_header_row(row):
    """ตั้งให้แถวนี้เป็น 'หัวตาราง' ที่พิมพ์ซ้ำทุกหน้าเมื่อตารางยาวข้ามหน้า
    (แก้ปัญหารายการเยอะแล้วหัวตารางหายในหน้าถัดไป)"""
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def _no_split_row(row):
    """ห้ามแถวนี้ถูกตัดแบ่งครึ่งข้ามหน้า (เนื้อหาในแถวอยู่หน้าเดียวกันเสมอ)"""
    trPr = row._tr.get_or_add_trPr()
    cs = OxmlElement("w:cantSplit")
    cs.set(qn("w:val"), "true")
    trPr.append(cs)


def _no_borders(table):
    """ลบเส้นขอบตาราง (ใช้กับตารางช่องลงนามให้ดูเป็นบล็อกสะอาด ไม่มีกรอบ)"""
    tblPr = table._element.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def _sign_table(doc, columns, *, after=6, gap=True):
    """ช่องลงนามแบบจัดคอลัมน์ด้วยตารางไร้เส้นขอบ (จัดบรรทัดตรงกันเป๊ะ)
    columns = [ [ (text, align), ... ], ... ] แต่ละคอลัมน์คือบล็อกลงนาม 1 ช่อง
    gap=False : ไม่เพิ่มบรรทัดว่างท้ายตาราง (ประหยัดพื้นที่ ให้เนื้อหาอยู่หน้าเดียว)
    """
    n = len(columns)
    table = doc.add_table(rows=1, cols=n)
    _no_borders(table)
    amap = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT}
    for cell, lines in zip(table.rows[0].cells, columns):
        first = True
        for text, align in lines:
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            # บรรทัดรูปลายเซ็น: ("__SIG__", "<path>") -> วางรูปกึ่งกลาง (PNG โปร่งใส ไม่ทับข้อความ)
            if text == "__SIG__" and align:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                try:
                    p.add_run().add_picture(align, height=Cm(1.35))
                except Exception:
                    pass
                continue
            p.alignment = amap[align]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            _csize(r, 16)
            r.font.name = THAI_FONT
            r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    if gap:
        doc.add_paragraph().paragraph_format.space_after = Pt(after)
    return table


def _cell_line(cell, segments, *, first=False, after=0):
    """เพิ่มบรรทัดในเซลล์ตาราง โดยกำหนดตัวหนาแยกช่วงได้
    segments = [(text, bold), ...]"""
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    for text, bold in segments:
        r = p.add_run(text)
        _csize(r, 16)
        _bcs(r, bold)
        r.font.name = THAI_FONT
        r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    return p


def _order_header_table(doc):
    """หัวใบสั่งซื้อ/จ้าง แบบ 2 คอลัมน์ (ไร้เส้น):
    ซ้าย = ข้อมูลผู้ขาย/ผู้รับจ้าง | ขวา = เลขที่/วันที่ + ข้อมูลส่วนราชการ
    จัดแนวด้วยตารางเพื่อให้ตรงกันแม้ที่อยู่ยาวหลายบรรทัด
    """
    table = doc.add_table(rows=1, cols=2)
    _no_borders(table)
    table.autofit = False
    table.allow_autofit = False
    left, right = table.rows[0].cells
    left.width = Cm(8.6)
    right.width = Cm(7.65)

    # ===== ซ้าย: ผู้ขาย/ผู้รับจ้าง =====
    _cell_line(left, [("{{ vendor_label }} : ", True), ("{{ vendor_name }}", False)], first=True)
    _cell_line(left, [("ที่อยู่ : ", True), ("{{ vendor_address }}", False)])
    _cell_line(left, [("โทรศัพท์ : ", True), ("{{ vendor_phone }}", False)])
    _cell_line(left, [("เลขประจำตัวผู้เสียภาษี : ", True), ("{{ vendor_tax_id }}", False)])

    # ===== ขวา: เลขที่/วันที่ + ส่วนราชการ (โรงเรียน) =====
    _cell_line(right, [("{{ order_kind }} เลขที่ : ", True), ("{{ order_no }}", False)], first=True)
    _cell_line(right, [("วันที่ : ", True), ("{{ order_date_thai }}", False)])
    _cell_line(right, [("ส่วนราชการ : ", True), ("{{ school_name }}", False)])
    _cell_line(right, [("ที่อยู่ : ", True), ("{{ school_address }}", False)])
    return table


def _attachment_table(doc):
    """ตารางรายละเอียดแนบท้าย 7 คอลัมน์ (หัวตาราง 2 ชั้น) ตามแบบฟอร์มมาตรฐาน
    ที่ | รายการ | จำนวน | หน่วย | ราคาที่สืบได้ (บาท) | [จำนวนและวงเงินที่ขอครั้งนี้: หน่วยละ | บาท]
    """
    table = doc.add_table(rows=2, cols=7)
    table.style = "Table Grid"
    r0 = table.rows[0].cells
    r1 = table.rows[1].cells
    # คอลัมน์ 0-4: ผสานสองแถวแนวตั้ง
    for i, lab in enumerate(["ที่", "รายการ", "จำนวน", "หน่วย", "ราคาที่สืบได้ (บาท)"]):
        c = r0[i].merge(r1[i])
        _set_cell(c, lab, bold=True, align="center")
    # คอลัมน์ 5-6: หัวรวม + หัวย่อย
    top = r0[5].merge(r0[6])
    _set_cell(top, "จำนวนและวงเงินที่ขอครั้งนี้", bold=True, align="center")
    _set_cell(r1[5], "หน่วยละ", bold=True, align="center")
    _set_cell(r1[6], "บาท", bold=True, align="center")
    _repeat_header_row(table.rows[0])      # หัว 2 ชั้นซ้ำทุกหน้า
    _repeat_header_row(table.rows[1])
    # loop ข้อมูล (docxtpl 3 แถว)
    _set_cell(table.add_row().cells[0], "{%tr for it in items %}")
    drow = table.add_row()
    d = drow.cells
    _no_split_row(drow)
    _set_cell(d[0], "{{ loop.index }}", align="center")
    _set_cell(d[1], "{{ it.name }}")
    _set_cell(d[2], "{{ it.qty }}", align="center")
    _set_cell(d[3], "{{ it.unit }}", align="center")
    _set_cell(d[4], "{{ it.unit_price }}", align="right")
    _set_cell(d[5], "{{ it.unit_price }}", align="right")
    _set_cell(d[6], "{{ it.amount }}", align="right")
    _set_cell(table.add_row().cells[0], "{%tr endfor %}")
    _summary_rows(table, label_cols=[0, 1, 2, 3, 4, 5], value_col=6, total_cols=7)
    # ความกว้างคงที่: ช่อง "รายการ" กว้างสุด (รวม ~16.25 ซม.)
    _fixed_cols(table, [Cm(1.0), Cm(5.6), Cm(1.4), Cm(1.4), Cm(2.2), Cm(2.1), Cm(2.55)])
    return table


def _table_indent(table, cm):
    """กำหนดระยะเยื้องซ้ายของทั้งตาราง (ให้ตารางกรรมการเยื้องเข้าเหมือนเนื้อหา)"""
    tblPr = table._element.tblPr
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), str(int(Cm(cm).twips)))
    ind.set(qn("w:type"), "dxa")
    tblPr.append(ind)


def _fixed_cols(table, widths):
    """กำหนดความกว้างคอลัมน์แบบคงที่ (ไม่ให้ Word ยืด-หดเอง) ให้ช่อง 'รายการ' กว้างสุด
    widths = list ของ Cm() เรียงตามคอลัมน์ รวมไม่ควรเกิน ~16.25 ซม. (พื้นที่พิมพ์ A4)"""
    table.autofit = False
    table.allow_autofit = False
    tbl = table._element
    tblPr = tbl.tblPr
    # layout แบบคงที่ (ใช้ความกว้างจาก tblGrid ไม่คำนวณจากเนื้อหา)
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    # ความกว้างรวมของตาราง
    total = sum(int(w.twips) for w in widths)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    # grid คอลัมน์ (ตัวกำหนดความกว้างจริงภายใต้ fixed layout)
    for g in tbl.findall(qn("w:tblGrid")):
        tbl.remove(g)
    grid = OxmlElement("w:tblGrid")
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(w.twips)))
        grid.append(col)
    tblPr.addnext(grid)
    # เผื่อ Word บางรุ่นอ่านจากความกว้างเซลล์ -> ตั้งให้แถวที่จำนวนเซลล์ตรงด้วย
    for row in table.rows:
        cells = row.cells
        if len(cells) == len(widths):
            for c, w in zip(cells, widths):
                c.width = w


def _member_table(doc, list_name, *, prefix="", member_var="m", indent=0.4):
    """ตารางรายชื่อกรรมการแบบไร้เส้นขอบ จัดเป็น 4 คอลัมน์ให้ตรงกัน:
    ลำดับ | ชื่อ | ตำแหน่ง ... | บทบาท (ประธาน/กรรมการ)
    list_name = ชื่อตัวแปร list ในแม่แบบ (เช่น 'inspect_members' / 'spec_members')
    """
    # ความกว้างรวม + ระยะเยื้อง ต้องไม่เกินพื้นที่พิมพ์ A4 (~16.25 ซม.)
    # เลขลำดับแคบ (ให้ชิดชื่อ) · คอลัมน์ตำแหน่งกว้างพอให้ "ตำแหน่ง ครูชำนาญการพิเศษ" อยู่บรรทัดเดียว
    widths = [Cm(0.85), Cm(4.8), Cm(5.6), Cm(4.5)]   # รวม 15.75 ซม.
    table = doc.add_table(rows=1, cols=4)
    _no_borders(table)
    table.autofit = False
    table.allow_autofit = False
    if indent:
        _table_indent(table, indent)

    def setw(cells):
        for c, w in zip(cells, widths):
            c.width = w

    op = table.rows[0].cells
    _set_cell(op[0], "{%%tr for %s in %s %%}" % (member_var, list_name))
    setw(op)
    # เลขลำดับกรรมการเป็น 1. 2. 3. (มีจุดท้าย) — ไม่ใช้รูปแบบ 2.1 2.2
    num_text = prefix + "{{ loop.index }}" + ("" if prefix else ".")
    d = table.add_row().cells
    _set_cell(d[0], num_text, align="left", size=16)
    _set_cell(d[1], "{{ %s.name }}" % member_var, size=16)
    _set_cell(d[2], "ตำแหน่ง {{ %s.position }}" % member_var, size=16)
    _set_cell(d[3], "{{ %s.role }}" % member_var, size=16)
    setw(d)
    cl = table.add_row().cells
    _set_cell(cl[0], "{%tr endfor %}")
    setw(cl)
    return table


def build_purchase_request():
    """แม่แบบ: บันทึกข้อความ รายงานขอซื้อ/จ้าง (+ รายละเอียดแนบท้าย)"""
    doc = Document()
    _font(doc)

    # ===== ตราครุฑ (ซ้าย) + "บันทึกข้อความ" (กึ่งกลาง) =====
    _krut_and_title(doc)
    # หัวเรื่อง: label ตัวหนา ตามระเบียบงานสารบรรณ
    _p_runs(doc, [("ส่วนราชการ  ", True), ("{{ school_office }}", False)])
    _p_runs(doc, [("ที่  ", True), ("{{ memo_no }}", False),
                  ("\t", False), ("วันที่ ", True), ("{{ request_date }}", False)],
            tab_cm=7)
    _p_runs(doc, [("เรื่อง  ", True), ("รายงานขอ{{ proc_type }}{{ subject }}", False)])
    # มาตรฐานงานสารบรรณ: ไม่มีเส้นขีดทึบคั่นหัวกระดาษกับเนื้อหา
    _p_runs(doc, [("เรียน  ", True), ("{{ director_office }}", False)])

    # ===== เนื้อหา (บีบระยะห่าง after=0 เพื่อให้ทั้งฉบับอยู่หน้าเดียว) =====
    _p(doc,
       "ด้วย {{ department }} มีความประสงค์จะ{{ proc_type }}{{ subject }} จำนวน {{ item_count }} รายการ "
       "จากเงิน{{ budget_source }} ตามโครงการ{{ project_name }} จำนวน {{ total_amount }} บาท "
       "({{ total_baht }}) รายละเอียดดังแนบ",
       align="justify", indent=1.25, after=0)
    _p(doc, LEGAL_INTRO, align="justify", indent=1.25, after=0)

    _p(doc, "1.  เหตุผลและความจำเป็นที่จะต้อง{{ proc_type }} คือ  {{ purpose }}", align="justify", indent=1.25, after=0)
    _p(doc, "2.  รายละเอียดงานที่จะ{{ proc_type }}คือ ... (รายละเอียดตามบันทึกที่แนบ)", indent=1.25, after=0)
    _p(doc, "3.  ราคากลางของพัสดุที่จะ{{ proc_type }} เป็นเงิน  {{ total_amount }}  บาท", indent=1.25, after=0)
    _p(doc, "4.  วงเงินที่จะ{{ proc_type }}ครั้งนี้ จำนวน {{ total_amount }} บาท ({{ total_baht }})", indent=1.25, after=0)
    _p(doc, "5.  กำหนดส่งมอบพัสดุภายใน {{ delivery_days }} วัน (การนับวันนับถัดจากวันลงนาม)", indent=1.25, after=0)
    _p(doc,
       "6.  {{ proc_type }}โดยวิธี{{ method }} เนื่องจากการจัดซื้อจัดจ้างพัสดุที่มีการผลิต จำหน่าย "
       "ก่อสร้าง หรือให้บริการทั่วไป และมีวงเงินการจัดซื้อจัดจ้างครั้งหนึ่งไม่เกิน 500,000 บาท "
       "ที่กำหนดในกฎกระทรวง", align="justify", indent=1.25, after=0)
    _p(doc, "7.  หลักเกณฑ์พิจารณาการคัดเลือกข้อเสนอ โดยใช้เกณฑ์ราคา", indent=1.25, after=0)
    _p(doc, "8.  ข้อเสนออื่น ๆ เห็นควรแต่งตั้งผู้ตรวจรับพัสดุ ตามเสนอ", indent=1.25, after=0)
    _p(doc, "จึงเรียนมาเพื่อโปรดพิจารณา", indent=1.25, bold=True, after=0)

    # ===== ผู้ตรวจรับ: คนเดียว (โดยอนุโลม) หรือ คณะกรรมการ =====
    _p(doc, "{%p if inspection_mode == 'single' %}")
    _p(doc,
       "อนุมัติให้แต่งตั้ง {{ inspector_name }} ตำแหน่ง {{ inspector_position }} เป็นผู้ตรวจรับ"
       "การจัด{{ proc_type }} โดยใช้บันทึกฉบับนี้แทนคำสั่งแต่งตั้งคณะกรรมการตรวจรับ โดยอนุโลม",
       align="justify", indent=1.25)
    _p(doc, "{%p else %}")
    _p(doc, "เห็นชอบในรายงานขอ{{ proc_type }} ดังกล่าวข้างต้น และอนุมัติให้แต่งตั้งคณะกรรมการตรวจรับ ดังนี้",
       align="justify", indent=1.25)
    _member_table(doc, "inspect_members")
    _p(doc, "{%p endif %}")

    # ===== ลงนาม: เจ้าหน้าที่ / หัวหน้าเจ้าหน้าที่ (จัดคอลัมน์ด้วยตารางไร้เส้น) =====
    # ใช้ gap=False + ระยะ before เล็กน้อย เพื่อประหยัดพื้นที่ให้เนื้อหาอยู่หน้าเดียว
    _sign_table(doc, [
        [("ลงชื่อ.....................................เจ้าหน้าที่", "center"),
         ("( {{ officer_name }} )", "center")],
        [("ลงชื่อ.....................................หัวหน้าเจ้าหน้าที่", "center"),
         ("( {{ head_officer_name }} )", "center")],
    ], gap=False)
    # ===== เห็นชอบ/อนุมัติ โดยผู้อำนวยการ =====
    _p(doc, "(   )  เห็นชอบ        (   )  อนุมัติ", align="center", before=4)
    _p(doc, "ลงชื่อ.................................................", align="center")
    _p(doc, "( {{ director_name }} )", align="center")
    _p(doc, "{{ director_office }}", align="center")
    _p(doc, "วันที่ {{ request_date }}", align="center")

    # ===== รายละเอียดแนบท้าย (ขึ้นหน้าใหม่) =====
    doc.add_page_break()
    _p(doc, "รายละเอียดแนบท้ายบันทึกข้อความ ที่ {{ memo_no }} ลงวันที่ {{ request_date }}",
       align="center", bold=True)
    _p(doc, "สำหรับจัด{{ proc_type }}พัสดุ จำนวน {{ item_count }} รายการ", align="center", bold=True)
    _p(doc, "{{ school_name }}", align="center", after=4)

    # ตารางรายละเอียดแนบท้าย 7 คอลัมน์ (หัว 2 ชั้น) + แถวสรุป VAT/รวม/คำอ่าน
    _attachment_table(doc)

    _p(doc, "", after=4)
    _sign_table(doc, [
        [("ลงชื่อ.....................................เจ้าหน้าที่", "center"),
         ("( {{ officer_name }} )", "center")],
        [("ลงชื่อ.....................................หัวหน้าเจ้าหน้าที่", "center"),
         ("( {{ head_officer_name }} )", "center")],
    ])

    TEMPLATES_DIR.mkdir(exist_ok=True)
    out = TEMPLATES_DIR / "รายงานขอซื้อ.docx"
    doc.save(str(out))
    return out


def build_inspection():
    """
    แม่แบบ: ใบตรวจรับพัสดุ (รวมท่อน 'เสนอผลตรวจรับ + ขออนุมัติเบิกจ่าย' ตามฟอร์มจริง)
    รองรับผู้ตรวจรับคนเดียว และคณะกรรมการตรวจรับ
    หมายเหตุ: วันที่ส่งมอบ/ตรวจรับ/ใบส่งของ เป็นช่องจุดไข่ปลาให้กรอกตอนตรวจรับจริง
    """
    doc = Document()
    _font(doc)

    _p(doc, "ใบตรวจรับพัสดุ", align="center", bold=True, size=20, after=2)
    _p(doc, "ตามระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560 ข้อ 175",
       align="center", bold=True, after=6)
    _p(doc, "เขียนที่ {{ school_name }}", align="right")
    _p(doc, "วันที่ {{ inspect_date_thai }}", align="right", after=6)

    _p(doc,
       "ตามที่{{ school_name }} ได้จัด{{ proc_type }}{{ subject }} จาก {{ vendor_name }} "
       "ตาม{{ order_kind }} เลขที่ {{ order_no }} ลงวันที่ {{ order_date_thai }} "
       "ครบกำหนดส่งมอบวันที่ {{ delivery_due_thai }} บัดนี้ {{ vendor_label }}ได้ส่งมอบพัสดุ "
       "ตามใบส่งของ เล่มที่ {{ delivery_note_book }} เลขที่ {{ delivery_note_no }} ลงวันที่ {{ inspect_date_thai }} "
       "{{ committee_word }}ได้ตรวจรับพัสดุแล้ว ปรากฏว่าถูกต้องครบถ้วนตาม{{ order_kind }}ทุกประการ "
       "โดยส่งมอบเกินกำหนดจำนวน {{ overdue_days }} วัน คิดค่าปรับในอัตราร้อยละ {{ penalty_rate }} ต่อวัน เป็นเงินทั้งสิ้น {{ fine_amount }} บาท "
       "จึงออกหนังสือสำคัญฉบับนี้ให้ไว้ {{ vendor_label }}ควรได้รับเงินเป็นจำนวนเงินทั้งสิ้น "
       "{{ total_amount }} บาท ({{ total_baht }}) ตาม{{ order_kind }}",
       align="justify", indent=1.25)
    _p(doc,
       "จึงขอเสนอรายงานต่อ{{ director_office }} เพื่อโปรดทราบ ตามนัยข้อ 175 (4) "
       "แห่งระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560",
       align="justify", indent=1.25, after=6)

    # ===== ลงนามผู้ตรวจรับ (คนเดียว / คณะกรรมการ) =====
    _p(doc, "{%p if inspection_mode == 'committee' %}")
    _p(doc, "{%p for m in inspect_members %}")
    _p(doc, "(ลงชื่อ) ........................................... {{ m.role }}", align="center")
    _p(doc, "( {{ m.name }} )", align="center")
    _p(doc, "{%p endfor %}")
    _p(doc, "{%p else %}")
    _p(doc, "(ลงชื่อ) ........................................... ผู้ตรวจรับพัสดุ", align="center")
    _p(doc, "( {{ inspector_name }} )", align="center")
    _p(doc, "{%p endif %}")

    # ===== เสนอผลการตรวจรับ + ขออนุมัติเบิกจ่าย =====
    _p(doc, "", after=6)
    _p_runs(doc, [("เรียน  ", True), ("{{ director_office }}", True)])
    _p(doc,
       "เพื่อโปรดทราบผลการตรวจรับพัสดุ ค่าจัด{{ proc_type }}{{ subject }} {{ committee_word }}"
       "ได้ดำเนินการตรวจรับพัสดุดังกล่าวเรียบร้อยแล้ว รายละเอียดตามใบตรวจรับพัสดุที่รายงานเสนอ "
       "และขออนุมัติเบิกจ่ายเงินให้{{ vendor_label }} เป็นเงิน {{ total_amount }} บาท ({{ total_baht }})",
       align="justify", indent=1.25, after=4)
    _p(doc, "ลงชื่อ..............................................เจ้าหน้าที่", align="center")
    _p(doc, "( {{ officer_name }} )", align="center", after=4)
    _p(doc, "ความเห็นของหัวหน้าเจ้าหน้าที่")
    _p(doc, "..............................................................................................................")
    _p(doc, "ลงชื่อ.............................................หัวหน้าเจ้าหน้าที่", align="center")
    _p(doc, "( {{ head_officer_name }} )", align="center", after=6)
    _p(doc, "คำสั่ง     (   ) ทราบ      (   ) อนุมัติ", indent=1.25)
    _p(doc, "ลงชื่อ.................................................", align="center")
    _p(doc, "( {{ director_name }} )", align="center")
    _p(doc, "{{ director_office }}", align="center")

    TEMPLATES_DIR.mkdir(exist_ok=True)
    out = TEMPLATES_DIR / "ใบตรวจรับพัสดุ.docx"
    doc.save(str(out))
    return out


def _summary_rows(table, *, label_cols, value_col, total_cols):
    """เติมแถวสรุปท้ายตาราง: รวมเป็นเงิน / ภาษีมูลค่าเพิ่ม / รวมเป็นเงินทั้งสิ้น + คำอ่านจำนวนเงิน
    label_cols = ดัชนีคอลัมน์ที่ใช้เขียนป้ายกำกับ (รวมกัน)
    value_col  = ดัชนีคอลัมน์ที่ใช้ใส่ตัวเลข
    total_cols = จำนวนคอลัมน์ทั้งหมด (สำหรับผสานแถวคำอ่าน)
    """
    def _label_row(text, value, *, bold=False):
        cells = table.add_row().cells
        merged = cells[label_cols[0]]
        for ci in label_cols[1:]:
            merged = merged.merge(cells[ci])
        _set_cell(merged, text, bold=True, align="center")
        _set_cell(cells[value_col], value, bold=bold, align="right")

    # แสดงรายการภาษีเฉพาะกรณีราคารวม VAT (ใช้ {%tr if%} ของ docxtpl)
    pre = table.add_row().cells
    _set_cell(pre[0], "{%tr if has_vat %}")
    _label_row("รวมเป็นเงิน", "{{ price_ex_vat }}")
    _label_row("ภาษีมูลค่าเพิ่ม", "{{ vat_amount }}")
    end = table.add_row().cells
    _set_cell(end[0], "{%tr endif %}")
    _label_row("รวมเป็นเงินทั้งสิ้น", "{{ total_amount }}", bold=True)
    # แถวคำอ่านจำนวนเงิน (ผสานทั้งแถว)
    baht = table.add_row().cells
    m = baht[0]
    for ci in range(1, total_cols):
        m = m.merge(baht[ci])
    _set_cell(m, "({{ total_baht }})", bold=True, align="center")


def _item_table(doc):
    """ตารางรายการพัสดุแบบ loop 3 แถว + แถวสรุป VAT/รวมทั้งสิ้น/คำอ่าน (ใช้ซ้ำหลายเอกสาร)"""
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    h = table.rows[0].cells
    _set_cell(h[0], "ลำดับ", bold=True, align="center")
    _set_cell(h[1], "รายการ", bold=True, align="center")
    _set_cell(h[2], "จำนวน/หน่วย", bold=True, align="center")
    _set_cell(h[3], "ราคาต่อหน่วย", bold=True, align="center")
    _set_cell(h[4], "จำนวนเงิน (บาท)", bold=True, align="center")
    _repeat_header_row(table.rows[0])      # หัวตารางซ้ำทุกหน้าเมื่อรายการเยอะ
    _set_cell(table.add_row().cells[0], "{%tr for it in items %}")
    drow = table.add_row()
    d = drow.cells
    _no_split_row(drow)                    # แต่ละรายการไม่ถูกตัดครึ่งข้ามหน้า
    _set_cell(d[0], "{{ loop.index }}", align="center")
    _set_cell(d[1], "{{ it.name }}")
    _set_cell(d[2], "{{ it.qty }} {{ it.unit }}", align="center")
    _set_cell(d[3], "{{ it.unit_price }}", align="right")
    _set_cell(d[4], "{{ it.amount }}", align="right")
    _set_cell(table.add_row().cells[0], "{%tr endfor %}")
    _summary_rows(table, label_cols=[0, 1, 2, 3], value_col=4, total_cols=5)
    # ความกว้างคงที่: ช่อง "รายการ" กว้างสุด (รวม ~16.25 ซม.)
    _fixed_cols(table, [Cm(1.2), Cm(7.0), Cm(2.6), Cm(2.6), Cm(2.85)])
    return table


def build_purchase_order():
    """แม่แบบ: ใบสั่งซื้อ / ใบสั่งจ้าง (ชื่อหัวเปลี่ยนตามประเภทอัตโนมัติ)
    Layout: โรงเรียน (ผู้สั่ง) ซ้าย + เลขที่/วันที่ ขวา → ผู้ขาย/รับจ้าง ด้านล่าง
    """
    doc = Document()
    _font(doc)

    # ตราครุฑกึ่งกลาง + ชื่อเอกสาร (ตามแบบฟอร์มจริง) — ครุฑใหญ่ขึ้นให้เด่น
    _krut_center(doc, height_cm=2.0)
    _p(doc, "{{ order_kind }}", align="center", bold=True, size=20, after=6)

    # ===== หัว 2 คอลัมน์: ซ้าย=ผู้ขาย | ขวา=ส่วนราชการ (จัดแนวด้วยตารางไร้เส้น) =====
    _order_header_table(doc)
    _p(doc, "", after=4)

    _p(doc,
       "ตามที่ {{ vendor_name }} ผู้เสนอราคาไว้ต่อ{{ school_name }} ซึ่งได้รับราคาและตกลง"
       "{{ proc_type }} ตามรายการดังต่อไปนี้", align="justify", indent=1.25)

    _item_table(doc)
    _p(doc, "", after=4)

    _p(doc, "การสั่ง{{ proc_type }} อยู่ภายใต้เงื่อนไขต่อไปนี้", indent=1.25)
    _p(doc, "1.  กำหนดส่งมอบภายใน {{ delivery_days }} วัน นับถัดจากวันที่{{ vendor_label }}ได้รับ{{ order_kind }}", indent=1.25)
    _p(doc, "2.  ครบกำหนดส่งมอบวันที่ {{ delivery_due_thai }}", indent=1.25)
    _p(doc, "3.  สถานที่ส่งมอบ {{ school_name }}", indent=1.25)
    _p(doc,
       "4.  สงวนสิทธิ์ค่าปรับกรณีส่งมอบเกินกำหนด โดยคิดค่าปรับเป็นรายวันในอัตราร้อยละ "
       "{{ penalty_rate }} ของราคาสิ่งของที่ยังไม่ได้รับมอบ", align="justify", indent=1.25)
    _p(doc,
       "5.  โรงเรียนสงวนสิทธิ์ที่จะไม่รับมอบ ถ้าปรากฏว่าสิ่งของนั้นมีลักษณะไม่ตรงตามรายการที่ระบุไว้ใน{{ order_kind }}",
       align="justify", indent=1.25, after=8)

    # ===== ลงนาม 2 ฝั่ง =====
    # ฝั่งผู้ขาย/รับจ้าง: ชื่อในวงเล็บ = เจ้าของร้าน (ถ้ามี) ไม่งั้นใช้ชื่อร้าน
    _sign_table(doc, [
        [("ลงชื่อ............................ผู้สั่ง{{ proc_type }}", "center"),
         ("( {{ signer_name }} )", "center"),
         ("{{ signer_position }}", "center"),
         ("วันที่ {{ order_date_thai }}", "center")],
        [("ลงชื่อ............................{{ vendor_label }}", "center"),
         ("( {{ vendor_signer }} )", "center"),   # เจ้าของร้าน/ผู้มีอำนาจ
         ("{{ vendor_name }}", "center"),           # ชื่อร้าน/บริษัท
         ("วันที่ {{ order_date_thai }}", "center")],
    ])

    TEMPLATES_DIR.mkdir(exist_ok=True)
    out = TEMPLATES_DIR / "ใบสั่งซื้อจ้าง.docx"
    doc.save(str(out))
    return out


def _krut_center(doc, *, height_cm=1.5):
    """ตราครุฑกึ่งกลางด้านบน (ค่าปริยาย 1.5 ซม. ใช้กับ 'คำสั่ง'/'ประกาศ')"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    krut = _krut_path()
    if krut:
        p.add_run().add_picture(str(krut), height=Cm(height_cm))
    else:
        r = p.add_run("[ครุฑ]")
        r.font.size = Pt(9)
        r.font.name = THAI_FONT


def _memo_head(doc, subject_expr, memo_no_expr="{{ memo_no }}", date_expr="{{ request_date }}"):
    """หัวบันทึกข้อความมาตรฐาน (ครุฑซ้าย + บันทึกข้อความ + label ตัวหนา ไม่มีเส้นคั่น)"""
    _krut_and_title(doc)
    _p_runs(doc, [("ส่วนราชการ  ", True), ("{{ school_office }}", False)])
    _p_runs(doc, [("ที่  ", True), (memo_no_expr, False),
                  ("\t", False), ("วันที่ ", True), (date_expr, False)], tab_cm=7)
    _p_runs(doc, [("เรื่อง  ", True), (subject_expr, False)])
    _p_runs(doc, [("เรียน  ", True), ("{{ director_office }}", False)])


def _signoff_officers(doc):
    """ช่องลงนาม เจ้าหน้าที่ + หัวหน้าเจ้าหน้าที่ (2 คอลัมน์ จัดด้วยตารางไร้เส้น)"""
    _p(doc, "", after=4)
    _sign_table(doc, [
        [("ลงชื่อ.....................................เจ้าหน้าที่", "center"),
         ("( {{ officer_name }} )", "center")],
        [("ลงชื่อ.....................................หัวหน้าเจ้าหน้าที่", "center"),
         ("( {{ head_officer_name }} )", "center")],
    ])


def _signoff_director(doc, *, with_approve=True):
    """ช่องลงนามผู้อำนวยการ (กึ่งกลาง)"""
    if with_approve:
        _p(doc, "(   )  เห็นชอบ        (   )  อนุมัติ", align="center", before=4)
    _p(doc, "(ลงชื่อ).........................................", align="center")
    _p(doc, "( {{ director_name }} )", align="center")
    _p(doc, "{{ director_office }}", align="center")
    _p(doc, "วันที่ {{ order_date_thai }}", align="center")


def build_result_report():
    """แม่แบบ: บันทึกข้อความ รายงานผลการพิจารณาและขออนุมัติสั่งซื้อ/สั่งจ้าง"""
    doc = Document()
    _font(doc)
    _memo_head(doc,
               "รายงานผลการพิจารณาและขออนุมัติสั่ง{{ proc_type }}",
               memo_no_expr="{{ result_memo_no }}",
               date_expr="{{ result_date_thai }}")
    _p(doc,
       "ตามที่{{ school_name }} เห็นชอบรายงานขอ{{ proc_type }}{{ subject }} จำนวน "
       "{{ total_amount }} บาท ({{ total_baht }}) ตามระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้าง"
       "และการบริหารพัสดุภาครัฐ พ.ศ. 2560 ข้อ 24 รายละเอียดดังแนบ",
       align="justify", indent=1.25)
    _p(doc,
       "ในการนี้เจ้าหน้าที่ได้เจรจาตกลงราคากับ {{ vendor_by }} ซึ่งมีอาชีพ{{ vendor_occupation }}แล้ว "
       "ปรากฏว่าเสนอราคาเป็นเงิน {{ total_amount }} บาท ({{ total_baht }}) ดังนั้นเพื่อให้เป็นไปตาม"
       "ระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560 ข้อ 79 "
       "จึงเห็นควร{{ proc_type }}จากผู้เสนอราคาดังกล่าว",
       align="justify", indent=1.25)
    _p(doc, "จึงเรียนมาเพื่อโปรดทราบและพิจารณา", indent=1.25)
    _p(doc,
       "อนุมัติให้สั่ง{{ proc_type }} จาก {{ vendor_name }} เป็น{{ vendor_label }} ในวงเงิน "
       "{{ total_amount }} บาท ({{ total_baht }}) กำหนดเวลาการส่งมอบ {{ delivery_days }} วัน",
       align="justify", indent=1.25)
    _p(doc, "ลงนามใน{{ order_kind }} ดังแนบ", indent=1.25)
    _signoff_officers(doc)
    _signoff_director(doc, with_approve=True)

    TEMPLATES_DIR.mkdir(exist_ok=True)
    out = TEMPLATES_DIR / "รายงานผลพิจารณา.docx"
    doc.save(str(out))
    return out


def build_inspect_command():
    """แม่แบบ: คำสั่งแต่งตั้งผู้ตรวจรับ (ครุฑกึ่งกลาง)"""
    doc = Document()
    _font(doc)
    _krut_center(doc)
    _p(doc, "คำสั่ง{{ school_name }}", align="center", bold=True, size=18, after=0)
    _p(doc, "ที่ {{ command_no }}", align="center", bold=True, after=0)
    _p(doc, "เรื่อง แต่งตั้งผู้ตรวจรับพัสดุ {{ subject }} โดยวิธี{{ method }}",
       align="center", bold=True, after=0)
    _p(doc, "─────────────────────", align="center", after=6)

    _p(doc,
       "ด้วย{{ school_name }} มีความประสงค์จัด{{ proc_type }}{{ subject }} โดยวิธี{{ method }} "
       "และเพื่อให้เป็นไปตามระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ "
       "พ.ศ. 2560", align="justify", indent=1.25)
    _p(doc, "จึงขอแต่งตั้งรายชื่อต่อไปนี้ เป็นคณะกรรมการตรวจรับพัสดุ โดยวิธี{{ method }} ดังนี้",
       indent=1.25)
    _member_table(doc, "inspect_members")
    _p(doc, "อำนาจและหน้าที่", bold=True, indent=1.25)
    _p(doc, "ทำการตรวจรับพัสดุให้เป็นไปตามเงื่อนไขของสัญญาหรือข้อตกลงนั้น",
       align="justify", indent=1.25)
    _p(doc, "ทั้งนี้ ตั้งแต่บัดนี้เป็นต้นไป", bold=True, indent=2.5, after=6)
    _p(doc, "สั่ง ณ วันที่ {{ command_date_official }}", align="center", after=24)
    _p(doc, "( {{ director_name }} )", align="center")
    _p(doc, "{{ director_office }}", align="center")

    TEMPLATES_DIR.mkdir(exist_ok=True)
    out = TEMPLATES_DIR / "คำสั่งแต่งตั้ง.docx"
    doc.save(str(out))
    return out


def build_quotation():
    """แม่แบบ: ใบเสนอราคา (จากผู้ขาย/ผู้รับจ้าง)"""
    doc = Document()
    _font(doc)
    _p(doc, "ใบเสนอราคา", align="center", bold=True, size=20, after=4)
    _p(doc, "วันที่ {{ quote_date_thai }}", align="right")
    _p(doc, "เรียน  ผู้อำนวยการ{{ school_name }}")
    _p(doc,
       "ข้าพเจ้า {{ vendor_name }} อยู่บ้านเลขที่ {{ vendor_address }} โทรศัพท์ {{ vendor_phone }} "
       "เลขประจำตัวผู้เสียภาษี {{ vendor_tax_id }} เป็นผู้มีคุณสมบัติครบถ้วนตามที่กำหนด "
       "และไม่เป็นผู้ทิ้งงานของทางราชการ", align="justify", indent=1.25)
    _p(doc, "ข้าพเจ้าขอเสนอราคา รวมทั้งบริการและกำหนดเวลาส่งมอบ ดังต่อไปนี้", indent=1.25)
    _item_table(doc)
    _p(doc, "", after=4)
    _p(doc, "ซึ่งเป็นราคาสุทธิ รวมทั้งภาษีอากรอื่น และค่าใช้จ่ายทั้งปวงไว้ด้วยแล้ว {{ vat_note }}",
       align="justify", indent=1.25)
    _p(doc, "{%p if has_vat %}")
    _p(doc, "(ราคาก่อนภาษี {{ price_ex_vat }} บาท + ภาษีมูลค่าเพิ่ม {{ vat_amount }} บาท)", indent=1.25)
    _p(doc, "{%p endif %}")
    _p(doc, "คำเสนอนี้จะยืนอยู่เป็นระยะเวลา 15 วัน นับตั้งแต่วันที่ยื่นใบเสนอราคา", indent=1.25)
    _p(doc, "กำหนดส่งมอบพัสดุภายใน {{ delivery_days }} วันทำการ นับตั้งแต่วันที่ยื่นใบเสนอราคา", indent=1.25, after=8)
    _sign_table(doc, [
        [("ลงชื่อ..........................ผู้ต่อรองราคา", "center"),
         ("( {{ officer_name }} )", "center")],
        [("ลงชื่อ..........................ผู้เสนอราคา", "center"),
         ("( {{ vendor_signer }} )", "center"),
         ("{{ vendor_name }}", "center")],
    ])
    _p(doc, "หมายเหตุ  ผู้เสนอราคาต้องเป็นเจ้าของหรือหุ้นส่วนผู้จัดการ หรือผู้รับมอบอำนาจ", before=2)
    TEMPLATES_DIR.mkdir(exist_ok=True)
    out = TEMPLATES_DIR / "ใบเสนอราคา.docx"
    doc.save(str(out))
    return out


def build_winner_announcement():
    """แม่แบบ: ประกาศผู้ชนะการเสนอราคา"""
    doc = Document()
    _font(doc)
    _p(doc, "ประกาศ{{ school_name }}", align="center", bold=True, size=18, after=0)
    _p(doc, "เรื่อง ประกาศผู้ชนะการเสนอราคา {{ proc_type }}{{ subject }} โดยวิธี{{ method }}",
       align="center", bold=True, after=0)
    _p(doc, "─────────────────────", align="center", after=6)
    _p(doc,
       "ตามที่{{ school_name }} ได้{{ proc_type }}{{ subject }} โดยวิธี{{ method }} "
       "การประกาศผู้ได้รับการคัดเลือกโดยวิธี{{ method }} นั้น", align="justify", indent=1.25)
    _p(doc,
       "{{ proc_type }}{{ subject }} จำนวน {{ item_count }} รายการ ผู้เสนอราคาที่ชนะการเสนอราคา "
       "ได้แก่ {{ vendor_name }} โดยเสนอราคาต่ำสุด เป็นจำนวนเงินทั้งสิ้น {{ total_amount }} บาท "
       "({{ total_baht }})", align="justify", indent=1.25)
    _p(doc, "รวมภาษีมูลค่าเพิ่มและภาษีอื่น ค่าขนส่ง ค่าจดทะเบียน และค่าใช้จ่ายอื่น ๆ ทั้งปวง",
       align="justify", indent=1.25, after=8)
    _p(doc, "ประกาศ ณ วันที่ {{ order_date_thai }}", align="center", after=12)
    _p(doc, "(ลงชื่อ).........................................", align="center")
    _p(doc, "( {{ director_name }} )", align="center")
    _p(doc, "{{ director_office }}", align="center")
    TEMPLATES_DIR.mkdir(exist_ok=True)
    out = TEMPLATES_DIR / "ประกาศผู้ชนะ.docx"
    doc.save(str(out))
    return out


def build_spec_committee():
    """แม่แบบ: บันทึกขออนุมัติแต่งตั้งคณะกรรมการกำหนดคุณลักษณะเฉพาะ/ราคากลาง (ข้อ 21)"""
    doc = Document()
    _font(doc)
    _memo_head(doc, "ขออนุมัติแต่งตั้งคณะกรรมการกำหนดรายละเอียดคุณลักษณะเฉพาะพัสดุ และราคากลาง",
               memo_no_expr="{{ spec_memo_no }}", date_expr="{{ spec_date_thai }}")
    _p(doc,
       "ด้วย{{ school_name }} จะดำเนินการ{{ proc_type }}{{ subject }} โดยใช้วงเงินงบประมาณ จำนวน "
       "{{ total_amount }} บาท ({{ total_baht }}) เพื่อให้การกำหนดรายละเอียดคุณลักษณะเฉพาะพัสดุ"
       "เป็นไปตามระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560 "
       "ข้อ 21 จึงขอแต่งตั้งคณะกรรมการ ประกอบด้วย", align="justify", indent=1.25)
    _member_table(doc, "spec_members")
    _p(doc, "โดยให้มีหน้าที่ (1) จัดทำรายละเอียดคุณลักษณะเฉพาะ และกำหนดหลักเกณฑ์การพิจารณา "
            "(2) จัดทำราคากลางของพัสดุตามแนวทางที่เกี่ยวข้อง", align="justify", indent=1.25)
    _p(doc, "จึงเรียนมาเพื่อโปรดพิจารณาอนุมัติ", indent=1.25)
    _signoff_officers(doc)
    _signoff_director(doc, with_approve=True)
    TEMPLATES_DIR.mkdir(exist_ok=True)
    out = TEMPLATES_DIR / "แต่งตั้งคุณลักษณะ.docx"
    doc.save(str(out))
    return out


def build_tor():
    """แม่แบบ: รายละเอียดคุณลักษณะเฉพาะของพัสดุ (TOR)"""
    doc = Document()
    _font(doc)
    _p(doc, "รายละเอียดคุณลักษณะเฉพาะของพัสดุที่จะ{{ proc_type }} (TOR)", align="center", bold=True, size=18)
    _p(doc, "ตามพระราชบัญญัติการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560 (ข้อ 21)",
       align="center", after=6)
    _p(doc, "1. ความเป็นมา/เหตุผลความจำเป็น", bold=True)
    _p(doc, "ด้วย{{ school_name }} มีความประสงค์จะ{{ proc_type }}{{ subject }} โดยวิธี{{ method }} "
            "เนื่องจาก {{ purpose }}", align="justify", indent=1.25)
    _p(doc, "2. ลักษณะการจัดหา", bold=True)
    _p(doc, "งาน{{ proc_type }}", indent=1.25)
    _p(doc, "3. รายละเอียดคุณลักษณะเฉพาะของพัสดุ", bold=True)
    _item_table(doc)
    _p(doc, "", after=4)
    _p(doc, "4. กำหนดเวลาส่งมอบ", bold=True)
    _p(doc, "ส่งมอบภายใน {{ delivery_days }} วัน", indent=1.25)
    _p(doc, "5. หลักเกณฑ์การพิจารณาคัดเลือกข้อเสนอ", bold=True)
    _p(doc, "พิจารณาตัดสินโดยใช้เกณฑ์ราคา", indent=1.25)
    _p(doc, "6. ราคากลางและที่มาของราคากลาง", bold=True)
    _p(doc, "ราคากลาง {{ total_amount }} บาท ({{ total_baht }}) ได้มาจาก {{ price_ref_source }}",
       align="justify", indent=1.25)
    _p(doc, "7. วงเงินงบประมาณ", bold=True)
    _p(doc, "ใช้งบประมาณจากเงิน{{ budget_source }} จำนวน {{ total_amount }} บาท ({{ total_baht }})",
       indent=1.25, after=8)
    _p(doc, "คณะกรรมการกำหนดคุณลักษณะเฉพาะและราคากลาง", bold=True)
    _member_table(doc, "spec_members")
    TEMPLATES_DIR.mkdir(exist_ok=True)
    out = TEMPLATES_DIR / "TOR.docx"
    doc.save(str(out))
    return out


def build_delivery_note():
    """แม่แบบ: ใบส่งมอบงาน + แจ้งหนี้ขอเบิกเงิน"""
    doc = Document()
    _font(doc)
    _p(doc, "ใบส่งมอบงาน", align="center", bold=True, size=20, after=4)
    _p(doc, "เขียนที่ {{ school_name }}", align="right")
    _p(doc, "วันที่ {{ delivery_date_thai }}", align="right", after=4)
    _p(doc, "เรื่อง  ส่งมอบงาน{{ proc_type }}และแจ้งหนี้ขอเบิกเงิน")
    _p(doc, "เรียน  {{ director_office }}")
    _p(doc,
       "ตามที่{{ school_name }} ได้ตกลงให้ข้าพเจ้า {{ vendor_name }} รับ{{ proc_type }}{{ subject }} "
       "โดยวิธี{{ method }} ตาม{{ order_kind }} เลขที่ {{ order_no }} ลงวันที่ {{ order_date_thai }} "
       "ในวงเงิน {{ total_amount }} บาท ({{ total_baht }}) นั้น", align="justify", indent=1.25)
    _p(doc,
       "บัดนี้ ข้าพเจ้าได้ปฏิบัติตาม{{ order_kind }} เสร็จเรียบร้อยแล้ว จึงขอส่งมอบงาน "
       "เพื่อตรวจรับและขอเบิกจ่ายเงิน จำนวน {{ total_amount }} บาท ({{ total_baht }})",
       align="justify", indent=1.25, after=8)
    _p(doc, "ขอแสดงความนับถือ", align="center")
    _p(doc, "(ลงชื่อ).........................................", align="center")
    _p(doc, "( {{ vendor_signer }} )", align="center")
    _p(doc, "{{ vendor_name }}", align="center")
    TEMPLATES_DIR.mkdir(exist_ok=True)
    out = TEMPLATES_DIR / "ใบส่งมอบงาน.docx"
    doc.save(str(out))
    return out


def _finance_table(doc):
    """ตารางสรุปการจ่ายเงิน (ไร้เส้น) จัดคอลัมน์ให้ตรงกัน:
    ป้ายกำกับ (ซ้าย) | จำนวนเงิน (ชิดขวา) | หน่วย/คำอ่าน (ซ้าย)
    """
    lines = [
        ("มูลค่าสินค้าหรือบริการ", "{{ goods_value }}", "บาท"),
        ("บวก ภาษีมูลค่าเพิ่ม", "{{ vat_value }}", "บาท"),
        ("จำนวนเงินที่ขอเบิกทั้งสิ้น", "{{ total_amount }}", "บาท"),
        ("หัก ภาษีเงินได้", "{{ wht_amount }}", "บาท"),
        ("หัก ค่าปรับ", "{{ fine_amount }}", "บาท"),
        ("คงเหลือจ่ายจริง เป็นเงิน", "{{ net_pay }}", "บาท  ({{ net_pay_baht }})"),
    ]
    table = doc.add_table(rows=len(lines), cols=3)
    _no_borders(table)
    table.autofit = False
    table.allow_autofit = False
    _table_indent(table, 1.6)
    widths = [Cm(6.0), Cm(3.2), Cm(6.5)]
    for i, (label, amount, unit) in enumerate(lines):
        cells = table.rows[i].cells
        bold = (i == len(lines) - 1)        # แถวสุดท้าย (คงเหลือจ่ายจริง) ตัวหนา
        _set_cell(cells[0], label, bold=bold, align="left", size=16)
        _set_cell(cells[1], amount, bold=bold, align="right", size=16)
        _set_cell(cells[2], unit, bold=bold, align="left", size=16)
        for c, w in zip(cells, widths):
            c.width = w
    return table


def build_disbursement():
    """แม่แบบ: บันทึกข้อความ รายงานผลการตรวจรับพัสดุและอนุมัติเบิกจ่ายเงิน
    (รวมความเห็นเจ้าหน้าที่การเงิน + รายละเอียดการคำนวณเงินที่จ่ายจริง)"""
    doc = Document()
    _font(doc)
    _memo_head(doc, "รายงานผลการตรวจรับพัสดุและอนุมัติเบิกจ่ายเงิน",
               memo_no_expr="{{ inspect_memo_no }}",
               date_expr="{{ inspect_date_thai }}")
    _p(doc,
       "ตามที่{{ school_name }} ได้จัด{{ proc_type }}{{ subject }} โดยมี {{ vendor_name }} เป็น{{ vendor_label }} "
       "ตาม{{ order_kind }}เลขที่ {{ order_no }} ลงวันที่ {{ order_date_thai }} ครบกำหนดส่งมอบวันที่ "
       "{{ delivery_due_thai }} รวมเป็นเงิน {{ total_amount }} บาท ({{ total_baht }}) นั้น",
       align="justify", indent=1.25)
    _p(doc,
       "บัดนี้ {{ vendor_label }}ได้ส่งมอบพัสดุเสร็จเรียบร้อยแล้ว ตามใบส่งของ/ใบกำกับภาษี/ใบเสร็จรับเงิน "
       "เล่มที่ {{ delivery_note_book }} เลขที่ {{ delivery_note_no }} ลงวันที่ {{ inspect_date_thai }} "
       "และ{{ committee_word }}ได้ทำการตรวจรับพัสดุเมื่อวันที่ {{ inspect_date_thai }} "
       "ไว้เป็นการถูกต้องครบถ้วนดังหลักฐานที่แนบ",
       align="justify", indent=1.25)
    _p(doc,
       "จึงเรียนมาเพื่อทราบผลการตรวจรับพัสดุ ตามนัยข้อ 175 (4) แห่งระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อ"
       "จัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560", align="justify", indent=1.25)
    _signoff_officers(doc)

    # ===== ความเห็นของเจ้าหน้าที่การเงิน + รายละเอียดการจ่าย =====
    _p(doc, "ความเห็นของเจ้าหน้าที่การเงิน", bold=True, indent=1.25, before=4)
    _p(doc, "ขออนุมัติจ่ายเงินให้แก่ {{ vendor_name }} รายละเอียดดังนี้", indent=1.25, after=2)
    # ตารางไร้เส้น 3 คอลัมน์ (ป้ายกำกับ | จำนวนเงินชิดขวา | หน่วย/คำอ่าน) จัดแนวตรงกันทุกแถว
    _finance_table(doc)
    _p(doc, "", after=2)
    _sign_table(doc, [
        [("ลงชื่อ.....................................เจ้าหน้าที่การเงิน", "center"),
         ("( {{ finance_officer_name }} )", "center")],
        [("", "center")],
    ])
    _signoff_director(doc, with_approve=True)
    TEMPLATES_DIR.mkdir(exist_ok=True)
    out = TEMPLATES_DIR / "รายงานเบิกจ่าย.docx"
    doc.save(str(out))
    return out


def build_all():
    TEMPLATES_DIR.mkdir(exist_ok=True)
    built = [build_purchase_request(), build_inspection(), build_purchase_order(),
             build_result_report(), build_inspect_command(),
             build_quotation(), build_winner_announcement(), build_spec_committee(),
             build_tor(), build_delivery_note(), build_disbursement()]
    return built


if __name__ == "__main__":
    for f in build_all():
        print("สร้างแม่แบบ:", f.name)
