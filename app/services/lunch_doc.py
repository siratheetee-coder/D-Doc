# -*- coding: utf-8 -*-
"""
lunch_doc.py
------------
ออกเอกสารจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ) รายงวด (สร้าง docx ตอนรันไทม์)
- render_installment_doc(inst, school, menus): ไฟล์ "งวด X" รวม 3 ส่วน
    (1) บันทึกรายงานผู้ควบคุม + ตารางเมนูรายวัน  (2) ใบส่งมอบงาน  (3) ใบตรวจรับงานจ้าง

อิงโครงสร้าง/ถ้อยคำจากไฟล์จริงที่โรงเรียนใช้ ใช้ helper ร่วมกับ build_templates
ดึงรายชื่อคณะกรรมการที่กรอกไว้ (LunchCommittee) มาลงในเอกสาร ถ้าไม่มีเว้นจุดไข่ปลาให้เซ็น
"""
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.services.doc_page import set_a4

from app.database import get_data_dir
from app.thai_utils import thai_date, bahttext
from app.services.build_templates import (
    _font, _p, _set_cell, _repeat_header_row, _no_split_row, _sign_table,
    _krut_and_title, _krut_center, _p_runs, _hr, _no_borders,
    THAI_FONT, _csize, _bcs,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH

_BLANK = "............................"


def _safe(text: str) -> str:
    for ch in '<>:"/\\|?*\n\r\t':
        text = text.replace(ch, "_")
    return text.strip()[:80]


def _school_disp(school) -> str:
    """ชื่อโรงเรียนแบบมีคำว่า 'โรงเรียน' นำหน้าเพียงครั้งเดียว
    (กันซ้ำ 'โรงเรียนโรงเรียนบ้านหินลาด' เมื่อชื่อที่บันทึกมี 'โรงเรียน' อยู่แล้ว)"""
    n = (getattr(school, "name", "") or "").strip()
    if not n:
        return "โรงเรียน"
    return n if n.startswith("โรงเรียน") else "โรงเรียน" + n


def _money(x) -> str:
    x = round(float(x or 0), 2)
    return f"{int(x):,}" if x == int(x) else f"{x:,.2f}"


def _save(doc, name: str) -> str:
    out_dir = get_data_dir() / "documents"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / (_safe(name) + ".docx")
    doc.save(str(out_path))
    return str(out_path)


def _begin(doc):
    """เริ่มเอกสาร: ถ้า doc=None สร้างใหม่ (own=True) ไม่งั้นต่อท้ายด้วย page break (own=False)"""
    if doc is None:
        doc = Document(); set_a4(doc)
        _font(doc)
        return doc, True
    if doc.paragraphs or doc.tables:   # เว้นหน้าเฉพาะเมื่อมีเนื้อหาก่อนหน้าแล้ว
        doc.add_page_break()
    return doc, False


def _finish(doc, own, name):
    return _save(doc, name) if own else doc


def _dnum(dt) -> str:
    return thai_date(dt) if dt else _BLANK


def _docnos(rnd) -> dict:
    """เลขที่/วันที่รายฉบับของรอบ (JSON) {kind:{field:value}}"""
    import json
    try:
        return json.loads(getattr(rnd, "doc_nos", "") or "{}") or {}
    except Exception:
        return {}


def _doc_no(rnd, kind, fallback=""):
    """เลขที่เอกสารรายฉบับ (doc_nos[kind]['no']) ไม่มีก็คืน fallback"""
    return ((_docnos(rnd).get(kind) or {}).get("no") or "").strip() or fallback


def _doc_dt(rnd, kind, field="date"):
    """วันที่ในฟิลด์ field ของเอกสาร kind (คืน datetime หรือ None)"""
    from datetime import datetime
    s = (_docnos(rnd).get(kind) or {}).get(field) or ""
    try:
        return datetime.fromisoformat(s) if s else None
    except Exception:
        return None


def _memo_ref2(rnd, kind):
    """(วันที่str, เลขที่) ของบันทึก/คำสั่งชนิด kind - fallback เลขตกลงจ้างเดิม"""
    dt = _doc_dt(rnd, kind, "date") or rnd.order_date
    no = _doc_no(rnd, kind, (rnd.order_no or "").strip())
    return _dnum(dt), no


def _menu_text(m) -> str:
    parts = [(m.main or "").strip()]
    if (m.dessert or "").strip():
        parts.append((m.dessert or "").strip())
    return "  ".join(p for p in parts if p)


_CHECK = "☐"        # ☐ ช่องติ๊ก
_RESULT_CRITERIA = ["ความสะอาด", "คุณภาพอาหาร", "ความทันเวลา", "ความเพียงพอ"]
_RESULT_OPTIONS = ["ดีมาก", "ดี", "พอใช้", "ปรับปรุง"]
_TH_NUM = ["๑", "๒", "๓", "๔", "๕", "๖", "๗", "๘", "๙", "๑๐"]


def _result_cell_text() -> str:
    """ข้อความคอลัมน์ 'ผลการดำเนินงาน' : 4 หัวข้อ x 4 ตัวเลือกให้ติ๊ก"""
    opts = " ".join(f"{_CHECK}{o}" for o in _RESULT_OPTIONS)
    return "\n".join(f"{c}\n{opts}" for c in _RESULT_CRITERIA)


def _rule(doc):
    """เส้นคั่นแนวนอนเต็มความกว้าง (ขอบล่างของย่อหน้าว่าง)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def _committee_cell_text(members, n=3, names=False) -> str:
    """ช่องลงชื่อกรรมการแบบแนวราบ เลขอารบิก: "1) ...........  2) ...........  3) ..........."
    names=True : แทรกชื่อในวงเล็บต่อท้ายแต่ละช่อง (ปกติ names=False เว้นจุดไข่ปลาให้เซ็น)"""
    count = max(n, len(members or [])) if names else n
    dots = "............................."
    segs = []
    for i in range(count):
        seg = f"{i + 1}) {dots}"
        if names:
            nm = members[i].name if (members and i < len(members) and members[i].name) else ""
            if nm:
                seg += f" ( {nm} )"
        segs.append(seg)
    return "    ".join(segs)


def _daily_table(doc, menus, committee=None):
    """ตารางควบคุมงานรายวัน: วัน เดือน ปี | รายการอาหารตาม TOR | ผลการดำเนินงาน (ติ๊ก) | คณะกรรมการควบคุมงาน"""
    widths = [Cm(3.0), Cm(4.0), Cm(4.2), Cm(5.5)]   # ขยายวันที่ให้อยู่บรรทัดเดียว + ช่องลงชื่อกว้าง
    headers = ["วัน เดือน ปี", "รายการอาหาร\nตามขอบเขตงาน TOR", "ผลการดำเนินงาน",
               "คณะกรรมการควบคุมงานจ้าง\nประกอบอาหาร"]
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    t.autofit = False
    hdr = t.rows[0]
    _repeat_header_row(hdr)
    for c, h, w in zip(hdr.cells, headers, widths):
        _set_cell(c, h, bold=True, align="center", size=13)
        c.width = w
    committee_txt = _committee_cell_text(committee, names=False)   # ไม่ใส่ชื่อ ลดความสูงตาราง
    for m in (menus if menus else [None] * 5):
        r = t.add_row()
        _no_split_row(r)
        _set_cell(r.cells[0], thai_date(m.date) if (m and m.date) else "", align="center", size=13)
        _set_cell(r.cells[1], _menu_text(m) if m else "", align="left", size=13)
        _set_cell(r.cells[2], _result_cell_text(), align="left", size=12)
        _set_cell(r.cells[3], committee_txt, align="left", size=12)
        for c, w in zip(r.cells, widths):
            c.width = w
    return t


def _inspect_table(doc, menus, committee=None):
    """ตารางใบตรวจรับพัสดุ: วัน เดือน ปี | รายการอาหาร | ลายมือชื่อผู้ส่งมอบงาน | ผู้ตรวจรับพัสดุ/คณะกรรมการ"""
    widths = [Cm(3.0), Cm(4.6), Cm(1.9), Cm(6.8)]   # ขยายตาราง + วันที่บรรทัดเดียว + ช่องลงชื่อกว้าง (แนวราบ)
    headers = ["วัน เดือน ปี", "รายการอาหาร", "ลายมือชื่อผู้\nส่งมอบงาน",
               "ผู้ตรวจรับงานจ้างหรือคณะกรรมการ\nตรวจรับงานจ้าง"]
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    t.autofit = False
    hdr = t.rows[0]
    _repeat_header_row(hdr)
    for c, h, w in zip(hdr.cells, headers, widths):
        _set_cell(c, h, bold=True, align="center", size=13)
        c.width = w
    committee_txt = _committee_cell_text(committee, names=False)   # ไม่ใส่ชื่อกรรมการตรวจรับ
    for m in (menus if menus else [None] * 5):
        r = t.add_row()
        _no_split_row(r)
        _set_cell(r.cells[0], thai_date(m.date) if (m and m.date) else "", align="center", size=13)
        _set_cell(r.cells[1], _menu_text(m) if m else "", align="left", size=13)
        _set_cell(r.cells[2], "", size=13)
        _set_cell(r.cells[3], committee_txt, align="left", size=12)
        for c, w in zip(r.cells, widths):
            c.width = w
    return t


def _menu_table3(doc, menus, third_header):
    """ตาราง 3 คอลัมน์ (09 ใบส่งมอบ / 10 ใบตรวจรับ): วัน เดือน ปี | รายการอาหาร | <third_header>"""
    widths = [Cm(3.0), Cm(7.0), Cm(6.2)]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.autofit = False
    hdr = t.rows[0]
    _repeat_header_row(hdr)
    for c, h, w in zip(hdr.cells, ["วัน เดือน ปี", "รายการอาหาร", third_header], widths):
        _set_cell(c, h, bold=True, align="center", size=14)
        c.width = w
    for m in (menus if menus else [None] * 5):
        r = t.add_row()
        _no_split_row(r)
        vals = [thai_date(m.date) if (m and m.date) else "", _menu_text(m) if m else "", ""]
        for c, v, w in zip(r.cells, vals, widths):
            _set_cell(c, v, size=14)
            c.width = w
    return t


def render_installment_doc(inst, school, menus) -> str:
    doc = Document(); set_a4(doc)
    _font(doc)
    rnd = inst.round
    vendor = rnd.vendor
    vname = vendor.name if vendor else _BLANK
    order_no = _doc_no(rnd, "hire-order", (getattr(rnd, "order_no", "") or "").strip() or _BLANK)
    sname = _school_disp(school)
    director = (school.director_name or "").strip()
    officer = (getattr(school, "officer_name", "") or "").strip()
    head_officer = (getattr(school, "head_officer_name", "") or "").strip()
    control_members = [m for m in getattr(rnd, "committees", []) if m.kind == "control"]
    inspect_members = [m for m in getattr(rnd, "committees", []) if m.kind == "inspect"]
    amount = _money(inst.amount or 0)
    amt_text = bahttext(inst.amount or 0)
    period = f"วันที่ {_dnum(inst.start_date)} ถึงวันที่ {_dnum(inst.end_date)} รวม {inst.days or ''} วัน"

    # ===== ส่วนที่ 1: บันทึกรายงานผู้ควบคุมและคณะกรรมการตรวจการจ้าง =====
    _p(doc, "บันทึกรายงานผู้ควบคุมและคณะกรรมการตรวจการจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ)",
       align="center", bold=True, size=16, after=4)
    _p(doc, f"เขียนที่ {sname}", align="right", after=0)
    _p(doc, f"วันที่ {_dnum(inst.inspect_date or inst.end_date)}", align="right", after=6)
    _p(doc, f"ตามที่{sname} ได้ตกลงจ้าง {vname} ประกอบอาหารกลางวัน (ปรุงสำเร็จ) "
            f"ให้นักเรียนรับประทาน งวดที่ {inst.seq} ระหว่าง{period} นั้น",
       align="justify", indent=1.25)
    _p(doc, "คณะกรรมการควบคุมงานและคณะกรรมการตรวจการจ้าง ขอรายงานผลการดำเนินงาน "
            "การประกอบอาหารกลางวันเป็นรายวัน ดังนี้", align="justify", indent=1.25, after=4)
    _daily_table(doc, menus, control_members)
    _p(doc, "", after=4)
    _p(doc, "ความเห็นของผู้อำนวยการสถานศึกษา : ทราบผลการดำเนินการประกอบอาหารกลางวัน",
       indent=1.25, after=10)
    _p(doc, "(ลงชื่อ)...........................................", align="center", after=3, line=1.15)
    _p(doc, f"( {director or _BLANK} )", align="center", after=3, line=1.15)
    _p(doc, f"ผู้อำนวยการ{sname}", align="center", after=6, line=1.15)

    # ===== ส่วนที่ 2: ใบส่งมอบงาน =====
    doc.add_page_break()
    _p(doc, "ใบส่งมอบงาน", align="center", bold=True, size=18, after=6)
    _p(doc, f"วันที่ {_dnum(inst.deliver_date or inst.end_date)}", align="right", after=6)
    _p(doc, f"ตามที่{sname} ได้ตกลงจ้างข้าพเจ้า {vname} ตามใบสั่งจ้าง เลขที่ {order_no} "
            f"เพื่อประกอบอาหารกลางวัน (ปรุงสำเร็จ) สำหรับนักเรียน งวดที่ {inst.seq} "
            f"ระหว่าง{period} นั้น", align="justify", indent=1.25)
    _p(doc, "บัดนี้ ข้าพเจ้าได้ดำเนินการประกอบอาหารเสร็จเรียบร้อยตามข้อกำหนดของงานแล้ว "
            "จึงขอส่งมอบงานตามเอกสารที่แนบมาพร้อมนี้", align="justify", indent=1.25, after=4)
    _menu_table3(doc, menus, "ผู้ส่งมอบงาน")
    _p(doc, f"ขอเบิกเงิน จำนวน {amount} บาท ({amt_text})", indent=1.25, before=4, after=14)
    _sign_table(doc, [
        [("(ลงชื่อ)...........................................ผู้ส่งมอบงาน", "center"),
         (f"( {vname} )", "center")],
    ])

    # ===== ส่วนที่ 3: ใบตรวจรับงานจ้าง =====
    doc.add_page_break()
    _p(doc, "ใบตรวจรับงานจ้าง", align="center", bold=True, size=18, after=4)
    _p(doc, f"เขียนที่ {sname}", align="right", after=0)
    _p(doc, f"วันที่ {_dnum(inst.inspect_date or inst.end_date)}", align="right", after=6)
    _p(doc, f"ตามที่{sname} ได้ตกลงจ้าง {vname} ประกอบอาหารกลางวัน (ปรุงสำเร็จ) "
            f"ให้นักเรียนรับประทาน ตามใบสั่งจ้าง เลขที่ {order_no} นั้น",
       align="justify", indent=1.25)
    _p(doc, f"บัดนี้ ผู้รับจ้างได้ส่งมอบงานทุกวันตามข้อตกลง และคณะกรรมการตรวจรับงานจ้าง "
            f"ได้ตรวจรับไว้ถูกต้องครบถ้วนแล้ว งวดที่ {inst.seq} ตามบันทึกข้อตกลงจ้างแล้ว ดังนี้",
       align="justify", indent=1.25, after=4)
    _inspect_table(doc, menus, inspect_members)
    _p(doc, "", after=4)
    _p(doc, "เรียน  ผู้อำนวยการ" + sname, indent=1.25, after=0)
    _p(doc, f"เพื่อโปรดทราบผลการตรวจรับงานจ้าง และขออนุมัติจ่ายเงินให้ผู้รับจ้าง งวดที่ {inst.seq} "
            f"ระหว่าง{period} เป็นเงิน {amount} บาท ({amt_text})",
       align="justify", indent=1.25, after=8)
    _sign_table(doc, [
        [("(ลงชื่อ)...........................................", "center"),
         (f"( {officer or _BLANK} )", "center"), ("เจ้าหน้าที่", "center")],
        [("(ลงชื่อ)...........................................", "center"),
         (f"( {head_officer or _BLANK} )", "center"), ("หัวหน้าเจ้าหน้าที่", "center")],
    ], gap=False)
    _rule(doc)
    # บล็อกความเห็น ผอ. + ลายเซ็น ให้อยู่หน้าเดียวกันเสมอ (กันชื่อ ผอ. หลุดไปหน้าใหม่ตอนงวดสั้น)
    end_paras = [
        _p(doc, "ความเห็นของผู้บริหารสถานศึกษา", indent=1.25, after=0),
        _p(doc, "(   ) ทราบผลการตรวจรับ          (   ) อนุมัติ", indent=1.5, after=12),
        _p(doc, "(ลงชื่อ)...........................................", align="center", after=3, line=1.15),
        _p(doc, f"( {director or _BLANK} )", align="center", after=3, line=1.15),
        _p(doc, f"ตำแหน่ง ผู้อำนวยการ{sname}", align="center", after=0, line=1.15),
    ]
    for p in end_paras[:-1]:
        p.paragraph_format.keep_with_next = True

    return _save(doc, f"งวดที่{inst.seq}_ปี{rnd.program.year}")


def _simple_table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    for c, h, w in zip(t.rows[0].cells, headers, widths):
        _set_cell(c, h, bold=True, align="center", size=14)
        c.width = w
    for row in rows:
        rc = t.add_row().cells
        for c, v, w in zip(rc, row, widths):
            _set_cell(c, v, size=14)
            c.width = w
    return t


def render_disburse_lunch_doc(inst, school, wht_rate=0.01) -> str:
    """เอกสารขอเบิกจ่ายรายงวด: บันทึกขออนุมัติ + ใบสำคัญรับเงิน + หนังสือรับรองหักภาษี ณ ที่จ่าย"""
    doc = Document(); set_a4(doc)
    _font(doc)
    rnd = inst.round
    prog = rnd.program
    vendor = rnd.vendor
    vname = vendor.name if vendor else _BLANK
    vaddr = (vendor.address if vendor else "") or _BLANK
    vtax = (vendor.tax_id if vendor else "") or _BLANK
    order_no = getattr(rnd, "order_no", None) or _BLANK
    sname = _school_disp(school)
    saddr = (school.address or "").strip()
    director = (school.director_name or "").strip() or _BLANK
    fin = (school.finance_officer_name or "").strip() or _BLANK
    fund = (prog.funding_org or "").strip() or "องค์กรปกครองส่วนท้องถิ่น"
    amt = round(float(inst.amount or 0), 2)
    wht = round(amt * float(wht_rate or 0), 2)
    net = round(amt - wht, 2)
    A, W, N = _money(amt), _money(wht), _money(net)
    period = (f"งวดที่ {inst.seq} ระหว่างวันที่ {_dnum(inst.start_date)} ถึงวันที่ "
              f"{_dnum(inst.end_date)} รวม {inst.days or ''} วัน")

    # ===== 1. บันทึกข้อความ ขออนุมัติเบิกจ่าย =====
    _memo_head(doc, school, [f"ขออนุมัติเบิกจ่ายเงินอุดหนุนอาหารกลางวัน รับจาก{fund}"],
               _dnum(inst.inspect_date or inst.end_date), order_no)
    _p(doc, f"ตามที่{sname}ได้จ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ) จำนวน {inst.days or ''} วัน "
            f"จาก {vname} จำนวนเงิน {A} บาท ({bahttext(amt)}) ตามใบสั่งจ้าง เลขที่ {order_no} "
            f"{period} จากเงินนอกงบประมาณ ประเภทเงินอุดหนุนอาหารกลางวันรับจาก{fund} นั้น",
       align="justify", indent=1.25)
    _p(doc, "บัดนี้ ผู้รับจ้างได้ส่งมอบอาหาร (ตามรายการอาหาร) ถูกต้องครบถ้วนแล้ว ตามนัยข้อ ๑๗๕ (๔) "
            "แห่งระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. ๒๕๖๐ "
            "เห็นควรเบิกจ่ายให้แก่ผู้รับจ้าง โดยมีรายละเอียด ดังนี้", align="justify", indent=1.25, after=4)
    from docx.enum.table import WD_TABLE_ALIGNMENT
    mt = doc.add_table(rows=0, cols=2); mt.style = "Table Grid"; mt.autofit = False
    _money_rows = [("จำนวนเงินขอเบิก", A), ("ภาษีมูลค่าเพิ่ม (ถ้ามี)", "-"),
                   ("มูลค่าสินค้า", "-"), ("หักภาษี ณ ที่จ่าย", W),
                   ("ค่าปรับ (ถ้ามี)", "-"), ("คงเหลือจ่ายจริง", N)]
    for k, (label, v) in enumerate(_money_rows):
        cells = mt.add_row().cells
        bold = (k == len(_money_rows) - 1)
        _set_cell(cells[0], label, align="left", size=14, bold=bold)
        _set_cell(cells[1], (f"{v} บาท" if v != "-" else "-"), align="right", size=14, bold=bold)
        cells[0].width = Cm(9.5); cells[1].width = Cm(5.0)
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    _p(doc, "", after=4)
    _p(doc, f"จึงเรียนมาเพื่อโปรดพิจารณาอนุมัติจ่ายเงิน (เงินอุดหนุนอาหารกลางวันรับจาก{fund}) "
            f"แก่ผู้รับจ้าง จำนวน {N} บาท ({bahttext(net)})", align="justify", indent=1.25, after=10)
    _sign_table(doc, [
        [("(ลงชื่อ)...........................................เจ้าหน้าที่การเงิน", "center"),
         (f"( {fin} )", "center")],
    ])
    _p(doc, "ความเห็นของผู้บริหารสถานศึกษา   (   ) อนุมัติ", indent=1.25, before=4, after=10)
    _p(doc, "(ลงชื่อ)...........................................", align="center", after=0)
    _p(doc, f"( {director} )", align="center", after=0)
    _p(doc, f"ผู้อำนวยการ{sname}", align="center", after=0)

    # ===== 2. ใบสำคัญรับเงิน =====
    doc.add_page_break()
    _p(doc, "ใบสำคัญรับเงิน", align="center", bold=True, size=18, after=4)
    _p(doc, f"{sname}  {saddr}", align="right", after=0)
    _p(doc, f"วันที่ {_dnum(inst.inspect_date or inst.end_date)}", align="right", after=6)
    _p(doc, f"ข้าพเจ้า {vname} บ้านเลขที่ {vaddr} ได้รับเงินจาก {sname} ดังรายการต่อไปนี้",
       align="justify", indent=1.25, after=4)
    _simple_table(doc, ["ลำดับที่", "รายการ", "จำนวนเงิน"],
                  [["1", f"ค่าจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ) {period}", A],
                   ["", "รวมเงิน", A]],
                  [Cm(1.6), Cm(10.4), Cm(4.0)])
    _p(doc, f"(ตัวอักษร)  ({bahttext(amt)})", indent=1.25, before=2, after=12)
    _sign_table(doc, [
        [("(ลงชื่อ)...........................................ผู้รับเงิน", "center"),
         (f"( {vname} )", "center")],
        [("(ลงชื่อ)...........................................ผู้จ่ายเงิน", "center"),
         (f"( {fin} )", "center")],
    ])

    # ===== 3. หนังสือรับรองการหักภาษี ณ ที่จ่าย =====
    doc.add_page_break()
    _p(doc, "หนังสือรับรองการหักภาษี ณ ที่จ่าย", align="center", bold=True, size=18, after=2)
    _p(doc, "ตามมาตรา ๕๐ ทวิ แห่งประมวลรัษฎากร", align="center", after=8)
    _p(doc, "ผู้มีหน้าที่หักภาษี ณ ที่จ่าย :", bold=True, after=0)
    _p(doc, f"ส่วนราชการ {sname}   เลขประจำตัวผู้เสียภาษี {getattr(school,'tax_id','') or _BLANK}", after=0)
    _p(doc, f"ที่อยู่ {saddr or _BLANK}", after=0)
    _p(doc, f"ขอรับรองว่าได้หักภาษี ณ ที่จ่าย ตามใบสั่งจ้าง เลขที่ {order_no}", after=6)
    _p(doc, "ผู้ถูกหักภาษี ณ ที่จ่าย :", bold=True, after=0)
    _p(doc, f"ชื่อ {vname}   เลขประจำตัวประชาชน {vtax}", after=0)
    _p(doc, f"ที่อยู่ {vaddr}", after=6)
    _simple_table(doc, ["ประเภทเงินได้ที่จ่าย", "วันที่จ่าย", "จำนวนเงินที่จ่าย", "ภาษีที่หัก"],
                  [["ค่าจ้างเหมาประกอบอาหารกลางวัน", _dnum(inst.inspect_date or inst.end_date), A, W],
                   ["รวม", "", A, W]],
                  [Cm(6.4), Cm(3.2), Cm(3.2), Cm(3.2)])
    _p(doc, f"รวมเงินภาษีที่หัก (ตัวอักษร)  ({bahttext(wht)})", indent=1.25, before=2, after=12)
    _p(doc, "(ลงชื่อ)...........................................ผู้จ่ายเงิน", align="center", after=0)
    _p(doc, f"( {director} )", align="center", after=0)
    _p(doc, f"ผู้อำนวยการ{sname}", align="center", after=0)

    return _save(doc, f"ขอเบิกจ่าย_งวดที่{inst.seq}_ปี{prog.year}")


def _student_tiers(prog):
    """แยกนักเรียนเป็น 2 กลุ่มตามใบสั่งจ้าง: (อนุบาล-ประถม) และ (มัธยม)"""
    t1 = sum((c.num_students or 0) for c in prog.classes
             if (c.level or "").startswith(("อ", "ป")))
    t2 = sum((c.num_students or 0) for c in prog.classes
             if (c.level or "").startswith("ม"))
    return t1, t2


def _cell_lines(cell, lines, *, align="left", size=13, bold=False):
    """เติมข้อความหลายบรรทัดในเซลล์ (แต่ละ element = 1 บรรทัด)"""
    amap = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT}
    cell.text = ""
    for i, ln in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = amap[align]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(ln)
        r.font.name = THAI_FONT
        _csize(r, size)
        _bcs(r, bold)
        r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)


def _order_item_table(doc, desc_head, tiers, total, baht):
    """ตารางรายการใบสั่งจ้าง (merge สวยตามแบบ):
    หัว 6 คอลัมน์ · แถวรายการ 1 แถว (รายการมีหัวเรื่อง+รายการย่อยหลายบรรทัด, คอลัมน์ตัวเลขเรียงตามชั้น)
    · แถวสรุป รวมเป็นเงิน/ภาษี/รวมทั้งสิ้น (merge คอลัมน์ซ้าย) · แถวตัวอักษร (merge ทั้งแถว)
    tiers = [(ชื่อชั้น, 'x คน', 'y บาท/วัน', 'z', 'จำนวนเงิน'), ...]"""
    widths = [Cm(1.2), Cm(6.3), Cm(1.8), Cm(2.4), Cm(1.6), Cm(2.2)]
    headers = ["ลำดับ", "รายการ", "จำนวน", "ราคาต่อหน่วย", "จำนวนวัน", "จำนวนเงิน (บาท)"]
    t = doc.add_table(rows=1, cols=6)
    t.style = "Table Grid"
    t.autofit = False
    for c, h, w in zip(t.rows[0].cells, headers, widths):
        _set_cell(c, h, bold=True, align="center", size=13)
        c.width = w
    # แถวรายการเดียว: รายการมีหัวเรื่อง + รายการย่อยตามชั้น, คอลัมน์ตัวเลขเรียงตามชั้น
    row = t.add_row().cells
    for c, w in zip(row, widths):
        c.width = w
    _set_cell(row[0], "1", align="center", size=13)
    desc = [desc_head] + [f"{i}. {tt[0]}" for i, tt in enumerate(tiers, 1)]
    _cell_lines(row[1], desc, align="left", size=13)
    _cell_lines(row[2], [tt[1] for tt in tiers], align="center", size=13)
    _cell_lines(row[3], [tt[2] for tt in tiers], align="center", size=13)
    _cell_lines(row[4], [tt[3] for tt in tiers], align="center", size=13)
    _cell_lines(row[5], [tt[4] for tt in tiers], align="right", size=13)

    def summary(label, amount, bold=False):
        r = t.add_row().cells
        for c, w in zip(r, widths):
            c.width = w
        merged = r[0].merge(r[1]).merge(r[2]).merge(r[3]).merge(r[4])
        _set_cell(merged, label, align="right", size=13, bold=bold)
        _set_cell(r[5], amount, align="right", size=13, bold=bold)

    summary("รวมเป็นเงิน", total)
    summary("ภาษีมูลค่าเพิ่ม", "-")
    summary("รวมเป็นเงินทั้งสิ้น", total, bold=True)
    # แถวตัวอักษร (merge ทั้งแถว)
    r = t.add_row().cells
    merged = r[0]
    for k in range(1, 6):
        merged = merged.merge(r[k])
    _set_cell(merged, f"(ตัวอักษร)  {baht}", align="center", size=13, bold=True)
    return t


def render_order_doc(rnd, school, doc=None) -> str:
    """ใบสั่งจ้างเหมาประกอบอาหารกลางวัน (สัญญา 1 รอบ) - รูปแบบตามแบบฟอร์มจริงของโรงเรียน (เลขอารบิก)"""
    doc, own = _begin(doc)
    prog = rnd.program
    vendor = rnd.vendor
    vname = vendor.name if vendor else _BLANK
    vaddr = (getattr(vendor, "address", "") or "").strip() if vendor else ""
    vphone = (getattr(vendor, "phone", "") or "").strip() if vendor else ""
    vtax = (getattr(vendor, "tax_id", "") or "").strip() if vendor else ""
    sname = _school_disp(school)
    saddr = (school.address or "").strip()
    director = (school.director_name or "").strip() or _BLANK
    order_no = _doc_no(rnd, "hire-order", (rnd.order_no or "").strip() or _BLANK)
    order_dt = _doc_dt(rnd, "hire-order", "date") or rnd.order_date
    total = round(float(rnd.amount or 0), 2)
    money, baht = _money(total), bahttext(total)
    rate = prog.rate_per_head or 0
    days = rnd.days or 0
    t1, t2 = _student_tiers(prog)
    insts = sorted(rnd.installments or [], key=lambda i: i.seq)
    n_inst = len(insts)
    per_days = (insts[0].days if insts else 0)
    ds, de = _dnum(rnd.start_date), _dnum(rnd.end_date)
    term = 1 if (rnd.start_date and rnd.start_date.month in (5, 6, 7, 8, 9, 10)) else 2

    _p(doc, "ใบสั่งจ้าง", align="center", bold=True, size=20, after=4)
    # หัว: ผู้รับจ้าง (ซ้าย) | ใบสั่งจ้าง เลขที่/วันที่ + โรงเรียน (ขวา)
    left = (f"ผู้รับจ้าง  {vname}\nที่อยู่ {vaddr or _BLANK}\n"
            f"โทรศัพท์ {vphone or '-'}\nเลขประจำตัวผู้เสียภาษี {vtax or '-'}")
    right = (f"ใบสั่งจ้าง เลขที่ {order_no}\nวันที่ {_dnum(order_dt)}\n{sname}\n{saddr}")
    _simple_table(doc, ["ผู้รับจ้าง", "ใบสั่งจ้าง"], [[left, right]], [Cm(8.5), Cm(7.0)])
    _p(doc, f"ตามที่ {vname} ได้เสนอราคาไว้ต่อ{sname} ซึ่งได้รับราคาและตกลงจ้าง ตามรายการดังต่อไปนี้",
       align="justify", indent=1.25, before=4, after=4)

    # ตารางรายการ (merge สวยตามแบบ) - หัวเรื่องงานจ้าง + รายการย่อยตามระดับชั้น
    desc_head = (f"ดำเนินการจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ) สำหรับนักเรียน{('ระดับอนุบาลถึงชั้นมัธยมศึกษาปีที่ 3' if t2 else 'ระดับอนุบาลถึงชั้นประถมศึกษาปีที่ 6')} "
                 f"ประจำภาคเรียนที่ {term} ปีการศึกษา {prog.year} รอบ {rnd.seq} "
                 f"ระหว่างวันที่ {ds} ถึงวันที่ {de} ดังนี้")
    tiers = []
    if t1:
        tiers.append(("ระดับอนุบาล-ระดับประถมศึกษา", f"{t1} คน", f"{_money(rate)} บาท/วัน", str(days), _money(t1 * rate * days)))
    if t2:
        tiers.append(("ระดับมัธยมศึกษา", f"{t2} คน", f"{_money(rate)} บาท/วัน", str(days), _money(t2 * rate * days)))
    if not tiers:
        tiers.append((f"นักเรียน {prog.total_students} คน", f"{prog.total_students} คน", f"{_money(rate)} บาท/วัน", str(days), money))
    _order_item_table(doc, desc_head, tiers, money, baht)

    _p(doc, "การสั่งจ้าง อยู่ภายใต้เงื่อนไขต่อไปนี้", bold=True, indent=0.5, before=4, after=0)
    _p(doc, f"1. กำหนดส่งมอบภายใน ตามงวดงาน {n_inst or '-'} งวดงาน งวดงานละ {per_days or '-'} วัน "
            f"รวม {days} วัน นับถัดจากวันที่ผู้รับจ้างได้รับใบสั่งจ้าง", align="justify", indent=1)
    _p(doc, "2. ครบกำหนดส่งมอบวันที่ - ตามงวดงานที่กำหนด", indent=1)
    _p(doc, f"3. สถานที่ส่งมอบ {sname}", indent=1)
    _p(doc, "4. ระยะเวลารับประกัน -", indent=1)
    _p(doc, "5. สงวนสิทธิ์ค่าปรับกรณีส่งมอบเกินกำหนด โดยคิดค่าปรับเป็นรายวันในอัตราร้อยละ 0.20 "
            "ของมูลค่าตามใบสั่งจ้าง", align="justify", indent=1)
    _p(doc, "6. โรงเรียนสงวนสิทธิ์ที่จะไม่รับมอบ ถ้าปรากฏว่างานจ้างนั้นมีลักษณะไม่ตรงตามรายการที่ระบุไว้"
            "ในใบสั่งจ้าง กรณีนี้ผู้รับจ้างจะต้องดำเนินการเปลี่ยนใหม่ให้ถูกต้องตามใบสั่งจ้างทุกประการ",
       align="justify", indent=1)
    _p(doc, "7. กรณีงานจ้าง ผู้รับจ้างจะต้องไม่เอางานทั้งหมดหรือแต่บางส่วนแห่งสัญญานี้ไปจ้างช่วงอีกทอดหนึ่ง "
            "เว้นแต่การจ้างช่วงงานแต่บางส่วนที่ได้รับอนุญาตเป็นหนังสือจากผู้ว่าจ้างแล้ว หากฝ่าฝืนผู้รับจ้างต้อง"
            "ชำระค่าปรับให้แก่ผู้ว่าจ้างเป็นจำนวนเงินในอัตราร้อยละ 10 (สิบ) ของวงเงินของงานที่จ้างช่วงตามสัญญา "
            "ทั้งนี้ ไม่ตัดสิทธิผู้ว่าจ้างในการบอกเลิกสัญญา", align="justify", indent=1)
    _p(doc, f"8. การส่งมอบงานและการจ่ายเงิน ผู้รับจ้างต้องสรุปรายการประกอบอาหารให้แก่ผู้ว่าจ้างเพื่อทำการ"
            f"เบิกจ่ายเงิน จำนวน {n_inst or '-'} งวด โดยมีรายละเอียด ดังนี้", align="justify", indent=1)
    if insts:
        for i in insts:
            amt = round(float(i.amount or 0), 2)
            last = " (งวดสุดท้าย)" if i.seq == n_inst else ""
            _p(doc, f"งวดที่ {i.seq}{last} จ่ายเป็นเงิน {_money(amt)} บาท ({bahttext(amt)}) เมื่อได้ดำเนินการ"
                    f"ส่งมอบงานงวดที่ {i.seq} ให้แล้วเสร็จภายใน {_dnum(i.end_date)}",
               align="justify", indent=1.25, after=0)
    else:
        _p(doc, "(ยังไม่ได้แบ่งงวด - เพิ่มงวดในหน้าจัดการงวด)", indent=1.25, after=0)
    _p(doc, "9. เงื่อนไขการสั่งจ้างและการสิ้นสุดของใบสั่งจ้าง", indent=1, before=2, after=0)
    _p(doc, f"9.1 ในการสั่งจ้างประกอบอาหารในครั้งนี้ กำหนดมูลค่าตามใบสั่งจ้างให้อยู่ภายในวงเงิน {money} บาท "
            f"({baht}) โดยมีระยะเวลาดำเนินการภายใน {days} วัน นับถัดจากวันลงนามในใบสั่งจ้าง",
       align="justify", indent=1.5, after=0)
    _p(doc, "9.2 การสั่งจ้างประกอบอาหารตามใบสั่งจ้าง จะสิ้นสุดลงเมื่อโรงเรียนเบิกจ่ายครบมูลค่าตามใบสั่งจ้าง "
            "หรือครบกำหนดระยะเวลาตามที่กำหนดในข้อ 9.1 แล้วแต่เงื่อนไขใดถึงก่อน", align="justify", indent=1.5)
    _p(doc, "10. เอกสารแนบท้ายใบสั่งจ้าง", indent=1, after=0)
    _p(doc, f"10.1 ขอบเขตของงาน การจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ) ประจำภาคเรียนที่ {term} "
            f"ปีการศึกษา {prog.year} รอบ {rnd.seq} (ระหว่างวันที่ {ds} ถึงวันที่ {de} จำนวน {days} วัน)",
       align="justify", indent=1.5, after=0)
    _p(doc, "10.2 ใบเสนอราคา", indent=1.5, after=0)
    _p(doc, "เอกสารแนบท้ายใบสั่งจ้างให้ถือเป็นเอกสารส่วนหนึ่งของใบสั่งจ้างฉบับนี้ กรณีเอกสารแนบท้ายขัดหรือแย้ง"
            "กับใบสั่งจ้างฉบับนี้ ให้ใช้ใบสั่งจ้างนี้บังคับ", align="justify", indent=1, after=14)
    _sign_table(doc, [
        [("(ลงชื่อ)...........................................ผู้สั่งจ้าง", "center"),
         (f"( {director} )", "center"),
         (f"ผู้อำนวยการ{sname}", "center"),
         (f"วันที่ {_dnum(order_dt)}", "center")],
        [("(ลงชื่อ)...........................................ผู้รับใบสั่งจ้าง", "center"),
         (f"( {vname} )", "center"),
         (f"วันที่ {_dnum(order_dt)}", "center"),
         ("", "center")],
    ])
    return _finish(doc, own, f"ใบสั่งจ้าง_รอบที่{rnd.seq}_ปี{prog.year}")


_COM_ORDER = [
    ("tor", "แต่งตั้งคณะกรรมการจัดทำขอบเขตของงาน (TOR) การจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ)",
     "จัดทำขอบเขตของงาน (TOR) การจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ) ให้ถูกต้องครบถ้วน"),
    ("control", "แต่งตั้งคณะกรรมการควบคุมงานจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ)",
     "ควบคุมงานจ้างเหมาประกอบอาหารกลางวัน ตรวจสอบคุณภาพ ความสะอาด และปริมาณอาหารเป็นรายวัน"),
    ("inspect", "แต่งตั้งคณะกรรมการตรวจรับการจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ)",
     "ตรวจรับงานจ้างเหมาประกอบอาหารกลางวัน ให้เป็นไปตามเงื่อนไขของสัญญาหรือข้อตกลง"),
]

# ชุดคำสั่งสำหรับซื้อวัตถุดิบเอง / จ้างแม่ครัว = ตรวจรับพัสดุ / ผู้ควบคุมประกอบอาหาร / ตรวจการประกอบอาหาร
_COM_ORDER_INGREDIENT = [
    ("inspect", "แต่งตั้งคณะกรรมการตรวจรับพัสดุ (วัตถุดิบเพื่อใช้ประกอบอาหารกลางวัน)",
     "ตรวจรับพัสดุ (วัตถุดิบ) ในการส่งมอบทุกครั้ง ให้ถูกต้องครบถ้วน ตามระเบียบฯ พ.ศ. 2560 ข้อ 175"),
    ("cook_control", "แต่งตั้งผู้ควบคุมรับผิดชอบในการประกอบอาหารกลางวัน",
     "ควบคุมรับผิดชอบการประกอบอาหารกลางวันให้ถูกสุขลักษณะ สะอาด และมีคุณค่าทางโภชนาการ"),
    ("food_inspect", "แต่งตั้งคณะกรรมการตรวจการประกอบอาหารกลางวัน",
     "ตรวจการประกอบอาหารกลางวันเป็นรายวัน ตรวจสอบความสะอาด คุณภาพ ปริมาณ และความทันเวลา"),
]


def _com_order_for(prog):
    """ชุดคำสั่งแต่งตั้งกรรมการตามวิธีดำเนินการของโครงการ + ชื่องานสำหรับข้อความนำ"""
    mode = getattr(prog, "operate_mode", "hire")
    if mode in ("ingredient", "person"):
        work = ("จัดซื้อวัตถุดิบเพื่อประกอบอาหารกลางวัน" if mode == "ingredient"
                else "จ้างเหมาประกอบอาหารกลางวัน (จ้างแม่ครัว)")
        return _COM_ORDER_INGREDIENT, work
    return _COM_ORDER, "จ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ)"


def render_committee_order_doc(rnd, school, doc=None) -> str:
    """คำสั่งแต่งตั้งคณะกรรมการ 3 ฉบับในไฟล์เดียว (TOR / ควบคุมงาน / ตรวจรับ)"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    director = (school.director_name or "").strip() or _BLANK
    period = (f"(ระหว่างวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)} "
              f"จำนวน {rnd.days or ''} วัน)")
    com_order, work = _com_order_for(prog)
    groups = {k: [m for m in rnd.committees if m.kind == k] for k, _, _ in com_order}
    # เลขที่/วันที่คำสั่งแยกรายฉบับ (doc_nos["cmd-<kind>"]) - แต่ละคำสั่งมีเลขของตัวเอง
    # fallback: เลขรวมเดิม doc_nos["committee"] -> command_no
    old_no = _doc_no(rnd, "committee", (getattr(rnd, "command_no", "") or "").strip())
    old_date = _doc_dt(rnd, "committee", "date") or getattr(rnd, "command_date", None) or rnd.order_date

    first = True
    for kind, subject, duty in com_order:
        members = groups.get(kind) or []
        cmd_no = _doc_no(rnd, f"cmd-{kind}", old_no)
        cmd_date = _doc_dt(rnd, f"cmd-{kind}", "date") or old_date
        if not first:
            doc.add_page_break()
        first = False
        _krut_center(doc)
        _p(doc, f"คำสั่ง{sname}", align="center", bold=True, size=18, after=0)
        _p(doc, f"ที่ {cmd_no or ('....../' + str(prog.year))}",
           align="center", bold=True, after=0)
        _p(doc, f"เรื่อง {subject}", align="center", bold=True, after=0)
        _p(doc, period, align="center", after=0)
        _p(doc, "─────────────────────", align="center", after=6)
        _p(doc, f"ด้วย{sname} จะดำเนินการ{work}ให้บริการแก่นักเรียน "
                "เพื่อให้การดำเนินการดังกล่าวเป็นไปด้วยความเรียบร้อย บังเกิดผลดีแก่ทางราชการ "
                "จึงอาศัยอำนาจตามระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ "
                "พ.ศ. ๒๕๖๐ แต่งตั้งบุคคลต่อไปนี้เป็นคณะกรรมการ", align="justify", indent=1.25, after=4)
        if members:
            for i, m in enumerate(members, 1):
                _p(doc, f"{i}. {m.name}        ตำแหน่ง {m.position}        {m.role}",
                   indent=1.5, after=0)
        else:
            for i in range(1, 4):
                _p(doc, f"{i}. ...........................................        ตำแหน่ง ..................        "
                        f"{'ประธานกรรมการ' if i == 1 else 'กรรมการ'}", indent=1.5, after=0)
        _p(doc, f"ให้คณะกรรมการที่ได้รับแต่งตั้ง {duty} และปฏิบัติหน้าที่ให้ถูกต้องตามระเบียบ"
                "ของทางราชการอย่างเคร่งครัด", align="justify", indent=1.25, before=4)
        _p(doc, "ทั้งนี้ ตั้งแต่บัดนี้เป็นต้นไป", bold=True, indent=1.25, after=6)
        _p(doc, f"สั่ง ณ วันที่ {_dnum(cmd_date)}", align="center", after=14)
        _p(doc, "(ลงชื่อ)...........................................", align="center", after=0)
        _p(doc, f"( {director} )", align="center", after=0)
        _p(doc, f"ผู้อำนวยการ{sname}", align="center", after=0)

    return _finish(doc, own, f"คำสั่งแต่งตั้งกรรมการ_รอบที่{rnd.seq}_ปี{prog.year}")


def _committee_lines(doc, members, fallback_n=3):
    """รายชื่อกรรมการเป็นตารางไร้เส้นขอบ ให้ ชื่อ/ตำแหน่ง/บทบาท ตรงคอลัมน์กัน"""
    data = list(members) if members else [None] * fallback_n
    widths = [Cm(1.0), Cm(6.25), Cm(4.8), Cm(4.2)]   # รวม 16.25 = พื้นที่พิมพ์ A4
    t = doc.add_table(rows=len(data), cols=4)
    _no_borders(t)
    for i, (row, m) in enumerate(zip(t.rows, data), 1):
        if m:
            name, pos, role = m.name, ("ตำแหน่ง " + (m.position or "ครู")), m.role
        else:
            name = "..........................................."
            pos = "ตำแหน่ง .................."
            role = "ประธานกรรมการ" if i == 1 else "กรรมการ"
        for c, v, w in zip(row.cells, [f"{i}.", name, pos, role], widths):
            _set_cell(c, v, size=16, align="left")
            c.width = w


def render_hire_report_doc(rnd, school, doc=None) -> str:
    """รายงานขอจ้างเหมาประกอบอาหารกลางวัน (บันทึกข้อความเปิดเรื่อง)"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    saddr = (school.address or "").strip()
    director = (school.director_name or "").strip() or _BLANK
    officer = (school.officer_name or "").strip() or _BLANK
    head = (school.head_officer_name or "").strip() or _BLANK
    fund = (prog.funding_org or "").strip() or "องค์กรปกครองส่วนท้องถิ่น"
    total = round(float(rnd.amount or 0), 2)
    rate = prog.rate_per_head or 0
    days = rnd.days or 0
    t1, t2 = _student_tiers(prog)
    period = (f"ระหว่างวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)} "
              f"จำนวน {days} วัน")

    dr = f"ประจำวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)}"
    _memo_head(doc, school,
               [f"รายงานขอจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ) ประจำปีการศึกษา {prog.year} "
                f"({dr} จำนวน {days} วัน)"],
               *_memo_ref2(rnd, "hire-report"))
    _p(doc, f"ด้วย{sname} จ้างเหมาประกอบอาหาร (ปรุงสำเร็จ) ให้แก่นักเรียนรับประทาน {dr} "
            f"ปีการศึกษา {prog.year} การจัดจ้างครั้งนี้ดำเนินการโดยวิธีเฉพาะเจาะจงตามมาตรา 56 (2) (ข) "
            "ประกอบหนังสือกระทรวงการคลัง ด่วนที่สุด ที่ กค (กวจ) 0405.2/ว 116 ลงวันที่ 12 มีนาคม 2562 "
            "ซึ่งมีรายละเอียดดังต่อไปนี้", align="justify", indent=1.25)
    _p(doc, "1. เหตุผลและความจำเป็นที่ต้องจ้าง", bold=True, indent=1.25, after=0)
    _p(doc, "เพื่อประกอบอาหารกลางวัน (ปรุงสำเร็จ) ให้แก่นักเรียนระดับอนุบาลจนถึงชั้นประถมศึกษาปีที่ 6",
       align="justify", indent=1.5)
    _p(doc, "2. ขอบเขตของงานที่จะจ้าง", bold=True, indent=1.25, after=0)
    _p(doc, f"การจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ) {dr} (รายละเอียดตามเอกสารแนบ)",
       align="justify", indent=1.5)
    _p(doc, "3. ราคากลางของงานที่จะจ้าง", bold=True, indent=1.25, after=0)
    _p(doc, f"เป็นเงิน {_money(total)} บาท ({bahttext(total)}) โดยมีแหล่งที่มาจาก{fund}", indent=1.5)
    _p(doc, "4. วงเงินที่จะจ้าง", bold=True, indent=1.25, after=0)
    _p(doc, f"เป็นเงิน {_money(total)} บาท ({bahttext(total)})", indent=1.5)
    _p(doc, "5. กำหนดเวลาที่ต้องการให้งานนั้นแล้วเสร็จ", bold=True, indent=1.25, after=0)
    _p(doc, f"ระยะเวลาการจ้าง จำนวน {days} วัน ตั้งแต่วันที่ {_dnum(rnd.start_date)} "
            f"ถึงวันที่ {_dnum(rnd.end_date)}", indent=1.5)
    _p(doc, "6. วิธีที่จะจ้าง และเหตุผลที่จะต้องจ้างโดยวิธีนั้น", bold=True, indent=1.25, after=0)
    _p(doc, "ดำเนินการด้วยวิธีเฉพาะเจาะจง เนื่องจากการจัดซื้อจัดจ้างพัสดุที่มีการผลิต จำหน่าย หรือให้บริการ "
            "ทั่วไป และมีวงเงินในการจัดซื้อจัดจ้างครั้งหนึ่งไม่เกินวงเงินตามที่กำหนดในกฎกระทรวง",
       align="justify", indent=1.5)
    _p(doc, "7. หลักเกณฑ์การพิจารณาคัดเลือกข้อเสนอ", bold=True, indent=1.25, after=0)
    _p(doc, "การพิจารณาคัดเลือกข้อเสนอโดยใช้เกณฑ์ราคา", indent=1.5)
    _p(doc, "8. การขออนุมัติแต่งตั้งคณะกรรมการต่าง ๆ", bold=True, indent=1.25, after=0)
    _p(doc, "8.1 แต่งตั้งผู้ควบคุมและคณะกรรมการตรวจการประกอบอาหาร", indent=1.5, after=0)
    _committee_lines(doc, [m for m in rnd.committees if m.kind == "control"])
    _p(doc, "8.2 คณะกรรมการตรวจรับงานจ้าง/ผู้ตรวจรับงานจ้าง", indent=1.5, before=2, after=0)
    _committee_lines(doc, [m for m in rnd.committees if m.kind == "inspect"])
    _p(doc, "โดยให้คณะกรรมการตรวจรับงานจ้าง/ผู้ตรวจรับงานจ้างที่ได้รับการแต่งตั้ง ปฏิบัติหน้าที่ตามระเบียบ"
            "กระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560 ข้อ 175",
       align="justify", indent=1.5)
    _p(doc, "จึงเรียนมาเพื่อโปรดพิจารณา หากเห็นชอบขอได้โปรด", indent=1.25, before=2, after=0)
    _p(doc, f"1. อนุมัติให้ดำเนินการจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ) {dr} ตามรายงานขอจ้างข้างต้น",
       align="justify", indent=1.5, after=0)
    _p(doc, "2. อนุมัติให้แต่งตั้งคณะกรรมการ ตามข้อ 8.1 และ 8.2", indent=1.5, after=10)
    _sign_table(doc, [[
        ("ลงชื่อ ..............................................เจ้าหน้าที่", "center"),
        (f"( {officer} )", "center")]])
    _p(doc, "ความเห็นของหัวหน้าเจ้าหน้าที่ ......................................................................",
       indent=1.25, before=2, after=8)
    _sign_table(doc, [[
        ("ลงชื่อ ..............................................หัวหน้าเจ้าหน้าที่", "center"),
        (f"( {head} )", "center")]])
    _p(doc, "คำสั่ง   เห็นชอบ / อนุมัติ / ลงนามแล้ว", align="center", bold=True, before=4, after=8)
    _sign_table(doc, [[
        ("ลงชื่อ ..............................................ผู้อำนวยการโรงเรียน", "center"),
        (f"( {director} )", "center")]])
    return _finish(doc, own, f"รายงานขอจ้าง_รอบที่{rnd.seq}_ปี{prog.year}")


def _memo_head(doc, school, subject_lines, date, doc_no=None):
    """หัวบันทึกข้อความมาตรฐานสารบรรณ: ครุฑ + หัวข้อตัวหนา (ส่วนราชการ/ที่/วันที่/เรื่อง/เรียน) + เส้นคั่น"""
    sname = _school_disp(school)
    office = (f"{sname}  " + (school.address or "").strip()).strip()
    _krut_and_title(doc)
    _p_runs(doc, [("ส่วนราชการ  ", True), (office, False)], after=0)
    _p_runs(doc, [("ที่  ", True), ((doc_no or "").strip() or _BLANK, False),
                  ("\t", False), ("วันที่  ", True), (date, False)], tab_cm=8, after=0)
    for i, s in enumerate(subject_lines):
        if i == 0:
            _p_runs(doc, [("เรื่อง  ", True), (s, False)], after=0)
        else:
            _p(doc, "        " + s, after=0)
    _p_runs(doc, [("เรียน  ", True), (f"ผู้อำนวยการ{sname}", False)], after=0)
    _hr(doc)


def _hdr_memo(doc, school, prog, subject_lines, date, doc_no=None):
    _memo_head(doc, school, subject_lines, date, doc_no)


def render_winner_doc(rnd, school, doc=None) -> str:
    """ประกาศผู้ชนะการเสนอราคา (จ้างเหมาอาหารกลางวัน)"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    director = (school.director_name or "").strip() or _BLANK
    vname = rnd.vendor.name if rnd.vendor else _BLANK
    total = round(float(rnd.amount or 0), 2)
    period = (f"ระหว่างวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)} "
              f"จำนวน {rnd.days or ''} วัน")
    dr = f"ประจำวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)} ({rnd.days or ''} วันทำการ)"
    _krut_center(doc)
    _p(doc, f"ประกาศ{sname}", align="center", bold=True, size=18, after=0)
    _p(doc, "เรื่อง ประกาศผู้ชนะการเสนอราคา สำหรับการจ้างประกอบอาหารกลางวัน (ปรุงสำเร็จ)",
       align="center", bold=True, after=0)
    _p(doc, f"{dr}", align="center", after=0)
    _p(doc, "โดยวิธีเฉพาะเจาะจง", align="center", after=0)
    _p(doc, "-------------------------------", align="center", after=6)
    _p(doc, f"ตามที่{sname} โดย{director} ได้มีโครงการจ้างประกอบอาหารกลางวัน (ปรุงสำเร็จ) "
            f"{dr} โดยวิธีเฉพาะเจาะจง นั้น", align="justify", indent=1.25)
    _p(doc, f"โครงการจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ) {dr} ผู้ได้รับการคัดเลือก ได้แก่ {vname} "
            f"โดยเสนอราคาเป็นเงินทั้งสิ้น {_money(total)} บาท ({bahttext(total)}) รวมภาษีมูลค่าเพิ่ม "
            "และภาษีอื่น ค่าขนส่ง ค่าจดทะเบียน และค่าใช้จ่ายอื่น ๆ ทั้งปวง",
       align="justify", indent=1.25, after=10)
    _p(doc, f"ประกาศ ณ วันที่ {_dnum(_doc_dt(rnd, 'winner', 'date') or rnd.order_date)}", align="center", after=14)
    _p(doc, "(ลงชื่อ)...........................................", align="center", after=0)
    _p(doc, f"( {director} )", align="center", after=0)
    _p(doc, f"ผู้อำนวยการ{sname}", align="center", after=0)
    return _finish(doc, own, f"ประกาศผู้ชนะ_รอบที่{rnd.seq}_ปี{prog.year}")


def render_result_doc(rnd, school, doc=None) -> str:
    """รายงานผลการพิจารณาและขออนุมัติสั่งจ้าง"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    director = (school.director_name or "").strip() or _BLANK
    officer = (school.officer_name or "").strip() or _BLANK
    head = (school.head_officer_name or "").strip() or _BLANK
    vname = rnd.vendor.name if rnd.vendor else _BLANK
    total = round(float(rnd.amount or 0), 2)
    period = f"ระหว่างวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)} จำนวน {rnd.days or ''} วัน"
    dr = f"ประจำวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)} ({rnd.days or ''} วันทำการ)"
    _hdr_memo(doc, school, prog,
              ["รายงานผลการพิจารณาและขออนุมัติสั่งจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ)",
               dr], *_memo_ref2(rnd, "result"))
    _p(doc, f"ตามที่ผู้อำนวยการ{sname} เห็นชอบให้ดำเนินการจ้างเหมาประกอบอาหารกลางวัน {dr} "
            f"โดยวิธีเฉพาะเจาะจง วงเงินงบประมาณ {_money(total)} บาท ({bahttext(total)}) นั้น เจ้าหน้าที่ได้"
            "เจรจาตกลงราคากับผู้ประกอบการโดยตรงตามระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้างและ"
            "การบริหารพัสดุภาครัฐ พ.ศ. 2560 ข้อ 79 แล้ว ขอรายงานผลการพิจารณา ดังนี้",
       align="justify", indent=1.25, after=4)
    _simple_table(doc,
                  ["รายการพิจารณา", "ผู้ชนะการเสนอราคา", "ราคาที่เสนอ\n(รวม VAT)", "ราคาที่ตกลงจ้าง\n(รวม VAT)"],
                  [[f"ดำเนินการจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ) {dr}",
                    vname, _money(total), _money(total)],
                   ["รวม", "", _money(total), _money(total)]],
                  [Cm(6.05), Cm(4.2), Cm(3), Cm(3)])   # รวม 16.25 = พื้นที่พิมพ์ A4
    _p(doc, f"จึงเห็นสมควรรับราคาจาก {vname} การจัดจ้างคราวนี้ไม่เกินวงเงินที่ประมาณไว้และไม่สูงกว่า"
            "ราคากลาง เจ้าหน้าที่ได้ต่อรองราคาแล้ว ผู้เสนอราคาไม่สามารถลดราคาลงได้อีกตามใบเสนอราคาที่แนบ "
            f"กำหนดส่งมอบภายในระยะเวลาที่กำหนด สถานที่ส่งมอบ ณ {sname}",
       align="justify", indent=1.25, before=4)
    _p(doc, "จึงเรียนมาเพื่อโปรดพิจารณาอนุมัติให้ดำเนินการจัดจ้างจากผู้ชนะการเสนอราคาดังกล่าว และลงนาม"
            "ในประกาศรายชื่อผู้ชนะการเสนอราคา และใบสั่งจ้าง ที่เสนอมาพร้อมนี้", align="justify", indent=1.25, after=12)
    _sign_table(doc, [
        [("ลงชื่อ ..........................................เจ้าหน้าที่", "center"),
         (f"( {officer} )", "center")],
        [("ลงชื่อ ..........................................หัวหน้าเจ้าหน้าที่", "center"),
         (f"( {head} )", "center")],
    ])
    _p(doc, "อนุมัติ/ลงนามแล้ว", align="center", bold=True, before=4, after=8)
    _p(doc, "(ลงชื่อ)...........................................", align="center", after=0)
    _p(doc, f"( {director} )", align="center", after=0)
    _p(doc, f"ผู้อำนวยการ{sname}", align="center", after=0)
    return _finish(doc, own, f"รายงานผลพิจารณา_รอบที่{rnd.seq}_ปี{prog.year}")


def render_tor_request_doc(rnd, school, doc=None) -> str:
    """บันทึกข้อความขออนุมัติแต่งตั้งคณะกรรมการจัดทำ TOR"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    officer = (school.officer_name or "").strip() or _BLANK
    head = (school.head_officer_name or "").strip() or _BLANK
    period = f"ระหว่างวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)} จำนวน {rnd.days or ''} วัน"
    _hdr_memo(doc, school, prog,
              ["ขออนุมัติแต่งตั้งคณะกรรมการจัดทำขอบเขตของงาน (TOR) "
               "งานจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ)"], *_memo_ref2(rnd, "tor-request"))
    _p(doc, "ข้อเท็จจริง", bold=True, indent=1.25, after=0)
    _p(doc, f"{sname} ดำเนินการจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ) ประจำปีการศึกษา {prog.year} "
            f"({period}) โดยวิธีเฉพาะเจาะจง สำหรับนักเรียนระดับชั้นอนุบาลถึงระดับชั้นมัธยมศึกษาปีที่ ๓ "
            "ในโรงเรียนขยายโอกาสทางการศึกษา", align="justify", indent=1.5)
    _p(doc, "ข้อเสนอและข้อพิจารณา", bold=True, indent=1.25, after=0)
    _p(doc, "เพื่อให้เป็นไปตามระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ "
            "พ.ศ. ๒๕๖๐ ข้อ ๒๑ เห็นควรแต่งตั้งคณะกรรมการจัดทำขอบเขตของงาน (TOR) ดังรายชื่อต่อไปนี้",
       align="justify", indent=1.5)
    _committee_lines(doc, [m for m in rnd.committees if m.kind == "tor"])
    _p(doc, "จึงเรียนมาเพื่อโปรดพิจารณา", indent=1.25, before=4, after=12)
    _sign_table(doc, [
        [("(ลงชื่อ)...........................................เจ้าหน้าที่", "center"),
         (f"( {officer} )", "center")],
        [("(ลงชื่อ)...........................................หัวหน้าเจ้าหน้าที่", "center"),
         (f"( {head} )", "center")],
    ])
    return _finish(doc, own, f"ขออนุมัติTOR_รอบที่{rnd.seq}_ปี{prog.year}")


def _tor_committee_signs(doc, members):
    """ช่องลงชื่อคณะกรรมการจัดทำ TOR (ประธาน/กรรมการ/กรรมการและเลขานุการ) ตามที่กรอกไว้
    ไม่มีข้อมูล -> เว้นจุดไข่ปลา 3 คน"""
    DOT = "..............................................."
    rows = []
    if members:
        for m in members:
            rows.append((m.role or "กรรมการ", m.name, m.position or "ครู"))
    else:
        rows = [("ประธานกรรมการ", "", "ครู"), ("กรรมการ", "", "ครู"),
                ("กรรมการและเลขานุการ", "", "ครู")]
    for role, name, pos in rows:
        _p(doc, f"(ลงชื่อ) {DOT} {role}", align="center", before=6, after=0)
        _p(doc, f"( {name.strip() if name else '............................................'} )",
           align="center", after=0)
        _p(doc, f"ตำแหน่ง {pos}", align="center", after=0)


def render_tor_doc(rnd, school, doc=None) -> str:
    """ขอบเขตของงาน (TOR) การจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ)
    รูปแบบตามแบบฟอร์มจริงของโรงเรียน (9 หัวข้อ + ลงชื่อคณะกรรมการจัดทำ TOR
    + บันทึกรายงานผลการจัดทำ TOR ต่อท้ายในไฟล์เดียว)"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    total = round(float(rnd.amount or 0), 2)
    money, baht = _money(total), bahttext(total)
    days = rnd.days or 0
    rate = prog.rate_per_head or 0
    students = prog.total_students
    fund = (prog.funding_org or "").strip() or "องค์กรปกครองส่วนท้องถิ่น"
    area = (getattr(school, "area_office", "") or "").strip() or \
        "สำนักงานเขตพื้นที่การศึกษาประถมศึกษา............ เขต ...."
    insts = sorted(rnd.installments or [], key=lambda i: i.seq)
    t1, t2 = _student_tiers(prog)
    ds, de = _dnum(rnd.start_date), _dnum(rnd.end_date)
    term = 1 if (rnd.start_date and rnd.start_date.month in (5, 6, 7, 8, 9, 10)) else 2
    lvl = ("ระดับชั้นอนุบาล ถึงระดับชั้นมัธยมศึกษาปีที่ 3 ในโรงเรียนขยายโอกาส"
           if t2 else "ระดับชั้นอนุบาล ถึงระดับชั้นประถมศึกษาปีที่ 6")
    tor_members = [m for m in rnd.committees if m.kind == "tor"]

    # ---------- หัวเรื่อง ----------
    _p(doc, "ขอบเขตของงาน (TOR) การจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ)",
       align="center", bold=True, size=18, after=0)
    _p(doc, f"{sname}  ภาคเรียนที่ {term} ประจำปีการศึกษา {prog.year} (รอบ {rnd.seq})",
       align="center", after=0)
    _p(doc, f"ระหว่างวันที่ {ds} ถึง วันที่ {de}", align="center", after=0)
    _p(doc, f"สังกัด{area}  กระทรวงศึกษาธิการ", align="center", after=6)

    # ---------- 1. ความเป็นมา ----------
    _p(doc, "1. ความเป็นมา", bold=True, indent=0.5, after=0)
    _p(doc, "ประเทศไทยได้ให้ความสำคัญกับการจัดการอาหาร จึงได้กำหนดยุทธศาสตร์การจัดการด้านอาหารของประเทศไทย "
            "ที่สอดรับกับยุทธศาสตร์ชาติ 20 ปี เชื่อมโยงการดำเนินงานทุกมิติที่เกี่ยวข้องกับอาหารสู่โภชนาการ "
            "เพื่อสุขภาพของผู้บริโภค โดยน้อมนำหลักปรัชญาของเศรษฐกิจพอเพียงมาเป็นหลักคิดพื้นฐานในการปฏิบัติ "
            "โดยมุ่งให้ประเทศไทยมีความมั่นคงทางด้านอาหารและโภชนาการ ซึ่งโครงการอาหารกลางวันในโรงเรียน "
            "เป็นโครงการหนึ่งที่มีความสำคัญ ช่วยส่งเสริมให้นักเรียนซึ่งอยู่ในวัยที่กำลังเจริญเติบโตมีสุขภาพ "
            "พลานามัยดีขึ้น (อ้างอิงจาก : คู่มือการดำเนินงานอาหารกลางวัน, สำนักงานคณะกรรมการการศึกษาขั้นพื้นฐาน)",
       align="justify", indent=1, after=0)
    _p(doc, f"ดังนั้น{sname} จึงถือเป็นหน้าที่ที่จะพัฒนาสุขภาพอนามัยของนักเรียน ซึ่งเป็นพื้นฐานสำคัญ "
            "และเป็นปัจจัยแรกที่จะส่งผลต่อคุณภาพการเรียนรู้ของนักเรียน จึงเสนอร่างขอบเขตของงาน "
            "(Term of Reference : TOR) โครงการอาหารกลางวันขึ้น", align="justify", indent=1)

    # ---------- 2. วัตถุประสงค์ ----------
    _p(doc, "2. วัตถุประสงค์", bold=True, indent=0.5, after=0)
    for s in ["1. เพื่อให้นักเรียนได้รับประทานอาหารกลางวัน ที่มีคุณค่า และเพียงพอต่อความต้องการของร่างกาย",
              "2. เพื่อช่วยเหลือนักเรียนที่ขาดแคลนและยากจน ให้ได้รับประทานอาหารกลางวันทุกคน",
              "3. เพื่อให้นักเรียนมีสุขนิสัยที่ดีในการรับประทานอาหาร",
              "4. เพื่อสนับสนุนกิจกรรมการเรียนการสอนกลุ่มการงานอาชีพ"]:
        _p(doc, s, indent=1, after=0)

    # ---------- 3. คุณสมบัติของผู้เสนอราคา ----------
    _p(doc, "3. คุณสมบัติของผู้เสนอราคา (กำหนดตามแบบที่กรมบัญชีกลางกำหนด)", bold=True, indent=0.5, after=0)
    _p(doc, "(1) คุณสมบัติตามกฎหมาย (มาตรา 64 แห่งพระราชบัญญัติการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ "
            "พ.ศ. 2560) ผู้จะเข้ายื่นข้อเสนอในการจัดซื้อจัดจ้างของหน่วยงานของรัฐ อย่างน้อยต้องมีคุณสมบัติ "
            "และไม่มีลักษณะต้องห้าม ดังต่อไปนี้", align="justify", indent=1, after=0)
    for s in ["(1.1) มีความสามารถตามกฎหมาย",
              "(1.2) ไม่เป็นบุคคลล้มละลาย",
              "(1.3) ไม่อยู่ระหว่างเลิกกิจการ",
              "(1.4) ไม่เป็นบุคคลซึ่งอยู่ระหว่างถูกระงับการยื่นข้อเสนอหรือทำสัญญากับหน่วยงานของรัฐ ตามมาตรา 106 วรรคสาม",
              "(1.5) ไม่เป็นบุคคลซึ่งถูกแจ้งเวียนชื่อให้เป็นผู้ทิ้งงานของหน่วยงานของรัฐ ตามมาตรา 109",
              "(1.6) คุณสมบัติหรือลักษณะต้องห้ามอื่นตามที่คณะกรรมการนโยบายประกาศกำหนดในราชกิจจานุเบกษา"]:
        _p(doc, s, indent=1.5, after=0)
    _p(doc, "(2) คุณสมบัติตามที่คณะกรรมการนโยบายการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐกำหนด "
            "(หนังสือด่วนที่สุด ที่ กค (กนบ) 0405.2/ว 410 ลงวันที่ 24 ตุลาคม 2560)",
       align="justify", indent=1, after=0)
    for s in ["1. มีความสามารถตามกฎหมาย    2. ไม่เป็นบุคคลล้มละลาย    3. ไม่อยู่ระหว่างเลิกกิจการ",
              "4. ไม่อยู่ระหว่างถูกระงับการยื่นข้อเสนอหรือทำสัญญากับหน่วยงานของรัฐไว้ชั่วคราว",
              "5. ไม่เป็นผู้ทิ้งงานตามบัญชีรายชื่อผู้ทิ้งงานในระบบเครือข่ายสารสนเทศของกรมบัญชีกลาง",
              "6. มีคุณสมบัติและไม่มีลักษณะต้องห้ามตามที่คณะกรรมการนโยบายฯ กำหนดในราชกิจจานุเบกษา",
              "7. เป็นบุคคลธรรมดาหรือนิติบุคคลผู้มีอาชีพรับจ้างงานที่จัดจ้างในครั้งนี้",
              "8. ไม่เป็นผู้มีผลประโยชน์ร่วมกันกับผู้ยื่นข้อเสนอรายอื่น",
              "9. ไม่เป็นผู้ได้รับเอกสิทธิ์หรือความคุ้มกันซึ่งอาจปฏิเสธไม่ยอมขึ้นศาลไทย"]:
        _p(doc, s, indent=1.5, after=0)

    # ---------- 4. ขอบเขตการดำเนินงาน ----------
    _p(doc, "4. ขอบเขตการดำเนินงาน", bold=True, indent=0.5, before=2, after=0)
    _p(doc, f"ผู้รับจ้างต้องเป็นผู้รับผิดชอบการประกอบอาหารกลางวัน (ปรุงสำเร็จ) ประจำภาคเรียนที่ {term} "
            f"ปีการศึกษา {prog.year} รอบ {rnd.seq} ระหว่างวันที่ {ds} ถึง วันที่ {de} จำนวน {days} วัน "
            f"ให้แก่นักเรียน{lvl} ภายในวงเงิน {money} บาท ({baht}) โดยผู้รับจ้างต้อง"
            "ดำเนินการประกอบอาหารกลางวัน (ปรุงสำเร็จ) ในทุกวันทำการ ตามรายการอาหารที่กำหนด",
       align="justify", indent=1, after=0)
    _p(doc, "ทั้งนี้ การประกอบอาหารในแต่ละวันจะต้องเป็นรายการตามเมนู Thai School Lunch โดยโรงเรียนจะกำหนดให้"
            "เจ้าหน้าที่อาหารกลางวันเป็นผู้แจ้งรายการอาหารที่ต้องการ ให้ผู้รับจ้างทราบก่อนถึงวันประกอบอาหาร "
            "(ปรุงสำเร็จ) อย่างน้อย 5 วันทำการ", align="justify", indent=1)

    # ---------- 5. การส่งมอบงาน ----------
    _p(doc, "5. การส่งมอบงาน", bold=True, indent=0.5, after=0)
    _p(doc, f"โรงเรียนกำหนดการส่งมอบงานออกเป็นจำนวน {len(insts) or '.......'} งวด โดยมีรายละเอียด ดังนี้",
       indent=1, after=0)
    if insts:
        for it in insts:
            _p(doc, f"งวดที่ {it.seq} ผู้รับจ้างต้องประกอบอาหารกลางวัน (ปรุงสำเร็จ) และสรุปรายการประกอบอาหาร "
                    f"ระหว่างวันที่ {_dnum(it.start_date)} ถึงวันที่ {_dnum(it.end_date)} จำนวน {it.days or ''} วัน",
               align="justify", indent=1.25, after=0)
    _p(doc, "หมายเหตุ : อาจมีการเปลี่ยนแปลงวันตามประกาศโรงเรียน ในการหยุดเรียนหรือประกาศเปิดเรียนชดเชย",
       indent=1, before=2, after=0)
    _p(doc, "โดยในการส่งมอบงานแต่ละงวด ผู้รับจ้างจะต้องเป็นผู้สรุปรายการประกอบอาหารเป็นรายวันในงวดงานนั้น ๆ "
            "ส่งให้แก่โรงเรียนเพื่อทำการเบิกจ่ายเงิน ซึ่งโรงเรียนจะชำระเงินให้แก่ผู้รับจ้างตามที่ได้ประกอบ"
            "อาหารจริงในแต่ละงวด เมื่อคณะกรรมการตรวจรับได้ดำเนินการตรวจรับไว้ถูกต้องครบถ้วนแล้ว",
       align="justify", indent=1)

    # ---------- 6. หลักเกณฑ์ ----------
    _p(doc, "6. หลักเกณฑ์การพิจารณาคัดเลือกข้อเสนอ", bold=True, indent=0.5, after=0)
    _p(doc, "เกณฑ์ราคา", indent=1)

    # ---------- 7. วงเงินงบประมาณ ----------
    _p(doc, "7. วงเงินงบประมาณ", bold=True, indent=0.5, after=0)
    _p(doc, f"ภาคเรียนที่ {term} ปีการศึกษา {prog.year} รอบ {rnd.seq} ระหว่างวันที่ {ds} ถึงวันที่ {de} "
            f"จำนวน {days} วัน ได้รับจัดสรรงบประมาณจาก{fund} จำนวน {money} บาท ({baht}) "
            "รายละเอียด ดังนี้", align="justify", indent=1, after=0)
    if t1:
        _p(doc, f"ระดับชั้นอนุบาล-ประถมศึกษา จำนวนนักเรียน {t1} คน ในอัตราคนละ {_money(rate)} บาท/ต่อวัน "
                f"จำนวน {days} วัน เป็นเงิน {_money(t1 * rate * days)} บาท", indent=1.25, after=0)
    if t2:
        _p(doc, f"ระดับมัธยมศึกษา จำนวนนักเรียน {t2} คน ในอัตราคนละ {_money(rate)} บาท/ต่อวัน "
                f"จำนวน {days} วัน เป็นเงิน {_money(t2 * rate * days)} บาท", indent=1.25, after=0)

    # ---------- 8. งวดงานและการจ่ายเงิน ----------
    _p(doc, "8. งวดงานและการจ่ายเงิน", bold=True, indent=0.5, before=2, after=0)
    _p(doc, f"โรงเรียนกำหนดการจ่ายเงินออกเป็น จำนวน {len(insts) or '.......'} งวด โดยมีรายละเอียด ดังนี้",
       indent=1, after=0)
    if insts:
        for it in insts:
            amt = round(float(it.amount or 0), 2)
            _p(doc, f"งวดที่ {it.seq} จ่ายเป็นเงิน {_money(amt)} บาท ({bahttext(amt)}) เมื่อผู้รับจ้างประกอบ"
                    "อาหาร (ปรุงสำเร็จ) และสรุปรายการประกอบอาหารให้แก่โรงเรียน หลังจากวันสุดท้ายที่ได้ส่งมอบ "
                    "และมีการตรวจรับอาหารเสร็จเรียบร้อย", align="justify", indent=1.25, after=0)

    # ---------- 9. ค่าปรับ ----------
    _p(doc, "9. ค่าปรับ 0.20", bold=True, indent=0.5, before=2, after=0)
    _p(doc, "อ้างอิงตามหนังสือคณะกรรมการวินิจฉัยปัญหาการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ "
            "ที่ กค (กวจ) 0405.2/ว 116 ลงวันที่ 12 มีนาคม 2562 เรื่องแนวทางปฏิบัติในการจัดซื้อวัตถุดิบ "
            "เพื่อใช้ในการประกอบอาหาร การจ้างบุคคลเพื่อประกอบอาหาร หรือการจ้างเหมาประกอบอาหาร (ปรุงสำเร็จ) "
            "ข้อ 4 การกำหนดค่าปรับในสัญญาหรือข้อตกลง", align="justify", indent=1, after=10)

    # ---------- ลงชื่อคณะกรรมการจัดทำ TOR ----------
    _tor_committee_signs(doc, tor_members)

    # ---------- บันทึกข้อความ รายงานผลการจัดทำ TOR (ต่อท้ายในไฟล์เดียว) ----------
    doc.add_page_break()
    rep_date, rep_no = _memo_ref2(rnd, "tor-request")
    _memo_head(doc, school,
               ["รายงานผลการจัดทำขอบเขตของงาน (TOR) การจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ)"],
               rep_date, rep_no)
    cmd_no = _doc_no(rnd, "cmd-tor", "......../" + str(prog.year))
    cmd_dt = _doc_dt(rnd, "cmd-tor", "date")
    _p(doc, f"ตามคำสั่ง{sname} ที่ {cmd_no} ลงวันที่ {_dnum(cmd_dt)} ได้แต่งตั้งคณะกรรมการกำหนดจัดทำ"
            "ขอบเขตของงาน (TOR) การจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ) "
            f"จำนวน 1 โครงการ งบประมาณ {money} บาท ({baht}) นั้น", align="justify", indent=1.25, before=6)
    _p(doc, "บัดนี้ คณะกรรมการจัดทำขอบเขตของงาน (TOR) การจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ) "
            "ได้พิจารณาดำเนินการเรียบร้อยแล้ว รายละเอียดปรากฏตามเอกสารแนบท้ายรายงานฉบับนี้",
       align="justify", indent=1.25)
    _p(doc, "จึงเรียนมาเพื่อโปรดพิจารณาอนุมัติ", indent=1.25, after=10)
    _tor_committee_signs(doc, tor_members)
    return _finish(doc, own, f"ขอบเขตของงาน_TOR_รอบที่{rnd.seq}_ปี{prog.year}")


def render_quotation_doc(rnd, school, doc=None) -> str:
    """ใบเสนอราคา (ผู้รับจ้างเสนอราคาจ้างเหมาประกอบอาหารกลางวัน)"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    total = round(float(rnd.amount or 0), 2)
    rate = prog.rate_per_head or 0
    days = rnd.days or 0
    students = prog.total_students
    v = rnd.vendor
    vname = v.name if v else _BLANK
    vaddr = (getattr(v, "address", "") or "").strip() if v else ""
    vphone = (getattr(v, "phone", "") or "").strip() if v else ""
    vowner = (getattr(v, "owner_name", "") or "").strip() if v else ""
    vtax = (getattr(v, "tax_id", "") or "").strip() if v else ""
    period = f"ระหว่างวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)} จำนวน {days} วัน"

    dr = f"ประจำวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)}"
    officer = (school.officer_name or "").strip() or _BLANK
    _p(doc, "ใบเสนอราคา", align="center", bold=True, size=20, after=6)
    _p(doc, f"วันที่  {_dnum(rnd.order_date)}", align="right", after=4)
    _p(doc, f"เรียน  ผู้อำนวยการ{sname}", after=4)
    _p(doc, f"1. ข้าพเจ้า {vowner or vname} ซึ่งเป็นผู้มีอำนาจลงนามผูกพันสถานประกอบการ คือ {vname} "
            f"ตั้งอยู่เลขที่ {vaddr or _BLANK} เลขประจำตัวผู้เสียภาษี {vtax or _BLANK} ซึ่งได้ศึกษา"
            f"ทำความเข้าใจขอบเขตของงานการจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ) {dr} ของ{sname} "
            "และรายละเอียดต่าง ๆ โดยตลอดและยอมรับข้อกำหนดและเงื่อนไขนั้นแล้ว รวมทั้งรับรองว่าข้าพเจ้า"
            "เป็นผู้มีคุณสมบัติครบถ้วนตามที่กำหนด และไม่เป็นผู้ทิ้งงานของทางราชการ",
       align="justify", indent=1.25, after=2)
    _p(doc, f"2. ข้าพเจ้าขอเสนอราคาจ้างเหมาประกอบอาหารกลางวัน (ปรุงสำเร็จ) {dr} เป็นเงินทั้งสิ้น "
            f"{_money(total)} บาท (ตัวอักษร {bahttext(total)}) ซึ่งเป็นราคาที่รวมภาษีมูลค่าเพิ่ม รวมทั้ง"
            "ภาษีอากรอื่น และค่าใช้จ่ายทั้งปวงไว้ด้วยแล้ว", align="justify", indent=1.25, after=2)
    _p(doc, "3. คำเสนอนี้จะยืนอยู่เป็นระยะเวลา 30 วัน นับตั้งแต่วันที่ได้ยื่นใบเสนอราคา",
       align="justify", indent=1.25, after=2)
    _p(doc, f"4. กำหนดส่งมอบ {dr} นับถัดจากวันลงนามใบสั่งจ้าง/ข้อตกลงจ้าง",
       align="justify", indent=1.25, after=14)
    # แยกคำว่าตำแหน่งเป็นบรรทัดใต้เส้นลงชื่อ กัน "ราคา" ตกบรรทัด
    _sign_table(doc, [
        [("ลงชื่อ ...........................................", "center"),
         ("ผู้เจรจาตกลงราคา", "center"),
         (f"( {officer} )", "center"),
         ("เจ้าหน้าที่", "center")],
        [("ลงชื่อ ...........................................", "center"),
         ("ผู้เสนอราคา", "center"),
         (f"( {vowner or vname} )", "center"),
         ("", "center")]])
    return _finish(doc, own, f"ใบเสนอราคา_รอบที่{rnd.seq}_ปี{prog.year}")


def render_contract_bundle(rnd, school) -> str:
    """ออกเอกสารต่อรอบทั้งชุดเป็นไฟล์ Word เดียว (เรียงตามลำดับงานจริง)"""
    doc = Document(); set_a4(doc)
    _font(doc)
    render_tor_request_doc(rnd, school, doc)        # บันทึกขออนุมัติ TOR
    render_committee_order_doc(rnd, school, doc)     # คำสั่งแต่งตั้งกรรมการ (3 ฉบับ)
    render_tor_doc(rnd, school, doc)                 # ขอบเขตของงาน (TOR)
    render_hire_report_doc(rnd, school, doc)         # รายงานขอจ้าง
    render_quotation_doc(rnd, school, doc)           # ใบเสนอราคา
    render_result_doc(rnd, school, doc)              # รายงานผลพิจารณา
    render_winner_doc(rnd, school, doc)              # ประกาศผู้ชนะ
    render_order_doc(rnd, school, doc)               # ใบสั่งจ้าง
    return _save(doc, f"ชุดเอกสารจ้างเหมา_รอบที่{rnd.seq}_ปี{rnd.program.year}")
