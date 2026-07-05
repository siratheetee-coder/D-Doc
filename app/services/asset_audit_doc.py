# -*- coding: utf-8 -*-
"""
asset_audit_doc.py — เอกสารตรวจสอบพัสดุประจำปี (ระเบียบกระทรวงการคลังฯ พ.ศ. 2560 ข้อ 213)
ชุดเดียวได้ 4 ส่วน: (1) บันทึกขอแต่งตั้งคณะกรรมการ (2) คำสั่งแต่งตั้ง
(3) บันทึกรายงานผลการตรวจสอบ (4) บัญชีครุภัณฑ์คงเหลือ (ดึงจากทะเบียนครุภัณฑ์)
สร้าง docx ตอนรันไทม์ ใช้ helper ร่วมกับ build_templates (ฟอนต์ TH Sarabun, ครุฑ)
ถ้อยคำอ้างระเบียบเป็นข้อความราชการสาธารณะ อิงไฟล์ตัวอย่างที่โรงเรียนใช้จริง
"""
from docx import Document
from docx.shared import Cm

from app.database import get_data_dir
from app.thai_utils import thai_date
from app.services.build_templates import (
    _font, _krut_and_title, _p, _p_runs, _sign_table, _set_cell, _hr,
    _repeat_header_row, _no_split_row, _no_borders,
)

_BLANK = "............................"


def _safe(text: str) -> str:
    for ch in '<>:"/\\|?*\n\r\t':
        text = text.replace(ch, "_")
    return text.strip()[:80]


def _save(doc, name: str) -> str:
    out_dir = get_data_dir() / "documents"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / (_safe(name) + ".docx")
    doc.save(str(out_path))
    return str(out_path)


def _office(school) -> str:
    return "  ".join(p for p in [(school.name or "").strip(),
                                 (school.address or "").strip()] if p)


def _director_line(school) -> str:
    name = (school.name or "").strip()
    return "ผู้อำนวยการ" + name if name.startswith("โรงเรียน") else "ผู้อำนวยการโรงเรียน"


def _members(ctx) -> list:
    """รายชื่อกรรมการ [{name, position, role}] (กรองที่มีชื่อ)"""
    return [m for m in ctx.get("members", []) if (m.get("name") or "").strip()]


def _member_lines(doc, members):
    """รายชื่อกรรมการเป็นตารางไร้เส้นขอบ ให้ ชื่อ/ตำแหน่ง/บทบาท ตรงคอลัมน์กัน"""
    data = list(members) if members else [None, None, None]
    widths = [Cm(1.0), Cm(6.6), Cm(4.8), Cm(4.2)]
    t = doc.add_table(rows=len(data), cols=4)
    _no_borders(t)
    for i, (row, m) in enumerate(zip(t.rows, data), 1):
        if m:
            name = (m.get("name") or "").strip()
            pos = "ตำแหน่ง " + ((m.get("position") or "ครู").strip())
            role = (m.get("role") or "กรรมการ").strip()
        else:
            name, pos = _BLANK, "ตำแหน่ง .................."
            role = "ประธานกรรมการ" if i == 1 else "กรรมการ"
        for c, v, w in zip(row.cells, [f"{i}.", name, pos, role], widths):
            _set_cell(c, v, size=16, align="left")
            c.width = w


def _memo_header(doc, school, subject, doc_no, date_txt):
    _krut_and_title(doc)
    _p_runs(doc, [("ส่วนราชการ  ", True), (_office(school), False)])
    _p_runs(doc, [("ที่  ", True), (doc_no or _BLANK, False),
                  ("\t", False), ("วันที่ ", True), (date_txt, False)], tab_cm=8)
    _p_runs(doc, [("เรื่อง  ", True), (subject, False)])
    _p_runs(doc, [("เรียน  ", True), (_director_line(school), False)])
    _hr(doc)


_LEGAL = ("อาศัยอำนาจตามความในข้อ 213 แห่งระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้างและ"
          "การบริหารพัสดุภาครัฐ พ.ศ. 2560 กำหนดให้ก่อนสิ้นเดือนกันยายนของทุกปี ให้หัวหน้า"
          "หน่วยงานของรัฐแต่งตั้งผู้รับผิดชอบในการตรวจสอบพัสดุซึ่งมิใช่เจ้าหน้าที่พัสดุ ทำการ"
          "ตรวจสอบการรับจ่ายพัสดุงวดตั้งแต่วันที่ 1 ตุลาคมปีก่อน จนถึงวันที่ 30 กันยายนปีปัจจุบัน "
          "และตรวจนับพัสดุคงเหลือ ณ วันสิ้นงวด ว่าการรับจ่ายถูกต้องหรือไม่ พัสดุคงเหลือมีอยู่ตรง"
          "ตามบัญชีหรือทะเบียนหรือไม่ มีพัสดุใดชำรุด เสื่อมสภาพ หรือสูญไปเพราะเหตุใด หรือพัสดุใด"
          "ไม่จำเป็นต้องใช้ในราชการต่อไป แล้วรายงานผลต่อผู้แต่งตั้งภายใน 30 วันทำการ")


def render_appoint_memo(school, ctx, doc=None):
    """(1) บันทึกข้อความ ขอแต่งตั้งคณะกรรมการตรวจสอบพัสดุประจำปี"""
    own = doc is None
    if own:
        doc = Document(); _font(doc)
    else:
        doc.add_page_break()
    year = ctx.get("year")
    head = (getattr(school, "head_officer_name", "") or "").strip() or _BLANK
    director = (school.director_name or "").strip() or _BLANK
    _memo_header(doc, school, f"การตรวจสอบพัสดุประจำปี ประจำปีงบประมาณ พ.ศ. {year}",
                 ctx.get("memo_no") or "", thai_date(ctx.get("date")))
    _p(doc, _LEGAL, align="justify", indent=1.25, after=2)
    _p(doc, f"ดังนั้น เพื่อให้การตรวจสอบการรับจ่ายพัสดุประจำปีงบประมาณ พ.ศ. {year} เป็นไปด้วย"
            "ความเรียบร้อยถูกต้องตามระเบียบดังกล่าวข้างต้น จึงขอแต่งตั้งบุคคลผู้มีรายนามต่อไปนี้"
            "เป็นคณะกรรมการตรวจสอบพัสดุประจำปี", align="justify", indent=1.25, after=2)
    _member_lines(doc, _members(ctx))
    _p(doc, "จึงเรียนมาเพื่อโปรดพิจารณา", indent=1.25, before=2, after=12)
    _sign_table(doc, [[
        ("ลงชื่อ ......................................", "center"),
        (f"( {head} )", "center"),
        ("หัวหน้าเจ้าหน้าที่พัสดุ", "center"),
    ]])
    _p(doc, "- ทราบ    - ดำเนินการ", indent=1.25, before=4, after=10)
    _sign_table(doc, [[
        ("ลงชื่อ ......................................", "center"),
        (f"( {director} )", "center"),
        (_director_line(school), "center"),
    ]])
    return _save(doc, f"ขอแต่งตั้งกรรมการตรวจสอบพัสดุ_ปีงบ{year}") if own else doc


def render_appoint_order(school, ctx, doc=None):
    """(2) คำสั่งโรงเรียน แต่งตั้งคณะกรรมการตรวจสอบพัสดุประจำปี"""
    own = doc is None
    if own:
        doc = Document(); _font(doc)
    else:
        doc.add_page_break()
    year = ctx.get("year")
    sname = (school.name or "โรงเรียน").strip()
    director = (school.director_name or "").strip() or _BLANK
    _krut_and_title(doc)
    _p(doc, f"คำสั่ง{sname}", align="center", bold=True, size=18, after=0)
    _p(doc, f"ที่ {ctx.get('order_no') or _BLANK}", align="center", after=0)
    _p(doc, f"เรื่อง แต่งตั้งคณะกรรมการตรวจสอบพัสดุประจำปี ประจำปีงบประมาณ พ.ศ. {year}",
       align="center", bold=True, after=0)
    _p(doc, "-----------------------------------", align="center", after=6)
    _p(doc, _LEGAL + " จึงแต่งตั้งบุคคลผู้มีรายนามข้างท้ายนี้เป็นคณะกรรมการตรวจสอบพัสดุประจำปี ดังนี้",
       align="justify", indent=1.25, after=2)
    _member_lines(doc, _members(ctx))
    _p(doc, "ให้คณะกรรมการที่ได้รับการแต่งตั้งตามคำสั่งนี้ ปฏิบัติหน้าที่ที่ได้รับมอบหมายให้บังเกิด"
            "ผลดีต่อทางราชการโดยเคร่งครัด", align="justify", indent=1.25, before=2, after=2)
    _p(doc, f"ทั้งนี้ ตั้งแต่วันที่ 1 ตุลาคม พ.ศ. {year} เป็นต้นไป", indent=1.25, after=1)
    _p(doc, f"สั่ง ณ วันที่ {thai_date(ctx.get('date'))}", indent=1.25, after=12)
    _sign_table(doc, [[
        ("ลงชื่อ ......................................", "center"),
        (f"( {director} )", "center"),
        (_director_line(school), "center"),
    ]])
    return _save(doc, f"คำสั่งแต่งตั้งกรรมการตรวจสอบพัสดุ_ปีงบ{year}") if own else doc


def render_result_memo(school, ctx, assets=None, doc=None):
    """(3) บันทึกข้อความ รายงานผลการตรวจสอบพัสดุประจำปี"""
    own = doc is None
    if own:
        doc = Document(); _font(doc)
    else:
        doc.add_page_break()
    year = ctx.get("year")
    director = (school.director_name or "").strip() or _BLANK
    damaged = (ctx.get("damaged_count") or "").strip() or "-"
    _memo_header(doc, school, f"รายงานผลการตรวจสอบพัสดุประจำปี ประจำปีงบประมาณ พ.ศ. {year}",
                 ctx.get("result_memo_no") or "", thai_date(ctx.get("date")))
    _p(doc, f"ตามคำสั่ง{(school.name or 'โรงเรียน').strip()} ที่ {ctx.get('order_no') or _BLANK} "
            f"เรื่อง แต่งตั้งคณะกรรมการตรวจสอบพัสดุประจำปี ประจำปีงบประมาณ พ.ศ. {year} "
            "ให้ดำเนินการตรวจสอบและตรวจนับวัสดุ/ครุภัณฑ์ที่คงเหลืออยู่ ณ วันสิ้นงวด นั้น",
       align="justify", indent=1.25, after=2)
    _p(doc, "บัดนี้ คณะกรรมการตามคำสั่งดังกล่าว ได้ดำเนินการตรวจสอบการรับ - จ่ายพัสดุเรียบร้อยแล้ว "
            "ผลการตรวจสอบสรุปได้ ดังนี้", align="justify", indent=1.25, after=2)
    _p(doc, "1. การตรวจสอบการรับ - จ่าย ได้ตรวจสอบเอกสารฝ่ายรับและเอกสารฝ่ายจ่าย บัญชีวัสดุ "
            "และทะเบียนครุภัณฑ์ ปรากฏว่าถูกต้องครบถ้วนตรงกัน", indent=1.25, after=1)
    _p(doc, f"2. การตรวจนับพัสดุคงเหลือ ณ วันที่ 30 กันยายน {year} ปรากฏว่ามีพัสดุคงเหลือ"
            "ตรงตามบัญชีและทะเบียน", indent=1.25, after=1)
    _p(doc, f"3. การตรวจสอบสภาพครุภัณฑ์ ปรากฏว่ามีครุภัณฑ์ชำรุด เสื่อมสภาพ จำนวน {damaged} รายการ",
       indent=1.25, after=2)
    _p(doc, "จึงเรียนมาเพื่อโปรดทราบและพิจารณาดำเนินการต่อไป", indent=1.25, after=12)
    members = _members(ctx)
    for m in (members or [{"name": _BLANK, "role": "ประธานกรรมการ"},
                          {"name": _BLANK, "role": "กรรมการ"},
                          {"name": _BLANK, "role": "กรรมการและเลขานุการ"}]):
        _sign_table(doc, [[
            (f"ลงชื่อ ...................................... {m.get('role','กรรมการ')}", "center"),
            (f"( {m.get('name', _BLANK)} )", "center"),
        ]])
    _p(doc, "- ทราบ    - รายงาน สตง.", indent=1.25, before=4, after=8)
    _sign_table(doc, [[
        ("ลงชื่อ ......................................", "center"),
        (f"( {director} )", "center"),
        (_director_line(school), "center"),
    ]])
    return _save(doc, f"รายงานผลตรวจสอบพัสดุ_ปีงบ{year}") if own else doc


def render_inventory(school, ctx, assets, doc=None):
    """(4) บัญชีครุภัณฑ์คงเหลือ (ดึงจากทะเบียนครุภัณฑ์ที่ยังใช้งาน)"""
    own = doc is None
    if own:
        doc = Document(); _font(doc)
    else:
        doc.add_page_break()
    year = ctx.get("year")
    _p(doc, f"บัญชีรายการครุภัณฑ์คงเหลือประจำปีงบประมาณ พ.ศ. {year}",
       align="center", bold=True, size=17, after=0)
    _p(doc, f"ณ วันที่ 30 กันยายน {year}", align="center", after=0)
    _p(doc, (school.name or "").strip(), align="center", after=6)
    headers = ["ที่", "เลขครุภัณฑ์", "รายการ", "จำนวน", "ราคาทุน (บาท)", "สภาพ"]
    widths = [Cm(1.2), Cm(3.6), Cm(6.5), Cm(1.8), Cm(3), Cm(2.4)]
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"; t.autofit = False
    _repeat_header_row(t.rows[0]); _no_split_row(t.rows[0])
    for c, h, w in zip(t.rows[0].cells, headers, widths):
        _set_cell(c, h, bold=True, align="center", size=14); c.width = w
    total = 0.0
    live = [a for a in (assets or []) if (a.status or "ใช้งาน") != "จำหน่ายแล้ว"]
    for i, a in enumerate(live, start=1):
        total += a.cost or 0
        vals = [str(i), a.asset_code or "-", a.name, f"{a.quantity or 1:g} {a.unit or ''}".strip(),
                f"{a.cost or 0:,.2f}", a.status or "ใช้งาน"]
        aligns = ["center", "left", "left", "center", "right", "center"]
        r = t.add_row(); _no_split_row(r)
        for c, v, w, al in zip(r.cells, vals, widths, aligns):
            _set_cell(c, v, align=al, size=13); c.width = w
    if not live:
        r = t.add_row()
        _set_cell(r.cells[0], "-", align="center", size=13)
    _p(doc, f"รวมครุภัณฑ์คงเหลือ {len(live)} รายการ เป็นเงิน {total:,.2f} บาท",
       bold=True, before=4, after=12)
    members = _members(ctx)
    for m in (members or [{"name": _BLANK, "role": "ประธานกรรมการ"}]):
        _sign_table(doc, [[
            (f"ลงชื่อ ...................................... {m.get('role','กรรมการ')}", "center"),
            (f"( {m.get('name', _BLANK)} )", "center"),
        ]])
    return _save(doc, f"บัญชีครุภัณฑ์คงเหลือ_ปีงบ{year}") if own else doc


def render_audit_bundle(school, ctx, assets) -> str:
    """ออกชุดเอกสารตรวจสอบพัสดุประจำปีทั้งชุดเป็นไฟล์เดียว"""
    doc = Document(); _font(doc)
    render_appoint_memo(school, ctx, doc)
    render_appoint_order(school, ctx, doc)
    render_result_memo(school, ctx, assets, doc)
    render_inventory(school, ctx, assets, doc)
    return _save(doc, f"ชุดตรวจสอบพัสดุประจำปี_ปีงบ{ctx.get('year')}")
