# -*- coding: utf-8 -*-
"""
office_doc.py
-------------
สร้างไฟล์ Word เอกสารงานธุรการจากข้อมูลจริง:
  - บันทึกข้อความ (OfficeMemo)
  - คำสั่งโรงเรียน (SchoolOrder)

ใช้ helper จัดรูปแบบร่วมกับ build_templates (ฟอนต์ TH Sarabun, ครุฑ, ตัวหนา complex script)
"""
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.doc_page import set_a4

from app.database import get_data_dir
from app.thai_utils import thai_date, thai_date_official
from app.services.build_templates import (
    _font, _p, _p_runs, _krut_and_title, _krut_center, _sign_table, _hr,
    _no_borders, _set_cell, THAI_FONT,
)


def _foot_line(par, text, size=15):
    """เพิ่มบรรทัดใน footer ด้วยฟอนต์ไทย (ชิดซ้าย)"""
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = par.add_run(text)
    run.font.name = THAI_FONT
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), THAI_FONT)


def _safe(text: str) -> str:
    for ch in '<>:"/\\|?*':
        text = text.replace(ch, "_")
    return text.strip()


def _inline_to_anchor(inline, dy_emu):
    """แปลง <wp:inline> (รูปแบบไหลตามข้อความ) เป็น <wp:anchor> แบบ 'อยู่หน้าข้อความ'
    (behindDoc=0, wrapNone) จัดกึ่งกลางคอลัมน์ + เลื่อนแนวตั้ง dy_emu (ลบ = ขึ้น)"""
    drawing = inline.getparent()
    anchor = OxmlElement("wp:anchor")
    for k, v in (("distT", "0"), ("distB", "0"), ("distL", "0"), ("distR", "0"),
                 ("simplePos", "0"), ("relativeHeight", "251658240"), ("behindDoc", "0"),
                 ("locked", "0"), ("layoutInCell", "1"), ("allowOverlap", "1")):
        anchor.set(k, v)
    sp = OxmlElement("wp:simplePos"); sp.set("x", "0"); sp.set("y", "0"); anchor.append(sp)
    ph = OxmlElement("wp:positionH"); ph.set("relativeFrom", "column")
    al = OxmlElement("wp:align"); al.text = "center"; ph.append(al); anchor.append(ph)
    pv = OxmlElement("wp:positionV"); pv.set("relativeFrom", "paragraph")
    off = OxmlElement("wp:posOffset"); off.text = str(int(dy_emu)); pv.append(off); anchor.append(pv)
    # ย้ายลูก extent/effectExtent มาก่อน แล้วแทรก wrapNone ตามด้วย docPr/graphic (ลำดับตาม schema)
    for tag in ("wp:extent", "wp:effectExtent"):
        el = inline.find(qn(tag))
        if el is not None:
            anchor.append(el)
    anchor.append(OxmlElement("wp:wrapNone"))
    for tag in ("wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
        el = inline.find(qn(tag))
        if el is not None:
            anchor.append(el)
    drawing.remove(inline)
    drawing.append(anchor)


def _float_signature(paragraph, signer_name, height_cm=1.45, dy_cm=-0.55):
    """ถ้าผู้ลงนามมีลายเซ็น -> วางรูปลายเซ็นแบบ 'อยู่หน้าข้อความ' ทับบนพารากราฟที่ระบุ
    (PNG โปร่งใส ลอยทับบรรทัด '(ลงชื่อ)' ไม่ดันข้อความ) คืน True ถ้าใส่แล้ว"""
    from app.services.signature import signature_path_for_current
    path = signature_path_for_current(signer_name)
    if not path:
        return False
    try:
        shape = paragraph.add_run().add_picture(path, height=Cm(height_cm))
        _inline_to_anchor(shape._inline, dy_cm * 360000)
        return True
    except Exception:
        return False


def _save_doc(doc, fname: str) -> str:
    """บันทึก .docx ลง data/documents; ถ้าเขียนไม่ได้ (path/permission บนเซิร์ฟเวอร์)
    ใช้โฟลเดอร์ชั่วคราวของระบบแทน แล้วคืน path (เสิร์ฟจากหน่วยความจำภายหลัง)"""
    import tempfile
    try:
        out_dir = get_data_dir() / "documents"
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / fname
        doc.save(str(out_path))
        return str(out_path)
    except OSError:
        out_path = Path(tempfile.gettempdir()) / fname
        doc.save(str(out_path))
        return str(out_path)


def _school_office(school) -> str:
    parts = [school.name or "", school.address or ""]
    return "  ".join(p for p in parts if p).strip()


def _director_office(school) -> str:
    name = (school.name or "").strip()
    if name.startswith("โรงเรียน"):
        return "ผู้อำนวยการ" + name
    return school.director_position or "ผู้อำนวยการโรงเรียน"


def _body_paragraphs(doc, body: str):
    """แตกเนื้อหาเป็นย่อหน้า (เว้นบรรทัด = ย่อหน้าใหม่) จัดชิดขอบแบบไทย เยื้องบรรทัดแรก"""
    for line in (body or "").split("\n"):
        line = line.strip()
        if line:
            _p(doc, line, align="justify", indent=1.25, after=2)


def render_memo(memo, school) -> str:
    """สร้างไฟล์ .docx บันทึกข้อความ คืนค่าที่อยู่ไฟล์"""
    doc = Document(); set_a4(doc)
    _font(doc)
    _krut_and_title(doc)   # ครุฑ + "บันทึกข้อความ"

    _p_runs(doc, [("ส่วนราชการ  ", True), (memo.from_dept or _school_office(school), False)])
    _p_runs(doc, [("ที่  ", True), (memo.memo_no or "", False),
                  ("\t", False), ("วันที่ ", True), (thai_date(memo.date), False)], tab_cm=7)
    _p_runs(doc, [("เรื่อง  ", True), (memo.subject or "", False)])
    _p_runs(doc, [("เรียน  ", True), (memo.to_person or _director_office(school), False)])
    _hr(doc)
    _body_paragraphs(doc, memo.body)

    # ลงนาม
    _p(doc, "", after=12)
    signer = (memo.signer_name or school.director_name or "")
    position = (memo.signer_position or _director_office(school))
    tbl = _sign_table(doc, [[
        ("ลงชื่อ.........................................", "center"),
        (f"( {signer} )", "center"),
        (position, "center"),
    ]])
    _float_signature(tbl.rows[0].cells[0].paragraphs[0], signer)

    fname = _safe(f"บันทึกข้อความ_{memo.memo_no or memo.id}_{memo.subject}") + ".docx"
    return _save_doc(doc, fname)


def render_order(order, school) -> str:
    """สร้างไฟล์ .docx คำสั่งโรงเรียน คืนค่าที่อยู่ไฟล์"""
    doc = Document(); set_a4(doc)
    _font(doc)
    _krut_center(doc, height_cm=1.8)
    _p(doc, "คำสั่ง" + (school.name or ""), align="center", bold=True, size=18, after=0)
    _p(doc, "ที่ " + (order.order_no or ""), align="center", bold=True, after=0)
    _p(doc, "เรื่อง " + (order.subject or ""), align="center", bold=True, after=0)
    _p(doc, "─────────────────────", align="center", after=6)

    _body_paragraphs(doc, order.body)

    _p(doc, "", after=6)
    _p(doc, "สั่ง ณ วันที่ " + thai_date_official(order.date), align="center", after=12)
    sign_p = _p(doc, "(ลงชื่อ).........................................", align="center")
    _float_signature(sign_p, school.director_name)
    _p(doc, f"( {school.director_name or ''} )", align="center")
    _p(doc, _director_office(school), align="center")

    fname = _safe(f"คำสั่ง_{order.order_no or order.id}_{order.subject}") + ".docx"
    return _save_doc(doc, fname)


_BLANK = ".........................................."


def render_official_letter(letter, school) -> str:
    """สร้างไฟล์ .docx หนังสือราชการภายนอก (ครุฑกลาง + ขอแสดงความนับถือ) คืนค่าที่อยู่ไฟล์"""
    doc = Document(); set_a4(doc)
    _font(doc)
    _krut_center(doc, height_cm=2.0)

    # บรรทัด "ที่ ..." (ซ้าย) + ชื่อ/ที่อยู่ส่วนราชการ (ขวา) แบบตารางไร้เส้น
    t = doc.add_table(rows=1, cols=2)
    _no_borders(t)
    t.autofit = False
    t.columns[0].width = Cm(7.5)
    t.columns[1].width = Cm(8.0)
    _set_cell(t.rows[0].cells[0], "ที่  " + (letter.doc_no or _BLANK), align="left", size=16)
    right = (school.name or "") + (("\n" + school.address) if school.address else "")
    _set_cell(t.rows[0].cells[1], right, align="right", size=16)

    _p(doc, "วันที่ " + thai_date_official(letter.date), align="center", after=4)
    _p_runs(doc, [("เรื่อง  ", True), (letter.subject or _BLANK, False)])
    _p_runs(doc, [("เรียน  ", True), (letter.to or _BLANK, False)])
    if (letter.ref or "").strip():
        _p_runs(doc, [("อ้างถึง  ", True), (letter.ref, False)])
    if (letter.enclosure or "").strip():
        _p_runs(doc, [("สิ่งที่ส่งมาด้วย  ", True), (letter.enclosure, False)])

    _p(doc, "", after=4)
    _body_paragraphs(doc, letter.body)

    # คำลงท้าย + ลงนาม (จัดกึ่งกลางครึ่งขวา)
    _p(doc, "", after=8)
    _p(doc, letter.closing or "ขอแสดงความนับถือ", align="center", after=10)
    signer = (letter.signer_name or school.director_name or "")
    position = (letter.signer_position or _director_office(school))
    tbl = _sign_table(doc, [[
        ("(ลงชื่อ).........................................", "center"),
        (f"( {signer} )", "center"),
        (position, "center"),
    ]])
    _float_signature(tbl.rows[0].cells[0].paragraphs[0], signer)

    # ส่วนราชการเจ้าของเรื่อง + เบอร์โทร -> ตรึงไว้ท้ายหน้า (footer)
    foot = doc.sections[0].footer
    foot.is_linked_to_previous = False
    _foot_line(foot.paragraphs[0] if foot.paragraphs else foot.add_paragraph(),
               school.name or "")
    _foot_line(foot.add_paragraph(), "โทร. ........................................")

    fname = _safe(f"หนังสือราชการ_{letter.doc_no or letter.id}_{letter.subject}") + ".docx"
    return _save_doc(doc, fname)
