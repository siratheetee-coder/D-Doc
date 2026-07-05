# -*- coding: utf-8 -*-
"""
finance_doc.py
--------------
สร้างไฟล์ Word "บันทึกข้อความ ขออนุมัติเบิกจ่ายเงินนอกงบประมาณ" (DisburseMemo)
ตามแบบมาตรฐานราชการ: เนื้อความ + รายละเอียดการคำนวณเงิน + บล็อกอนุมัติของผู้อำนวยการ
ใช้ helper จัดรูปแบบร่วมกับ build_templates (ฟอนต์ TH Sarabun, ครุฑ, ตัวหนา complex script)
"""
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.database import get_data_dir
from app.thai_utils import thai_date, bahttext
from app.services.build_templates import (
    _font, _p, _p_runs, _krut_and_title, _sign_table, _hr,
)


def _safe(text: str) -> str:
    for ch in '<>:"/\\|?*':
        text = text.replace(ch, "_")
    return text.strip()


def _school_office(school) -> str:
    parts = [school.name or "", school.address or ""]
    return "  ".join(p for p in parts if p).strip()


def _director_office(school) -> str:
    name = (school.name or "").strip()
    if name.startswith("โรงเรียน"):
        return "ผู้อำนวยการ" + name
    return school.director_position or "ผู้อำนวยการโรงเรียน"


def _fmt(v) -> str:
    return "{:,.2f}".format(float(v or 0))


def _no_borders(table):
    tbl = table._tbl
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tbl.tblPr.append(borders)


def _set_cell(cell, text, *, bold=False, align="left", size=16):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "TH Sarabun New"
    r._element.rPr.rFonts.set(qn("w:cs"), "TH Sarabun New")


def _money_breakdown(doc, lines):
    """ตารางไร้เส้น 3 คอลัมน์: ป้ายกำกับ | จำนวนเงิน(ชิดขวา) | บาท/คำอ่าน"""
    table = doc.add_table(rows=len(lines), cols=3)
    _no_borders(table)
    table.autofit = False
    widths = [Cm(6.5), Cm(3.5), Cm(6.5)]
    for i, (label, amount, tail) in enumerate(lines):
        cells = table.rows[i].cells
        bold = (i == len(lines) - 1)
        _set_cell(cells[0], label, bold=bold, align="left")
        _set_cell(cells[1], amount, bold=bold, align="right")
        _set_cell(cells[2], tail, bold=bold, align="left")
        for c, w in zip(cells, widths):
            c.width = w
    return table


def render_disburse(memo, school) -> str:
    """สร้างไฟล์ .docx บันทึกข้อความขออนุมัติเบิกจ่ายเงินนอกงบประมาณ คืนค่าที่อยู่ไฟล์"""
    doc = Document()
    _font(doc)
    _krut_and_title(doc)   # ครุฑ + "บันทึกข้อความ"

    amount = float(memo.amount or 0)
    vat = float(memo.vat or 0)
    wht = float(memo.wht or 0)
    fine = float(memo.fine or 0)
    goods = round(amount - vat, 2)
    net = round(amount - wht - fine, 2)
    src = (memo.budget_source or (memo.account.name if memo.account else "") or "เงินอุดหนุน")
    kind = memo.proc_kind or "จัดซื้อ"
    payee = memo.payee or "..............................."

    _p_runs(doc, [("ส่วนราชการ  ", True), (_school_office(school), False)])
    _p_runs(doc, [("ที่  ", True), (memo.memo_no or "", False),
                  ("\t", False), ("วันที่ ", True), (thai_date(memo.date), False)], tab_cm=8)
    _p_runs(doc, [("เรื่อง  ", True),
                  (f"ขออนุมัติเบิกจ่ายเงินนอกงบประมาณ ประเภท{src}", False)])
    _p_runs(doc, [("เรียน  ", True), (_director_office(school), False)])
    _hr(doc)

    # ย่อหน้านำ
    _p(doc,
       f"ตามที่{school.name or 'โรงเรียน'} ได้ดำเนินการ{kind}{memo.subject or ''} "
       f"เพื่อใช้ในโครงการ/กิจกรรม{(' ' + memo.note) if (memo.note or '').strip() else '..............................'} "
       f"(รายละเอียดตามที่แนบ) ตามใบส่งสินค้า/ใบกำกับภาษี เลขที่.................. ลงวันที่.................. "
       f"จาก {payee} จำนวนเงิน {_fmt(amount)} บาท ({bahttext(amount)}) "
       f"และในการนี้โรงเรียนขอใช้เงินนอกงบประมาณ ประเภท{src} นั้น",
       align="justify", indent=1.25, after=2)
    _p(doc,
       "บัดนี้ ผู้ขายได้ส่งมอบงานถูกต้องครบถ้วนแล้ว ตามระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้าง"
       "และการบริหารพัสดุภาครัฐ พ.ศ. 2560", align="justify", indent=1.25, after=2)
    _p(doc, "เห็นควรเบิกจ่ายให้แก่ผู้ขาย ตามใบส่งของ โดยมีรายละเอียดดังนี้",
       align="justify", indent=1.25, after=2)

    _money_breakdown(doc, [
        ("จำนวนเงินขอเบิก", _fmt(amount), "บาท"),
        ("ภาษีมูลค่าเพิ่ม", _fmt(vat), "บาท"),
        ("มูลค่าสินค้า", _fmt(goods), "บาท"),
        ("หัก ภาษี ณ ที่จ่าย", _fmt(wht), "บาท"),
        ("ค่าปรับ", _fmt(fine), "บาท"),
        ("คงเหลือจ่ายจริง", _fmt(net), f"บาท  ({bahttext(net)})"),
    ])
    _p(doc, "", after=2)
    _p(doc, "จึงเรียนมาเพื่อโปรดอนุมัติ", align="justify", indent=1.25, after=10)

    # ลงนามเจ้าหน้าที่การเงิน
    fin = (school.finance_officer_name or school.officer_name or "").strip()
    _sign_table(doc, [[
        ("ลงชื่อ.........................................เจ้าหน้าที่การเงิน", "center"),
        (f"( {fin} )", "center"),
    ]])

    # บล็อกอนุมัติของผู้อำนวยการ
    _p(doc, "", after=6)
    _p(doc, "1.  ทราบ", indent=1.25, after=1)
    _p(doc, f"2.  อนุมัติให้จ่ายเงิน {_fmt(net)} บาท ({bahttext(net)})", indent=1.25, after=10)
    _sign_table(doc, [[
        ("ลงชื่อ.........................................", "center"),
        (f"( {(school.director_name or '').strip()} )", "center"),
        (_director_office(school), "center"),
    ]])
    _p(doc, "วันที่.........................................", align="center", after=2)

    out_dir = get_data_dir() / "documents"
    out_dir.mkdir(exist_ok=True)
    fname = _safe(f"ขออนุมัติเบิกจ่าย_{memo.memo_no or memo.id}_{memo.subject}") + ".docx"
    out_path = out_dir / fname
    doc.save(str(out_path))
    return str(out_path)
