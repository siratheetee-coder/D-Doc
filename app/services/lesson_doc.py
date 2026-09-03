# -*- coding: utf-8 -*-
"""
lesson_doc.py - แบบบันทึกการตรวจแผนการจัดการเรียนรู้ (ใบรับรอง)
ออกคู่กับไฟล์แผนที่ครูอัปโหลด · แปะลายเซ็นหัวหน้าฝ่ายวิชาการ + ผอ. อัตโนมัติ
ฟอนต์ TH SarabunPSK · กระดาษ A4
"""
import io
from pathlib import Path

from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.database import get_data_dir
from app.services.super_doc import _doc, _logo, _p, _font, _safe, FONT
from app.services.signature import signature_path_for
from app.thai_utils import _THAI_MONTHS


def _thai_date(dt) -> str:
    """คืนวันที่แบบไทย เช่น '9 มิถุนายน 2569' (ว่างถ้าไม่มี)"""
    if not dt:
        return ""
    return f"{dt.day} {_THAI_MONTHS[dt.month]} {dt.year + 543}"


def _sign_block(doc, db, name, position, date_str):
    """ช่องลงนาม: ลายเซ็น (แปะจากทะเบียนถ้ามี) + (ชื่อ) + ตำแหน่ง + วันที่"""
    name = (name or "").strip()
    # ลายเซ็น (ภาพ) กึ่งกลาง - ถ้าไม่มีเว้นบรรทัดว่างไว้เซ็นมือ
    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_before = Pt(6); sp.paragraph_format.space_after = Pt(0)
    path = signature_path_for(db, name) if name else None
    if path:
        try:
            sp.add_run().add_picture(path, height=Cm(1.2))
        except Exception:
            _font(sp.add_run("ลงชื่อ ......................................"), 16)
    else:
        _font(sp.add_run("ลงชื่อ ......................................"), 16)
    _p(doc, f"( {name} )" if name else "( ...................................... )",
       align="center", size=16, after=0)
    _p(doc, position or "", align="center", size=16, after=0)
    _p(doc, f"วันที่ {date_str}" if date_str else "วันที่ ........./........./.........",
       align="center", size=16, after=6)


def render_lesson_plan_cert(db, plan, school) -> str:
    """สร้างแบบบันทึกการตรวจแผน (docx) คืน path ไฟล์"""
    from app.models import Person
    doc = _doc(top=1.6, bottom=1.6)
    _logo(doc, school, height_cm=1.6, after=2)
    _p(doc, school.name or "", align="center", bold=True, size=18, after=0)
    _p(doc, "แบบบันทึกการตรวจแผนการจัดการเรียนรู้", align="center", bold=True, size=18, after=10)

    teacher = plan.teacher or (db.get(Person, plan.person_id) if plan.person_id else None)
    term = plan.term if plan.term in (1, 2) else "....."
    _p(doc, f"ชื่อครูผู้สอน  {(teacher.name if teacher else '') or '.....'}", size=16, after=4)
    _p(doc, f"เรื่อง/หน่วยการจัดการเรียนรู้  {plan.title or '.....'}", size=16, after=4)
    _p(doc, f"ภาคเรียนที่ {term}  ปีการศึกษา {plan.year or '.....'}", size=16, after=4)
    _p(doc, f"วันที่ส่งแผน  {_thai_date(plan.submitted_at) or '.....'}", size=16, after=6)

    # ---- ลายเซ็นครูผู้จัดทำ/เสนอแผน (แปะให้ถ้ามีในทะเบียน) ----
    _sign_block(doc, db, (teacher.name if teacher else ""),
                "ครูผู้สอน (ผู้เสนอแผน)", _thai_date(plan.submitted_at))

    # ---- ผลการตรวจของหัวหน้าฝ่ายวิชาการ ----
    _p(doc, "ความเห็นหัวหน้ากลุ่มบริหารงานวิชาการ", bold=True, size=16, after=2)
    _p(doc, (plan.comment or "ตรวจแผนการจัดการเรียนรู้แล้ว เห็นควรเสนอผู้อำนวยการเพื่อพิจารณา"),
       size=16, after=2, align="justify")
    academic = db.get(Person, plan.academic_by) if plan.academic_by else None
    academic_name = academic.name if academic else (school.academic_head_name or "")
    _sign_block(doc, db, academic_name,
                "หัวหน้ากลุ่มบริหารงานวิชาการ", _thai_date(plan.reviewed_at))

    # ---- ความเห็น/การอนุมัติของ ผอ. ----
    _p(doc, "ความเห็นผู้อำนวยการโรงเรียน", bold=True, size=16, after=2)
    _p(doc, (plan.director_comment or "อนุมัติ"), size=16, after=2, align="justify")
    director = db.get(Person, plan.director_by) if plan.director_by else None
    director_name = director.name if director else (school.director_name or "")
    _sign_block(doc, db, director_name,
                school.director_position or "ผู้อำนวยการโรงเรียน", _thai_date(plan.director_at))

    out = get_data_dir() / "generated"
    out.mkdir(exist_ok=True)
    fname = _safe(f"ตรวจแผน_{(teacher.name if teacher else '')}_{plan.title}")[:80] + ".docx"
    path = out / fname
    doc.save(str(path))
    return str(path)
