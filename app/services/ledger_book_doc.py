# -*- coding: utf-8 -*-
"""
ledger_book_doc.py
------------------
เอกสารบัญชีการเงินแบบราชการ (แบบ สตง.):
  1) สมุดเงินสด (Cash Book)         -> render_cash_book / build_cash_book_xlsx
  2) บัญชีแยกประเภท เต็มรูปแบบ        -> render_general_ledger / build_ledger_xlsx
คอลัมน์เงินใช้ เดบิต(รับ) / เครดิต(จ่าย) / คงเหลือ ตามหลักบัญชีเงินสดของหน่วยงานภาครัฐ
สร้าง docx/xlsx ตอนรันไทม์ (ไม่พึ่ง template) แล้วคืน path
"""
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from app.services.doc_page import set_a4

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

from app.database import get_data_dir
from app.thai_utils import thai_date, _THAI_MONTHS

THAI_FONT = "TH Sarabun New"
_thin = Side(style="thin")
_BORDER = Border(left=_thin, right=_thin, top=_thin, bottom=_thin)
_HEAD = PatternFill("solid", fgColor="DCFCE7")
_TOTAL = PatternFill("solid", fgColor="F1F5F9")
_SUBTOT = PatternFill("solid", fgColor="FEF9C3")


def _safe(text: str) -> str:
    for ch in '<>:"/\\|?*':
        text = text.replace(ch, "_")
    return text.strip()


def _fmt(v):
    return "{:,.2f}".format(v) if v else ("0.00" if v == 0 else "-")


def _fmt0(v):
    """ช่องว่างถ้าเป็น 0/None (สำหรับช่องเดบิต/เครดิตที่ไม่มีค่า)"""
    return "{:,.2f}".format(v) if v else ""


def be_year(dt) -> int:
    return (dt.year + 543) if dt else 0


# ---------------- Word helpers ----------------
def _p(doc, text="", *, align="center", bold=False, size=14, after=2):
    p = doc.add_paragraph()
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = THAI_FONT
    r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    r._element.rPr.rFonts.set(qn("w:ascii"), THAI_FONT)
    r._element.rPr.rFonts.set(qn("w:hAnsi"), THAI_FONT)
    return p


def _set_cell(cell, text, *, bold=False, align="left", size=13, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = THAI_FONT
    r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    r._element.rPr.rFonts.set(qn("w:ascii"), THAI_FONT)
    r._element.rPr.rFonts.set(qn("w:hAnsi"), THAI_FONT)
    if fill is not None:
        shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {
            qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): fill})
        cell._tc.get_or_add_tcPr().append(shd)


def _landscape(doc):
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.left_margin = sec.right_margin = Cm(1.5)
    sec.top_margin = sec.bottom_margin = Cm(1.5)
    base = doc.styles["Normal"]
    base.font.name = THAI_FONT
    base.font.size = Pt(13)
    base._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)


def _sign_block(doc, school):
    _p(doc, "", after=10)
    officer = (school.finance_officer_name or school.officer_name or "").strip()
    _p(doc, "ลงชื่อ.............................................ผู้จัดทำ", align="center", after=0)
    _p(doc, f"( {officer} )", align="center", after=0)
    _p(doc, "เจ้าหน้าที่การเงินและบัญชี", align="center", after=10)
    _p(doc, "ลงชื่อ.............................................ผู้ตรวจสอบ", align="center", after=0)
    _p(doc, f"( {(school.director_name or '').strip()} )", align="center", after=0)
    director_pos = ("ผู้อำนวยการ" + school.name) if (school.name or "").startswith("โรงเรียน") \
        else (school.director_position or "ผู้อำนวยการโรงเรียน")
    _p(doc, director_pos, align="center", after=2)


# ---------------- สมุดเงินสด (Word) ----------------
def render_cash_book(school, fiscal_year, rows, opening, totals, scope_name="ทุกบัญชี") -> str:
    """rows: list ของ dict {date, desc, ref, debit(รับ), credit(จ่าย), balance, subtotal(bool), month}"""
    doc = Document(); set_a4(doc)
    _landscape(doc)
    _p(doc, "สมุดเงินสด", align="center", bold=True, size=18, after=0)
    _p(doc, (school.name or ""), align="center", bold=True, size=15, after=0)
    _p(doc, f"ประจำปีงบประมาณ {fiscal_year}   ({scope_name})", align="center", size=14, after=6)

    headers = ["ว/ด/ป", "รายการ", "เลขที่เอกสาร", "รับ (เดบิต)", "จ่าย (เครดิต)", "คงเหลือ"]
    widths = [Cm(2.6), Cm(9.5), Cm(3.4), Cm(3.4), Cm(3.4), Cm(3.4)]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for c, (h, w) in enumerate(zip(headers, widths)):
        _set_cell(table.rows[0].cells[c], h, bold=True, align="center", fill="DCFCE7")
        table.rows[0].cells[c].width = w

    # แถวยอดยกมา
    op = table.add_row().cells
    _set_cell(op[0], "", align="center")
    _set_cell(op[1], "ยอดยกมา", bold=True)
    _set_cell(op[2], "")
    _set_cell(op[3], ""); _set_cell(op[4], "")
    _set_cell(op[5], _fmt(opening), bold=True, align="right")
    for c, w in enumerate(widths):
        op[c].width = w

    for row in rows:
        cells = table.add_row().cells
        sub = row.get("subtotal")
        fill = "FEF9C3" if sub else None
        _set_cell(cells[0], row.get("date", ""), align="center", fill=fill)
        _set_cell(cells[1], row.get("desc", ""), bold=bool(sub), fill=fill)
        _set_cell(cells[2], row.get("ref", ""), align="center", fill=fill)
        _set_cell(cells[3], _fmt0(row.get("debit")), align="right", bold=bool(sub), fill=fill)
        _set_cell(cells[4], _fmt0(row.get("credit")), align="right", bold=bool(sub), fill=fill)
        _set_cell(cells[5], _fmt(row.get("balance")) if row.get("balance") is not None else "",
                  align="right", bold=bool(sub), fill=fill)
        for c, w in enumerate(widths):
            cells[c].width = w

    # แถวรวมทั้งสิ้น
    tr = table.add_row().cells
    _set_cell(tr[0], "", fill="F1F5F9")
    _set_cell(tr[1], "รวมทั้งสิ้น", bold=True, align="right", fill="F1F5F9")
    _set_cell(tr[2], "", fill="F1F5F9")
    _set_cell(tr[3], _fmt(totals.get("debit")), bold=True, align="right", fill="F1F5F9")
    _set_cell(tr[4], _fmt(totals.get("credit")), bold=True, align="right", fill="F1F5F9")
    _set_cell(tr[5], _fmt(totals.get("balance")), bold=True, align="right", fill="F1F5F9")
    for c, w in enumerate(widths):
        tr[c].width = w

    _sign_block(doc, school)
    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    out = out_dir / (_safe(f"สมุดเงินสด_ปีงบ{fiscal_year}_{scope_name}") + ".docx")
    doc.save(str(out))
    return str(out)


# ---------------- สมุดเงินสดแบบราชการ: แยกรับ/จ่าย + 3 งบ (Word) ----------------
_FUND_COLS = ["เงินงบประมาณ", "เงินรายได้แผ่นดิน", "เงินนอกงบประมาณ"]


def _cb_section(doc, title, total_label, rows, *, open_by_fund=None):
    """สร้าง 1 ตารางของสมุดเงินสด (ด้านรับ หรือ ด้านจ่าย) คืน dict ยอดรวมแต่ละคอลัมน์"""
    _p(doc, title, align="left", bold=True, size=13.5, after=2)
    headers = ["วันที่", "เลขที่เอกสาร", "รายการ", "เงินสด"] + _FUND_COLS
    widths = [Cm(2.3), Cm(3.0), Cm(8.2), Cm(3.2), Cm(3.2), Cm(3.4), Cm(3.4)]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for c, (h, w) in enumerate(zip(headers, widths)):
        _set_cell(table.rows[0].cells[c], h, bold=True, align="center", fill="DCFCE7")
        table.rows[0].cells[c].width = w
    tot = {"เงินสด": 0.0}
    for f in _FUND_COLS:
        tot[f] = 0.0

    def _row(date, ref, desc, cash, fund_amt: dict, *, bold=False, fill=None):
        cells = table.add_row().cells
        _set_cell(cells[0], date, align="center", fill=fill)
        _set_cell(cells[1], ref, align="center", fill=fill)
        _set_cell(cells[2], desc, bold=bold, fill=fill)
        _set_cell(cells[3], _fmt0(cash), align="right", bold=bold, fill=fill)
        for i, f in enumerate(_FUND_COLS):
            _set_cell(cells[4 + i], _fmt0(fund_amt.get(f, 0)), align="right", bold=bold, fill=fill)
        for c, w in enumerate(widths):
            cells[c].width = w

    if open_by_fund is not None:
        op_cash = sum(open_by_fund.get(f, 0) for f in _FUND_COLS)
        _row("", "", "ยอดยกมา", op_cash, open_by_fund, bold=True)
        tot["เงินสด"] += op_cash
        for f in _FUND_COLS:
            tot[f] += open_by_fund.get(f, 0)
    for r in rows:
        amt = r["amount"] or 0
        _row(r["date"], r["ref"], r["desc"], amt, {r["fund"]: amt})
        tot["เงินสด"] += amt
        tot[r["fund"]] = tot.get(r["fund"], 0) + amt
    # แถวรวม
    cells = table.add_row().cells
    _set_cell(cells[0], "", fill="F1F5F9"); _set_cell(cells[1], "", fill="F1F5F9")
    _set_cell(cells[2], total_label, bold=True, align="right", fill="F1F5F9")
    _set_cell(cells[3], _fmt(tot["เงินสด"]), bold=True, align="right", fill="F1F5F9")
    for i, f in enumerate(_FUND_COLS):
        _set_cell(cells[4 + i], _fmt(tot[f]), bold=True, align="right", fill="F1F5F9")
    for c, w in enumerate(widths):
        cells[c].width = w
    return tot


def render_cash_book_fund(school, fiscal_year, scope_name, open_by_fund, receipts, payments) -> str:
    """สมุดเงินสดแบบราชการ แนวนอน: ด้านรับ (เดบิตเงินสด) + ด้านจ่าย (เครดิตเงินสด)
    เครดิต/เดบิตแยกตามประเภทเงิน (งบประมาณ/รายได้แผ่นดิน/นอกงบประมาณ)"""
    doc = Document(); set_a4(doc)
    _landscape(doc)
    _p(doc, "สมุดเงินสด", align="center", bold=True, size=18, after=0)
    _p(doc, (school.name or ""), align="center", bold=True, size=15, after=0)
    _p(doc, f"ประจำปีงบประมาณ {fiscal_year}   ({scope_name})", align="center", size=14, after=8)

    rin = _cb_section(doc, "ด้านรับ  (เดบิต = เงินสด · เครดิตแยกตามประเภทเงิน)",
                      "รวมด้านรับ", receipts, open_by_fund=open_by_fund)
    _p(doc, "", after=6)
    rout = _cb_section(doc, "ด้านจ่าย  (เครดิต = เงินสด · เดบิตแยกตามประเภทเงิน)",
                       "รวมด้านจ่าย", payments)

    _p(doc, "", after=4)
    carry_cash = rin["เงินสด"] - rout["เงินสด"]
    per_fund = "   ".join(f"{f} {_fmt(rin[f] - rout[f])}" for f in _FUND_COLS)
    _p(doc, f"เงินสดคงเหลือยกไป  {_fmt(carry_cash)}  บาท", align="right", bold=True, size=14, after=0)
    _p(doc, f"คงเหลือแยกประเภทเงิน:  {per_fund}", align="right", size=12.5, after=2)

    _sign_block(doc, school)
    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    out = out_dir / (_safe(f"สมุดเงินสด_ปีงบ{fiscal_year}_{scope_name}") + ".docx")
    doc.save(str(out))
    return str(out)


# ---------------- บัญชีแยกประเภท เต็มรูปแบบ (Word) ----------------
def render_general_ledger(school, account, fiscal_year, rows, opening) -> str:
    """rows: list ของ dict {date, desc, ref, debit, credit, balance, side, subtotal, month}"""
    doc = Document(); set_a4(doc)
    _landscape(doc)
    _p(doc, "บัญชีแยกประเภททั่วไป", align="center", bold=True, size=18, after=0)
    _p(doc, (school.name or ""), align="center", bold=True, size=15, after=0)
    _p(doc, f"ชื่อบัญชี  {account.name}          ประจำปีงบประมาณ {fiscal_year}",
       align="center", size=14, after=6)

    headers = ["ว/ด/ป", "รายการ", "เลขที่เอกสาร", "เดบิต", "เครดิต", "ดุล", "คงเหลือ"]
    widths = [Cm(2.6), Cm(8.8), Cm(3.2), Cm(3.2), Cm(3.2), Cm(1.6), Cm(3.4)]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for c, (h, w) in enumerate(zip(headers, widths)):
        _set_cell(table.rows[0].cells[c], h, bold=True, align="center", fill="DCFCE7")
        table.rows[0].cells[c].width = w

    op = table.add_row().cells
    _set_cell(op[0], "", align="center")
    _set_cell(op[1], "ยอดยกมา", bold=True)
    _set_cell(op[2], ""); _set_cell(op[3], ""); _set_cell(op[4], "")
    _set_cell(op[5], "เดบิต" if opening >= 0 else "เครดิต", align="center")
    _set_cell(op[6], _fmt(abs(opening)), bold=True, align="right")
    for c, w in enumerate(widths):
        op[c].width = w

    for row in rows:
        cells = table.add_row().cells
        sub = row.get("subtotal")
        fill = "FEF9C3" if sub else None
        _set_cell(cells[0], row.get("date", ""), align="center", fill=fill)
        _set_cell(cells[1], row.get("desc", ""), bold=bool(sub), fill=fill)
        _set_cell(cells[2], row.get("ref", ""), align="center", fill=fill)
        _set_cell(cells[3], _fmt0(row.get("debit")), align="right", bold=bool(sub), fill=fill)
        _set_cell(cells[4], _fmt0(row.get("credit")), align="right", bold=bool(sub), fill=fill)
        _set_cell(cells[5], row.get("side", ""), align="center", fill=fill)
        _set_cell(cells[6], _fmt(row.get("balance")) if row.get("balance") is not None else "",
                  align="right", bold=bool(sub), fill=fill)
        for c, w in enumerate(widths):
            cells[c].width = w

    _sign_block(doc, school)
    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    out = out_dir / (_safe(f"บัญชีแยกประเภท_{account.name}_ปีงบ{fiscal_year}") + ".docx")
    doc.save(str(out))
    return str(out)


# ---------------- Excel ----------------
def _xcell(ws, r, c, val, *, bold=False, align="left", money=False, fill=None, wrap=False):
    cell = ws.cell(row=r, column=c, value=val)
    cell.font = Font(name=THAI_FONT, bold=bold, size=13)
    cell.alignment = Alignment(horizontal=align, vertical="center", wrap_text=wrap)
    cell.border = _BORDER
    if money:
        cell.number_format = "#,##0.00"
    if fill:
        cell.fill = fill
    return cell


def _cashbook_sheet(ws, title, scope_name, fiscal_year, rows, opening, totals):
    ws.merge_cells("A1:F1")
    t = ws.cell(1, 1, f"{title}  ประจำปีงบประมาณ {fiscal_year}  ({scope_name})")
    t.font = Font(name=THAI_FONT, bold=True, size=16)
    t.alignment = Alignment(horizontal="center", vertical="center")
    heads = ["ว/ด/ป", "รายการ", "เลขที่เอกสาร", "รับ (เดบิต)", "จ่าย (เครดิต)", "คงเหลือ"]
    for c, h in enumerate(heads, 1):
        _xcell(ws, 3, c, h, bold=True, align="center", fill=_HEAD)
    _xcell(ws, 4, 1, "")
    _xcell(ws, 4, 2, "ยอดยกมา", bold=True)
    _xcell(ws, 4, 3, ""); _xcell(ws, 4, 4, "", money=True); _xcell(ws, 4, 5, "", money=True)
    _xcell(ws, 4, 6, opening, bold=True, money=True, align="right")
    r = 5
    for row in rows:
        sub = row.get("subtotal")
        fill = _SUBTOT if sub else None
        _xcell(ws, r, 1, row.get("date", ""), align="center", fill=fill)
        _xcell(ws, r, 2, row.get("desc", ""), bold=bool(sub), fill=fill, wrap=True)
        _xcell(ws, r, 3, row.get("ref", ""), align="center", fill=fill)
        _xcell(ws, r, 4, row.get("debit") or None, money=True, align="right", bold=bool(sub), fill=fill)
        _xcell(ws, r, 5, row.get("credit") or None, money=True, align="right", bold=bool(sub), fill=fill)
        _xcell(ws, r, 6, row.get("balance"), money=True, align="right", bold=bool(sub), fill=fill)
        r += 1
    _xcell(ws, r, 1, "", fill=_TOTAL); _xcell(ws, r, 2, "รวมทั้งสิ้น", bold=True, align="right", fill=_TOTAL)
    _xcell(ws, r, 3, "", fill=_TOTAL)
    _xcell(ws, r, 4, totals.get("debit"), bold=True, money=True, align="right", fill=_TOTAL)
    _xcell(ws, r, 5, totals.get("credit"), bold=True, money=True, align="right", fill=_TOTAL)
    _xcell(ws, r, 6, totals.get("balance"), bold=True, money=True, align="right", fill=_TOTAL)
    for i, w in enumerate([13, 46, 16, 16, 16, 16], 1):
        ws.column_dimensions[chr(64 + i)].width = w


def build_cash_book_xlsx(fiscal_year, rows, opening, totals, scope_name="ทุกบัญชี") -> str:
    wb = Workbook(); ws = wb.active; ws.title = "สมุดเงินสด"
    _cashbook_sheet(ws, "สมุดเงินสด", scope_name, fiscal_year, rows, opening, totals)
    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    path = out_dir / (_safe(f"สมุดเงินสด_ปีงบ{fiscal_year}_{scope_name}") + ".xlsx")
    wb.save(str(path))
    return str(path)


def _cb_fund_section_xlsx(ws, r, title, total_label, rows, open_by_fund=None):
    """เขียน 1 ส่วนสมุดเงินสด (ด้านรับ/ด้านจ่าย) แบบ 3 งบ ลงชีต คืน (แถวถัดไป, ยอดรวม)"""
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=7)
    tt = ws.cell(r, 1, title); tt.font = Font(name=THAI_FONT, bold=True, size=13)
    tt.alignment = Alignment(horizontal="left", vertical="center")
    r += 1
    heads = ["วันที่", "เลขที่เอกสาร", "รายการ", "เงินสด"] + _FUND_COLS
    for c, h in enumerate(heads, 1):
        _xcell(ws, r, c, h, bold=True, align="center", fill=_HEAD)
    r += 1
    tot = {"เงินสด": 0.0}
    for f in _FUND_COLS:
        tot[f] = 0.0

    def _wrow(date, ref, desc, cash, fund_amt: dict, *, bold=False, fill=None):
        _xcell(ws, r, 1, date, align="center", fill=fill)
        _xcell(ws, r, 2, ref, align="center", fill=fill)
        _xcell(ws, r, 3, desc, bold=bold, fill=fill, wrap=True)
        _xcell(ws, r, 4, cash or None, money=True, align="right", bold=bold, fill=fill)
        for i, f in enumerate(_FUND_COLS):
            _xcell(ws, r, 5 + i, fund_amt.get(f) or None, money=True, align="right", bold=bold, fill=fill)

    if open_by_fund is not None:
        op_cash = sum(open_by_fund.get(f, 0) for f in _FUND_COLS)
        _wrow("", "", "ยอดยกมา", op_cash, open_by_fund, bold=True)
        tot["เงินสด"] += op_cash
        for f in _FUND_COLS:
            tot[f] += open_by_fund.get(f, 0)
        r += 1
    for row in rows:
        amt = row["amount"] or 0
        _wrow(row["date"], row["ref"], row["desc"], amt, {row["fund"]: amt})
        tot["เงินสด"] += amt
        tot[row["fund"]] = tot.get(row["fund"], 0) + amt
        r += 1
    _xcell(ws, r, 1, "", fill=_TOTAL); _xcell(ws, r, 2, "", fill=_TOTAL)
    _xcell(ws, r, 3, total_label, bold=True, align="right", fill=_TOTAL)
    _xcell(ws, r, 4, tot["เงินสด"], bold=True, money=True, align="right", fill=_TOTAL)
    for i, f in enumerate(_FUND_COLS):
        _xcell(ws, r, 5 + i, tot[f], bold=True, money=True, align="right", fill=_TOTAL)
    return r + 2, tot


def build_cash_book_fund_xlsx(fiscal_year, scope_name, open_by_fund, receipts, payments) -> str:
    """สมุดเงินสดแบบราชการ (Excel): แยกด้านรับ/จ่าย + คอลัมน์ เงินสด + 3 งบ"""
    wb = Workbook(); ws = wb.active; ws.title = "สมุดเงินสด"
    ws.merge_cells("A1:G1")
    t = ws.cell(1, 1, f"สมุดเงินสด  ประจำปีงบประมาณ {fiscal_year}  ({scope_name})")
    t.font = Font(name=THAI_FONT, bold=True, size=16)
    t.alignment = Alignment(horizontal="center", vertical="center")
    r = 3
    r, rin = _cb_fund_section_xlsx(ws, r, "ด้านรับ  (เดบิต = เงินสด · เครดิตแยกตามประเภทเงิน)",
                                   "รวมด้านรับ", receipts, open_by_fund)
    r, rout = _cb_fund_section_xlsx(ws, r, "ด้านจ่าย  (เครดิต = เงินสด · เดบิตแยกตามประเภทเงิน)",
                                    "รวมด้านจ่าย", payments)
    # คงเหลือยกไป แยกประเภทเงิน
    _xcell(ws, r, 3, "เงินสดคงเหลือยกไป", bold=True, align="right", fill=_TOTAL)
    _xcell(ws, r, 4, rin["เงินสด"] - rout["เงินสด"], bold=True, money=True, align="right", fill=_TOTAL)
    for i, f in enumerate(_FUND_COLS):
        _xcell(ws, r, 5 + i, rin[f] - rout[f], bold=True, money=True, align="right", fill=_TOTAL)
    for i, w in enumerate([13, 16, 40, 15, 15, 16, 16], 1):
        ws.column_dimensions[chr(64 + i)].width = w
    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    path = out_dir / (_safe(f"สมุดเงินสด_ปีงบ{fiscal_year}_{scope_name}") + ".xlsx")
    wb.save(str(path))
    return str(path)


def build_ledger_xlsx(account, fiscal_year, rows, opening) -> str:
    wb = Workbook(); ws = wb.active; ws.title = "บัญชีแยกประเภท"[:31]
    ws.merge_cells("A1:G1")
    t = ws.cell(1, 1, f"บัญชีแยกประเภททั่วไป  บัญชี {account.name}  ปีงบประมาณ {fiscal_year}")
    t.font = Font(name=THAI_FONT, bold=True, size=16)
    t.alignment = Alignment(horizontal="center", vertical="center")
    heads = ["ว/ด/ป", "รายการ", "เลขที่เอกสาร", "เดบิต", "เครดิต", "ดุล", "คงเหลือ"]
    for c, h in enumerate(heads, 1):
        _xcell(ws, 3, c, h, bold=True, align="center", fill=_HEAD)
    _xcell(ws, 4, 1, ""); _xcell(ws, 4, 2, "ยอดยกมา", bold=True)
    _xcell(ws, 4, 3, ""); _xcell(ws, 4, 4, "", money=True); _xcell(ws, 4, 5, "", money=True)
    _xcell(ws, 4, 6, "เดบิต" if opening >= 0 else "เครดิต", align="center")
    _xcell(ws, 4, 7, abs(opening), bold=True, money=True, align="right")
    r = 5
    for row in rows:
        sub = row.get("subtotal")
        fill = _SUBTOT if sub else None
        _xcell(ws, r, 1, row.get("date", ""), align="center", fill=fill)
        _xcell(ws, r, 2, row.get("desc", ""), bold=bool(sub), fill=fill, wrap=True)
        _xcell(ws, r, 3, row.get("ref", ""), align="center", fill=fill)
        _xcell(ws, r, 4, row.get("debit") or None, money=True, align="right", bold=bool(sub), fill=fill)
        _xcell(ws, r, 5, row.get("credit") or None, money=True, align="right", bold=bool(sub), fill=fill)
        _xcell(ws, r, 6, row.get("side", ""), align="center", fill=fill)
        _xcell(ws, r, 7, row.get("balance"), money=True, align="right", bold=bool(sub), fill=fill)
        r += 1
    for i, w in enumerate([13, 42, 16, 15, 15, 8, 16], 1):
        ws.column_dimensions[chr(64 + i)].width = w
    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    path = out_dir / (_safe(f"บัญชีแยกประเภท_{account.name}_ปีงบ{fiscal_year}") + ".xlsx")
    wb.save(str(path))
    return str(path)
