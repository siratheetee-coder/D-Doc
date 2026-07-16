# -*- coding: utf-8 -*-
"""
cash_report.py
--------------
สร้างไฟล์ Word "รายงานเงินคงเหลือประจำวัน" (แบบมาตรฐานราชการ)
ดึงหมวด/บัญชีจากที่ผู้ใช้สร้างเอง + แยกคอลัมน์ เงินสด/เงินฝากธนาคาร/เงินฝากส่วนราชการผู้เบิก
ตามที่ระบุไว้ที่แต่ละบัญชี (deposit_type)
"""
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.services.doc_page import set_a4

from app.database import get_data_dir
from app.thai_utils import _THAI_MONTHS

THAI_FONT = "TH Sarabun New"
DEPOSIT_TYPES = {"cash": "เงินสด", "bank": "เงินฝากธนาคาร", "agency": "เงินฝากส่วนราชการผู้เบิก"}


def _safe(text: str) -> str:
    for ch in '<>:"/\\|?*':
        text = text.replace(ch, "_")
    return text.strip()


def be_date_parts(dt):
    """คืน (วัน, ชื่อเดือนไทย, ปีพ.ศ.) จาก datetime"""
    if not dt:
        return ("........", "............", "........")
    return (dt.day, _THAI_MONTHS[dt.month], dt.year + 543)


def _fmt(v):
    return "{:,.2f}".format(v) if v else "-"


def _set_cell(cell, text, *, bold=False, align="left", size=14, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = THAI_FONT
    r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    r._element.rPr.rFonts.set(qn("w:ascii"), THAI_FONT)
    r._element.rPr.rFonts.set(qn("w:hAnsi"), THAI_FONT)
    if fill is not None:
        tcpr = cell._tc.get_or_add_tcPr()
        shd = tcpr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): fill})
        tcpr.append(shd)


def _p(doc, text="", *, align="left", bold=False, size=14, after=2):
    p = doc.add_paragraph()
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = THAI_FONT
    r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    return p


def render_cash_report(school, rows, totals, as_of) -> str:
    """rows: list ของ dict {name, header(bool), indent(bool), cash, bank, agency, total}
    totals: dict {cash, bank, agency, total}"""
    doc = Document(); set_a4(doc)
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(1.5)
    base = doc.styles["Normal"]
    base.font.name = THAI_FONT
    base.font.size = Pt(14)
    base._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)

    _p(doc, "รายงานเงินคงเหลือประจำวัน " + (school.name or ""), align="center", bold=True, size=17, after=0)
    d, mon, be = be_date_parts(as_of)
    _p(doc, f"ประจำวันที่ {d} เดือน {mon} พ.ศ. {be}", align="center", size=15, after=6)

    headers = ["ประเภท", "เงินสด", "เงินฝากธนาคาร", "เงินฝากส่วนราชการผู้เบิก", "รวม", "หมายเหตุ"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    widths = [Cm(5.4), Cm(2.4), Cm(2.6), Cm(2.9), Cm(2.6), Cm(2.1)]   # รวม 18.0 = พื้นที่พิมพ์ A4 (21.0 - ขอบ 1.5x2)
    for c, (h, w) in enumerate(zip(headers, widths)):
        _set_cell(table.rows[0].cells[c], h, bold=True, align="center")
        table.rows[0].cells[c].width = w

    for row in rows:
        cells = table.add_row().cells
        lvl = row.get("level", 0)
        kind = row.get("kind", "leaf")
        name = ("    " * lvl) + row["name"]
        bold = kind in ("group", "sub")
        fill = "DCFCE7" if kind == "group" else ("F1F5F9" if kind == "sub" else None)
        _set_cell(cells[0], name, bold=bold, fill=fill)
        _set_cell(cells[1], _fmt(row.get("cash")), align="right", bold=bold, fill=fill)
        _set_cell(cells[2], _fmt(row.get("bank")), align="right", bold=bold, fill=fill)
        _set_cell(cells[3], _fmt(row.get("agency")), align="right", bold=bold, fill=fill)
        _set_cell(cells[4], _fmt(row.get("total")), align="right", bold=bold, fill=fill)
        _set_cell(cells[5], "", fill=fill)
        for c, w in enumerate(widths):
            cells[c].width = w

    # แถวรวม
    tcells = table.add_row().cells
    _set_cell(tcells[0], "รวม", bold=True, align="center")
    _set_cell(tcells[1], _fmt(totals.get("cash")), bold=True, align="right")
    _set_cell(tcells[2], _fmt(totals.get("bank")), bold=True, align="right")
    _set_cell(tcells[3], _fmt(totals.get("agency")), bold=True, align="right")
    _set_cell(tcells[4], _fmt(totals.get("total")), bold=True, align="right")
    _set_cell(tcells[5], "")
    for c, w in enumerate(widths):
        tcells[c].width = w

    # ลงนาม
    _p(doc, "", after=8)
    officer = (school.finance_officer_name or school.officer_name or "").strip()
    _p(doc, "ลงชื่อ.............................................ผู้จัดทำรายงาน", align="center", after=0)
    _p(doc, f"( {officer} )", align="center", after=0)
    _p(doc, "เจ้าหน้าที่การเงิน", align="center", after=8)

    _p(doc, "คณะกรรมการเก็บรักษาเงิน ได้ตรวจสอบนับเงินสดคงเหลือประจำวันถูกต้อง ตามรายการข้างต้นแล้ว "
            "และได้นำเงินสดเก็บรักษาไว้ในตู้นิรภัยเป็นที่เรียบร้อยแล้ว", align="left", after=8)
    _p(doc, "(ลงชื่อ)......................................กรรมการ      "
            "(ลงชื่อ)......................................กรรมการ      "
            "(ลงชื่อ)......................................กรรมการ", align="center", after=10)

    _p(doc, "ลงชื่อ.............................................", align="center", after=0)
    _p(doc, f"( {(school.director_name or '').strip()} )", align="center", after=0)
    director_pos = ("ผู้อำนวยการ" + school.name) if (school.name or "").startswith("โรงเรียน") \
        else (school.director_position or "ผู้อำนวยการโรงเรียน")
    _p(doc, director_pos, align="center", after=2)

    out_dir = get_data_dir() / "documents"
    out_dir.mkdir(exist_ok=True)
    out = out_dir / (_safe(f"รายงานเงินคงเหลือประจำวัน_{be}-{as_of.month:02d}-{as_of.day:02d}" if as_of else "รายงานเงินคงเหลือประจำวัน") + ".docx")
    doc.save(str(out))
    return str(out)
