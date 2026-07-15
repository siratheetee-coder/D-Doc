# -*- coding: utf-8 -*-
"""
lunch_receipt.py — ใบสำคัญรับเงิน (หลักฐานการรับเงินอุดหนุนอาหารกลางวันจาก อปท.)
ออกจากรายการ "รับ" ในบัญชีรับ-จ่ายอาหารกลางวัน
"""
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.database import get_data_dir
from app.thai_utils import thai_date, bahttext

THAI_FONT = "TH Sarabun New"


def _safe(text: str) -> str:
    for ch in '<>:"/\\|?*':
        text = text.replace(ch, "_")
    return text.strip()


def _p(doc, text="", *, align="left", bold=False, size=15, after=6, indent=0.0):
    p = doc.add_paragraph()
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = THAI_FONT
    r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    return p


def render_lunch_receipt(school, program, ledger) -> str:
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(1.5)
    base = doc.styles["Normal"]; base.font.name = THAI_FONT; base.font.size = Pt(15)
    base._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)

    _p(doc, "ใบสำคัญรับเงิน", align="center", bold=True, size=20, after=2)
    _p(doc, (school.name or ""), align="center", bold=True, size=16, after=0)
    _p(doc, "โครงการอาหารกลางวัน", align="center", size=14, after=10)

    d = ledger.date
    _p(doc, f"วันที่ {thai_date(d) if d else '.................................'}", align="right", after=10)

    amount = ledger.amount or 0
    org = (getattr(program, "funding_org", "") or "").strip() or "องค์กรปกครองส่วนท้องถิ่น"
    detail = (ledger.detail or "").strip() or "เงินอุดหนุนโครงการอาหารกลางวัน"
    ref = (ledger.ref or "").strip()

    _p(doc, f"ข้าพเจ้า ในนาม {school.name or '..............................'} "
            f"ได้รับเงินจาก {org} เป็นจำนวนเงิน {amount:,.2f} บาท "
            f"({bahttext(amount)}) เป็นค่า {detail}"
            + (f" งวด/อ้างอิงที่ {ref}" if ref else "")
            + " ไว้เป็นการถูกต้องเรียบร้อยแล้ว",
       align="justify", size=15, after=6, indent=1.25)
    _p(doc, "จึงได้ลงลายมือชื่อไว้เป็นหลักฐาน", align="justify", size=15, after=24, indent=1.25)

    fin = (getattr(school, "finance_officer_name", "") or getattr(school, "officer_name", "") or "").strip()
    director = (getattr(school, "director_name", "") or "").strip()

    # ลงนาม 2 ฝั่ง: ผู้รับเงิน (การเงิน) | ผู้อนุมัติ (ผอ.)
    tbl = doc.add_table(rows=1, cols=2)
    for cell, (line1, name, pos) in zip(tbl.rows[0].cells, [
            ("(ลงชื่อ)..............................ผู้รับเงิน", fin, "เจ้าหน้าที่การเงิน"),
            ("(ลงชื่อ)..............................ผู้อนุมัติ", director, None)]):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, txt in enumerate([line1, f"( {name or '.......................................'} )", pos or "ผู้อำนวยการโรงเรียน"]):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(txt); r.font.size = Pt(15); r.font.name = THAI_FONT
            r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)

    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    out = out_dir / (_safe(f"ใบสำคัญรับเงินอาหารกลางวัน_{ledger.id}") + ".docx")
    doc.save(str(out))
    return str(out)
