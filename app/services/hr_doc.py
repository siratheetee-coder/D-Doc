# -*- coding: utf-8 -*-
"""
hr_doc.py — เอกสารงานบุคคล (Word)
- render_leave_form  : ใบลา (ป่วย/กิจ/พักผ่อน ฯลฯ)
- render_certificate : หนังสือรับรอง (การเป็นบุคลากร / เงินเดือน)
"""
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.database import get_data_dir
from app.thai_utils import thai_date

THAI_FONT = "TH Sarabun New"
_BLANK = "................................"


def _safe(text: str) -> str:
    for ch in '<>:"/\\|?*':
        text = text.replace(ch, "_")
    return text.strip()


def _p(doc, text="", *, align="left", bold=False, size=15, after=6, indent=0.0):
    p = doc.add_paragraph()
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = THAI_FONT
    r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    return p


def _doc():
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(1.5)
    base = doc.styles["Normal"]; base.font.name = THAI_FONT; base.font.size = Pt(15)
    base._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    return doc


def _director_pos(school):
    return ("ผู้อำนวยการ" + school.name) if (school.name or "").startswith("โรงเรียน") \
        else (getattr(school, "director_position", "") or "ผู้อำนวยการโรงเรียน")


def _dt(d):
    return thai_date(d) if d else _BLANK


# ---------------- ใบลา ----------------
def render_leave_form(school, person, record, type_label) -> str:
    doc = _doc()
    _p(doc, "ใบลา" + type_label.replace("ลา", "", 1), align="center", bold=True, size=20, after=10)

    _p(doc, f"เขียนที่ {school.name or ''}", align="right", after=0)
    _p(doc, f"วันที่ {_dt(record.created_at)}", align="right", after=8)

    _p(doc, f"เรื่อง  ขอ{type_label}", after=2)
    _p(doc, f"เรียน  {_director_pos(school)}", after=6)

    name = person.name or _BLANK
    pos = person.position or "ครู"
    _p(doc, f"ข้าพเจ้า {name} ตำแหน่ง {pos} "
            f"ขอ{type_label}ตั้งแต่วันที่ {_dt(record.start_date)} ถึงวันที่ {_dt(record.end_date)} "
            f"มีกำหนด {record.days:g} วัน "
            + (f"เนื่องจาก {record.reason} " if (record.reason or '').strip() else "")
            + "ในระหว่างลาข้าพเจ้าจะติดต่อได้ที่ "
            + ((record.contact or '').strip() or _BLANK),
       align="justify", size=15, after=6, indent=1.25)
    _p(doc, "จึงเรียนมาเพื่อโปรดพิจารณาอนุญาต", align="justify", after=20, indent=1.25)

    _p(doc, "ขอแสดงความนับถือ", align="center", after=2)
    _p(doc, "(ลงชื่อ)...................................ผู้ลา", align="center", after=0)
    _p(doc, f"( {name} )", align="center", after=14)

    # ความเห็น/คำสั่งผู้บังคับบัญชา
    _p(doc, "คำสั่ง", bold=True, after=2)
    _p(doc, "☐ อนุญาต            ☐ ไม่อนุญาต", indent=1.25, after=14)
    _p(doc, "(ลงชื่อ)...................................", align="center", after=0)
    _p(doc, f"( {(getattr(school, 'director_name', '') or '').strip() or _BLANK} )", align="center", after=0)
    _p(doc, _director_pos(school), align="center", after=2)

    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    out = out_dir / (_safe(f"ใบลา_{name}_{record.id}") + ".docx")
    doc.save(str(out))
    return str(out)


# ---------------- หนังสือรับรองบุคลากร (ฟอร์มราชการ: ครุฑ + ที่ + ที่อยู่) ----------------
def _hdr_cell(cell, lines, *, align="left", size=15):
    cell.text = ""
    for i, txt in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(txt); r.font.size = Pt(size); r.font.name = THAI_FONT
        r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)


def render_certificate(school, person) -> str:
    """หนังสือรับรองบุคลากร — ฟอร์มราชการ (ครุฑ / ที่ / ที่อยู่โรงเรียน) + ลายเซ็น ผอ."""
    from app.services.build_templates import _krut_path
    from app.services.office_doc import _float_signature
    from app.thai_utils import _THAI_MONTHS
    import datetime as _d

    from docx.enum.table import WD_ALIGN_VERTICAL

    doc = _doc()
    # ---- หัวกระดาษแถวเดียว: ที่ (ซ้าย) | ครุฑ (กลาง) | ชื่อ+ที่อยู่โรงเรียน (ขวา) ----
    addr = [school.name or ""]
    if (getattr(school, "address", "") or "").strip():
        addr.append(school.address.strip())
    dp = " ".join(x for x in [
        ("อำเภอ" + school.district) if (getattr(school, "district", "") or "").strip() else "",
        ("จังหวัด" + school.province) if (getattr(school, "province", "") or "").strip() else "",
    ] if x)
    if dp:
        addr.append(dp)
    prefix = (getattr(school, "doc_prefix", "") or "").strip()

    t = doc.add_table(rows=1, cols=3)
    t.autofit = False
    left, mid, right = t.rows[0].cells
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _hdr_cell(left, [f"ที่  {prefix} ............/............".replace("  ", " ").strip()])
    mid.text = ""
    mp = mid.paragraphs[0]
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_after = Pt(0)
    krut = _krut_path()
    if krut:
        mp.add_run().add_picture(str(krut), height=Cm(2.0))
    _hdr_cell(right, addr, size=13.5)
    # ซ้าย+กลางกว้างขึ้น -> ครุฑอยู่กลางหน้า และบล็อกที่อยู่เลื่อนไปทางขวา
    left.width = Cm(6.0); mid.width = Cm(3.2); right.width = Cm(6.8)
    _p(doc, "", after=12)

    name = person.name or _BLANK
    pos = (person.position or "ครู").strip()
    rank = (person.rank or "").strip()
    sal = person.salary or 0

    area = (getattr(school, "area_office", "") or "").strip()
    sangkad = f"สังกัด{area} " if area else "สังกัด................................................ "
    body = (f"หนังสือฉบับนี้ให้ไว้เพื่อรับรองว่า {name} ปัจจุบันเป็นข้าราชการครูและ"
            f"บุคลากรทางการศึกษา ตำแหน่ง{pos}"
            + (f" วิทยฐานะ{rank}" if rank else "")
            + f" {school.name or 'โรงเรียน'} {sangkad}"
              f"สำนักงานคณะกรรมการการศึกษาขั้นพื้นฐาน กระทรวงศึกษาธิการ "
            + (f"รับเงินเดือนในอัตรา {sal:,.2f} บาท ({_baht(sal)}) " if sal else
               "รับเงินเดือนในอัตรา ........................ บาท (........................................บาทถ้วน) ")
            + "เริ่มรับราชการ ตั้งแต่วันที่ "
            + (thai_date(person.start_date) if person.start_date else "........................................")
            + " จนถึงปัจจุบันจริง")
    _p(doc, body, align="justify", size=15, after=6, indent=1.25)
    _p(doc, "หนังสือรับรองฉบับนี้ออกให้เพื่อนำไป"
            "................................................................ เท่านั้น",
       align="justify", size=15, after=14, indent=1.25)

    now = _d.datetime.now()
    _p(doc, f"ให้ไว้  ณ  วันที่ {now.day} เดือน{_THAI_MONTHS[now.month]} พ.ศ. {now.year + 543}",
       align="center", size=15, after=0)
    _p(doc, "", after=0, size=15)   # เว้นที่ให้ลายเซ็น ผอ.
    _p(doc, "", after=0, size=15)

    # ---- ลงนาม ผอ. (วางลายเซ็นจริงถ้ามี) ----
    director = (getattr(school, "director_name", "") or "").strip()
    sign_p = _p(doc, "(ลงชื่อ)..............................................", align="center", after=0)
    _float_signature(sign_p, director)
    _p(doc, f"( {director or _BLANK} )", align="center", after=0)
    _p(doc, _director_pos(school), align="center", after=2)

    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    out = out_dir / (_safe(f"หนังสือรับรองบุคลากร_{name}") + ".docx")
    doc.save(str(out))
    return str(out)


# ---------------- คำสั่งไปราชการ ----------------
def render_travel_order(school, person, record) -> str:
    doc = _doc()
    _p(doc, "คำสั่ง" + (school.name or "โรงเรียน"), align="center", bold=True, size=17, after=0)
    _p(doc, f"ที่  {record.doc_no or '............/............'}", align="center", size=14, after=0)
    _p(doc, "เรื่อง  ให้ข้าราชการครูและบุคลากรทางการศึกษาไปราชการ", align="center", bold=True, size=15, after=8)
    _p(doc, "───────────────────", align="center", after=8)

    name = person.name or _BLANK
    pos = person.position or "ครู"
    subject = (record.subject or "").strip() or "ปฏิบัติราชการ"
    place = (record.place or "").strip() or _BLANK
    _p(doc, f"ด้วย {school.name or 'โรงเรียน'} มีความจำเป็นต้องให้บุคลากรไปราชการเพื่อ {subject} "
            f"จึงอาศัยอำนาจตามความในมาตราที่เกี่ยวข้อง แต่งตั้งให้บุคคลดังต่อไปนี้ไปราชการ",
       align="justify", size=15, after=6, indent=1.25)
    _p(doc, f"{name} ตำแหน่ง {pos} "
            f"ไปราชการ ณ {place} "
            f"ตั้งแต่วันที่ {_dt(record.start_date)} ถึงวันที่ {_dt(record.end_date)} "
            f"รวม {record.days:g} วัน"
            + (f" โดยเบิกค่าใช้จ่ายในการเดินทางไปราชการ จำนวน {record.budget:,.2f} บาท" if (record.budget or 0) else ""),
       align="justify", size=15, after=6, indent=1.25)
    _p(doc, "ทั้งนี้ ให้ผู้ได้รับแต่งตั้งปฏิบัติหน้าที่ที่ได้รับมอบหมายด้วยความเรียบร้อย เกิดผลดีแก่ทางราชการ",
       align="justify", size=15, after=8, indent=1.25)
    _p(doc, f"สั่ง ณ วันที่ {_dt(record.doc_date)}", align="center", after=24)

    _p(doc, "(ลงชื่อ)...................................", align="center", after=0)
    _p(doc, f"( {(getattr(school, 'director_name', '') or '').strip() or _BLANK} )", align="center", after=0)
    _p(doc, _director_pos(school), align="center", after=2)

    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    out = out_dir / (_safe(f"คำสั่งไปราชการ_{name}_{record.id}") + ".docx")
    doc.save(str(out))
    return str(out)


# ---------------- ก.พ.7 (ทะเบียนประวัติ) ----------------
def _tc(cell, text, *, bold=False, align="left", fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(13); r.font.name = THAI_FONT
    r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    if fill:
        tcpr = cell._tc.get_or_add_tcPr()
        tcpr.append(tcpr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): fill}))


def render_kp7(school, person) -> str:
    doc = _doc()
    _p(doc, "ทะเบียนประวัติบุคลากร (ก.พ.7)", align="center", bold=True, size=18, after=0)
    _p(doc, (school.name or ""), align="center", bold=True, size=15, after=8)

    # ข้อมูลส่วนตัว
    info = [
        ("ชื่อ-นามสกุล", person.name or _BLANK),
        ("ประเภท/ตำแหน่ง", f"{person.person_type or '-'} · {person.position or '-'}"),
        ("วิทยฐานะ/ระดับ", person.rank or "-"),
        ("เลขประจำตัวประชาชน", person.id_card or "-"),
        ("วันเดือนปีเกิด", thai_date(person.birthdate) if person.birthdate else "-"),
        ("วันบรรจุ/เริ่มปฏิบัติงาน", thai_date(person.start_date) if person.start_date else "-"),
        ("เงินเดือนปัจจุบัน", f"{(person.salary or 0):,.2f} บาท" if person.salary else "-"),
        ("โทรศัพท์ / อีเมล", f"{person.phone or '-'} / {person.email or '-'}"),
    ]
    t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
    for label, val in info:
        cells = t.add_row().cells
        _tc(cells[0], label, bold=True, fill="F1F5F9"); _tc(cells[1], val)
        cells[0].width = Cm(5.5); cells[1].width = Cm(11)
    _p(doc, "", after=6)

    # ประวัติการดำรงตำแหน่ง/วิทยฐานะ
    _p(doc, "ประวัติการดำรงตำแหน่ง / เลื่อนวิทยฐานะ", bold=True, size=14, after=2)
    rt = doc.add_table(rows=1, cols=4); rt.style = "Table Grid"
    for c, h in enumerate(["วันที่", "ตำแหน่ง", "วิทยฐานะ/ระดับ", "เลขที่คำสั่ง"]):
        _tc(rt.rows[0].cells[c], h, bold=True, align="center", fill="CCFBF1")
    rows = sorted(person.rank_history, key=lambda r: (r.date or __import__('datetime').datetime.min))
    for r in rows:
        cells = rt.add_row().cells
        _tc(cells[0], thai_date(r.date) if r.date else "-", align="center")
        _tc(cells[1], r.position or "-"); _tc(cells[2], r.rank or "-")
        _tc(cells[3], r.doc_no or "-", align="center")
    if not rows:
        _tc(rt.add_row().cells[0], "-")
    _p(doc, "", after=6)

    # เครื่องราชอิสริยาภรณ์
    _p(doc, "เครื่องราชอิสริยาภรณ์ที่ได้รับ", bold=True, size=14, after=2)
    dt = doc.add_table(rows=1, cols=3); dt.style = "Table Grid"
    for c, h in enumerate(["ปี พ.ศ.", "ชั้นตรา", "เลขที่ประกาศ"]):
        _tc(dt.rows[0].cells[c], h, bold=True, align="center", fill="CCFBF1")
    decs = sorted(person.decorations, key=lambda d: (d.year or 0))
    for d in decs:
        cells = dt.add_row().cells
        _tc(cells[0], d.year or "-", align="center"); _tc(cells[1], d.name or "-")
        _tc(cells[2], d.ref or "-", align="center")
    if not decs:
        _tc(dt.add_row().cells[0], "-")

    _p(doc, "", after=16)
    _p(doc, "(ลงชื่อ)...................................ผู้จัดทำ/รับรอง", align="center", after=0)
    _p(doc, f"( {(getattr(school, 'director_name', '') or '').strip() or _BLANK} )", align="center", after=0)
    _p(doc, _director_pos(school), align="center", after=2)

    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    out = out_dir / (_safe(f"กพ7_{person.name}") + ".docx")
    doc.save(str(out))
    return str(out)


def _baht(v):
    try:
        from app.thai_utils import bahttext
        return bahttext(v)
    except Exception:
        return ""
