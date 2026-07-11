# -*- coding: utf-8 -*-
"""
proc_alt_doc.py
---------------
สร้างไฟล์ Word เอกสารจัดซื้อ "วิธีพิเศษ" แบบย่อ (สร้าง docx ตรงตอนรันไทม์)
- render_w804(proc, school)    : ว.804 ซื้อไม่เกิน 50,000 (ชุด 4 ส่วน)
- render_w119_t1(proc, school) : ว.119 ตารางที่ 1 ซื้อพัสดุไม่เกิน 10,000 (บันทึกเดียว)
- render_w119_t2(proc, school) : ว.119 ตารางที่ 2 ค่าบริหาร/ฝึกอบรม (บันทึกเดียว)

อ้างถ้อยคำ/โครงสร้างจากไฟล์ตัวอย่างราชการที่โรงเรียนใช้จริง ใช้ helper ร่วมกับ build_templates
(ฟอนต์ TH Sarabun, ครุฑ, จัดกระจายแบบไทย) เหมือน finance_doc.py / asset_dispose_doc.py
"""
from pathlib import Path

from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.database import get_data_dir
from app.thai_utils import thai_date, bahttext
from app.services.build_templates import (
    _font, _krut_and_title, _hr, _p, _p_runs, _sign_table, _set_cell,
    _repeat_header_row, _no_split_row, THAI_FONT, _csize, _bcs,
)

_ALIGN = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
          "right": WD_ALIGN_PARAGRAPH.RIGHT}


def _cell_ml(cell, lines, *, size=15):
    """เขียนหลายบรรทัดในเซลล์เดียว (ฟอนต์ไทย) — lines = [(text, align, bold), ...]"""
    cell.text = ""
    for i, (text, align, bold) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = _ALIGN[align]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(text)
        r.font.name = THAI_FONT
        _csize(r, size)
        _bcs(r, bold)
        r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)


def _opinion_box(doc, head_officer, approver):
    """กล่องมีเส้นขอบ 2 คอลัมน์ (ตาม template ว.804):
    ซ้าย = ความเห็นของหัวหน้าเจ้าหน้าที่ + ช่องลงนาม
    ขวา = ☐ เห็นชอบ / ☐ ลงนามแล้ว + ช่องลงนาม"""
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.autofit = False
    left, right = t.rows[0].cells
    left.width = Cm(8.0)
    right.width = Cm(8.0)
    _line = "ลงชื่อ......................................."
    _date = "วันที่......................................."
    _cell_ml(left, [
        ("ความเห็นของหัวหน้าเจ้าหน้าที่", "left", True),
        ("", "left", False),
        (_line, "center", False),
        (f"( {head_officer} )", "center", False),
        (_date, "center", False),
    ])
    _cell_ml(right, [
        ("☐  เห็นชอบ", "left", False),
        ("☐  ลงนามแล้ว", "left", False),
        ("", "left", False),
        (_line, "center", False),
        (f"( {approver} )", "center", False),
        (_date, "center", False),
    ])
    return t

_BLANK = "............................"
_BLANK_S = "................"


def _safe(text: str) -> str:
    for ch in '<>:"/\\|?*\n\r\t':
        text = text.replace(ch, "_")
    return text.strip()[:80]


def _school_office(school) -> str:
    parts = [school.name or "", school.address or ""]
    return "  ".join(p for p in parts if p).strip()


def _director_office(school) -> str:
    name = (school.name or "").strip()
    if name.startswith("โรงเรียน"):
        return "ผู้อำนวยการ" + name
    return getattr(school, "director_position", None) or "ผู้อำนวยการโรงเรียน"


def _director_line(school) -> str:
    name = (school.name or "").strip()
    return "ผู้อำนวยการ" + name if name.startswith("โรงเรียน") else "ผู้อำนวยการโรงเรียน"


def _money(x) -> str:
    x = round(float(x or 0), 2)
    return f"{int(x):,}" if x == int(x) else f"{x:,.2f}"


def _project(proc) -> str:
    return (proc.project_name or (proc.project.name if getattr(proc, "project", None) else "")
            or proc.subject or "").strip()


def _items(proc) -> list:
    return list(proc.items or [])


def _total(proc) -> float:
    items = _items(proc)
    if proc.total_amount:
        return float(proc.total_amount)
    return float(sum((i.quantity or 0) * (i.unit_price or 0) for i in items))


def _extra(proc) -> dict:
    """อ่านฟิลด์เสริมเฉพาะรูปแบบจาก proc.case_extra (JSON)"""
    import json
    try:
        return json.loads(proc.case_extra) if getattr(proc, "case_extra", "") else {}
    except Exception:
        return {}


def _x(extra: dict, key: str, blank: str = _BLANK_S) -> str:
    """ค่าฟิลด์เสริม หรือจุดไข่ปลาถ้าเว้นว่าง"""
    v = (extra.get(key) or "").strip()
    return v or blank


def _xdate(extra: dict, key: str, blank: str = _BLANK_S) -> str:
    """ค่าฟิลด์เสริมชนิดวันที่ -> แปลงเป็นวันที่ไทยเต็ม (เช่น ๒๒ มิถุนายน พ.ศ. ๒๕๖๙)"""
    from app.thai_utils import parse_be_date
    v = (extra.get(key) or "").strip()
    if not v:
        return blank
    d = parse_be_date(v)
    return thai_date(d) if d else v


def _save(doc, name: str) -> str:
    out_dir = get_data_dir() / "documents"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / (_safe(name) + ".docx")
    doc.save(str(out_path))
    return str(out_path)


def _grid(doc, headers, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    hdr = t.rows[0]
    _repeat_header_row(hdr)
    _no_split_row(hdr)
    for c, h, w in zip(hdr.cells, headers, widths):
        _set_cell(c, h, bold=True, align="center", size=14)
        c.width = w
    return t


def _grow(t, values, widths, aligns, *, bold=False):
    r = t.add_row()
    _no_split_row(r)
    for c, v, w, a in zip(r.cells, values, widths, aligns):
        _set_cell(c, v, align=a, size=14, bold=bold)
        c.width = w
    return r


def _header(doc, *, subject_line: str, school, doc_no="", doc_date=None,
            via_head=False):
    """หัวบันทึกข้อความ: ครุฑ + ส่วนราชการ/ที่/วันที่/เรื่อง/เรียน"""
    _krut_and_title(doc)
    if isinstance(doc_date, str):
        date_txt = doc_date.strip() or _BLANK
    else:
        date_txt = thai_date(doc_date) if doc_date else _BLANK
    _p_runs(doc, [("ส่วนราชการ  ", True), (_school_office(school), False)])
    _p_runs(doc, [("ที่  ", True), (doc_no or _BLANK, False),
                  ("\t", False), ("วันที่ ", True), (date_txt, False)], tab_cm=8)
    _p_runs(doc, [("เรื่อง  ", True), (subject_line, False)])
    rian = _director_line(school) + ("  ผ่าน หัวหน้าเจ้าหน้าที่" if via_head else "")
    _p_runs(doc, [("เรียน  ", True), (rian, False)])
    _hr(doc)


# ============================================================
# ว.804 — ซื้อไม่เกิน 50,000 บาท (ชุด 4 ส่วน)
# ============================================================
_W804_REF = ("หนังสือคณะกรรมการวินิจฉัยปัญหาการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ "
             "ด่วนที่สุด ที่ กค (กวจ) 0405.2/ว 804 ลงวันที่ 12 พฤศจิกายน 2568 "
             "เรื่อง แนวทางปฏิบัติสำหรับการจัดซื้อวงเงินไม่เกิน 50,000 บาท")


def render_w804(proc, school) -> str:
    doc = _font_doc()
    items = _items(proc)
    total = _total(proc)
    subject = (proc.subject or "").strip()
    fy = proc.fiscal_year or ""
    officer = (getattr(school, "officer_name", "") or "").strip()
    n = len(items) or 1
    ex = _extra(proc)

    # ---------- ส่วนที่ 1: รายงานขอซื้อ ----------
    _header(doc, subject_line=f"รายงานขอซื้อ{subject} โดยวิธีเฉพาะเจาะจง",
            school=school, doc_no=(proc.memo_no or "").strip(),
            doc_date=proc.request_date, via_head=True)
    _p(doc, f"ด้วย {school.name or 'โรงเรียน'} มีความประสงค์จะซื้อ{subject} "
            f"จำนวน {n} รายการ สำหรับโครงการ/กิจกรรม{_project(proc)} โดยวิธีเฉพาะเจาะจง "
            f"ประจำปีงบประมาณ พ.ศ. {fy} ซึ่งมีรายละเอียดดังต่อไปนี้",
       align="justify", indent=1.25, after=2)
    _p(doc, "๑. รายละเอียดคุณลักษณะเฉพาะของพัสดุและขอบเขตงาน : รายละเอียดตามคุณลักษณะเฉพาะ"
            "ของพัสดุหรือขอบเขตของงานที่แนบ", align="justify", indent=1.25, after=2)
    _p(doc, f"๒. วงเงินที่จะซื้อ : วงเงินงบประมาณ {_money(total)} บาท ({bahttext(total)}) "
            f"โดยใช้จ่ายจากเงินงบประมาณรายจ่ายประจำปีงบประมาณ พ.ศ. {fy}",
       align="justify", indent=1.25, after=2)
    _p(doc, "๓. ราคากลางของพัสดุที่จะซื้อ : ให้ถือวงเงินตามข้อ ๒ เป็นราคากลางของพัสดุที่จะซื้อในครั้งนี้",
       align="justify", indent=1.25, after=2)
    _p(doc, f"๔. ผู้รับผิดชอบในการจัดซื้อ : มอบหมายให้ {officer or _BLANK} "
            f"ตำแหน่ง {_BLANK_S} เป็นผู้มีหน้าที่ดำเนินการจัดซื้อกับผู้ประกอบการที่มีอาชีพขายพัสดุ"
            "โดยตรง และเป็นผู้ตรวจรับพัสดุในครั้งนี้", align="justify", indent=1.25, after=2)
    _p(doc, f"การดำเนินการจัดซื้อครั้งนี้ เป็นไปตาม{_W804_REF}",
       align="justify", indent=1.25, after=2)
    _p(doc, "จึงเรียนมาเพื่อโปรดพิจารณา หากเห็นชอบได้โปรด", align="justify", indent=1.25, after=2)
    _p(doc, "๑. เห็นชอบรายละเอียดคุณลักษณะเฉพาะของพัสดุหรือขอบเขตของงานตามข้อ ๑",
       align="justify", indent=1.25, after=2)
    _p(doc, "๒. อนุมัติให้ดำเนินการ และมอบหมายผู้รับผิดชอบดำเนินการจัดซื้อ ตามรายละเอียด"
            "ในรายงานขอซื้อดังกล่าวข้างต้น", align="justify", indent=1.25, after=10)
    _sign_table(doc, [[
        ("ลงชื่อ.........................................", "center"),
        (f"( {officer or _BLANK} )", "center"),
        ("เจ้าหน้าที่", "center"),
    ]])
    _p(doc, "", after=4)
    head_officer = (getattr(school, "head_officer_name", "") or "").strip() or _BLANK
    director = (getattr(school, "director_name", "") or "").strip() or _BLANK
    _opinion_box(doc, head_officer, director)

    # ---------- ส่วนที่ 2: รายละเอียดคุณลักษณะเฉพาะ (แนบท้าย) ----------
    doc.add_page_break()
    _p(doc, "รายละเอียดคุณลักษณะเฉพาะของพัสดุที่จะดำเนินการซื้อ", align="center", bold=True, after=2)
    _p(doc, f"งานจัดซื้อวัสดุโครงการ/กิจกรรม{_project(proc)} จำนวน {n} รายการ",
       align="center", after=1)
    _p(doc, f"แนบท้ายบันทึกข้อความ ที่ {(proc.memo_no or '').strip() or _BLANK_S} "
            f"ลงวันที่ {thai_date(proc.request_date) if proc.request_date else _BLANK_S}",
       align="center", after=4)
    w = [Cm(1.5), Cm(10.5), Cm(2.5), Cm(2.5)]
    t = _grid(doc, ["ลำดับ", "ชื่อพัสดุ และรายละเอียดคุณลักษณะเฉพาะ", "จำนวน", "หน่วยนับ"], w)
    for i, it in enumerate(items, 1):
        _grow(t, [str(i), (it.name or "").strip(), _money(it.quantity), (it.unit or "").strip()],
              w, ["center", "left", "center", "center"])
    if not items:
        _grow(t, ["", "", "", ""], w, ["center", "left", "center", "center"])
    _p(doc, "", after=8)
    _sign_table(doc, [[
        ("ผู้รับผิดชอบจัดทำคุณลักษณะเฉพาะ", "center"),
        ("ลงชื่อ..................................................................", "center"),
        (f"( {officer or _BLANK} )", "center"),
    ]])

    # ---------- ส่วนที่ 3: รายงานสรุปผล + ขออนุมัติเบิกจ่าย ----------
    doc.add_page_break()
    _header(doc, subject_line=f"รายงานสรุปผลการจัดซื้อ{subject} โดยวิธีเฉพาะเจาะจง และอนุมัติเบิกจ่าย",
            school=school, doc_no=ex.get("report2_no", "").strip() or (proc.memo_no or "").strip(),
            doc_date=(ex.get("report2_date", "").strip() or proc.request_date))
    _p(doc, f"ตามที่ {school.name or 'โรงเรียน'} เห็นชอบให้ดำเนินการซื้อ{subject} "
            f"จำนวน {n} รายการ สำหรับโครงการ/กิจกรรม{_project(proc)} โดยวิธีเฉพาะเจาะจง "
            f"ประจำปีงบประมาณ พ.ศ. {fy} นั้น", align="justify", indent=1.25, after=2)
    _p(doc, f"เจ้าหน้าที่ผู้รับผิดชอบได้ดำเนินการจัดซื้อกับผู้ประกอบการโดยตรงตาม{_W804_REF} "
            "ขอรายงานผลการดำเนินการ ดังนี้", align="justify", indent=1.25, after=2)
    _p_runs(doc, [("ชื่อผู้ประกอบการ บริษัท/ห้าง/ร้าน  ", False),
                  ((proc.vendor.name if proc.vendor else _BLANK), False)], indent=1.25)
    w2 = [Cm(7.5), Cm(2.0), Cm(2.0), Cm(2.7), Cm(2.8)]
    t2 = _grid(doc, ["รายการจัดซื้อ", "จำนวน", "หน่วยนับ", "ราคาต่อหน่วย", "จำนวนเงิน"], w2)
    for it in items:
        amt = (it.quantity or 0) * (it.unit_price or 0)
        _grow(t2, [(it.name or "").strip(), _money(it.quantity), (it.unit or "").strip(),
                   _money(it.unit_price), _money(amt)], w2,
              ["left", "center", "center", "right", "right"])
    rtot = t2.add_row()
    _no_split_row(rtot)
    _set_cell(rtot.cells[0], "รวมเป็นเงิน", bold=True, align="center", size=14)
    rtot.cells[0].merge(rtot.cells[3])
    _set_cell(rtot.cells[4], _money(total), bold=True, align="right", size=14)
    _p(doc, "**ซึ่งเป็นราคาที่รวมภาษีมูลค่าเพิ่มและค่าใช้จ่ายทั้งปวงไว้แล้วด้วย",
       align="justify", indent=0, after=2, size=14)
    _p(doc, "การจัดซื้อคราวนี้ไม่เกินวงเงินที่ประมาณไว้ และเห็นว่าเป็นราคาที่เหมาะสม และได้ใช้"
            f"เงินสดสำรองจ่าย ชำระให้ผู้ขายเรียบร้อยแล้ว ตามหลักฐานการรับเงิน ใบเสร็จรับเงิน "
            f"เล่มที่ {_x(ex,'receipt_book')} เลขที่ {_x(ex,'receipt_no')} ลงวันที่ {_xdate(ex,'receipt_date')} "
            f"และได้รับมอบพัสดุไว้ครบถ้วนเรียบร้อยแล้ว เมื่อวันที่ {_xdate(ex,'deliver_date')}",
       align="justify", indent=1.25, after=8)
    _sign_table(doc, [[
        ("ลงชื่อ…………………………………………ผู้ตรวจรับพัสดุ (ผู้ได้รับมอบหมาย)", "center"),
        (f"( {officer or _BLANK} )", "center"),
    ]])
    _p(doc, "จึงเรียนมาเพื่อโปรดทราบและขออนุมัติเบิกจ่ายต่อไป", align="justify", indent=1.25, after=10)
    _sign_table(doc, [[
        ("ลงชื่อ…………………………………………เจ้าหน้าที่", "center"),
        (f"( {officer or _BLANK} )", "center"),
    ]])
    _p(doc, "☐  ทราบ", after=1)
    _p(doc, "☐  อนุมัติ", after=10)
    _sign_table(doc, [[
        ("ลงชื่อ…………………………………………", "center"),
        (f"( {(getattr(school, 'director_name', '') or '').strip() or _BLANK} )", "center"),
        ("ตำแหน่ง " + _director_office(school), "center"),
        ("วันที่........................................", "center"),
    ]])

    # ---------- ส่วนที่ 4: ใบติดใบเสร็จรับเงิน ----------
    doc.add_page_break()
    _p(doc, "ใบติดใบเสร็จรับเงิน", align="center", bold=True, size=18, after=2)
    # เว้นพื้นที่ว่างกลางหน้าไว้สำหรับติดใบเสร็จจริง แล้วดันข้อความ+ลายเซ็นไปอยู่ล่างหน้า
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Cm(17)
    _p(doc, f"จำนวนเงินตามใบเสร็จรับเงิน เล่มที่ {_x(ex,'receipt_book')} เลขที่ {_x(ex,'receipt_no')} "
            f"ลงวันที่ {_xdate(ex,'receipt_date')}", align="justify", indent=1.25, after=2)
    _p(doc, f"ข้าพเจ้าได้ทดรองจ่ายไปก่อนแล้ว เป็นจำนวนเงิน {_money(total)} บาท ({bahttext(total)}) "
            f"และได้รับมอบพัสดุไว้ครบถ้วนถูกต้องแล้ว ข้าพเจ้าขอเบิกเงินจัดซื้อ{subject} "
            f"โดยวิธีเฉพาะเจาะจง จำนวนเงิน {_money(total)} บาท ({bahttext(total)})",
       align="justify", indent=1.25, after=10)
    _sign_table(doc, [[
        ("ลงชื่อ..............................................ผู้ขอเบิก", "center"),
        (f"( {officer or _BLANK} )", "center"),
        ("ตำแหน่ง.......................................................", "center"),
    ]])

    return _save(doc, f"ว804_{(proc.memo_no or proc.id)}_{subject}")


# ============================================================
# ว.119 — ตารางที่ 1 / ตารางที่ 2 (โครงสร้างร่วม)
# ============================================================
_W119_REF = ("หนังสือด่วนที่สุด ที่ กค (กวจ) 0405.2/ว 119 ลงวันที่ 7 มีนาคม 2561 "
             "เรื่อง แนวทางการปฏิบัติในการดำเนินการจัดหาพัสดุที่เกี่ยวกับค่าใช้จ่ายในการบริหารงาน "
             "ค่าใช้จ่ายในการฝึกอบรม การจัดงาน และการประชุมของหน่วยงานของรัฐ")


def render_w119_t1(proc, school) -> str:
    """ว.119 ตารางที่ 1 — ซื้อพัสดุไม่เกิน 10,000 บาท (มีตารางรายการ + ลงนาม 3 ฝ่าย)"""
    doc = _font_doc()
    items = _items(proc)
    total = _total(proc)
    officer = (getattr(school, "officer_name", "") or "").strip()
    n = len(items) or 1
    ex = _extra(proc)
    _header(doc, subject_line="รายงานขอความเห็นชอบการจัดซื้อจัดจ้าง และขออนุมัติเบิกจ่ายเงิน",
            school=school, doc_no=(proc.memo_no or "").strip(), doc_date=proc.request_date)
    _p(doc, f"ด้วยข้าพเจ้า {officer or _BLANK} ได้รับอนุมัติให้ดำเนินการตามกิจกรรม"
            f"{_project(proc)} โครงการ{_project(proc)} และได้ดำเนินการจัดซื้อจัดจ้าง "
            f"จำนวน {n} รายการ เพื่อใช้ในกิจกรรมดังกล่าว โดยมีรายละเอียด ดังนี้",
       align="justify", indent=1.25, after=4)
    w = [Cm(1.2), Cm(6.3), Cm(1.6), Cm(2.3), Cm(2.6), Cm(2.5)]
    t = _grid(doc, ["ที่", "รายการ", "จำนวน", "ราคาต่อหน่วย", "จำนวนเงิน", "หลักฐานลำดับที่"], w)
    for i, it in enumerate(items, 1):
        amt = (it.quantity or 0) * (it.unit_price or 0)
        _grow(t, [str(i), (it.name or "").strip(), _money(it.quantity),
                  _money(it.unit_price), _money(amt), str(i)], w,
              ["center", "left", "center", "right", "right", "center"])
    rtot = t.add_row()
    _no_split_row(rtot)
    _set_cell(rtot.cells[0], "รวมทั้งสิ้น", bold=True, align="center", size=14)
    rtot.cells[0].merge(rtot.cells[3])
    _set_cell(rtot.cells[4], _money(total), bold=True, align="right", size=14)
    _set_cell(rtot.cells[5], "", align="center", size=14)
    _p(doc, f"รวมทั้งสิ้น (ตัวอักษร) {bahttext(total)}", indent=1.25, after=4, size=15)
    _p(doc, f"ทั้งนี้ การดำเนินการจัดซื้อดังกล่าว เป็นการดำเนินการตาม{_W119_REF} "
            "ตามตารางที่ 1 รายการค่าใช้จ่ายที่เป็นการจัดซื้อจัดจ้าง ฯ",
       align="justify", indent=1.25, after=2)
    _p(doc, "จึงเรียนมาเพื่อโปรดพิจารณาให้ความเห็นชอบ และให้ถือรายงานนี้เป็นหลักฐานการตรวจรับพัสดุ "
            "โดยอนุโลม", align="justify", indent=1.25, after=10)
    _sign_table(doc, [[
        ("ผู้รับผิดชอบกิจกรรม", "center"),
        ("(ลงชื่อ).....................................................", "center"),
        (f"( {officer or _BLANK} )", "center"),
        ("ตำแหน่ง...........................................", "center"),
    ]])
    # ความเห็นงานการเงิน
    fin = (getattr(school, "finance_officer_name", "") or "").strip()
    _p(doc, "ความเห็นของงานการเงิน", bold=True, indent=1.25, after=1, size=15)
    _p(doc, f"เรียน {_director_line(school)}", indent=1.25, after=1)
    _p(doc, "โปรดพิจารณา", indent=1.25, after=1)
    _p(doc, "๑. ให้ความเห็นชอบการจัดซื้อจัดจ้างดังกล่าวข้างต้น", indent=1.5, after=1)
    _p(doc, f"๒. อนุมัติให้จ่ายเงิน จำนวน {_money(total)} บาท ({bahttext(total)}) "
            f"ให้แก่ {_x(ex,'advance_payer',_BLANK)} ผู้ทดรองจ่าย/ผู้ยืมเงิน", indent=1.5, after=8)
    _sign_table(doc, [[
        ("(ลงชื่อ)............................................ เจ้าหน้าที่การเงิน", "center"),
        (f"( {fin or _BLANK} )", "center"),
    ]])
    # ความเห็นงานแผน/พัสดุ
    _p(doc, "ความเห็นของงานแผนงานและงานพัสดุ", bold=True, indent=1.25, after=1, size=15)
    _bk = (ex.get("budget_kind") or "").strip()
    _p(doc, "( ) ใช้งบตามกิจกรรม/โครงการที่อนุมัติไว้ตามแผน   งบ "
            + (_bk if _bk else "( ) รายหัว  ( ) 15 ปี  ( ) อื่น ๆ .............................."),
       indent=1.25, after=8)
    _sign_table(doc, [
        [("(ลงชื่อ).............................................เจ้าหน้าที่พัสดุ", "center"),
         (f"( {officer or _BLANK} )", "center")],
        [("(ลงชื่อ).............................................หัวหน้าเจ้าหน้าที่", "center"),
         (f"( {(getattr(school, 'head_officer_name', '') or '').strip() or _BLANK} )", "center")],
    ])
    # ความเห็น ผอ.
    _p(doc, "ความเห็นของผู้อำนวยการโรงเรียน", bold=True, indent=1.25, after=1, size=15)
    _p(doc, "( ) อนุมัติ   ( ) ไม่อนุมัติ เนื่องจาก ...............................................................",
       indent=1.25, after=8)
    _sign_table(doc, [[
        ("(ลงชื่อ) .......................................................................", "center"),
        (f"( {(getattr(school, 'director_name', '') or '').strip() or _BLANK} )", "center"),
        (_director_office(school), "center"),
        ("วันที่ .................................................", "center"),
    ]])
    return _save(doc, f"ว119ตาราง1_{(proc.memo_no or proc.id)}_{(proc.subject or '').strip()}")


def render_w119_t2(proc, school) -> str:
    """ว.119 ตารางที่ 2 — ค่าบริหาร/ฝึกอบรม (ความเรียง + ลงนาม 2 ฝ่าย)"""
    doc = _font_doc()
    total = _total(proc)
    n = len(_items(proc)) or 1
    ex = _extra(proc)
    officer = (getattr(school, "officer_name", "") or "").strip()
    vendor = proc.vendor.name if proc.vendor else _BLANK
    dept = (proc.department or "").strip() or (school.name or "หน่วยงาน")
    _header(doc, subject_line="รายงานขอความเห็นชอบการจัดซื้อจัดจ้าง",
            school=school, doc_no=(proc.memo_no or "").strip(), doc_date=proc.request_date)
    _p(doc, f"ด้วย {dept} ได้ดำเนินการจัดซื้อจัดจ้าง{(proc.subject or '').strip()} "
            f"จำนวน {n} รายการ รวมเป็นจำนวนเงินทั้งสิ้น {_money(total)} บาท ({bahttext(total)}) "
            f"และประสงค์จะรายงานขอความเห็นชอบในการดำเนินการจัดซื้อจัดจ้างในครั้งนี้ "
            f"ตามหลักฐานการจัดซื้อจัดจ้างเป็น ใบส่งของ/ใบแจ้งหนี้/ใบเสร็จรับเงิน/ใบสำคัญรับเงิน "
            f"ของ {vendor} เล่มที่ {_x(ex,'receipt_book')} เลขที่ {_x(ex,'receipt_no')} "
            f"วันที่ {_xdate(ex,'receipt_date')} เป็นเงิน {_money(total)} บาท ({bahttext(total)})",
       align="justify", indent=1.25, after=2)
    _p(doc, f"ทั้งนี้ การดำเนินการจัดซื้อจัดจ้างดังกล่าวนี้ เป็นไปตาม{_W119_REF} "
            "ตามตารางที่ 2 กรณีจัดซื้อจัดจ้างพัสดุที่เกี่ยวกับค่าใช้จ่ายในการบริหารงานที่มีวงเงิน"
            "การจัดซื้อจัดจ้างครั้งหนึ่งไม่เกิน 10,000 บาท", align="justify", indent=1.25, after=2)
    _p(doc, "จึงเรียนมาเพื่อโปรดพิจารณา", align="justify", indent=1.25, after=2)
    _p(doc, "๑. ให้ความเห็นชอบในการดำเนินการจัดซื้อจัดจ้าง ตามรายละเอียดข้างต้น และให้ถือรายงานนี้"
            "เป็นการตรวจรับพัสดุโดยอนุโลม", indent=1.25, after=2)
    _p(doc, f"๒. อนุมัติให้จ่ายเงินจำนวน {_money(total)} บาท ({bahttext(total)}) ให้แก่ {vendor}",
       indent=1.25, after=10)
    _sign_table(doc, [[
        ("ลงชื่อ ...................................... เจ้าหน้าที่ผู้รับผิดชอบ", "center"),
        (f"( {officer or _BLANK} )", "center"),
    ]])
    fin_head = (getattr(school, "finance_head_name", "") or
                getattr(school, "finance_officer_name", "") or "").strip()
    _sign_table(doc, [
        [("ความเห็นหัวหน้างานการเงิน", "left"),
         ("(   ) เห็นสมควรตามที่เสนอ", "left"),
         ("(   ) ไม่เห็นสมควร เนื่องจาก..............................", "left"),
         ("ลงชื่อ................................หัวหน้างานการเงิน", "left"),
         (f"( {fin_head or _BLANK} )", "left")],
        [("ผู้มีอำนาจลงนาม", "left"),
         ("(   ) เห็นชอบ/อนุมัติ", "left"),
         ("(   ) ไม่เห็นชอบ/ไม่อนุมัติ", "left"),
         ("ลงชื่อ.........................................", "left"),
         (f"( {(getattr(school, 'director_name', '') or '').strip() or _BLANK} )", "left"),
         (_director_office(school), "left")],
    ])
    return _save(doc, f"ว119ตาราง2_{(proc.memo_no or proc.id)}_{(proc.subject or '').strip()}")


# ============================================================
# ข้อ 79 วรรค 2 — กรณีจำเป็นเร่งด่วน (ดำเนินการไปก่อน) — บันทึกเดียวจบ
# ============================================================
def render_clause79(proc, school) -> str:
    """ข้อ 79 วรรคสอง — รายงานขอความเห็นชอบ (จำเป็นเร่งด่วน) อนุมัติโดยหัวหน้าเจ้าหน้าที่"""
    doc = _font_doc()
    total = _total(proc)
    n = len(_items(proc)) or 1
    ex = _extra(proc)
    officer = (getattr(school, "officer_name", "") or "").strip()
    head = (getattr(school, "head_officer_name", "") or "").strip()
    vendor = proc.vendor.name if proc.vendor else _BLANK
    who = officer or _BLANK
    purpose = (proc.purpose or "").strip()
    if purpose.startswith("เพื่อ"):        # กัน "เพื่อเพื่อ..." เมื่อ purpose ขึ้นต้นด้วยเพื่ออยู่แล้ว
        purpose = purpose[len("เพื่อ"):].strip()
    purpose_txt = ("เพื่อ" + purpose + " ") if purpose else ""
    _header(doc, subject_line="รายงานขอความเห็นชอบการจัดซื้อจัดจ้างตามระเบียบฯ ข้อ 79 วรรคสอง",
            school=school, doc_no=(proc.memo_no or "").strip(), doc_date=proc.request_date)
    _p(doc, f"ด้วย {who} มีความจำเป็นต้องซื้อ{(proc.subject or '').strip()} {purpose_txt}"
            f"จำนวน {n} รายการ เป็นจำนวนเงิน {_money(total)} บาท ({bahttext(total)}) "
            f"จาก {vendor} ตาม (ใบเสร็จรับเงิน/ใบส่งของ) เล่มที่ {_x(ex,'receipt_book')} "
            f"เลขที่ {_x(ex,'receipt_no')} วันที่ {_xdate(ex,'receipt_date')} "
            "เนื่องจากเป็นกรณีที่มีความจำเป็นเร่งด่วนที่เกิดขึ้นโดยไม่ได้คาดหมายไว้ก่อนและไม่อาจ"
            "ดำเนินการตามปกติได้ทัน จึงได้ดำเนินการจัดซื้อ/จัดจ้างไปก่อน ทั้งนี้ตามระเบียบกระทรวง"
            "การคลังว่าด้วยการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560 ข้อ 79 วรรคสอง",
       align="justify", indent=1.25, after=2)
    _p(doc, "จึงเรียนมาเพื่อโปรดพิจารณาให้ความเห็นชอบ และให้ถือรายงานนี้เป็นหลักฐานการตรวจรับพัสดุ"
            "ตามระเบียบฯ ข้อ 79 วรรคสอง", align="justify", indent=1.25, after=10)
    _sign_table(doc, [[
        ("ลงชื่อ ...................................... เจ้าหน้าที่ผู้รับผิดชอบ", "center"),
        (f"( {who} )", "center"),
    ]])
    _p(doc, "เรียน หัวหน้าเจ้าหน้าที่", indent=1.25, after=1)
    _p(doc, "โปรดพิจารณา", indent=1.25, after=1)
    _p(doc, "๑. ให้ความเห็นชอบการจัดซื้อ/จัดจ้างดังกล่าวข้างต้น", indent=1.5, after=1)
    _p(doc, f"๒. อนุมัติให้จ่ายเงินจำนวน {_money(total)} บาท ({bahttext(total)})", indent=1.5, after=10)
    _sign_table(doc, [[
        ("ลงชื่อ .........................................", "center"),
        (f"( {head or _BLANK} )", "center"),
        ("หัวหน้าเจ้าหน้าที่", "center"),
    ]])
    return _save(doc, f"ข้อ79วรรค2_{(proc.memo_no or proc.id)}_{(proc.subject or '').strip()}")


def _font_doc():
    from docx import Document
    doc = Document()
    _font(doc)
    return doc


# ตารางส่งออก: kind -> ฟังก์ชัน
RENDERERS = {
    "w804": render_w804,
    "w119t1": render_w119_t1,
    "w119t2": render_w119_t2,
    "clause79": render_clause79,
}
