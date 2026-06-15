"""
docgen.py
---------
สร้างไฟล์เอกสาร Word (.docx) จากข้อมูลในฐานข้อมูล

เฟส 1 สร้าง 2 เอกสาร:
1. รายงานขอซื้อขอจ้าง (บันทึกข้อความ ตามระเบียบฯ 2560 ข้อ 22)
2. ใบตรวจรับพัสดุ

หมายเหตุ: เฟสนี้เราสร้างเอกสารด้วยโค้ดโดยตรง (python-docx) เพื่อให้รันได้ทันที
เฟสถัดไปจะเปลี่ยนไปใช้ไฟล์เทมเพลต .docx จริงของโรงเรียน (โฟลเดอร์ doc_templates/)
"""
from pathlib import Path
from datetime import datetime

from docx import Document as DocxDocument
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.database import get_data_dir
from app.models import Procurement, School
from app.thai_utils import bahttext, thai_date

# ฟอนต์มาตรฐานหนังสือราชการไทย
THAI_FONT = "TH Sarabun New"
FONT_SIZE = 16  # pt


def _output_dir() -> Path:
    """โฟลเดอร์เก็บไฟล์เอกสารที่สร้าง: data/documents/"""
    d = get_data_dir() / "documents"
    d.mkdir(exist_ok=True)
    return d


def _set_thai_font(doc: DocxDocument):
    """ตั้งฟอนต์เริ่มต้นของเอกสารเป็น TH Sarabun New ขนาด 16"""
    style = doc.styles["Normal"]
    style.font.name = THAI_FONT
    style.font.size = Pt(FONT_SIZE)
    # ตั้งฟอนต์สำหรับอักขระภาษาไทยโดยเฉพาะ (east-asia / complex-script)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(_qn("w:cs"), THAI_FONT)
    rfonts.set(_qn("w:ascii"), THAI_FONT)
    rfonts.set(_qn("w:hAnsi"), THAI_FONT)


def _qn(tag: str):
    """ช่วยอ้างชื่อ tag แบบ namespace ของ python-docx"""
    from docx.oxml.ns import qn
    return qn(tag)


def _add_para(doc, text="", *, align=None, bold=False, size=FONT_SIZE, indent=None):
    """เพิ่มย่อหน้าพร้อมตั้งค่ารูปแบบให้สะดวก"""
    p = doc.add_paragraph()
    if indent is not None:
        p.paragraph_format.first_line_indent = Cm(indent)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = THAI_FONT
    run._element.rPr.rFonts.set(_qn("w:cs"), THAI_FONT)
    return p


def _safe_filename(text: str) -> str:
    """ทำชื่อไฟล์ให้ปลอดภัย (ตัดอักขระต้องห้ามใน Windows)"""
    for ch in '<>:"/\\|?*':
        text = text.replace(ch, "_")
    return text.strip()


def generate_purchase_request(db, procurement: Procurement, school: School) -> str:
    """
    สร้าง 'รายงานขอซื้อขอจ้าง' (บันทึกข้อความ)
    คืนค่า: ที่อยู่ไฟล์ .docx ที่สร้าง
    """
    doc = DocxDocument()
    _set_thai_font(doc)

    # ---- หัวบันทึกข้อความ ----
    _add_para(doc, "บันทึกข้อความ", align="center", bold=True, size=20)
    _add_para(doc, f"ส่วนราชการ  {school.name or '..............................'}")
    _add_para(
        doc,
        f"ที่  {school.doc_prefix or 'ศธ'}  {procurement.doc_no or '......./.......'}"
        f"          วันที่  {thai_date(procurement.request_date)}",
    )
    _add_para(
        doc,
        f"เรื่อง  รายงานขอ{procurement.proc_type}{procurement.subject} "
        f"โดยวิธี{procurement.method}",
        bold=True,
    )
    _add_para(doc, "เรียน  " + (school.director_position or "ผู้อำนวยการโรงเรียน"))
    _add_para(doc, "")

    # ---- เนื้อหา ----
    total = procurement.total_amount or 0
    _add_para(
        doc,
        f"ด้วย {school.name or 'โรงเรียน'} มีความจำเป็นต้อง{procurement.proc_type}"
        f"{procurement.subject} เนื่องจาก {procurement.purpose or '-'} "
        f"จึงขอรายงานขอ{procurement.proc_type} ดังรายการต่อไปนี้",
        align="justify", indent=1.25,
    )

    # ---- ตารางรายการพัสดุ ----
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["ลำดับ", "รายการ", "จำนวน", "ราคา/หน่วย", "จำนวนเงิน (บาท)"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = THAI_FONT
                r.font.size = Pt(FONT_SIZE)

    for idx, item in enumerate(procurement.items, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = item.name
        row[2].text = f"{item.quantity:g} {item.unit}"
        row[3].text = f"{item.unit_price:,.2f}"
        row[4].text = f"{item.amount:,.2f}"

    # แถวรวมเงิน
    sum_row = table.add_row().cells
    sum_row[1].text = "รวมเป็นเงินทั้งสิ้น"
    sum_row[4].text = f"{total:,.2f}"

    _add_para(doc, "")
    _add_para(
        doc,
        f"รวมเป็นเงินทั้งสิ้น {total:,.2f} บาท ({bahttext(total)}) "
        f"โดยใช้เงินจาก {procurement.budget_source or '............'}",
        indent=1.25,
    )
    _add_para(
        doc,
        f"จึงเรียนมาเพื่อโปรดพิจารณาอนุมัติให้ดำเนินการ{procurement.proc_type}"
        f"โดยวิธี{procurement.method} ตามรายการข้างต้น",
        align="justify", indent=1.25,
    )

    # ---- ลงชื่อ ----
    _add_para(doc, "")
    _add_para(doc, "ลงชื่อ ......................................... เจ้าหน้าที่พัสดุ", align="right")
    _add_para(doc, f"({school.supply_officer or '...............................'})", align="right")
    _add_para(doc, "")
    _add_para(doc, "ความเห็น/คำสั่ง ผู้อำนวยการ  ☐ อนุมัติ   ☐ ไม่อนุมัติ")
    _add_para(doc, "")
    _add_para(doc, "ลงชื่อ .........................................", align="right")
    _add_para(doc, f"({school.director_name or '...............................'})", align="right")
    _add_para(doc, school.director_position or "ผู้อำนวยการโรงเรียน", align="right")

    # ---- บันทึกไฟล์ ----
    fname = _safe_filename(f"รายงานขอซื้อขอจ้าง_{procurement.doc_no or procurement.id}_{procurement.subject}.docx")
    path = _output_dir() / fname
    doc.save(str(path))
    return str(path)


def generate_inspection_report(db, procurement: Procurement, school: School) -> str:
    """
    สร้าง 'ใบตรวจรับพัสดุ'
    คืนค่า: ที่อยู่ไฟล์ .docx ที่สร้าง
    """
    doc = DocxDocument()
    _set_thai_font(doc)

    _add_para(doc, "ใบตรวจรับพัสดุ", align="center", bold=True, size=20)
    _add_para(doc, school.name or "โรงเรียน......................", align="center")
    _add_para(doc, "")
    _add_para(
        doc,
        f"ตามที่ {school.name or 'โรงเรียน'} ได้ดำเนินการ{procurement.proc_type}"
        f"{procurement.subject} โดยวิธี{procurement.method} "
        f"ตามหนังสือเลขที่ {procurement.doc_no or '......'} "
        f"จากผู้ขาย/ผู้รับจ้าง {procurement.vendor.name if procurement.vendor else '............'} นั้น",
        align="justify", indent=1.25,
    )
    _add_para(
        doc,
        "คณะกรรมการตรวจรับพัสดุ/ผู้ตรวจรับพัสดุ ได้ตรวจรับพัสดุแล้ว "
        "ปรากฏว่าถูกต้องครบถ้วนตามรายการดังนี้",
        align="justify", indent=1.25,
    )

    # ตารางรายการ
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["ลำดับ", "รายการ", "จำนวน", "จำนวนเงิน (บาท)"]):
        c = table.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
    for idx, item in enumerate(procurement.items, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = item.name
        row[2].text = f"{item.quantity:g} {item.unit}"
        row[3].text = f"{item.amount:,.2f}"

    total = procurement.total_amount or 0
    _add_para(doc, "")
    _add_para(
        doc,
        f"รวมเป็นเงินทั้งสิ้น {total:,.2f} บาท ({bahttext(total)}) "
        "เห็นควรเบิกจ่ายเงินให้แก่ผู้ขาย/ผู้รับจ้างต่อไป",
        indent=1.25,
    )

    _add_para(doc, "")
    for label in ["ลงชื่อ ......................................... ประธานกรรมการ",
                  "ลงชื่อ ......................................... กรรมการ",
                  "ลงชื่อ ......................................... กรรมการ"]:
        _add_para(doc, label, align="center")

    fname = _safe_filename(f"ใบตรวจรับพัสดุ_{procurement.doc_no or procurement.id}_{procurement.subject}.docx")
    path = _output_dir() / fname
    doc.save(str(path))
    return str(path)


# ทะเบียนชนิดเอกสารที่สร้างได้ -> ฟังก์ชันที่ใช้สร้าง
DOC_GENERATORS = {
    "รายงานขอซื้อขอจ้าง": generate_purchase_request,
    "ใบตรวจรับพัสดุ": generate_inspection_report,
}
