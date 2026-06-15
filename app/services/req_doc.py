# -*- coding: utf-8 -*-
"""
req_doc.py
----------
สร้างไฟล์ Word 'ใบเบิกวัสดุ' จากข้อมูลใบเบิก (Requisition)
ใช้ helper จัดรูปแบบร่วมกับ build_templates (ฟอนต์ TH Sarabun, ตัวหนา complex script)
"""
from pathlib import Path

from docx import Document

from app.database import get_data_dir
from app.thai_utils import thai_date
from app.services.build_templates import _font, _p, _set_cell, THAI_FONT


def _safe(text: str) -> str:
    for ch in '<>:"/\\|?*':
        text = text.replace(ch, "_")
    return text.strip()


def render_requisition(req, school) -> str:
    """สร้างไฟล์ .docx ใบเบิกวัสดุ คืนค่าที่อยู่ไฟล์"""
    doc = Document()
    _font(doc)

    _p(doc, "ใบเบิกวัสดุ", align="center", bold=True, size=20, after=2)
    _p(doc, school.name or "", align="center", bold=True, after=8)

    _p(doc, "เลขที่ {{x}}".replace("{{x}}", req.req_no or str(req.id)) +
            "          วันที่ " + thai_date(req.date), after=2)
    _p(doc, "ผู้ขอเบิก " + (req.requester or "...................................") +
            "          ฝ่าย/งาน " + (req.department or "..............................."), after=2)
    _p(doc, "เพื่อใช้ในงาน " + (req.purpose or "..............................................................."),
       after=6)

    # ตารางรายการ
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    h = table.rows[0].cells
    _set_cell(h[0], "ลำดับ", bold=True, align="center")
    _set_cell(h[1], "รายการวัสดุ", bold=True, align="center")
    _set_cell(h[2], "จำนวน", bold=True, align="center")
    _set_cell(h[3], "หน่วย", bold=True, align="center")
    for i, it in enumerate(req.items, start=1):
        c = table.add_row().cells
        _set_cell(c[0], str(i), align="center")
        _set_cell(c[1], it.name or "")
        _set_cell(c[2], f"{(it.qty or 0):g}", align="center")
        _set_cell(c[3], it.unit or "", align="center")
    # เติมแถวว่างให้ครบอย่างน้อย 6 แถว (ฟอร์มกรอกมือ)
    for _ in range(max(0, 6 - len(req.items))):
        c = table.add_row().cells
        for j in range(4):
            _set_cell(c[j], "")

    _p(doc, "", after=10)

    # ลงนาม 3 ช่อง: ผู้เบิก / ผู้จ่าย / ผู้อนุมัติ
    sign = doc.add_table(rows=1, cols=3)
    cols = [
        ("ลงชื่อ..........................ผู้ขอเบิก", "( " + (req.requester or "") + " )"),
        ("ลงชื่อ..........................ผู้จ่ายวัสดุ", "( " + (school.officer_name or "") + " )"),
        ("ลงชื่อ..........................ผู้อนุมัติ", "( " + (school.director_name or "") + " )"),
    ]
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    # ลบเส้นขอบตารางลงนาม
    from app.services.build_templates import _no_borders
    _no_borders(sign)
    for cell, (line1, line2) in zip(sign.rows[0].cells, cols):
        for k, txt in enumerate((line1, line2)):
            para = cell.paragraphs[0] if k == 0 else cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = para.add_run(txt)
            r.font.name = THAI_FONT
            from docx.shared import Pt
            from docx.oxml.ns import qn
            r.font.size = Pt(15)
            r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)

    out_dir = get_data_dir() / "documents"
    out_dir.mkdir(exist_ok=True)
    fname = _safe(f"ใบเบิกวัสดุ_{req.req_no or req.id}") + ".docx"
    out_path = out_dir / fname
    doc.save(str(out_path))
    return str(out_path)
