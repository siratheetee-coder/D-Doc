# -*- coding: utf-8 -*-
"""
super_doc.py - เอกสารงานนิเทศภายในสถานศึกษา (Word) ให้ตรงแบบฟอร์มต้นฉบับ
- render_classroom_visit : แบบการเยี่ยมชั้นเรียน (10 ข้อ + ข้อเสนอแนะ + เกณฑ์แปลผล)
- render_supervision     : แบบบันทึกการนิเทศการจัดการเรียนรู้ (25 ข้อ 4 ด้าน + บันทึก 4 หัวข้อ)
ฟอนต์ TH SarabunPSK ตามต้นฉบับ · กระดาษ Letter (21.59x27.94 ซม.)
"""
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn

from app.database import get_data_dir

FONT = "TH SarabunPSK"
CHK = "✓"          # เครื่องหมายถูก


def _safe(text: str) -> str:
    for ch in '<>:"/\\|?*':
        text = text.replace(ch, "_")
    return text.strip()


VISIT_ITEMS = [
    "ครูเข้าสอนตรงเวลา",
    "ครูมีสื่อ อุปกรณ์การสอน พร้อมใช้งาน",
    "ครูแต่งกายเหมาะสมกับสภาพความเป็นครู",
    "ครูจัดบรรยากาศในห้องเรียนให้พร้อมต่อการเรียนรู้",
    "ครูพูดด้วยน้ำเสียงน่าฟังและเร้าความสนใจของนักเรียน",
    "ครูควบคุมดูแลการจัดห้องเรียนให้มีบรรยากาศแห่งการเรียนรู้",
    "นักเรียนกระตือรือร้นที่จะเรียนรู้",
    "นักเรียนมีความสนุกสนาน ร่าเริง แจ่มใส",
    "นักเรียนสนใจปฏิบัติกิจกรรมที่ได้รับมอบหมาย",
    "นักเรียนมีระเบียบวินัยดี และมีมารยาทเรียบร้อย",
]

SUP_DOMAINS = [
    ("1. ด้านความสามารถในการจัดทำแผนการจัดการเรียนรู้", [
        "การวางแผนการสอนที่มีประสิทธิภาพ",
        "แผนการจัดการเรียนรู้ถูกต้อง เป็นขั้นตอน และครบองค์ประกอบ",
        "แผนการจัดการเรียนรู้มีกิจกรรมที่ทำให้นักเรียนเกิดการเรียนรู้",
        "แผนการจัดการเรียนรู้มีการจัดหาสื่อที่เหมาะสมกับการเรียนรู้ของนักเรียน",
        "แผนการจัดการเรียนรู้มีการวัดและประเมินผลผู้เรียนได้อย่างเหมาะสม",
    ]),
    ("2. ด้านความสามารถในการจัดการเรียนรู้", [
        "ใช้เทคนิคต่าง ๆ ที่ทำให้นักเรียนทุกคนมีส่วนร่วมในชั้นเรียน",
        "เลือกใช้สื่อ เทคโนโลยีและอุปกรณ์การสอนที่เหมาะสม",
        "มีการประเมินนักเรียนระหว่างเรียน",
        "อธิบายเนื้อหาบทเรียนได้อย่างชัดเจน",
        "มีความสามารถในการควบคุมชั้นเรียนเมื่อทำกิจกรรม",
        "มีการจัดกิจกรรมการเรียนรู้ที่เน้นการพัฒนาการคิด ได้อภิปราย ซักถาม และแสดงความคิดเห็น",
        "มีการปรับเนื้อหา กิจกรรมในขณะจัดการเรียนรู้เพื่อให้เหมาะสมตามสถานการณ์หรือให้ทันเวลาที่เหลือ",
        "มีกิจกรรมการเรียนการสอนที่เชื่อมโยงหรือบูรณาการกับชีวิตประจำวัน สอดแทรกคุณธรรม จริยธรรมระหว่างเรียน",
        "ใช้ภาษาพูดและภาษาเขียนได้ถูกต้อง เหมาะสม",
    ]),
    ("3. ด้านความสามารถในการประเมินผล", [
        "วัดและประเมินผลด้วยวิธีการที่หลากหลาย",
        "วัดและประเมินผลสอดคล้องกับมาตรฐานการเรียนรู้ ตัวชี้วัด และจุดประสงค์การเรียนรู้",
        "มีเกณฑ์การวัดและประเมินผลที่ชัดเจน",
        "ให้ข้อมูลย้อนกลับแก่นักเรียนเพื่อการปรับปรุงหรือพัฒนา",
        "มีผลงาน ชิ้นงาน ภาระงาน ซึ่งเป็นหลักฐานการเรียนรู้",
    ]),
    ("4. ด้านความสามารถในการจัดสภาพแวดล้อมในชั้นเรียน", [
        "จัดสภาพห้องเรียนได้อย่างเหมาะสม และเอื้อต่อการเรียนรู้ของนักเรียน",
        "สร้างปฏิสัมพันธ์เชิงบวกในชั้นเรียน",
        "จัดชั้นเรียนให้มีความปลอดภัย ไม่เสี่ยงต่อการเกิดอุบัติเหตุในระหว่างการจัด การเรียนการสอน",
        "มีความสามารถในการควบคุมชั้นเรียน",
        "ชี้แจงกฎกติกาหรือข้อตกลงในการเรียน",
        "มีการดูแลพฤติกรรมของนักเรียนในชั้นเรียนอย่างใกล้ชิด",
    ]),
]


def _font(run, size, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:cs"), FONT)


def _p(doc, text="", *, align="left", bold=False, size=16, after=0, before=0):
    p = doc.add_paragraph()
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    pf = p.paragraph_format
    pf.space_after = Pt(after); pf.space_before = Pt(before)
    _font(p.add_run(text), size, bold)
    return p


def _cell(cell, text, *, size=16, bold=False, align="center", valign=True):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    _font(p.add_run(str(text)), size, bold)
    if valign:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _dots(value, n):
    """กรอกแล้ว = แสดงค่าล้วน (ไม่มีจุดต่อท้าย) · เว้นว่าง = จุดไข่ปลาให้เขียน"""
    v = ("" if value is None else str(value)).strip()
    return v if v else ("." * n)


def _paren(value, n):
    """ชื่อในวงเล็บ: กรอกแล้ว = ( ชื่อ ) ไม่มีจุดต่อท้าย · ว่าง = ( ....... )"""
    v = ("" if value is None else str(value)).strip()
    return f"( {v} )" if v else f"( {'.' * n} )"


def _legend(doc, rows, ncol, size=15):
    """เกณฑ์คะแนน 1-5 จัดด้วยตารางล่องหน (ไม่มีเส้น) ให้คอลัมน์ตรงกัน"""
    t = doc.add_table(rows=len(rows), cols=ncol)
    for ri, cells in enumerate(rows):
        for ci in range(ncol):
            _cell(t.rows[ri].cells[ci], cells[ci] if ci < len(cells) else "",
                  align="left", size=size, valign=False)
    return t


def _doc(top=1.5):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.59), Cm(27.94)     # Letter (ตามต้นฉบับ)
    sec.left_margin = sec.right_margin = Cm(2.54)
    sec.top_margin = Cm(top); sec.bottom_margin = Cm(1.5)
    base = doc.styles["Normal"]; base.font.name = FONT; base.font.size = Pt(16)
    base._element.rPr.rFonts.set(qn("w:cs"), FONT)
    return doc


def _parse_scores(csv, n):
    out = []
    for x in (csv or "").split(","):
        x = x.strip()
        out.append(int(x) if x.isdigit() else 0)
    out += [0] * (n - len(out))
    return out[:n]


def _avg(vals):
    v = [x for x in vals if x]
    return round(sum(v) / len(v), 2) if v else None


def _rating_header(t, ncols):
    """แถวหัวตาราง 2 แถว: ที่ | รายการ | ระดับการปฏิบัติ(รวม 5..1) | [หมายเหตุ]"""
    r0, r1 = t.rows[0].cells, t.rows[1].cells
    r0[0].merge(r1[0]); _cell(r0[0], "ที่", bold=True)
    r0[1].merge(r1[1]); _cell(r0[1], "รายการ", bold=True)
    lvl = r0[2]
    for j in range(3, 7):
        lvl = lvl.merge(r0[j])
    _cell(lvl, "ระดับการปฏิบัติ", bold=True)
    for j, lab in enumerate(("5", "4", "3", "2", "1")):
        _cell(r1[2 + j], lab, bold=True)
    if ncols == 8:
        r0[7].merge(r1[7]); _cell(r0[7], "หมายเหตุ", bold=True)


def _item_row(cells, no, text, score, ncols):
    _cell(cells[0], no)
    _cell(cells[1], text, align="left")
    for j in range(5):
        _cell(cells[2 + j], CHK if score == (5 - j) else "")
    if ncols == 8:
        _cell(cells[7], "")


# ============================ แบบการเยี่ยมชั้นเรียน ============================
def render_classroom_visit(school, record) -> str:
    from app.thai_utils import thai_date
    person = record.person
    scores = _parse_scores(record.scores, 10)
    doc = _doc(top=1.5)

    _p(doc, "แบบการเยี่ยมชั้นเรียน (CLASSROOM VISTATION)", align="center", bold=True, after=0)
    _p(doc, f"ภาคเรียนที่ {_dots(record.term, 8)} ปีการศึกษา {_dots(record.year, 14)}",
       align="center", bold=True, after=6)

    _p(doc, "ตอนที่ 1 ข้อมูลทั่วไป", bold=True)
    _p(doc, f"1. ผู้รับผิดชอบชั้นเรียน (ชื่อ-นามสกุล) {_dots(person.name if person else '', 40)}")
    _p(doc, f"2. กลุ่มสาระการเรียนรู้/วิชา {_dots(record.subject_group, 55)}")
    _p(doc, f"   เรื่องที่สอน {_dots(record.topic, 40)} ระดับชั้น {_dots(record.grade_level, 12)}")
    _p(doc, f"   วันที่ {_dots(thai_date(record.visit_date) if record.visit_date else '', 45)} "
            f"คาบที่ {_dots(record.period, 8)} เวลา {_dots(record.visit_time, 14)}")
    _p(doc, f"3. ผู้เยี่ยมชั้นเรียน (ชื่อ- นามสกุล) {_dots(record.visitor_name, 45)}", after=6)

    _p(doc, "ตอนที่ 2 การเยี่ยมชั้นเรียน", bold=True)
    _p(doc, "คำชี้แจง  แบบการเยี่ยมชั้นเรียนนี้ เป็นแบบเยี่ยมการจัดการเรียนการสอนของครูในแต่ละรายวิชาที่สอน "
            "โดยผู้บริหารหรือผู้เยี่ยมชั้นเรียน และบันทึกข้อมูลจากการเยี่ยมชั้นเรียนโดยทำเครื่องหมายถูก "
            f"( {CHK} ) ในแบบประเมินทุกข้อ", align="justify")
    _legend(doc, [["5 = ปฏิบัติได้ระดับดีมาก", "4 = ปฏิบัติได้ระดับดี", "3 = ปฏิบัติได้ระดับปานกลาง"],
                  ["2 = ปฏิบัติได้ระดับพอใช้", "1 = ควรปรับปรุงแก้ไข", ""]], 3)
    _p(doc, "", after=2)

    # ตาราง 14 แถว x 8 คอลัมน์
    t = doc.add_table(rows=2 + 10 + 2, cols=8); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _rating_header(t, 8)
    for i, item in enumerate(VISIT_ITEMS):
        _item_row(t.rows[2 + i].cells, i + 1, item, scores[i], 8)
    avg = _avg(scores)
    for ri, label in ((12, "ค่าเฉลี่ย"), (13, "รวมเฉลี่ยทุกด้าน")):
        cs = t.rows[ri].cells
        m = cs[0].merge(cs[1]); _cell(m, label, bold=True, align="center")
        mv = cs[2]
        for j in range(3, 7):
            mv = mv.merge(cs[j])
        _cell(mv, ("%.2f" % avg) if avg is not None else "", bold=True)
        _cell(cs[7], "")
    _widths_visit(t)

    _p(doc, "ตอนที่ 3 ข้อเสนอแนะ", bold=True, before=6)
    _suggest_lines(doc, record.suggestion, 6, indent=1.0)

    _p(doc, f"\t\t\t\t(ลงชื่อ) {'.' * 40} ผู้เยี่ยมชั้นเรียน", before=6)
    _p(doc, f"\t\t\t\t         {_paren(record.visitor_name, 40)}")
    _p(doc, f"\t\t\t\t         ตำแหน่ง {'.' * 35}", after=8)

    _p(doc, "เกณฑ์การแปลความหมาย")
    for rng, mean in (("4.51 – 5.00", "ปฏิบัติได้ระดับดีมาก"), ("3.51 – 4.50", "ปฏิบัติได้ระดับดี"),
                      ("2.51 – 3.50", "ปฏิบัติได้ระดับปานกลาง"), ("1.51 – 2.50", "ปฏิบัติได้ระดับพอใช้"),
                      ("1.00 – 1.50", "ควรปรับปรุงแก้ไข")):
        _p(doc, f"\t\tระดับคะแนน\t{rng}\tหมายถึง\t{mean}")

    return _save(doc, f"แบบการเยี่ยมชั้นเรียน_{person.name if person else ''}")


def _widths_visit(t):
    w = [Cm(0.9), Cm(8.4), Cm(1.1), Cm(1.1), Cm(1.1), Cm(1.1), Cm(1.1), Cm(1.7)]
    for row in t.rows:
        for c, cw in zip(row.cells, w):
            c.width = cw


def _suggest_lines(doc, text, min_lines, indent=0.0):
    """กรอกแล้ว = แสดงข้อความล้วน (ย่อหน้าเข้า) · ว่าง = เส้นจุดว่างให้เขียน"""
    text = (text or "").strip()
    if text:
        for ln in text.splitlines() or [text]:
            p = _p(doc, ln)
            if indent:
                p.paragraph_format.left_indent = Cm(indent)
    else:
        for _ in range(min_lines):
            _p(doc, "." * 118)


# ============================ แบบบันทึกการนิเทศการจัดการเรียนรู้ ============================
def render_supervision(school, record) -> str:
    from app.thai_utils import thai_date
    person = record.person
    pos = getattr(person, "position", "") or "" if person else ""
    rank = getattr(person, "rank", "") or "" if person else ""
    scores = _parse_scores(record.scores, 25)
    doc = _doc(top=1.75)

    _p(doc, "แบบบันทึกการนิเทศการจัดการเรียนรู้", align="center", bold=True, size=18, after=4)

    _p(doc, "ตอนที่ 1  ข้อมูลทั่วไปของผู้รับการนิเทศ", bold=True)
    _p(doc, "คำชี้แจง  โปรดเติมข้อความลงในช่องว่างที่กำหนดให้", bold=True)
    _p(doc, f"ชื่อผู้รับการนิเทศ {_dots(person.name if person else '', 30)} "
            f"ตำแหน่ง {_dots(pos, 12)} วิทยฐานะ {_dots(rank, 16)}")
    _p(doc, f"กลุ่มสาระการเรียนรู้ {_dots(record.subject_group, 60)}")
    _p(doc, f"รายวิชาที่สอน {_dots(record.subject_taught, 25)} รหัสวิชา {_dots(record.subject_code, 14)} "
            f"ชั้น {_dots(record.grade_class, 10)}")
    _p(doc, f"นิเทศครั้งที่ {_dots(record.round_no, 5)} วัน  เดือน  ปีที่รับการนิเทศ "
            f"{_dots(thai_date(record.sup_date) if record.sup_date else '', 40)}", after=6)

    _p(doc, "ตอนที่ 2  แบบประเมินสมรรถนะการจัดการเรียนรู้ของผู้รับการนิเทศ", bold=True)
    _p(doc, "คำชี้แจง  ให้ผู้นิเทศสังเกตกระบวนการจัดการเรียนรู้ของผู้รับการนิเทศทั้ง 4 ด้าน")
    _p(doc, f"\t  แล้วทำเครื่องหมาย {CHK} ในช่องที่มีการปฏิบัติมากที่สุดถึงน้อยที่สุด โดยใช้เกณฑ์ดังนี้")
    _legend(doc, [["5 หมายถึง มากที่สุด", "4 หมายถึง มาก", "3 หมายถึง ปานกลาง",
                   "2 หมายถึง น้อย", "1 หมายถึง น้อยที่สุด"]], 5, size=14)
    _p(doc, "", after=2)

    # จำนวนแถว: หัว 2 + (ต่อด้าน: 1 หัวข้อด้าน + N ข้อ + 1 เฉลี่ย) + 1 รวมเฉลี่ยทุกด้าน
    nrows = 2 + sum(2 + len(items) for _, items in SUP_DOMAINS) + 1
    t = doc.add_table(rows=nrows, cols=7); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _rating_header(t, 7)
    ri = 2
    no = 1
    dom_avgs = []
    for dom_title, items in SUP_DOMAINS:
        cs = t.rows[ri].cells                      # แถวหัวข้อด้าน (รวมทุกคอลัมน์)
        m = cs[0]
        for j in range(1, 7):
            m = m.merge(cs[j])
        _cell(m, dom_title, bold=True, align="left")
        ri += 1
        dvals = []
        for item in items:
            _item_row(t.rows[ri].cells, no, item, scores[no - 1], 7)
            dvals.append(scores[no - 1]); no += 1; ri += 1
        cs = t.rows[ri].cells                       # แถวเฉลี่ยของด้าน
        m = cs[0].merge(cs[1]); _cell(m, "เฉลี่ย", bold=True)
        av = _avg(dvals); dom_avgs.append(av)
        mv = cs[2]
        for j in range(3, 7):
            mv = mv.merge(cs[j])
        _cell(mv, ("%.2f" % av) if av is not None else "", bold=True)
        ri += 1
    cs = t.rows[ri].cells                            # รวมเฉลี่ยทุกด้าน
    m = cs[0].merge(cs[1]); _cell(m, "รวมเฉลี่ยทุกด้าน", bold=True)
    allv = _avg([a for a in dom_avgs if a is not None])
    mv = cs[2]
    for j in range(3, 7):
        mv = mv.merge(cs[j])
    _cell(mv, ("%.2f" % allv) if allv is not None else "", bold=True)
    _widths_sup(t)

    _p(doc, "ตอนที่ 3 ผู้นิเทศบันทึกเพิ่มเติมการนิเทศการจัดการเรียนรู้", bold=True, before=6)
    for n, (label, val) in enumerate([
        ("1.  สิ่งที่พบจากการสังเกตการจัดการเรียนรู้ในชั้นเรียนของผู้รับการนิเทศ", record.note_found),
        ("2.  การสะท้อนความคิดจากการจัดการเรียนรู้ในชั้นเรียนของผู้รับการนิเทศ", record.note_reflect),
        ("3.  ความประทับใจหรือจุดเด่นในการจัดการเรียนรู้ครั้งนี้", record.note_impress),
        ("4.  สิ่งที่ควรปรับปรุงหรือพัฒนา", record.note_improve),
    ]):
        _p(doc, label)
        _suggest_lines(doc, val, 2, indent=1.0)

    _p(doc, f"(ลงชื่อ) {'.' * 40} ผู้รับการนิเทศ", align="center", before=8)
    _p(doc, _paren(person.name if person else '', 35), align="center")
    _p(doc, f"(ลงชื่อ) {'.' * 40} ผู้นิเทศ", align="center", before=6)
    _p(doc, _paren(record.supervisor_name, 35), align="center")

    return _save(doc, f"แบบบันทึกการนิเทศ_{person.name if person else ''}")


def _widths_sup(t):
    w = [Cm(0.9), Cm(9.6), Cm(1.1), Cm(1.1), Cm(1.1), Cm(1.1), Cm(1.1)]
    for row in t.rows:
        for c, cw in zip(row.cells, w):
            c.width = cw


def _save(doc, fname):
    out_dir = get_data_dir() / "documents"; out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / (_safe(fname) + ".docx")
    doc.save(str(out))
    return str(out)
