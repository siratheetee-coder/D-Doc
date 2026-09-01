# -*- coding: utf-8 -*-
"""
gov_forms.py - กรอกข้อมูลลงแบบฟอร์มราชการต้นฉบับ (.docx) โดยคงเลย์เอาต์เป๊ะ
วิธีทำ: เปิดไฟล์แม่แบบที่ฝังไว้ (app/data/forms/*) แล้ว "แทรกค่า" ต่อท้าย run ที่รู้ตำแหน่ง
- ใบลา: leave_form.docx (แบบ ก.พ. ป่วย/คลอดบุตร/กิจส่วนตัว)
ครูพิมพ์รายละเอียดในระบบ -> ระบบเติมลงแม่แบบให้ -> ดาวน์โหลดไปพิมพ์/เซ็น
"""
import copy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.run import Run

from app.database import get_data_dir
from app.thai_utils import _THAI_MONTHS

_FORM_DIR = Path(__file__).resolve().parent.parent / "data" / "forms"
CHK = "✓"


def _safe(text: str) -> str:
    for ch in '<>:"/\\|?*':
        text = text.replace(ch, "_")
    return text.strip()


def _ins_after(run, text):
    """แทรก run ใหม่ (คัดลอกฟอนต์/ขนาดจาก run เดิม) ต่อท้าย run ที่ให้มา พร้อมข้อความ text"""
    new_r = copy.deepcopy(run._element)
    # ล้างเนื้อหาเดิมใน clone (ข้อความ/แท็บ/สัญลักษณ์/ขึ้นบรรทัด) เหลือแต่ rPr (ฟอนต์)
    for tag in ("w:t", "w:tab", "w:sym", "w:br", "w:cr"):
        for el in new_r.findall(qn(tag)):
            new_r.remove(el)
    rpr = new_r.find(qn("w:rPr"))       # ค่าที่กรอกต้องไม่มีเส้นใต้ (กันจุดไข่ปลาติดมากับ anchor)
    if rpr is not None:
        for u in rpr.findall(qn("w:u")):
            rpr.remove(u)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_r.append(t)
    run._element.addnext(new_r)
    return new_r


def _strip_u(run):
    """ลบเส้นใต้ (w:u) ของ run เดียว"""
    rpr = run._element.rPr
    if rpr is not None:
        for u in rpr.findall(qn("w:u")):
            rpr.remove(u)


def _first_tab(p, cm_pos):
    """ย้ายตำแหน่ง tab stop แรกของย่อหน้า (cm) เพื่อจัดแนวให้ตรงกับบรรทัดอื่น"""
    pPr = p._element.find(qn("w:pPr"))
    if pPr is None:
        return
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        return
    tab = tabs.find(qn("w:tab"))
    if tab is not None:
        tab.set(qn("w:pos"), str(int(cm_pos * 567)))   # 1cm = 567 twips


def _strip_following(p, idx):
    """ลบเส้นไข่ปลา (underline dotted) เฉพาะ 'ช่องว่างที่ตามหลัง' run idx (จนเจอข้อความจริง)
    -> เอาจุดออกเฉพาะฟิลด์ที่กรอกค่าแล้ว โดยไม่แตะฟิลด์อื่นบนบรรทัดเดียวกัน"""
    for r in p.runs[idx + 1:]:
        t = r.text
        if t == "" or t.strip() != "":     # เจอกล่อง/สัญลักษณ์ หรือข้อความจริง -> หยุด
            break
        rpr = r._element.rPr
        if rpr is not None:
            for u in rpr.findall(qn("w:u")):
                rpr.remove(u)


def _fill(p, pairs):
    """แทรกหลายค่าในย่อหน้าเดียว: pairs = [(run_index, text), ...]
    - ข้ามค่าว่าง (ปล่อยเส้นไข่ปลาไว้ให้เขียนมือ)
    - ค่าที่กรอก -> ลบเส้นไข่ปลาที่ตามหลังฟิลด์นั้น
    - แทรกจาก index มาก -> น้อย กันการเลื่อน index"""
    for idx, text in sorted(pairs, key=lambda x: -x[0]):
        if not (text or "").strip():
            continue
        try:
            _strip_u(p.runs[idx])          # ลบจุดของ run สมอ (ช่องว่างก่อนค่า)
            _strip_following(p, idx)
            _ins_after(p.runs[idx], text)
        except Exception:
            pass


def _insstrip(p, idx, text):
    """แทรกค่าเดี่ยว + ลบเส้นไข่ปลาที่ตามหลัง (สำหรับจุดที่ไม่ได้ใช้ _fill)"""
    _strip_u(p.runs[idx])
    _strip_following(p, idx)
    _ins_after(p.runs[idx], text)


def _full_date(dt):
    """วันที่แบบเต็ม '1 กันยายน 2569' (ว่างถ้าไม่มี)"""
    if not dt:
        return ""
    return f"{dt.day} {_THAI_MONTHS[dt.month]} {dt.year + 543}"


def _set_para(p, text, size=16):
    """แทนที่ทั้งย่อหน้าด้วยข้อความไหลปกติ (ไม่มีแท็บ) - กันข้อความยาวดันขอบ/ล้น"""
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return run


def _fill_paren(p, name):
    """ใส่ชื่อในวงเล็บให้ ')' ชิดชื่อ (ไม่เว้นช่องยาว) : '( ชื่อ )'"""
    i = _find_run(p, "(")
    if i is None:
        return
    for r in p.runs[i + 1:]:            # ยุบแท็บ/ช่องว่างระหว่าง ( ) ให้ ') ' ชิดชื่อ
        if r.text == ")":
            break
        if r.text != "" and r.text.strip() == "":
            r.text = ""
    _ins_after(p.runs[i], " " + (name or "").strip() + " ")


def _dparts(dt):
    """คืน (วัน, ชื่อเดือนไทย, ปีพ.ศ.) จาก datetime หรือ ('','','')"""
    if not dt:
        return "", "", ""
    return str(dt.day), _THAI_MONTHS[dt.month], str(dt.year + 543)


def _find_run(p, text):
    """คืน index ของ run แรกที่ข้อความตรง (หรือ None)"""
    for i, r in enumerate(p.runs):
        if r.text == text:
            return i
    return None


def _stamp_sig(run, db, name, height_cm=1.0):
    """แปะรูปลายเซ็นของ name (ถ้ามีในทะเบียน) ต่อท้าย run · คืน True ถ้าแปะได้"""
    if db is None or not (name or "").strip():
        return False
    try:
        from app.services.signature import signature_path_for
        path = signature_path_for(db, name)
    except Exception:
        path = None
    if not path:
        return False
    new_r = OxmlElement("w:r")
    run._element.addnext(new_r)
    try:
        Run(new_r, run._parent).add_picture(path, height=Cm(height_cm))
        return True
    except Exception:
        new_r.getparent().remove(new_r)
        return False


def _fill_date_slashes(run, d, m, y):
    """เติมวันที่ลงรูปแบบ 'วันที่ __/__/__' ที่อยู่ใน run เดียว โดยคงความยาวเดิม (กันเลย์เอาต์เลื่อน)"""
    t = run.text
    if t.count("/") < 2:
        return
    i1 = t.index("/"); i2 = t.index("/", i1 + 1)
    seg0, seg1, seg2 = t[:i1], t[i1 + 1:i2], t[i2 + 1:]

    def center(seg, val):
        val = str(val or "")
        w = len(seg)
        if not val or w <= 0:
            return seg
        if len(val) >= w:
            return val[:w]
        left = (w - len(val)) // 2
        return " " * left + val + " " * (w - len(val) - left)

    core = seg0.rstrip()
    trail = seg0[len(core):]
    run.text = core + center(trail, d) + "/" + center(seg1, m) + "/" + center(seg2, y)


def _set_tc_text(tc, text):
    """ตั้งข้อความในเซลล์ตาราง (กึ่งกลาง) - ล้าง run เดิมในย่อหน้าแรกแล้วใส่ค่าใหม่"""
    p = tc.find(qn("w:p"))
    if p is None:
        return
    for r in p.findall(qn("w:r")):
        p.remove(r)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = str(text)
    r.append(t); p.append(r)


def _strip_dots(p):
    """ลบเส้นไข่ปลา (underline dotted บน run แท็บ) ในย่อหน้า - ใช้กับบรรทัดที่กรอกค่าแล้ว"""
    for r in p.runs:
        rpr = r._element.rPr
        if rpr is not None:
            for u in rpr.findall(qn("w:u")):
                rpr.remove(u)


def _set_font_all(doc, font):
    """ตั้งฟอนต์ทั้งเอกสาร (รวม run ในกล่องข้อความ/ตาราง) + สไตล์ Normal
    ไม่กระทบสัญลักษณ์ (w:sym) เพราะ sym ใช้ฟอนต์ของตัวเอง"""
    for r in doc.element.body.iter(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr"); r.insert(0, rpr)
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
        for a in ("w:ascii", "w:hAnsi", "w:cs"):
            rf.set(qn(a), font)
    try:
        st = doc.styles["Normal"]
        st.font.name = font
        st._element.rPr.rFonts.set(qn("w:cs"), font)
    except Exception:
        pass


def _fill_stats_tables(doc, stats):
    """เติมตารางสถิติการลา (ทุกตารางในกล่องข้อความ) · stats={'sick':(ก่อน,ครั้งนี้,รวม),...}"""
    order = ["sick", "personal", "maternity"]     # แถว 1,2,3
    def fmt(v):
        return f"{v:g}" if v else ""
    for tbl in doc.element.body.iter(qn("w:tbl")):
        rows = tbl.findall(qn("w:tr"))
        for ri, key in zip((1, 2, 3), order):
            if ri >= len(rows):
                continue
            cells = rows[ri].findall(qn("w:tc"))
            vals = [fmt(v) for v in stats.get(key, (0, 0, 0))]
            for ci, val in zip((1, 2, 3), vals):
                if val and ci < len(cells):
                    _set_tc_text(cells[ci], val)


def _leave_stats(db, person_id, year, cur_type, cur_days, exclude_id=None):
    """รวมวันลาแต่ละประเภทในปี (พ.ศ.) จากทะเบียนวันลา · คืน {type:(ก่อนหน้า,ครั้งนี้,รวม)}"""
    from collections import defaultdict
    before = defaultdict(float)
    if db is not None and person_id and year:
        from app.models import LeaveRecord
        q = db.query(LeaveRecord).filter_by(person_id=person_id, year=year)
        for r in q.all():
            if exclude_id and r.id == exclude_id:
                continue
            before[_LEAVE_NORM.get((r.leave_type or "").strip(), "personal")] += (r.days or 0)
    out = {}
    for t in ("sick", "personal", "maternity"):
        b = before.get(t, 0)
        this = (cur_days or 0) if t == cur_type else 0
        out[t] = (b, this, b + this)
    return out


def _last_leave(db, person_id, cur_type, before_date, exclude_id=None):
    """รายการลาประเภทเดียวกันครั้งก่อนหน้า (ล่าสุดก่อน before_date) หรือ None"""
    if db is None or not person_id:
        return None
    from app.models import LeaveRecord
    q = db.query(LeaveRecord).filter_by(person_id=person_id)
    if exclude_id:
        q = q.filter(LeaveRecord.id != exclude_id)
    rows = [r for r in q.all()
            if _LEAVE_NORM.get((r.leave_type or "").strip(), "personal") == cur_type]
    if before_date:
        rows = [r for r in rows if r.start_date and r.start_date.date() < _as_date(before_date)]
    rows.sort(key=lambda r: (r.start_date or _MIN), reverse=True)
    return rows[0] if rows else None


import datetime as _dtmod
_MIN = _dtmod.datetime.min


def _as_date(dt):
    return dt.date() if hasattr(dt, "date") else dt


# normalize ชนิดลา รับได้ทั้งคีย์ทะเบียน (sick/personal/maternity) และป้ายไทย (ลาป่วย/ลากิจ...)
_LEAVE_NORM = {
    "sick": "sick", "ป่วย": "sick", "ลาป่วย": "sick",
    "personal": "personal", "กิจ": "personal", "กิจส่วนตัว": "personal",
    "ลากิจ": "personal", "ลากิจส่วนตัว": "personal",
    "maternity": "maternity", "คลอด": "maternity", "คลอดบุตร": "maternity",
    "ลาคลอด": "maternity", "ลาคลอดบุตร": "maternity",
}
# ชนิดลา -> ป่วย=ย่อหน้า9 · กิจส่วนตัว=10 · คลอดบุตร=11 (เช็กบ็อกซ์อยู่ run2)
_LEAVE_PARA = {"sick": 9, "personal": 10, "maternity": 11}
_LEAVE_LABEL = {"sick": "ป่วย", "personal": "กิจส่วนตัว", "maternity": "คลอดบุตร"}


def render_leave_official(school, person, record, db=None, approver=None,
                          checker=None, checker_date=None, approve_date=None,
                          write_date=None, work_group="") -> str:
    """กรอกใบลา (แบบ ก.พ.) จากไฟล์แม่แบบ คืน path .docx
    db: session โรงเรียน (ใช้คำนวณสถิติการลา/ลาครั้งสุดท้าย + หาลายเซ็น)
    checker: Person ผู้ตรวจสอบ (หัวหน้าบุคคล) · checker_date: วันที่ตรวจสอบ
    approver: Person ผอ. · approve_date: วันที่อนุมัติ (ติ๊ก 'อนุญาต' + ชื่อ/ตำแหน่ง/วันที่/ลายเซ็น)"""
    doc = Document(str(_FORM_DIR / "leave_form.docx"))
    sec = doc.sections[0]
    sec.top_margin = Cm(1.0); sec.bottom_margin = Cm(1.0)   # TH Sarabun New สูงกว่าเดิม เผื่อ 1 หน้า
    P = doc.paragraphs

    name = (person.name if person else "") or ""
    position = (getattr(person, "position", "") or "ครู")
    sch = school.name or ""
    area = (getattr(school, "area_office", "") or "").strip()   # สำนักงานเขต -> ช่องสังกัด
    wg = (work_group or "").strip()
    group_val = (wg + " " + sch).strip() if wg else sch          # กลุ่มงาน + ชื่อโรงเรียน
    addressee = ("ผู้อำนวยการ" + sch) if sch else (
        getattr(school, "director_position", "") or "ผู้อำนวยการโรงเรียน")

    typ_norm = _LEAVE_NORM.get((record.leave_type or "").strip(), "personal")
    typ_label = _LEAVE_LABEL[typ_norm]

    wd = write_date or getattr(record, "created_at", None) or getattr(record, "submitted_at", None)
    wdd, wmon, wyy = _dparts(wd)
    sd_d, sd_m, sd_y = _dparts(record.start_date)
    ed_d, ed_m, ed_y = _dparts(record.end_date)
    days = f"{(record.days or 0):g}"
    reason = (record.reason or "").strip()
    contact = (record.contact or "").strip()

    wmon_num = str(wd.month) if wd else ""

    # ---- หัวเอกสาร ----
    _fill(P[2], [(2, "  " + sch)])                              # เขียนที่
    try:
        P[2].runs[1].text = ""                                 # ตัดช่องว่างนำ ให้ตรงกับบรรทัด 'วันที่'
    except Exception:
        pass
    _fill(P[3], [(1, " " + wdd), (4, " " + wmon), (9, " " + wyy)])  # วันที่ เดือน พ.ศ.
    _fill(P[4], [(1, "ขอลา" + typ_label)])                     # เรื่อง
    _fill(P[5], [(1, addressee)])                              # เรียน
    # ข้าพเจ้า + ตำแหน่ง + กลุ่มงาน(+รร.) + สังกัด(เขต) รวมเป็นข้อความไหลต่อเนื่อง
    # (ต่อบรรทัดล่างขึ้นมา กันเหลือที่ว่างเยอะ + กันข้อความยาวดันขอบ) · P8 จะลบทิ้งท้าย
    body_line = ("ข้าพเจ้า " + name + "   ตำแหน่ง " + position
                 + "   กลุ่มงาน " + group_val + ("   สังกัด " + area if area else ""))
    _set_para(P[7], body_line, size=16)
    P[7].paragraph_format.first_line_indent = Cm(1.25)

    # ---- ชนิดการลา (ติ๊กในกล่อง run2) + เหตุผล (run5, เฉพาะป่วย/กิจ) ----
    # ติ๊กโดย "แทนที่อักขระกล่องเดิม" ด้วยเครื่องหมายถูก ไม่แทรก run ใหม่
    # (การแทรกจะเพิ่มความกว้าง ดันคำว่า 'เนื่องจาก' ไปชิดขอบ เหตุผลเลยล้นหน้า)
    sel_para = _LEAVE_PARA[typ_norm]
    try:
        box_run = P[sel_para].runs[2]
        for sym in box_run._element.findall(qn("w:sym")):   # ล้างกล่อง Wingdings เดิม
            box_run._element.remove(sym)
        box_run.text = CHK
    except Exception:
        pass
    if reason and sel_para in (9, 10):
        _fill(P[sel_para], [(5, " " + reason)])

    # ---- ช่วงวันลา + จำนวนวัน ----
    _fill(P[12], [(1, " " + sd_d), (3, " " + sd_m), (5, " " + sd_y),
                  (8, " " + ed_d), (10, " " + ed_m), (12, " " + ed_y), (16, days)])
    # ---- ที่อยู่ติดต่อระหว่างลา ----
    _fill(P[14], [(15, " " + contact)])
    if contact:                       # ลบเส้นไข่ปลาบรรทัดต่อจากที่อยู่ติดต่อ (P15) เมื่อกรอกแล้ว
        _strip_dots(P[15])

    # ---- ลงชื่อผู้ลา ----
    _first_tab(P[19], 9.5)                                     # จัดชื่อให้ตรงกับ ตำแหน่ง/วันที่
    _fill_paren(P[19], name)                                   # ( ชื่อ ) ให้ ) ชิดชื่อ
    _fill(P[20], [(3, " " + position)])                        # ตำแหน่ง
    if wd:                                                     # วันที่ = '1 กันยายน 2569'
        for k in (4, 5, 6, 7, 8):
            try:
                P[21].runs[k].text = ""
            except Exception:
                pass
        _insstrip(P[21], 3, " " + _full_date(wd))

    # ---- ตารางสถิติการลา + การลาครั้งสุดท้าย (ดึงจากทะเบียนวันลา) ----
    year = getattr(record, "year", None)
    if not year and record.start_date:
        year = record.start_date.year + 543
    exclude_id = record.id if getattr(record, "year", None) is not None else None
    _fill_stats_tables(doc, _leave_stats(db, getattr(person, "id", None), year,
                                         typ_norm, (record.days or 0), exclude_id))
    last = _last_leave(db, getattr(person, "id", None), typ_norm, record.start_date, exclude_id)
    if last is not None:
        l_sd, l_sm, l_sy = _dparts(last.start_date)
        l_ed, l_em, l_ey = _dparts(last.end_date)
        _fill(P[13], [(13, " " + l_sd), (15, " " + l_sm), (20, " " + l_sy)])
        _fill(P[14], [(11, f"{(last.days or 0):g}"), (8, " " + l_ey),
                      (2, " " + l_em), (0, " " + l_ed)])

    # ---- ผู้ตรวจสอบ (หัวหน้าบุคคล) ----
    if checker is not None:
        cname = (getattr(checker, "name", "") or "").strip()
        cpos = (getattr(checker, "position", "") or "").strip()
        i = _find_run(P[27], "(ลงชื่อ)")
        if i is not None:
            if cname:
                _insstrip(P[27], i, " " + cname + " ")
            _stamp_sig(P[27].runs[i], db, cname)
        i = _find_run(P[28], "ง")            # ตำแหน่ง (แยก run เป็น 'ตำแหน่'+'ง')
        if i is not None and cpos:
            _insstrip(P[28], i, " " + cpos)
        if checker_date:                     # วันที่ผู้ตรวจสอบ = '2 กันยายน 2569'
            try:
                r2 = P[29].runs[2]
                r2.text = "่  " + _full_date(checker_date)
                if r2._element.rPr is not None:
                    for u in r2._element.rPr.findall(qn("w:u")):
                        r2._element.rPr.remove(u)
            except Exception:
                pass

    # ---- คำสั่งอนุญาต + ลงนาม ผอ. (เมื่ออนุมัติแล้ว) ----
    if approver is not None:
        dname = (getattr(approver, "name", "") or "").strip()
        dpos = (getattr(school, "director_position", "") or "ผู้อำนวยการโรงเรียน")
        # ติ๊ก 'อนุญาต' - แทนที่กล่องก่อนคำว่า 'อนุญาต' ในตำแหน่งเดิม (ไม่ดันเลย์เอาต์)
        i = _find_run(P[27], "อนุญาต")
        if i is not None and i >= 1:
            try:
                bx = P[27].runs[i - 1]
                for sym in bx._element.findall(qn("w:sym")):
                    bx._element.remove(sym)
                bx.text = CHK
            except Exception:
                pass
        # ลายเซ็น ผอ. หลัง '(ลงชื่อ)' คอลัมน์ขวา (p28)
        i = _find_run(P[28], "(ลงชื่อ)")
        if i is not None:
            _stamp_sig(P[28].runs[i], db, dname)
        # ชื่อ ผอ. ในวงเล็บ (p29 คอลัมน์ขวา) - ) ชิดชื่อ
        _fill_paren(P[29], dname)
        # ตำแหน่ง ผอ. (p30)
        i = _find_run(P[30], "ตำแหน่ง")
        if i is not None:
            _insstrip(P[30], i, " " + dpos)
        # วันที่อนุมัติ ผอ. (p31) = '3 กันยายน 2569'
        if approve_date:
            i = _find_run(P[31], "วันที่")
            if i is not None:
                for k in range(i + 1, i + 8):
                    try:
                        if P[31].runs[k].text.strip() in ("", "/"):
                            P[31].runs[k].text = ""
                    except Exception:
                        break
                _insstrip(P[31], i, " " + _full_date(approve_date))

    # ---- เอา '/' ออกจากช่องวันที่ผู้ตรวจสอบ/ผอ ที่ยังว่าง (เหลือจุดไข่ปลาไว้เขียน) ----
    try:
        r = P[29].runs[2]
        if "/" in r.text:
            r.text = r.text.replace("/", " ")
    except Exception:
        pass
    try:
        i = _find_run(P[31], "วันที่")
        if i is not None:
            for k in range(i + 1, len(P[31].runs)):
                if "/" in P[31].runs[k].text:
                    P[31].runs[k].text = P[31].runs[k].text.replace("/", " ")
    except Exception:
        pass

    # ---- จัดบรรทัดหัวเอกสาร ----
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _AL
    from docx.shared import Cm as _Cm, Pt as _Pt
    P[2].alignment = _AL.LEFT                       # เขียนที่ ให้ตรงกับ วันที่ (บรรทัดถัดไป)
    try:
        P[7].paragraph_format.first_line_indent = _Cm(1.25)   # ย่อหน้า 'ข้าพเจ้า'
    except Exception:
        pass

    # ---- เปลี่ยนฟอนต์ทั้งไฟล์เป็น TH Sarabun New (รวมตัวเลขในตารางสถิติ) ----
    _set_font_all(doc, "TH Sarabun New")

    # ระยะห่างบรรทัด 1.0 (single) ตามต้นฉบับ · ตัดเว้นวรรคย่อหน้า
    for p in doc.paragraphs:
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.space_before = _Pt(0); pf.space_after = _Pt(0)
    # ลบบรรทัดกลุ่มงาน/สังกัดเดิม (P8) - ถูกรวมเข้าบรรทัด 'ข้าพเจ้า' (P7) แล้ว
    try:
        P[8]._element.getparent().remove(P[8]._element)
    except Exception:
        pass
    # ตัดเฉพาะย่อหน้าเว้นว่างตรงกลาง (P16) · คงบรรทัดว่างใต้หัวเรื่อง (P1) และก่อน 'ข้าพเจ้า' (P6)
    for idx in (16,):
        try:
            if not P[idx].text.strip():
                P[idx]._element.getparent().remove(P[idx]._element)
        except Exception:
            pass

    # ตัดย่อหน้าว่างท้ายเอกสาร (แม่แบบเว้นไว้เต็มหน้าพอดี) กันข้อความที่เติมดันตกหน้า 2
    for p in reversed(doc.paragraphs):
        if p.text.strip() == "":
            p._element.getparent().remove(p._element)
        else:
            break

    out = get_data_dir() / "documents"; out.mkdir(exist_ok=True)
    path = out / (_safe(f"ใบลา_{name}_{getattr(record,'id','')}") + ".docx")
    doc.save(str(path))
    return str(path)
