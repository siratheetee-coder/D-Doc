# -*- coding: utf-8 -*-
"""
lunch_ingredient_doc.py - เอกสารการจัดซื้อวัตถุดิบเพื่อประกอบอาหารกลางวัน (รูปแบบ 1)
ตามคู่มือการดำเนินงานโครงการอาหารกลางวัน สพฐ. (วงเงินไม่เกิน 500,000 บาท) แบบยืมเงิน->ส่งใช้
อ้างถ้อยคำ/โครงสร้างจากไฟล์ตัวอย่างที่โรงเรียนใช้จริง (Lunch examples/1 ...)

ชุดเอกสาร (ต่อรอบ/เดือน):
  02 บันทึกขออนุมัติยืมเงิน   03 สัญญายืมเงิน (ฟอร์ม)   04 แบบประมาณการค่าใช้จ่าย
  05 ใบจัดซื้อวัสดุเครื่องบริโภค (4 ส่วน)   06 ใบรับรายงานวัตถุดิบ (ฟอร์ม)
  07 ใบเสร็จรับเงิน (ฟอร์ม)   08 ใบรับรองการจ่ายเงิน (ฟอร์ม)
  09 บันทึกรายงานผู้ควบคุมและคณะกรรมการตรวจการประกอบอาหาร   10 ขออนุมัติเบิกจ่ายส่งใช้เงินยืม
"""
from datetime import datetime
from docx import Document
from docx.shared import Cm

from app.services.doc_page import set_a4

from app.thai_utils import bahttext
from app.services.build_templates import (
    _font, _p, _sign_table, _set_cell, _repeat_header_row, _no_split_row,
    _krut_and_title, _p_runs, _hr, _shrink_body_font, _csize, _bcs, THAI_FONT,
    _no_borders,
)
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _box_cell(cell, lines, *, size=14, before=0, after=1, line=None):
    """เติมข้อความหลายบรรทัดในเซลล์ตาราง (แต่ละ tuple = (text, align, bold))
    ใช้ทำฟอร์มแบบมีช่องกรอบ (เช่น สัญญายืมเงิน แบบ 8500)
    before/after = ระยะห่างบน/ล่างแต่ละบรรทัด (Pt) · line = ระยะบรรทัด (เช่น 1.5)"""
    amap = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT}
    cell.text = ""
    for i, (txt, al, bold) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = amap[al]
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(before)
        if line:
            p.paragraph_format.line_spacing = line
        r = p.add_run(txt)
        r.font.name = THAI_FONT
        _csize(r, size)
        _bcs(r, bold)
        r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
from app.services.lunch_doc import (_money, _dnum, _save, _simple_table, _daygroup_table, _BLANK,
                                    _memo_head, _committee_lines, _school_disp)

_THAI_MONTHS = ["", "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน",
                "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"]


def _month_year(dt) -> str:
    if not dt:
        return "................"
    return f"{_THAI_MONTHS[dt.month]} {dt.year + 543}"


def _fund(prog) -> str:
    return (prog.funding_org or "").strip() or "องค์กรปกครองส่วนท้องถิ่น"


def _docnos(rnd) -> dict:
    """เลขที่/วันที่รายฉบับเอกสารของรอบ (JSON) {kind:{no,date(ISO)}}"""
    import json
    try:
        return json.loads(getattr(rnd, "doc_nos", "") or "{}") or {}
    except Exception:
        return {}


def _memo_ref(rnd, kind=None):
    """เลขที่/วันที่ของเอกสารชนิด kind (จากทะเบียนกลาง) - คืน (วันที่str, เลขที่)
    ลำดับ: เลขรายฉบับ (doc_nos[kind]) -> เลขบันทึกรวมเดิม -> เลขตกลงจ้าง"""
    ent = _docnos(rnd).get(kind or "", {}) if kind else {}
    no = (ent.get("no") or "").strip() or (getattr(rnd, "memo_no", "") or "").strip() or (rnd.order_no or "").strip()
    dt = None
    if ent.get("date"):
        try:
            dt = datetime.fromisoformat(ent["date"])
        except Exception:
            dt = None
    dt = dt or getattr(rnd, "memo_date", None) or rnd.order_date
    return _dnum(dt), no


def _doc_no(rnd, kind, fallback=""):
    """เลขที่เอกสารรายฉบับ (doc_nos[kind]['no']) - ไม่มีก็คืน fallback"""
    return ((_docnos(rnd).get(kind) or {}).get("no") or "").strip() or fallback


def _doc_dt(rnd, kind, field="date"):
    """วันที่ในฟิลด์ field ของเอกสาร kind (คืน datetime หรือ None)"""
    s = (_docnos(rnd).get(kind) or {}).get(field) or ""
    try:
        return datetime.fromisoformat(s) if s else None
    except Exception:
        return None


def _borrower(school, prog=None):
    """ผู้ยืมเงิน/ผู้จ่ายเงิน = เจ้าหน้าที่โครงการอาหารกลางวัน (ถ้าตั้งไว้ในโครงการ)
    ถ้าไม่ได้ตั้ง จึงใช้เจ้าหน้าที่พัสดุเป็นค่าตั้งต้น (อาจเป็นคนละคนกัน)"""
    name = ((getattr(prog, "lunch_officer", "") or "").strip()
            or (getattr(school, "officer_name", "") or "").strip() or _BLANK)
    return name, "เจ้าหน้าที่โครงการอาหารกลางวัน"


def _round_ingredients(rnd):
    """วัตถุดิบของโครงการที่อยู่ในช่วงวันของรอบนี้ (เรียงตามวัน/ลำดับ)"""
    prog = rnd.program
    s = rnd.start_date.date() if rnd.start_date else None
    e = rnd.end_date.date() if rnd.end_date else None
    out = []
    for ig in sorted(prog.ingredients, key=lambda x: (x.date or datetime.max, x.seq, x.id)):
        d = ig.date.date() if ig.date else None
        if s and e and d and (d < s or d > e):
            continue
        out.append(ig)
    return out


def _committee(rnd, kind):
    """รายชื่อกรรมการชนิดที่ระบุ (เรียงตามลำดับ)"""
    return [m for m in sorted(rnd.committees, key=lambda x: x.seq) if m.kind == kind]


def _crit_cell():
    """เกณฑ์ประเมินการประกอบอาหาร (ให้ติ๊กในเอกสาร)"""
    return ("ความสะอาด ☐ดีมาก ☐ดี ☐พอใช้ ☐ปรับปรุง\n"
            "คุณภาพอาหาร ☐ดีมาก ☐ดี ☐พอใช้ ☐ปรับปรุง\n"
            "ความทันเวลา ☐ดีมาก ☐ดี ☐พอใช้ ☐ปรับปรุง\n"
            "ความเพียงพอ ☐ดีมาก ☐ดี ☐พอใช้ ☐ปรับปรุง")


def _round_menu_days(rnd):
    """เมนูรายวันในช่วงของรอบ (เรียงตามวัน)"""
    prog = rnd.program
    return [m for m in sorted(prog.menus, key=lambda x: (x.date or datetime.max))
            if m.date and rnd.start_date and rnd.end_date and rnd.start_date <= m.date <= rnd.end_date]


def _begin(doc):
    if doc is None:
        d = Document(); set_a4(d); _font(d); return d, True
    if doc.paragraphs or doc.tables:
        doc.add_page_break()
    return doc, False


def _finish(doc, own, name):
    return _save(doc, name) if own else doc


def render_borrow_memo(rnd, school, doc=None) -> str:
    """02 บันทึกขออนุมัติยืมเงินอุดหนุนอาหารกลางวัน"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    saddr = (school.address or "").strip()
    fund = _fund(prog)
    bname, bpos = _borrower(school, rnd.program)
    fin = (getattr(school, "finance_officer_name", "") or "").strip() or _BLANK
    director = (school.director_name or "").strip() or _BLANK
    students = prog.total_students
    days = rnd.days or 0
    rate = prog.rate_per_head or 0
    total = round(float(rnd.amount or 0), 2)

    _memo_head(doc, school, [f"ขออนุมัติยืมเงิน (เงินอุดหนุนอาหารกลางวันรับจาก{fund})"],
               *_memo_ref(rnd, "borrow"))
    _p(doc, f"ด้วยข้าพเจ้า {bname} ตำแหน่ง {bpos} มีความประสงค์ขอยืมเงิน (เงินอุดหนุนอาหารกลางวัน"
            f"รับจาก{fund}) สำหรับเป็นค่าใช้จ่ายอาหารกลางวันให้นักเรียนระดับอนุบาลถึงประถมศึกษาปีที่ 6 "
            f"จำนวน {students} คน ระหว่างวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)} "
            f"รวมระยะเวลา {days} วัน เป็นเงิน {_money(total)} บาท (ตัวอักษร {bahttext(total)}) "
            f"({students} คน x {days} วัน x {_money(rate)} บาท) ตามสัญญาการยืมเงินและประมาณการดังแนบ "
            "และขอรับรองว่า ข้าพเจ้าไม่มีหนี้ผูกพันเกี่ยวกับเงินยืมกับทางราชการแต่อย่างใด",
       align="justify", indent=1.25, after=2)
    _p(doc, "จึงเรียนมาเพื่อโปรดพิจารณาอนุมัติ", indent=1.25, after=12)
    _sign_table(doc, [
        [("(ลงชื่อ)....................................................ผู้ยืม", "center"),
         (f"( {bname} )", "center"),
         (f"ตำแหน่ง {bpos}", "center")]])
    _p(doc, "", after=4)
    _simple_table(doc, ["ความคิดเห็นเจ้าหน้าที่การเงิน", "คำสั่ง/การสั่งการ"],
                  [[f"ได้ตรวจสอบสัญญาการยืมเงินและเอกสารประกอบแล้วถูกต้องตามระเบียบ เห็นควรอนุมัติ"
                    f"ให้ยืมเงินให้แก่ผู้ยืม\n\n(ลงชื่อ)....................เจ้าหน้าที่การเงิน\n( {fin} )",
                    f"(  ) ทราบ  (  ) อนุมัติ  (  ) ลงนามในสัญญาการยืมเงิน\n\n"
                    f"(ลงชื่อ)....................ผู้อำนวยการโรงเรียน\n( {director} )"]],
                  [Cm(8.2), Cm(7.8)])
    return _finish(doc, own, f"บันทึกขออนุมัติยืมเงิน_รอบที่{rnd.seq}_ปี{prog.year}")


def render_estimate(rnd, school, doc=None) -> str:
    """04 แบบประมาณการค่าใช้จ่าย (แนบท้ายสัญญายืมเงิน)"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    bname, bpos = _borrower(school, rnd.program)
    students = prog.total_students
    days = rnd.days or 0
    rate = prog.rate_per_head or 0
    total = round(float(rnd.amount or 0), 2)

    _p(doc, "แบบประมาณการค่าใช้จ่าย", align="center", bold=True, size=18, after=0)
    _p(doc, f"{sname}", align="center", after=0)
    _p(doc, f"แนบท้ายสัญญาเงินยืมเลขที่ {_doc_no(rnd, 'loan', (rnd.order_no or '').strip() or '........./.........')} "
            f"ลงวันที่ {_dnum(_doc_dt(rnd, 'loan', 'date') or rnd.order_date)}", align="center", after=6)
    _simple_table(doc, ["รายการ", "จำนวนเงิน"],
                  [[f"ประมาณการค่าอาหารกลางวัน ประจำเดือน {_month_year(rnd.start_date)} "
                    f"จำนวน {students} คน x อัตราวันละ {_money(rate)} บาท x จำนวน {days} วัน",
                    _money(total)],
                   [f"รวมจำนวนเงินทั้งสิ้น (ตัวอักษร {bahttext(total)})", _money(total)]],
                  [Cm(11.5), Cm(4.5)])
    _p(doc, "", after=12)
    _sign_table(doc, [
        [("(ลงชื่อ)..................................................ผู้ประมาณการ/ผู้ยืม", "center"),
         (f"( {bname} )", "center"),
         (f"ตำแหน่ง {bpos}", "center")]])
    return _finish(doc, own, f"แบบประมาณการค่าใช้จ่าย_รอบที่{rnd.seq}_ปี{prog.year}")


def render_purchase_report(rnd, school, doc=None) -> str:
    """03 บันทึกข้อความ รายงานขอซื้อวัตถุดิบ (7 ข้อ + ข้อ 8 แต่งตั้งกรรมการ 3 ชุด)
    ตรงตามคู่มืออาหารกลางวัน สพฐ. (จัดซื้อวัตถุดิบ) - ดึงกรรมการ 3 ชุดที่ตั้งไว้มาแสดง"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    officer = (getattr(school, "officer_name", "") or "").strip() or _BLANK
    head = (getattr(school, "head_officer_name", "") or "").strip() or _BLANK
    director = (school.director_name or "").strip() or _BLANK
    students = prog.total_students
    days = rnd.days or 0
    total = round(float(rnd.amount or 0), 2)
    ds, de = _dnum(rnd.start_date), _dnum(rnd.end_date)
    month = _month_year(rnd.start_date)
    _memo_head(doc, school,
               [f"รายงานขอซื้อวัตถุดิบเพื่อใช้ประกอบอาหารประจำเดือน {month}",
                f"(ระหว่างวันที่ {ds} ถึงวันที่ {de})"],
               *_memo_ref(rnd, "report"))
    _p(doc, f"ด้วย{sname} จัดซื้อวัตถุดิบเพื่อใช้ในการประกอบอาหารให้นักเรียนรับประทาน ระหว่างวันที่ "
            f"{ds} ถึงวันที่ {de} รวม {days} วัน การจัดซื้อครั้งนี้ดำเนินการโดยวิธีเฉพาะเจาะจง ตามพระราชบัญญัติ"
            "การจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560 มาตรา 56 (2) (ข) และหนังสือคณะกรรมการ"
            "วินิจฉัยปัญหาการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ ด่วนที่สุด ที่ กค (กวจ) 0405.2/ว 116 "
            "ลงวันที่ 12 มีนาคม พ.ศ. 2562 ซึ่งมีรายละเอียดดังต่อไปนี้", align="justify", indent=1.25, after=2)
    _p(doc, "1. เหตุผลและความจำเป็นที่ต้องจัดซื้อ", bold=True, indent=1.25, after=0)
    _p(doc, f"เพื่อใช้ในการประกอบอาหารให้นักเรียนรับประทานในมื้อกลางวัน สำหรับนักเรียน จำนวน {students} คน",
       indent=1.5, after=2)
    _p(doc, "2. รายละเอียดคุณลักษณะเฉพาะของพัสดุ", bold=True, indent=1.25, after=0)
    _p(doc, "รายการวัตถุดิบและคุณลักษณะเฉพาะของวัตถุดิบ รายละเอียดตามเอกสารแนบท้าย", indent=1.5, after=2)
    _p(doc, "3. ราคากลางของพัสดุที่จะซื้อ", bold=True, indent=1.25, after=0)
    _p(doc, f"เป็นเงิน {_money(total)} บาท ({bahttext(total)})", indent=1.5, after=2)
    _p(doc, "4. วงเงินที่จะซื้อ", bold=True, indent=1.25, after=0)
    _p(doc, f"เป็นเงิน {_money(total)} บาท ({bahttext(total)})", indent=1.5, after=2)
    _p(doc, "5. กำหนดเวลาที่ต้องการพัสดุ", bold=True, indent=1.25, after=0)
    _p(doc, f"กำหนดส่งมอบระหว่างวันที่ {ds} ถึงวันที่ {de} ทุกวันทำการ เวลา 06.00 น.", indent=1.5, after=2)
    _p(doc, "6. วิธีที่จะซื้อและเหตุผลที่ต้องซื้อโดยวิธีเฉพาะเจาะจง", bold=True, indent=1.25, after=0)
    _p(doc, "ดำเนินการโดยวิธีเฉพาะเจาะจง เนื่องจากการจัดซื้อจัดจ้างพัสดุที่มีการผลิต จำหน่าย หรือให้บริการทั่วไป "
            "และมีวงเงินในการจัดซื้อจัดจ้างครั้งหนึ่งไม่เกินวงเงินตามที่กำหนดในกฎกระทรวง", align="justify", indent=1.5, after=2)
    _p(doc, "7. หลักเกณฑ์การพิจารณาคัดเลือกข้อเสนอ", bold=True, indent=1.25, after=0)
    _p(doc, "การพิจารณาคัดเลือกข้อเสนอโดยใช้เกณฑ์ราคา", indent=1.5, after=2)
    _p(doc, "8. การอนุมัติแต่งตั้งบุคคลหรือคณะกรรมการ ดังนี้", bold=True, indent=1.25, after=0)
    _p(doc, "8.1 คณะกรรมการตรวจรับพัสดุ ประกอบด้วย", indent=1.5, after=0)
    _committee_lines(doc, _committee(rnd, "inspect"))
    _p(doc, "8.2 ผู้ควบคุมรับผิดชอบในการประกอบอาหาร ได้แก่", indent=1.5, before=2, after=0)
    _committee_lines(doc, _committee(rnd, "cook_control"), fallback_n=1)
    _p(doc, "8.3 คณะกรรมการตรวจการประกอบอาหาร ประกอบด้วย", indent=1.5, before=2, after=0)
    _committee_lines(doc, _committee(rnd, "food_inspect"))
    _p(doc, "จึงเรียนมาเพื่อโปรดพิจารณา หากเห็นชอบขอได้โปรดอนุมัติให้ดำเนินการตามรายละเอียดในรายงานขอซื้อ"
            "ดังกล่าวข้างต้น", align="justify", indent=1.25, before=2, after=10)
    _sign_table(doc, [
        [("(ลงชื่อ)..............................................เจ้าหน้าที่พัสดุ", "center"),
         (f"( {officer} )", "center")],
        [("(ลงชื่อ)..............................................หัวหน้าเจ้าหน้าที่", "center"),
         (f"( {head} )", "center")]])
    _p(doc, "คำสั่ง   อนุมัติตามเสนอ", align="center", bold=True, before=4, after=8)
    _sign_table(doc, [
        [("(ลงชื่อ)..............................................ผู้อำนวยการโรงเรียน", "center"),
         (f"( {director} )", "center")]])
    return _finish(doc, own, f"รายงานขอซื้อวัตถุดิบ_รอบที่{rnd.seq}_ปี{prog.year}")


def _purchase_day_block(doc, rnd, school, day_dt, day_ings):
    """ใบจัดซื้อวัสดุเครื่องบริโภค 1 วัน (ฟอร์ม 4 ส่วน + ตารางวัสดุของวันนั้น)
    เติมให้ครบ: วันที่=วันนั้น · ผู้จัดทำ=เจ้าหน้าที่อาหารกลางวัน · ประเภทเงิน+จำนวนเงินจากวัตถุดิบวันนั้น"""
    from docx.shared import Cm as _Cm
    prog = rnd.program
    sname = _school_disp(school)
    officer = (getattr(school, "officer_name", "") or "").strip() or _BLANK
    head = (getattr(school, "head_officer_name", "") or "").strip() or _BLANK
    director = (school.director_name or "").strip() or _BLANK
    fin = (getattr(school, "finance_officer_name", "") or "").strip() or _BLANK
    bname, bpos = _borrower(school, prog)            # เจ้าหน้าที่โครงการอาหารกลางวัน
    fund = (getattr(prog, "funding_org", "") or "").strip()
    fund_txt = f"เงินอุดหนุนอาหารกลางวัน จาก อบต. {fund}" if fund else "เงินอุดหนุนอาหารกลางวัน"
    D = "..................................."
    daystr = _dnum(day_dt) if day_dt else "................................"
    total = sum((ig.quantity or 0) * (ig.unit_price or 0) for ig in day_ings)
    total_txt = _money(total) if day_ings else "....................."
    total_words = bahttext(total) if day_ings else D
    dkey = (day_dt.date() if day_dt else None)
    menu_main = next((m.main or "" for m in prog.menus if m.date and m.date.date() == dkey), "") if dkey else ""

    def nm_or(members, i):
        return (members[i].name if (members and i < len(members) and members[i].name) else D)
    ctrl = _committee(rnd, "cook_control") + _committee(rnd, "food_inspect")
    insp = _committee(rnd, "inspect")

    def sign(role, name=None, date_val=None):
        b = [(f"(ลงชื่อ) {D}", "center", False),
             (f"( {name} )" if name else f"( {D} )", "center", False),
             (role, "center", False)]
        if date_val is not None:
            b.append((f"วันที่ {date_val}", "center", False))
        return b

    _p(doc, "ใบจัดซื้อวัสดุเครื่องบริโภค วงเงินไม่เกิน 500,000 บาท", align="center", bold=True, size=17, after=0)
    _p(doc, f"{sname}", align="center", after=0)
    _p(doc, f"ประจำวันที่ {daystr}", align="center", after=4)

    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = "Table Grid"
    tbl.autofit = False
    for row in tbl.rows:
        row.cells[0].width = _Cm(7.75)
        row.cells[1].width = _Cm(7.75)

    _box_cell(tbl.cell(0, 0), [("ส่วนที่ 1 รายงานขอซื้อ", "left", True)], size=13)
    _box_cell(tbl.cell(0, 1), [("ส่วนที่ 3 ใบรับรองแทนใบเสร็จ (บค. ........./..........)", "left", True)], size=13)
    _box_cell(tbl.cell(1, 0), [
        (f"ด้วย {sname} ขอจัดซื้อวัสดุเครื่องบริโภคตามรายการต่อไปนี้ เพื่อประกอบอาหารให้แก่นักเรียน"
         f"รับประทาน ในวันที่ {daystr} การจัดซื้อครั้งนี้ดำเนินการโดยวิธีเฉพาะเจาะจง"
         "ตามมาตรา 56 (2) (ข) ประกอบหนังสือกระทรวงการคลัง ด่วนที่สุด ที่ กค (กวจ) 0405.2/ว 116 "
         "ลงวันที่ 12 มีนาคม 2562", "left", False),
        ("", "left", False),
    ] + sign("ผู้จัดทำรายการ", bname, date_val=daystr), size=12)
    _box_cell(tbl.cell(1, 1), [
        (f"ข้าพเจ้า {bname}", "left", False),
        (f"ตำแหน่ง {bpos}", "left", False),
        (f"ได้จ่ายเงินจำนวน {total_txt} บาท", "left", False),
        (f"ตัวอักษร ( {total_words} )", "left", False),
        ("โดยไม่อาจเรียกใบเสร็จรับเงินจากผู้รับเงินได้ ตามรายการที่แนบ", "left", False),
        ("", "left", False),
    ] + sign("ผู้จ่ายเงิน", bname, date_val=daystr), size=12)

    _box_cell(tbl.cell(2, 0), [("ส่วนที่ 2 การอนุมัติการจัดซื้อ", "left", True)], size=13)
    _box_cell(tbl.cell(2, 1), [("ส่วนที่ 4 ผลการตรวจรับและอนุมัติการจ่ายเงิน", "left", True)], size=13)
    _box_cell(tbl.cell(3, 0), [
        (f"เรียน  ผู้อำนวยการ{sname}", "left", False),
        ("เพื่อโปรดทราบและ", "left", False),
        ("1. เห็นชอบตามรายงานขอซื้อ", "left", False),
        ("2. แต่งตั้งผู้ควบคุมและคณะกรรมการตรวจการประกอบอาหาร", "left", False),
        (f"     2.1 {nm_or(ctrl, 0)}", "left", False),
        (f"     2.2 {nm_or(ctrl, 1)}", "left", False),
        (f"     2.3 {nm_or(ctrl, 2)}", "left", False),
        ("3. แต่งตั้งผู้ตรวจรับหรือคณะกรรมการตรวจรับพัสดุ", "left", False),
        (f"     3.1 {nm_or(insp, 0)}  ประธาน", "left", False),
        (f"     3.2 {nm_or(insp, 1)}  กรรมการ", "left", False),
        (f"     3.3 {nm_or(insp, 2)}  กรรมการ", "left", False),
    ] + sign("เจ้าหน้าที่", officer) + sign("หัวหน้าเจ้าหน้าที่", head, date_val=daystr) + [
        ("", "left", False),
        ("อนุมัติตามเสนอข้อ 1 และ 2", "center", True),
        ("", "left", False),
    ] + sign("ผู้อำนวยการโรงเรียน", director, date_val=daystr), size=12)
    _box_cell(tbl.cell(3, 1), [
        (f"เรียน  ผู้อำนวยการ{sname}", "left", False),
        ("เพื่อโปรดทราบ พัสดุตามรายการข้างต้นได้ทำการตรวจรับไว้เป็นการถูกต้อง ครบถ้วนแล้ว", "left", False),
        (f"(ลงชื่อ) {D} ประธานกรรมการตรวจรับ", "left", False),
        (f"(ลงชื่อ) {D} กรรมการ", "left", False),
        (f"(ลงชื่อ) {D} กรรมการ", "left", False),
    ] + sign("เจ้าหน้าที่", officer) + [
        ("ได้ตรวจสอบหลักฐานแล้วถูกต้อง จึงขออนุมัติเบิกจ่ายเงิน", "left", False),
        (f"ประเภท {fund_txt} จำนวนเงิน {total_txt} บาท", "left", False),
    ] + sign("เจ้าหน้าที่การเงิน", fin, date_val=daystr) + [
        ("ทราบ อนุมัติตามรายการที่ขอเบิก และจ่ายเงินได้", "center", True),
    ] + sign("ผู้อำนวยการโรงเรียน", director, date_val=daystr)
      + sign("ผู้รับเงิน", bname) + sign("ผู้จ่ายเงิน", bname), size=12)

    # ตารางวัสดุเครื่องบริโภคของวันนี้ (หัวเรื่อง center + ชื่อรายการอาหารเป็นแถวหัวตาราง center)
    _p(doc, "รายละเอียดแนบท้ายใบจัดซื้อเครื่องบริโภค วงเงินไม่เกิน 500,000 บาท",
       align="center", bold=True, before=6, after=2, size=14)
    body = []
    for ig in day_ings:
        amt = (ig.quantity or 0) * (ig.unit_price or 0)
        body.append([ig.name or "", f"{_money(ig.quantity)} {ig.unit or ''}".strip(),
                     _money(ig.unit_price), _money(amt)])
    if body:
        body.append(["", "", "รวมเป็นเงินทั้งสิ้น", _money(total)])
    else:
        body = [["", "", "", ""] for _ in range(6)]
    _simple_table(doc, ["วัสดุเครื่องบริโภค", "จำนวนหน่วย", "ราคา/หน่วย", "จำนวนเงิน"],
                  body, [Cm(6.2), Cm(3.0), Cm(3.0), Cm(3.0)],
                  title_row=("รายการอาหาร: " + menu_main) if menu_main else None)


def render_purchase_form(rnd, school, doc=None) -> str:
    """05 ใบจัดซื้อวัสดุเครื่องบริโภค วงเงินไม่เกิน 500,000 บาท - ออกแยกรายวัน (1 วัน 1 ใบ)
    เติมวันที่/ผู้จัดทำ/ประเภทเงิน/จำนวนเงิน จากวัตถุดิบที่กรอกในแต่ละวัน"""
    doc, own = _begin(doc)
    prog = rnd.program
    days = _ings_by_day(rnd)
    if not days:                                # ยังไม่กรอกวัตถุดิบ -> ใบเปล่า 1 ใบ
        _purchase_day_block(doc, rnd, school, None, [])
    else:
        for i, (d, igs) in enumerate(days):
            if i > 0:
                doc.add_page_break()
            day_dt = igs[0].date if igs else None
            _purchase_day_block(doc, rnd, school, day_dt, igs)
    return _finish(doc, own, f"ใบจัดซื้อวัสดุเครื่องบริโภค_รอบที่{rnd.seq}_ปี{prog.year}")


def render_material_report_form(rnd, school, doc=None) -> str:
    """06 ใบรับรายงานวัตถุดิบและปริมาณการจัดซื้อ (ฟอร์มเปล่าตามโปรแกรม Thai School Lunch)"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    _p(doc, "ใบรับรายงานวัตถุดิบและปริมาณการจัดซื้อ", align="center", bold=True, size=17, after=0)
    _p(doc, "(ตามโปรแกรม Thai School Lunch หรือปรับใช้ตามหลักโภชนาการ)", align="center", after=0)
    _p(doc, f"{sname}", align="center", after=0)
    _p(doc, f"ระหว่างวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)}", align="center", after=6)
    # ดึงวัตถุดิบรายวันมาเติมให้ (merge วัน/เมนู ต่อวัน ให้ดูสวย) - ถ้ายังไม่กรอกให้เป็นฟอร์มเปล่า
    menus_by_date = {m.date.date(): (m.main or "") for m in prog.menus if m.date}
    groups, total = [], 0.0
    for d, rows in _ings_by_day(rnd):
        items = []
        for ig in rows:
            amt = (ig.quantity or 0) * (ig.unit_price or 0); total += amt
            items.append([ig.name or "", _money(ig.quantity), ig.unit or "",
                          _money(ig.unit_price), _money(amt)])
        groups.append((_dnum(rows[0].date) if rows else "", menus_by_date.get(d, ""), items))
    headers = ["วันที่", "เมนูอาหาร", "ส่วนประกอบ", "จำนวน", "หน่วย", "ราคา/หน่วย", "จำนวนเงิน"]
    widths = [Cm(3.5), Cm(3.2), Cm(2.6), Cm(1.2), Cm(1.3), Cm(1.8), Cm(2.1)]
    if groups:
        _daygroup_table(doc, headers, widths, groups, total=total)
        _p(doc, f"รวมเป็นเงินทั้งสิ้น (ตัวอักษร) {bahttext(total)}", align="right", before=2, after=2, size=13)
    else:
        _daygroup_table(doc, headers, widths, [("", "", [["", "", "", "", ""]]) for _ in range(6)])
    bname, bpos = _borrower(school, rnd.program)
    _p(doc, "", after=16)   # เว้นที่ให้เซ็น
    _sign_table(doc, [
        [("(ลงชื่อ)....................................ผู้จัดทำรายงาน", "center"),
         (f"( {bname} )", "center"),
         (f"ตำแหน่ง {bpos}", "center")]])
    return _finish(doc, own, f"ใบรับรายงานวัตถุดิบ_รอบที่{rnd.seq}_ปี{prog.year}")


def render_receipt_form(rnd, school, doc=None) -> str:
    """08 ใบรับรองการจ่ายเงินค่าวัตถุดิบ (ตารางราคากลางที่จัดซื้อ) ตรงตามคู่มืออาหารกลางวัน สพฐ.
    ดึงวัตถุดิบ+เมนูรายวันมาเติม + ลงชื่อ เจ้าหน้าที่โครงการอาหารกลางวัน/ผู้จ่ายเงิน"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    bname, bpos = _borrower(school, rnd.program)
    _p(doc, "ใบรับรองการจ่ายเงินค่าวัตถุดิบ", align="center", bold=True, size=18, after=0)
    _p(doc, f"{sname}  ระหว่างวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)}",
       align="center", after=6)
    ings = _round_ingredients(rnd)
    menus_by_date = {m.date.date(): (m.main or "") for m in prog.menus if m.date}
    body, prev_d, total = [], None, 0.0
    for ig in ings:
        d = ig.date.date() if ig.date else None
        amt = (ig.quantity or 0) * (ig.unit_price or 0)
        total += amt
        body.append([menus_by_date.get(d, "") if d != prev_d else "",
                     ig.name or "", _money(ig.quantity), ig.unit or "", _money(amt)])
        prev_d = d
    if body:
        body.append(["", "รวมเงิน", "", "", _money(total)])
    else:
        body = [["", "", "", "", ""] for _ in range(6)]
    _simple_table(doc,
                  ["รายการอาหาร", "วัตถุดิบ/เครื่องปรุง", "จำนวน", "หน่วยนับ", "ราคากลางที่จัดซื้อ (บาท)"],
                  body,
                  [Cm(4.6), Cm(5.0), Cm(1.6), Cm(1.7), Cm(3.35)])   # ขยายรายการอาหาร/วัตถุดิบ ลดราคา
    if body and total:
        _p(doc, f"รวมเป็นเงินทั้งสิ้น (ตัวอักษร) {bahttext(total)}", align="right", before=2, after=2, size=13)
    _p(doc, "", after=16)   # เว้นที่ให้เซ็น
    _sign_table(doc, [
        [("(ลงชื่อ)............................................เจ้าหน้าที่โครงการอาหารกลางวัน/ผู้จ่ายเงิน", "center"),
         (f"( {bname} )", "center"),
         (f"ตำแหน่ง {bpos}", "center")]])
    _p(doc, "หมายเหตุ กรณีไม่อาจเรียกใบเสร็จรับเงินจากผู้ขายได้ ให้ใช้ใบรับรองการจ่ายเงินแทน "
            "(อ้างอิงระเบียบกระทรวงการคลังฯ พ.ศ. 2562 ข้อ 48)", indent=0.5, before=8, size=13)
    return _finish(doc, own, f"ใบรับรองการจ่ายเงินค่าวัตถุดิบ_รอบที่{rnd.seq}_ปี{prog.year}")


def render_control_report(rnd, school, doc=None) -> str:
    """09 บันทึกรายงานผู้ควบคุมและคณะกรรมการตรวจการประกอบอาหารกลางวัน"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    bname, _ = _borrower(school, rnd.program)
    director = (school.director_name or "").strip() or _BLANK
    _p(doc, "บันทึกรายงานผู้ควบคุมและคณะกรรมการตรวจการประกอบอาหารกลางวัน",
       align="center", bold=True, size=17, after=4)
    _p(doc, f"เขียนที่ {sname}", align="right", after=0)
    _p(doc, f"วันที่ {_dnum(rnd.end_date)}", align="right", after=6)
    _p(doc, f"ตามที่{sname} ได้มอบหมายให้ {bname} จัดซื้อวัตถุดิบและประกอบอาหารกลางวันให้นักเรียน"
            f"รับประทาน ระหว่างวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)} บัดนี้ "
            "ได้ดำเนินการประกอบอาหารทุกวันตามที่กำหนด ผู้ควบคุมและคณะกรรมการตรวจการประกอบอาหารกลางวัน "
            "ขอรายงานผลการดำเนินงาน ดังนี้", align="justify", indent=1.25, after=4)
    # ช่องขวา = จุดไข่ปลาเว้นว่างให้ผู้ควบคุม/กรรมการเซ็น 3 ช่อง (เว้นที่ให้เซ็นด้วย space before)
    days = _round_menu_days(rnd)
    # สร้างตารางเอง: ช่องผลตรวจฟอนต์ 12 ให้ทุกหัวข้ออยู่บรรทัดเดียว (ไม่ตก)
    hdr = ["วัน เดือน ปี", "รายการอาหาร", "ผลการตรวจสอบ", "ผู้ควบคุม/กรรมการตรวจการ"]
    W = [Cm(2.1), Cm(2.35), Cm(7.95), Cm(3.85)]   # ขยายช่องผู้ควบคุม
    crit = [(f"{h}  ☐ดีมาก ☐ดี ☐พอใช้ ☐ปรับปรุง", "left", False)
            for h in ("ความสะอาด", "คุณภาพอาหาร", "ความทันเวลา", "ความเพียงพอ")]
    signs = [("(ลงชื่อ) ...............................", "left", False) for _ in range(3)]   # 3 ช่อง
    tb = doc.add_table(rows=1, cols=4)
    tb.style = "Table Grid"
    tb.autofit = False
    tb.allow_autofit = False
    for c, h, w in zip(tb.rows[0].cells, hdr, W):
        _set_cell(c, h, bold=True, align="center", size=14)
        c.width = w
    for m in (days or [None] * 5):
        rc = tb.add_row().cells
        _set_cell(rc[0], _dnum(m.date) if m else "", size=14)
        _set_cell(rc[1], (m.main or "") if m else "", size=14)
        _box_cell(rc[2], crit, size=12)
        _box_cell(rc[3], signs, size=13, before=10, after=6)   # เว้นที่ให้เซ็นแต่ละช่อง
        for c, w in zip(rc, W):
            c.width = w
    _p(doc, "(  ) ทราบผลการดำเนินการประกอบอาหารกลางวัน", indent=1.25, before=4, after=10)
    _p(doc, "ลงชื่อ......................................................", align="center", after=0)
    _p(doc, f"( {director} )", align="center", after=0)
    _p(doc, f"ตำแหน่ง ผู้อำนวยการ{sname}", align="center", after=0)
    return _finish(doc, own, f"บันทึกรายงานตรวจการประกอบอาหาร_รอบที่{rnd.seq}_ปี{prog.year}")


def render_repay_memo(rnd, school, doc=None) -> str:
    """10 บันทึกขออนุมัติเบิกจ่ายเงินเพื่อส่งใช้เงินยืม"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    saddr = (school.address or "").strip()
    fund = _fund(prog)
    bname, bpos = _borrower(school, rnd.program)
    fin = (getattr(school, "finance_officer_name", "") or "").strip() or _BLANK
    director = (school.director_name or "").strip() or _BLANK
    total = round(float(rnd.amount or 0), 2)

    _memo_head(doc, school, [f"ขออนุมัติเบิกจ่ายเงินเพื่อส่งใช้เงินยืม (เงินอุดหนุนอาหารกลางวันรับจาก{fund})"],
               *_memo_ref(rnd, "repay"))
    _p(doc, f"ตามที่อนุมัติให้ {bname} ผู้ยืมเงินโครงการอาหารกลางวัน ยืมเงิน (เงินอุดหนุนอาหารกลางวัน"
            f"รับจาก{fund}) เพื่อเป็นค่าใช้จ่ายอาหารกลางวันให้นักเรียนรับประทาน จำนวนเงิน {_money(total)} "
            f"บาท (ตัวอักษร {bahttext(total)}) ตามสัญญาการยืมเงินที่ {_doc_no(rnd, 'loan', (rnd.order_no or '').strip() or _BLANK)} "
            f"ลงวันที่ {_dnum(_doc_dt(rnd, 'loan', 'date') or rnd.order_date)} นั้น", align="justify", indent=1.25, after=2)
    _p(doc, "บัดนี้ ได้ดำเนินการตามวัตถุประสงค์แล้ว ขอส่งใช้หลักฐาน และเงินสด (ถ้ามี) ดังนี้",
       align="justify", indent=1.25, after=0)
    _p(doc, f"1. หลักฐานค่าอาหารกลางวัน\t\tจำนวน {_money(total)} บาท", indent=1.5, after=0)
    _p(doc, "2. เงินสด (ถ้ามี)\t\t\tจำนวน - บาท", indent=1.5, after=0)
    _p(doc, f"รวมเป็นเงิน {_money(total)} บาท", indent=1.5, after=2)
    _p(doc, f"จึงเรียนมาเพื่อโปรดทราบ และอนุมัติเบิกจ่ายเงิน (เงินอุดหนุนอาหารกลางวันรับจาก{fund}) "
            f"จำนวน {_money(total)} บาท (ตัวอักษร {bahttext(total)})", align="justify", indent=1.25, after=12)
    _sign_table(doc, [
        [("(ลงชื่อ)....................................................ผู้ยืม", "center"),
         (f"( {bname} )", "center"),
         (f"ตำแหน่ง {bpos}", "center")]])
    _p(doc, "", after=4)
    _simple_table(doc, ["ความคิดเห็นเจ้าหน้าที่การเงิน", "คำสั่ง/การสั่งการ"],
                  [[f"ได้ตรวจสอบหลักฐานและเอกสารประกอบการส่งใช้เงินยืมแล้วถูกต้องครบถ้วนตามระเบียบ "
                    f"เห็นควรอนุมัติเบิกจ่ายเงิน\n\n(ลงชื่อ)....................เจ้าหน้าที่การเงิน\n( {fin} )",
                    f"(  ) ทราบ  (  ) อนุมัติ\n\n(ลงชื่อ)....................ผู้อำนวยการโรงเรียน\n( {director} )"]],
                  [Cm(8.2), Cm(7.8)])
    return _finish(doc, own, f"ขออนุมัติเบิกจ่ายส่งใช้เงินยืม_รอบที่{rnd.seq}_ปี{prog.year}")


def render_reimburse_summary(rnd, school, doc=None) -> str:
    """ตัวอย่าง 14 ใบสรุปเบิกเงินชดเชยเงินยืมโครงการอาหารกลางวัน
    (ใช้เฉพาะกรณีส่งใช้คืนเงินยืมเป็นภาคเรียน)"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    bname, bpos = _borrower(school, rnd.program)
    fin = (getattr(school, "finance_officer_name", "") or "").strip() or _BLANK
    director = (school.director_name or "").strip() or _BLANK
    total = round(float(rnd.amount or 0), 2)
    money = _money(total)
    DOT = "..............................................."

    _p(doc, "ใบสรุปเบิกเงินชดเชยเงินยืมโครงการอาหารกลางวัน", align="center", bold=True, size=18, after=0)
    _p(doc, f"{sname}", align="center", after=0)
    _p(doc, f"วันที่ {_dnum(rnd.end_date)}", align="center", after=6)
    # ตารางรายการค่าอาหารแต่ละงวด + รวม
    insts = sorted(rnd.installments or [], key=lambda i: i.seq)
    rows = []
    if insts:
        for k, it in enumerate(insts, 1):
            rows.append([str(k), f"ค่าอาหารประจำวันที่ {_dnum(it.end_date)}", _money(it.amount or 0), ""])
    else:
        for k in range(1, 6):
            rows.append([str(k), "ค่าอาหารประจำวันที่ ...................................", "", ""])
    rows.append(["", "รวม", _money(total), ""])
    _simple_table(doc, ["ลำดับที่", "รายการ", "จำนวนเงิน", "หมายเหตุ"],
                  rows, [Cm(1.6), Cm(7.4), Cm(3.0), Cm(3.5)])
    _p(doc, "ข้าพเจ้าขอเบิกเงินเพื่อชดเชยเงินยืม", indent=1.25, before=6, after=2)
    _p(doc, f"( / )  เงินอุดหนุนโครงการอาหารกลางวัน\t\tเป็นเงิน {money} บาท", indent=1.5, after=0)
    _p(doc, "(   )  กองทุนเพื่อโครงการอาหารกลางวัน\t\tเป็นเงิน ............................ บาท", indent=1.5, after=0)
    _p(doc, "(   )  บำรุงการศึกษา-โครงการอาหารกลางวัน\t\tเป็นเงิน ............................ บาท", indent=1.5, after=2)
    _p(doc, "ซึ่งจ่ายเป็นค่าอาหารตามใบสำคัญรับเงินที่แนบ", indent=1.25, after=12)
    _sign_table(doc, [
        [(f"ลงชื่อ {DOT} ผู้ขอเบิก", "center"),
         (f"( {bname} )", "center")]])
    _p(doc, "ได้ตรวจสอบรายการจ่ายตามใบสำคัญที่ขอเบิก จำนวน ................ ฉบับ ถูกต้องแล้ว",
       indent=1.25, before=10, after=10)
    _sign_table(doc, [
        [(f"ลงชื่อ {DOT} เจ้าหน้าที่การเงิน", "center"),
         (f"( {fin} )", "center"),
         (f"วันที่ {DOT}", "center")]])
    _p(doc, "ทราบและอนุมัติจ่ายเงิน", indent=1.25, before=10, after=10)
    _sign_table(doc, [
        [(f"ลงชื่อ {DOT} หัวหน้าสถานศึกษา", "center"),
         (f"( {director} )", "center"),
         (f"วันที่ {DOT}", "center")]])
    _p(doc, "*** ใช้เอกสารนี้เฉพาะกรณีส่งใช้คืนเงินยืมเป็นภาคเรียน", size=13, before=12)
    return _finish(doc, own, f"ใบสรุปเบิกเงินชดเชยเงินยืม_รอบที่{rnd.seq}_ปี{prog.year}")


def _ings_by_day(rnd):
    """จัดกลุ่มวัตถุดิบตามวัน คืน [(date, [ings]), ...] เรียงตามวัน"""
    from itertools import groupby
    ings = sorted(_round_ingredients(rnd), key=lambda x: (x.date or datetime.max, x.seq, x.id))
    out = []
    for d, grp in groupby(ings, key=lambda x: (x.date.date() if x.date else None)):
        out.append((d, list(grp)))
    return out


def render_purchase_attach(rnd, school, doc=None) -> str:
    """หน้าแนบท้ายใบจัดซื้อ: ตารางรายการอาหาร + วัสดุเครื่องบริโภครายวัน (ส่วนที่แยกจากฟอร์ม 4 ส่วน)"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    menus = {m.date.date(): (m.main or "") for m in prog.menus if m.date}
    _p(doc, "รายการอาหารและวัสดุเครื่องบริโภค (แนบท้ายใบจัดซื้อ)", align="center", bold=True, size=16, after=0)
    _p(doc, f"{sname}  ระหว่างวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)}",
       align="center", after=6)
    groups, tot = [], 0.0
    for d, rows in _ings_by_day(rnd):
        items = []
        for ig in rows:
            amt = (ig.quantity or 0) * (ig.unit_price or 0); tot += amt
            items.append([ig.name or "", f"{_money(ig.quantity)} {ig.unit or ''}".strip(),
                          _money(ig.unit_price), _money(amt)])
        groups.append((_dnum(rows[0].date) if rows else "", menus.get(d, ""), items))
    headers = ["วัน เดือน ปี", "รายการอาหาร", "วัสดุเครื่องบริโภค", "จำนวนหน่วย", "ราคา/หน่วย", "จำนวนเงิน"]
    widths = [Cm(2.6), Cm(3.0), Cm(3.4), Cm(2.2), Cm(2.1), Cm(2.2)]
    if groups:
        _daygroup_table(doc, headers, widths, groups, total=tot)
    else:
        _daygroup_table(doc, headers, widths, [("", "", [["", "", "", ""]]) for _ in range(6)])
    _p(doc, "", after=10)
    _sign_table(doc, [
        [("(ลงชื่อ)...........................................ผู้จัดทำรายการ", "center")]])
    return _finish(doc, own, f"แนบท้ายใบจัดซื้อ_รอบที่{rnd.seq}_ปี{prog.year}")


def render_ingredient_deliver(rnd, school, doc=None) -> str:
    """ใบส่งมอบวัตถุดิบเพื่อใช้ประกอบอาหารกลางวัน (รายวัน วันละหน้า) - ผู้ขายส่งมอบวัตถุดิบ"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    seller = (getattr(rnd.vendor, "name", "") or "").strip() if rnd.vendor else ""
    bname, _bpos = _borrower(school, prog)
    DN = "..............................................."
    days = _ings_by_day(rnd) or [(None, [])]
    first = True
    for d, rows in days:
        if not first:
            doc.add_page_break()
        first = False
        _p(doc, "ใบส่งมอบวัตถุดิบเพื่อใช้ประกอบอาหารกลางวัน", align="center", bold=True, size=17, after=0)
        _p(doc, f"{sname}", align="center", after=6)
        body = [[ig.name or "", f"{_money(ig.quantity)} {ig.unit or ''}".strip(), ""] for ig in rows]
        while len(body) < 8:
            body.append(["", "", ""])
        _simple_table(doc, ["รายการวัตถุดิบ", "จำนวน", "หมายเหตุ"],
                      body, [Cm(8.5), Cm(3.5), Cm(3.5)])
        _p(doc, "", after=10)
        _sign_table(doc, [
            [(f"(ลงชื่อ) {DN}", "center"), (f"( {seller or DN} )", "center"),
             ("ผู้ส่ง", "center"), (f"วันที่ {_dnum2(d)}", "center")],
            [(f"(ลงชื่อ) {DN}", "center"), (f"( {bname} )", "center"),
             ("ผู้รับ", "center"), (f"วันที่ {_dnum2(d)}", "center")]])
    return _finish(doc, own, f"ใบส่งมอบวัตถุดิบ_รอบที่{rnd.seq}_ปี{prog.year}")


def render_inspect_detail(rnd, school, doc=None) -> str:
    """ใบแสดงรายละเอียดการตรวจรับพัสดุ (แนบท้ายใบตรวจรับ) - รายวัน + ช่องลงชื่อกรรมการ 3 คน"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    students = prog.total_students or ""
    menus = {m.date.date(): (m.main or "") for m in prog.menus if m.date}
    days = _round_menu_days(rnd)
    _p(doc, "ใบแสดงรายละเอียดการตรวจรับพัสดุ", align="center", bold=True, size=17, after=0)
    _p(doc, f"แนบท้ายใบตรวจรับพัสดุ ลงวันที่ {_dnum(rnd.end_date)}", align="center", after=0)
    _p(doc, f"{sname}", align="center", after=6)
    sign3 = "1) .....................  2) .....................  3) ....................."
    body = []
    for m in days:
        body.append([_dnum(m.date), (m.main or ""), f"{students} ชุด", sign3, ""])
    if not body:
        body = [["", "", f"{students} ชุด", sign3, ""] for _ in range(6)]
    _simple_table(doc, ["วัน เดือน ปี", "รายการอาหาร", "จำนวน", "ลายมือชื่อกรรมการตรวจรับ", "หมายเหตุ"],
                  body, [Cm(2.6), Cm(3.4), Cm(1.6), Cm(5.9), Cm(2.0)])
    return _finish(doc, own, f"ใบแสดงรายละเอียดการตรวจรับ_รอบที่{rnd.seq}_ปี{prog.year}")


def render_reimburse_advance(rnd, school, doc=None) -> str:
    """ใบสรุปเบิกเงินทดรองจ่ายโครงการอาหารกลางวัน (สรุปค่าอาหารรายวัน + ขอเบิกชดเชยเงินทดรอง)"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    bname, _bpos = _borrower(school, prog)
    fin = (getattr(school, "finance_officer_name", "") or "").strip() or _BLANK
    director = (school.director_name or "").strip() or _BLANK
    DN = "..............................................."
    _p(doc, "ใบสรุปเบิกเงินทดรองจ่ายโครงการอาหารกลางวัน", align="center", bold=True, size=18, after=0)
    _p(doc, f"{sname}", align="center", after=0)
    _p(doc, f"วันที่ {_dnum(rnd.end_date)}", align="center", after=6)
    rows, tot = [], 0.0
    for k, (d, igs) in enumerate(_ings_by_day(rnd), 1):
        amt = sum((ig.quantity or 0) * (ig.unit_price or 0) for ig in igs); tot += amt
        rows.append([str(k), f"ค่าอาหารประจำวันที่ {_dnum2(d)}", _money(amt), ""])
    if not rows:
        rows = [[str(k), "ค่าอาหารประจำวันที่ ...................................", "", ""] for k in range(1, 6)]
    total = round(float(rnd.amount or 0), 2) if not tot else round(tot, 2)
    fuel = round(float(getattr(prog, "fuel_cost", 0) or getattr(rnd, "fuel_cost", 0) or 0), 2)  # ค่าเชื้อเพลิง (จากงบโครงการ)
    if fuel:
        rows.append([str(len(rows) + 1), "ค่าเชื้อเพลิงประกอบอาหาร (แก๊ส/ถ่าน)", _money(fuel), ""])
        total = round(total + fuel, 2)
    rows.append(["", "รวม", _money(total), ""])
    _simple_table(doc, ["ที่", "รายการ", "จำนวนเงิน", "หมายเหตุ"],
                  rows, [Cm(1.6), Cm(7.4), Cm(3.0), Cm(3.5)])
    _p(doc, f"(ตัวอักษร {bahttext(total)})", indent=0.5, before=2, after=6)
    _p(doc, "ข้าพเจ้าขอเบิกเงินเพื่อชดเชยเงินทดรองจ่าย", indent=1.25, after=2)
    _p(doc, f"( / )  เงินอุดหนุนโครงการอาหารกลางวัน\t\tเป็นเงิน {_money(total)} บาท", indent=1.5, after=0)
    _p(doc, "(   )  กองทุนเพื่อโครงการอาหารกลางวัน\t\tเป็นเงิน ............................ บาท", indent=1.5, after=0)
    _p(doc, "(   )  บำรุงการศึกษา-โครงการอาหารกลางวัน\t\tเป็นเงิน ............................ บาท", indent=1.5, after=2)
    _p(doc, "ซึ่งจ่ายเป็นค่าอาหารตามใบสำคัญรับเงินที่แนบ", indent=1.25, after=12)
    _sign_table(doc, [[(f"ลงชื่อ {DN} ผู้ขอเบิก", "center"), (f"( {bname} )", "center")]])
    _p(doc, "ได้ตรวจสอบรายการจ่ายตามใบสำคัญที่ขอเบิก จำนวน ................ ฉบับ ถูกต้องแล้ว",
       indent=1.25, before=10, after=10)
    _sign_table(doc, [[(f"ลงชื่อ {DN} เจ้าหน้าที่การเงิน", "center"), (f"( {fin} )", "center"),
                       (f"วันที่ {DN}", "center")]])
    _p(doc, "ทราบและอนุมัติจ่ายเงิน", indent=1.25, before=10, after=10)
    _sign_table(doc, [[(f"ลงชื่อ {DN} หัวหน้าสถานศึกษา", "center"), (f"( {director} )", "center"),
                       (f"วันที่ {DN}", "center")]])
    return _finish(doc, own, f"ใบสรุปเบิกเงินทดรองจ่าย_รอบที่{rnd.seq}_ปี{prog.year}")


def render_wht_cook(rnd, school, doc=None) -> str:
    """ใบรับรองการหักภาษี ณ ที่จ่าย (ค่าจ้างประกอบอาหารกลางวัน - จ้างแม่ครัว)"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    saddr = (school.address or "").strip()
    stax = (getattr(school, "tax_id", "") or "").strip() or "-"
    v = rnd.vendor
    vname = (v.name if v else "") or _BLANK
    vtax = (getattr(v, "tax_id", "") or "").strip() if v else ""
    vaddr = (getattr(v, "address", "") or "").strip() if v else ""
    fin = (getattr(school, "finance_officer_name", "") or "").strip() or _BLANK
    order_no = _doc_no(rnd, "order", (rnd.order_no or "").strip() or _BLANK)
    order_dt = _doc_dt(rnd, "order", "date") or rnd.order_date
    total = round(float(rnd.amount or 0), 2)
    tax = round(total * 0.01, 2)
    tb, ts = int(total), int(round((total - int(total)) * 100))
    xb, xs = int(tax), int(round((tax - int(tax)) * 100))

    _p(doc, "ใบรับรองการหักภาษี ณ ที่จ่าย", align="center", bold=True, size=18, after=4)
    _p(doc, f"ส่วนราชการ {sname}  เลขประจำตัวผู้เสียภาษี {stax}", after=0)
    _p(doc, f"ที่อยู่ {saddr or _BLANK}", after=4)
    _p(doc, f"ขอรับรองว่าได้หักภาษี ณ ที่จ่ายตามใบสั่งจ้าง เลขที่ {order_no}  ลงวันที่ {_dnum(order_dt)}",
       indent=1.25, after=0)
    _p(doc, f"ชื่อผู้ถูกหัก {vname}  เลขประจำตัวผู้เสียภาษี {vtax or '-'}", after=0)
    _p(doc, f"ที่อยู่ {vaddr or _BLANK}", after=6)
    _simple_table(doc,
                  ["ประเภทเงินได้ที่จ่าย", "วัน เดือน ปี ที่จ่าย", "จำนวนเงินที่จ่าย", "ภาษีที่หัก"],
                  [["ค่าจ้างประกอบอาหารกลางวัน", _dnum(rnd.end_date), f"{_money(total)}", f"{_money(tax)}"],
                   ["รวมเงินภาษีที่หักและนำส่ง", "", f"{_money(total)}", f"{_money(tax)}"]],
                  [Cm(6.0), Cm(3.5), Cm(3.0), Cm(3.0)])
    _p(doc, f"รวมเงินภาษี (ตัวอักษร) ({bahttext(tax)})", indent=0.5, before=2, after=12)
    _sign_table(doc, [
        [(f"(ลงชื่อ) ...........................................", "center"),
         (f"( {fin} )", "center"), ("เจ้าหน้าที่การเงิน", "center")]])
    return _finish(doc, own, f"หนังสือรับรองหักภาษี_รอบที่{rnd.seq}_ปี{prog.year}")


def _dnum2(d):
    """แปลง date (ไม่ใช่ datetime) เป็นข้อความวันที่ไทย"""
    if not d:
        return "................................"
    from datetime import datetime as _dt
    return _dnum(_dt(d.year, d.month, d.day))


def _committee_signs(doc, members):
    """ลายเซ็นคณะกรรมการตรวจรับ (ประธาน/กรรมการ) - ไม่มีชื่อ = เส้นจุด 3 คน"""
    seq = members if members else [None, None, None]
    for i, mem in enumerate(seq):
        role = (mem.role if mem else "") or ("ประธานกรรมการ" if i == 0 else "กรรมการ")
        _p(doc, f"(ลงชื่อ) ........................................... {role}", align="center", before=6, after=0)
        _p(doc, (f"( {mem.name} )" if mem else "(...........................................)"),
           align="center", after=0)


def render_purchase_summary(rnd, school, doc=None) -> str:
    """สรุปรายการจัดซื้อวัตถุดิบ (เพื่อประกอบการตรวจรับพัสดุ) - ดึงวัตถุดิบ+กรรมการตรวจรับมาเติม"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    ings = _round_ingredients(rnd)
    total = sum((i.quantity or 0) * (i.unit_price or 0) for i in ings)
    _p(doc, "สรุปรายการจัดซื้อวัตถุดิบเพื่อใช้ประกอบอาหาร", align="center", bold=True, size=17, after=0)
    _p(doc, f"{sname}  ระหว่างวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)}",
       align="center", after=4)
    _p(doc, "เรียน  ประธานคณะกรรมการตรวจรับพัสดุ/ผู้ตรวจรับพัสดุ", indent=1, after=2)
    _p(doc, "ขอรายงานรายการจัดซื้อวัตถุดิบเพื่อใช้ประกอบอาหาร เพื่อประกอบการตรวจรับพัสดุ ดังนี้",
       indent=1.25, after=2)
    body = [[i.name or "", _money(i.quantity), i.unit or "", _money(i.unit_price),
             _money((i.quantity or 0) * (i.unit_price or 0)), ""] for i in ings]
    if body:
        body.append(["รวมเป็นเงินทั้งสิ้น", "", "", "", _money(total), ""])
    else:
        body = [["", "", "", "", "", ""] for _ in range(5)]
    _simple_table(doc, ["วัตถุดิบ/เครื่องปรุง", "จำนวน", "หน่วยนับ", "ราคาต่อหน่วย", "จำนวนเงิน", "หมายเหตุ"],
                  body, [Cm(5.0), Cm(1.8), Cm(1.8), Cm(2.6), Cm(2.6), Cm(2.45)])
    _p(doc, "ได้ตรวจสอบรายการวัตถุดิบดังกล่าวแล้ว รายการ จำนวน ถูกต้องครบถ้วน",
       indent=1.25, before=4, after=8)
    _committee_signs(doc, _committee(rnd, "inspect"))
    return _finish(doc, own, f"สรุปรายการจัดซื้อวัตถุดิบ_รอบที่{rnd.seq}_ปี{prog.year}")


def render_inspection_note(rnd, school, doc=None) -> str:
    """ใบตรวจรับพัสดุ (ตามข้อ 175 ระเบียบฯ พ.ศ. 2560) - ดึงกรรมการตรวจรับ + เจ้าหน้าที่พัสดุ"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    officer = (getattr(school, "officer_name", "") or "").strip() or _BLANK
    days = len(_round_menu_days(rnd)) or (rnd.days or 0)
    _p(doc, "ใบตรวจรับพัสดุ", align="center", bold=True, size=18, after=2)
    _p(doc, f"เขียนที่ {sname}", align="right", after=0)
    _p(doc, f"วันที่ {_dnum(rnd.end_date)}", align="right", after=6)
    _p(doc, f"ตามรายงานสรุปการตรวจรับวัตถุดิบเพื่อใช้ประกอบอาหารกลางวัน ในระหว่างวันที่ "
            f"{_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)} รวม {days} วัน รายละเอียดตามแนบ "
            "คณะกรรมการตรวจรับพัสดุ ได้ตรวจรับและให้ถือว่าพัสดุ", align="justify", indent=1.25, after=2)
    _p(doc, "( / )  ถูกต้อง          (   )  ไม่ถูกต้อง  จำนวน .............. รายการ", indent=1.5, after=2)
    _p(doc, f"จึงรายงานต่อผู้อำนวยการ{sname} เพื่อโปรดทราบผลการตรวจรับ ตามนัยข้อ 175 "
            "แห่งระเบียบกระทรวงการคลังว่าด้วยการจัดซื้อจัดจ้างและการบริหารพัสดุภาครัฐ พ.ศ. 2560",
       align="justify", indent=1.25, after=8)
    _committee_signs(doc, _committee(rnd, "inspect"))
    _p(doc, f"เรียน  ผู้อำนวยการ{sname}", indent=1, before=10, after=0)
    _p(doc, "     เพื่อโปรดทราบ คณะกรรมการตรวจรับพัสดุได้ดำเนินการตรวจรับพัสดุเรียบร้อยแล้ว",
       indent=1.25, after=6)
    _p(doc, "(ลงชื่อ) ........................................... เจ้าหน้าที่พัสดุ", align="center", after=0)
    _p(doc, f"( {officer} )", align="center", after=0)
    return _finish(doc, own, f"ใบตรวจรับพัสดุ_รอบที่{rnd.seq}_ปี{prog.year}")


def _memo_head_to(doc, school, subject_lines, date, doc_no, addressee):
    """หัวบันทึกข้อความ (สารบรรณ) แบบระบุผู้รับเอง (เช่น เรียน ประธานกรรมการตรวจรับ)"""
    sname = _school_disp(school)
    office = (f"{sname}  " + (school.address or "").strip()).strip()
    _krut_and_title(doc)
    _p_runs(doc, [("ส่วนราชการ  ", True), (office, False)], after=0)
    _p_runs(doc, [("ที่  ", True), ((doc_no or "").strip() or _BLANK, False),
                  ("\t", False), ("วันที่  ", True), (date, False)], tab_cm=8, after=0)
    for i, s in enumerate(subject_lines):
        _p_runs(doc, [("เรื่อง  ", True), (s, False)], after=0) if i == 0 else _p(doc, "        " + s, after=0)
    _p_runs(doc, [("เรียน  ", True), (addressee, False)], after=0)
    _hr(doc)


def render_menu_list(rnd, school, doc=None) -> str:
    """01 รายการอาหารกลางวัน (ตาม Thai School Lunch) - ตารางเมนูรายวันในช่วงรอบ"""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    bname, _ = _borrower(school, rnd.program)
    _p(doc, "รายการอาหารกลางวัน (ตาม Thai School Lunch)", align="center", bold=True, size=17, after=0)
    _p(doc, f"{sname}  ระหว่างวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)}",
       align="center", after=6)
    days = _round_menu_days(rnd)
    if days:
        body = [[_dnum(m.date), " ".join(x for x in [(m.main or ""), (m.dessert or "")] if x), ""] for m in days]
    else:
        body = [["", "", ""] for _ in range(10)]
    _simple_table(doc, ["วัน เดือน ปี", "รายการอาหาร", "หมายเหตุ"], body,
                  [Cm(4.0), Cm(9.65), Cm(2.6)])   # ขยายคอลัมน์วันที่ไม่ให้บีบ
    _p(doc, "", after=14)   # เว้นที่ให้เจ้าหน้าที่เซ็น
    _sign_table(doc, [
        [("(ลงชื่อ)............................................เจ้าหน้าที่โครงการอาหารกลางวัน/ผู้จัดทำ", "center"),
         (f"( {bname} )", "center")]])
    return _finish(doc, own, f"รายการอาหารกลางวัน_รอบที่{rnd.seq}_ปี{prog.year}")


def render_inspect_notify(rnd, school, doc=None) -> str:
    """บันทึกข้อความ การตรวจรับพัสดุ (แจ้งประธานกรรมการตรวจรับให้ปฏิบัติหน้าที่)
    ตรงตามคู่มืออาหารกลางวัน สพฐ."""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    officer = (getattr(school, "officer_name", "") or "").strip() or _BLANK
    director = (school.director_name or "").strip() or _BLANK
    insp = _committee(rnd, "inspect")
    chair = insp[0].name if insp else _BLANK
    others = " และ ".join(m.name for m in insp[1:]) if len(insp) > 1 else "....................................."
    month = _month_year(rnd.start_date)
    _memo_head_to(doc, school,
                  [f"การตรวจรับพัสดุ รายการวัตถุดิบเพื่อใช้ในการประกอบอาหารกลางวันประจำเดือน {month}",
                   f"ระหว่างวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)}"],
                  *_memo_ref(rnd, "inspect-notify"),
                  "ประธานคณะกรรมการตรวจรับพัสดุ/ผู้ตรวจรับพัสดุ")
    _p(doc, f"ด้วยผู้อำนวยการ{sname} ได้แต่งตั้ง {chair} ประธานกรรมการตรวจรับพัสดุ/ผู้ตรวจรับพัสดุ "
            f"พร้อมด้วย {others} เป็นกรรมการตรวจรับพัสดุ เพื่อทำหน้าที่ตรวจรับวัตถุดิบเพื่อใช้ในการประกอบอาหาร"
            "ในการส่งมอบทุกครั้ง โดยให้จัดทำรายงานผลการตรวจรับเสนอหัวหน้าหน่วยงานของรัฐ (ผู้อำนวยการโรงเรียน) "
            "เป็นรายสัปดาห์หรือรายเดือน ตามความเหมาะสมแล้วแต่กรณี ทั้งนี้ ประธานคณะกรรมการตรวจรับพัสดุอาจ"
            "มอบหมายให้กรรมการคนหนึ่งคนใดทำหน้าที่ตรวจรับพัสดุเบื้องต้นในแต่ละครั้งที่มีการส่งมอบก็ได้ และให้"
            "กรรมการผู้มีหน้าที่จัดทำบันทึกสรุปรายการวัตถุดิบที่ตรวจรับในแต่ละครั้ง แล้วรวบรวมส่งมอบให้เจ้าหน้าที่"
            "พัสดุเก็บรวบรวมเสนอต่อคณะกรรมการตรวจรับพัสดุเป็นรายสัปดาห์หรือรายเดือนต่อไป",
       align="justify", indent=1.25, after=2)
    _p(doc, "สำหรับรายการวัตถุดิบเพื่อใช้ในการประกอบอาหารในแต่ละวัน เจ้าหน้าที่จะได้ส่งให้คณะกรรมการตรวจรับ"
            "พัสดุ/ผู้ตรวจรับพัสดุ ทราบต่อไป", align="justify", indent=1.25, after=2)
    _p(doc, "จึงเรียนมาเพื่อโปรดทราบ", indent=1.25, after=14)
    _sign_table(doc, [
        [("(ลงชื่อ)............................................", "center"),
         (f"( {officer} )", "center"),
         ("เจ้าหน้าที่พัสดุ", "center")]])
    return _finish(doc, own, f"บันทึกแจ้งตรวจรับพัสดุ_รอบที่{rnd.seq}_ปี{prog.year}")


def render_inspect_assign(rnd, school, doc=None) -> str:
    """ตัวอย่าง 4 บันทึกข้อความ การมอบหมายการตรวจรับวัตถุดิบเพื่อใช้ในการประกอบอาหารกลางวัน
    วันที่ตรวจรับ = วันสุดท้ายของแต่ละงวด (ต่อกรรมการแต่ละคน) · ตารางไร้เส้น ซ้าย ชื่อ/ขวา บทบาท+วันที่"""
    from docx.shared import Cm as _Cm
    doc, own = _begin(doc)
    prog = rnd.program
    insp = _committee(rnd, "inspect")
    insts = sorted(rnd.installments or [], key=lambda i: i.seq)
    dates = [_dnum(i.end_date) for i in insts] or ["............................." for _ in range(3)]
    DN = "..........................................."
    TH = ["๑", "๒", "๓", "๔", "๕", "๖", "๗", "๘"]
    members = insp if insp else [None, None, None]

    _memo_head_to(doc, school,
                  ["การมอบหมายการตรวจรับวัตถุดิบเพื่อใช้ในการประกอบอาหารกลางวัน"],
                  *_memo_ref(rnd, "inspect-assign"),
                  "คณะกรรมการตรวจรับพัสดุ")
    _p(doc, "ขอมอบหมายให้คณะกรรมการตรวจรับพัสดุ เพื่อทำหน้าที่ตรวจรับวัตถุดิบเพื่อใช้ในการประกอบอาหาร"
            "กลางวันเบื้องต้นในแต่ละครั้งที่มีการส่งมอบ ดังนี้", align="justify", indent=1.25, after=4)
    # ตารางไร้เส้น: ซ้าย = (ลำดับ) ชื่อ + ทำหน้าที่ตรวจรับพัสดุ | ขวา = บทบาท + วันที่ (วันสุดท้ายของแต่ละงวด)
    t = doc.add_table(rows=len(members), cols=2)
    t.autofit = False
    for i, m in enumerate(members):
        nm = (m.name if m else "") or ""
        role = "ประธานกรรมการตรวจรับพัสดุ" if i == 0 else "กรรมการตรวจรับพัสดุ"
        num = TH[i] if i < len(TH) else str(i + 1)
        _box_cell(t.cell(i, 0),
                  [(f"({num}) นาย/นาง/นางสาว {nm or DN}", "left", False),
                   ("      ทำหน้าที่ตรวจรับพัสดุ", "left", False)], size=14, after=2)
        _box_cell(t.cell(i, 1),
                  [(role, "left", False)] + [(f"วันที่ {d}", "left", False) for d in dates],
                  size=14, after=0)
        t.cell(i, 0).width = _Cm(7.5)
        t.cell(i, 1).width = _Cm(8.0)
    _no_borders(t)
    _p(doc, "และให้กรรมการที่ได้รับมอบหมายจัดทำบันทึกสรุปรายการวัตถุดิบที่ตรวจรับในแต่ละครั้ง "
            "ส่งมอบให้เจ้าหน้าที่เป็นรายสัปดาห์/รายเดือน ต่อไป", align="justify", indent=1.25, before=4, after=2)
    _p(doc, "จึงเรียนมาเพื่อโปรดทราบ", indent=1.25, after=10)
    _p(doc, f"(ลงชื่อ) {DN} ประธานคณะกรรมการตรวจรับพัสดุ", align="center", after=10)
    _p(doc, "ทราบ", after=2)
    for i, m in enumerate(members):
        role = "ประธานกรรมการตรวจรับพัสดุ" if i == 0 else "กรรมการตรวจรับพัสดุ"
        _p(doc, f"(ลงชื่อ) {DN} {role}", after=2)
    return _finish(doc, own, f"มอบหมายการตรวจรับวัตถุดิบ_รอบที่{rnd.seq}_ปี{prog.year}")


def render_inspect_report(rnd, school, doc=None) -> str:
    """บันทึกข้อความ รายงานการตรวจรับพัสดุ (เจ้าหน้าที่พัสดุ เสนอ ผอ. แนบใบตรวจรับ)
    ตรงตามคู่มืออาหารกลางวัน สพฐ."""
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    officer = (getattr(school, "officer_name", "") or "").strip() or _BLANK
    head = (getattr(school, "head_officer_name", "") or "").strip() or _BLANK
    director = (school.director_name or "").strip() or _BLANK
    days = len(_round_menu_days(rnd)) or (rnd.days or 0)
    month = _month_year(rnd.start_date)
    _memo_head(doc, school,
               [f"รายงานการตรวจรับพัสดุ รายการวัตถุดิบเพื่อใช้ในการประกอบอาหารกลางวันประจำเดือน {month}",
                f"ระหว่างวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)}"],
               *_memo_ref(rnd, "inspect-report"))
    _p(doc, f"ตามคำสั่ง{sname} เรื่อง แต่งตั้งคณะกรรมการตรวจรับพัสดุ ผู้ควบคุมรับผิดชอบในการประกอบ"
            "อาหาร และคณะกรรมการตรวจการประกอบอาหาร ในส่วนของคณะกรรมการตรวจรับพัสดุ มีหน้าที่ตรวจรับ"
            "อาหารสด อาหารแห้งรายวัน ซึ่งเจ้าหน้าที่โครงการอาหารกลางวันจัดซื้อตามรายการอาหารที่กำหนดในแต่ละวัน "
            "โดยให้จัดทำรายงานการตรวจรับเสนอผู้อำนวยการเป็นรายสัปดาห์ นั้น", align="justify", indent=1.25, after=2)
    _p(doc, f"เจ้าหน้าที่โครงการอาหารกลางวันได้จัดซื้อวัตถุดิบและส่งมอบให้คณะกรรมการตรวจรับพัสดุ ในระหว่าง"
            f"วันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)} จำนวน {days} วัน ตามแบบตรวจรับ"
            "ที่แนบมาพร้อมหนังสือนี้", align="justify", indent=1.25, after=2)
    _p(doc, "จึงเรียนมาเพื่อโปรดทราบ", indent=1.25, after=12)
    _sign_table(doc, [
        [("(ลงชื่อ)............................................", "center"),
         (f"( {officer} )", "center"),
         ("เจ้าหน้าที่พัสดุ", "center")]])
    _p(doc, "ความเห็นของหัวหน้าเจ้าหน้าที่พัสดุ   เรียน ผู้อำนวยการโรงเรียน เพื่อโปรดทราบ",
       indent=1.25, before=4, after=8)
    _sign_table(doc, [
        [("(ลงชื่อ)............................................", "center"),
         (f"( {head} )", "center"),
         ("หัวหน้าเจ้าหน้าที่พัสดุ", "center")]])
    _p(doc, "คำสั่ง   ทราบ", align="center", bold=True, before=4, after=8)
    _sign_table(doc, [
        [("(ลงชื่อ)............................................ผู้อำนวยการโรงเรียน", "center"),
         (f"( {director} )", "center")]])
    return _finish(doc, own, f"รายงานการตรวจรับพัสดุ_รอบที่{rnd.seq}_ปี{prog.year}")


def render_loan_contract(rnd, school, doc=None) -> str:
    """04 สัญญาการยืมเงิน (แบบ 8500) หน้า + หลัง (รายการส่งใช้เงินยืม)
    ตรงตามคู่มืออาหารกลางวัน สพฐ. - เติมชื่อผู้ยืม/จำนวนเงิน/ประมาณการให้"""
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    bname, bpos = _borrower(school, rnd.program)
    director = (school.director_name or "").strip() or _BLANK
    fund = _fund(prog)
    students = prog.total_students
    days = rnd.days or 0
    rate = prog.rate_per_head or 0
    total = round(float(rnd.amount or 0), 2)
    money, baht = _money(total), bahttext(total)
    month = _month_year(rnd.start_date)
    # เลขที่/วันที่ของสัญญายืมเงิน กรอกที่หน้าจัดการงวด (doc_nos["loan"]) - หลายช่อง
    oid = _doc_no(rnd, "loan", (rnd.order_no or "").strip() or _BLANK)
    borrow_dt = _doc_dt(rnd, "loan", "date") or rnd.order_date         # วันที่ยืม (ผู้ยืมลงชื่อ)
    present_dt = _doc_dt(rnd, "loan", "present_date")                  # วันที่เสนอ จนท.การเงิน/ผอ.อนุมัติ
    receive_dt = _doc_dt(rnd, "loan", "receive_date") or borrow_dt     # วันที่ได้รับเงิน
    od = _dnum(borrow_dt)
    present_s = _dnum(present_dt) if present_dt else ".................."
    receive_s = _dnum(receive_dt)
    DOT = "..............................................."
    PROMISE = ("ข้าพเจ้าสัญญาว่าจะปฏิบัติตามระเบียบของทางราชการทุกประการ และจะนำใบสำคัญคู่จ่ายที่ถูกต้อง "
               "พร้อมทั้งเงินเหลือจ่าย (ถ้ามี) ส่งใช้ภายในกำหนดไว้ในระเบียบการเบิกจ่ายเงินจากคลัง คือภายใน 30 วัน "
               "นับแต่วันที่ได้รับเงินยืมนี้ ถ้าข้าพเจ้าไม่ส่งใช้ตามกำหนด ข้าพเจ้ายินยอมให้หักเงินเดือน ค่าจ้าง "
               "หรือเงินอื่นใดที่ข้าพเจ้าพึงได้รับจากทางราชการ ชดใช้จำนวนเงินที่ยืมไปจนครบถ้วนได้ทันที")

    # ---------- หน้า (front) : ฟอร์มแบบ 8500 (ตารางมีกรอบ = ช่องๆ) ----------
    _p(doc, "แบบ 8500", align="right", size=13, after=0)
    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = "Table Grid"
    tbl.autofit = False
    tbl.allow_autofit = False
    LW, RW = Cm(11.0), Cm(5.25)
    for row in tbl.rows:
        row.cells[0].width = LW
        row.cells[1].width = RW
    # R0
    _box_cell(tbl.cell(0, 0), [("สัญญาการยืมเงิน", "center", True),
                               (f"ยื่นต่อ ผู้อำนวยการ{sname}", "left", False)], line=1.5)
    _box_cell(tbl.cell(0, 1), [(f"เลขที่ {oid}", "left", False),
                               (f"วันครบกำหนด {_dnum(rnd.end_date)}", "left", False)], line=1.5)
    # R1 (เต็มความกว้าง)
    c = tbl.cell(1, 0).merge(tbl.cell(1, 1))
    _box_cell(c, [(f"ข้าพเจ้า {bname}  ตำแหน่ง {bpos}", "left", False),
                  (f"สังกัด {sname}  {(school.address or '').strip()}", "left", False),
                  (f"มีความประสงค์ขอยืมเงินจาก เงินอุดหนุนอาหารกลางวันรับจาก{fund}", "left", False),
                  (f"เพื่อเป็นค่าใช้จ่ายในการประกอบอาหารกลางวันให้นักเรียนรับประทาน ประจำเดือน {month} "
                   f"จำนวน {students} คน x วันละ {_money(rate)} บาท x {days} วัน  ดังรายละเอียดต่อไปนี้", "left", False)], line=1.5)
    # R2: (ตัวอักษร) ... รวมเงิน (บาท) | ยอดเงิน
    _box_cell(tbl.cell(2, 0), [(f"(ตัวอักษร) {baht}          รวมเงิน (บาท)", "left", False)], line=1.5)
    _box_cell(tbl.cell(2, 1), [(money, "center", True)], size=15, line=1.5)
    # R3: คำสัญญา + ลงชื่อผู้ยืม
    c = tbl.cell(3, 0).merge(tbl.cell(3, 1))
    _box_cell(c, [(PROMISE, "left", False),
                  (f"ลงชื่อ {DOT}ผู้ยืม   ( {bname} )   วันที่ {od}", "left", False)], line=1.5)
    # R4: เสนอ + คำอนุมัติ
    c = tbl.cell(4, 0).merge(tbl.cell(4, 1))
    _box_cell(c, [("เสนอ  ผู้อำนวยการโรงเรียน", "left", True),
                  (f"ได้ตรวจสอบแล้วเห็นสมควรอนุมัติให้ยืมตามใบยืมฉบับนี้ได้ จำนวน {money} บาท ({baht})", "left", False),
                  (f"ลงชื่อ {DOT}เจ้าหน้าที่การเงิน   วันที่ {present_s}", "left", False),
                  ("คำอนุมัติ  อนุมัติให้ยืมตามเงื่อนไขข้างต้นได้", "left", True),
                  (f"ลงชื่อ {DOT}ผู้อนุมัติ   ( {director} )   วันที่ {present_s}", "left", False)], line=1.5)
    # R5: ใบรับเงิน
    c = tbl.cell(5, 0).merge(tbl.cell(5, 1))
    _box_cell(c, [("ใบรับเงิน", "left", True),
                  (f"ได้รับเงินยืมจำนวน {money} บาท ({baht}) ไปเป็นการถูกต้องแล้ว", "left", False),
                  (f"ลงชื่อ {DOT}ผู้รับเงิน   ( {bname} )   วันที่ {receive_s}", "left", False)], line=1.5)

    # ---------- หลัง (back) : รายการส่งใช้เงินยืม ----------
    doc.add_page_break()
    _p(doc, "รายการส่งใช้เงินยืม", align="center", bold=True, size=17, after=6)
    hdr = ["ครั้งที่", "วัน เดือน ปี", "เงินสดหรือใบสำคัญ", "จำนวนเงิน", "คงค้าง",
           "ลายมือชื่อผู้รับเงิน", "ใบรับเลขที่"]
    widths = [Cm(1.5), Cm(2.6), Cm(3.0), Cm(2.3), Cm(2.3), Cm(2.65), Cm(1.9)]
    t = doc.add_table(rows=13, cols=7)
    t.style = "Table Grid"
    hcells = t.rows[0].cells
    for c, v, w in zip(hcells, hdr, widths):
        _set_cell(c, v, size=14, align="center"); c.width = w
    for row in t.rows[1:]:
        row.height = Cm(0.9); row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for c, w in zip(row.cells, widths):
            c.width = w
    _p(doc, "หมายเหตุ  (1) ยื่นต่อ ผู้อำนวยการโรงเรียน  (2) ระบุชื่อส่วนราชการที่จ่ายเงิน  "
            "(3) ระบุวัตถุประสงค์ที่จะนำเงินยืมไปใช้จ่าย  (4) เสนอต่อผู้มีอำนาจอนุมัติ",
       before=6, size=13)
    if own:   # ออกเดี่ยว: ย่อเนื้อความเหลือ 14 ให้หน้า+หลังพอดี 2 หน้า (ตามแบบ 8500)
        _shrink_body_font(doc, 14)
    return _finish(doc, own, f"สัญญาการยืมเงิน_รอบที่{rnd.seq}_ปี{prog.year}")


def render_food_photos(rnd, school, doc=None) -> str:
    """10 รูปภาพอาหารกลางวันในโรงเรียน - ตารางช่องติดรูปอาหารรายวัน (เมนูดึงมาให้)"""
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    doc, own = _begin(doc)
    prog = rnd.program
    sname = _school_disp(school)
    _p(doc, "รูปภาพอาหารกลางวันในโรงเรียน", align="center", bold=True, size=18, after=0)
    _p(doc, f"{sname}  ประจำเดือน {_month_year(rnd.start_date)}", align="center", after=0)
    _p(doc, f"ระหว่างวันที่ {_dnum(rnd.start_date)} ถึงวันที่ {_dnum(rnd.end_date)}", align="center", after=6)
    days = _round_menu_days(rnd)
    items = [(f"{_dnum(m.date)}\n{(m.main or '')}") for m in days] or ["" for _ in range(6)]
    # ตาราง 2 คอลัมน์ ต่อแถว: ป้ายวัน/เมนู อยู่บน + ช่องรูปสูง
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    for i in range(0, len(items), 2):
        pair = items[i:i + 2]
        # แถวป้าย
        lab = t.add_row().cells
        for j in range(2):
            _set_cell(lab[j], pair[j] if j < len(pair) else "", size=14, align="center")
            lab[j].width = Cm(8.1)
        # แถวช่องติดรูป (สูง)
        ph = t.add_row()
        ph.height = Cm(5.2); ph.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for j in range(2):
            _set_cell(ph.cells[j], "(ติดรูปอาหาร)" if (j < len(pair) and pair[j]) else "", size=12, align="center")
            ph.cells[j].width = Cm(8.1)
    return _finish(doc, own, f"รูปภาพอาหารกลางวัน_รอบที่{rnd.seq}_ปี{prog.year}")


def render_ingredient_bundle(rnd, school) -> str:
    """ออกชุดเอกสารซื้อวัตถุดิบทั้งชุดเป็นไฟล์เดียว (เรียงตามลำดับงานจริง)"""
    doc = Document(); set_a4(doc); _font(doc)
    # เรียงตามลำดับคู่มืออาหารกลางวัน สพฐ. (จัดซื้อวัตถุดิบ)
    render_menu_list(rnd, school, doc)            # 1  รายการอาหาร (Thai School Lunch)
    render_purchase_report(rnd, school, doc)      # 2  รายงานขอซื้อ + แต่งตั้งกรรมการ 3 ชุด
    render_purchase_form(rnd, school, doc)        # 3  ใบจัดซื้อวัสดุ 4 ส่วน + หน้าแนบท้าย
    render_borrow_memo(rnd, school, doc)          # 4  ขออนุมัติยืมเงิน
    render_loan_contract(rnd, school, doc)        # 5  สัญญาการยืมเงิน (แบบ 8500) หน้า+หลัง
    render_estimate(rnd, school, doc)             # 6  แบบประมาณการ
    render_inspect_notify(rnd, school, doc)       # 7  แจ้งประธานกรรมการตรวจรับ (การตรวจรับพัสดุ)
    render_inspect_assign(rnd, school, doc)       # 8  การมอบหมายการตรวจรับวัตถุดิบ (ตัวอย่าง 4)
    render_ingredient_deliver(rnd, school, doc)   # 9  ใบส่งมอบวัตถุดิบ (รายวัน)
    render_material_report_form(rnd, school, doc) # 10 ใบรับรายงานวัตถุดิบ
    render_purchase_summary(rnd, school, doc)     # 11 สรุปรายการจัดซื้อ (เพื่อตรวจรับ)
    render_inspection_note(rnd, school, doc)      # 12 ใบตรวจรับพัสดุ (ข้อ 175)
    render_inspect_detail(rnd, school, doc)       # 13 ใบแสดงรายละเอียดการตรวจรับ (แนบท้าย)
    render_receipt_form(rnd, school, doc)         # 14 ใบรับรองการจ่ายเงินค่าวัตถุดิบ
    render_control_report(rnd, school, doc)       # 15 รายงานการประกอบอาหาร
    render_food_photos(rnd, school, doc)          # 16 รูปภาพอาหารกลางวัน
    render_repay_memo(rnd, school, doc)           # 17 อนุมัติเบิกจ่ายส่งใช้เงินยืม
    render_reimburse_summary(rnd, school, doc)    # 18 ใบสรุปเบิกเงินชดเชยเงินยืม (กรณีส่งใช้รายภาคเรียน)
    return _save(doc, f"ชุดเอกสารซื้อวัตถุดิบ_รอบที่{rnd.seq}_ปี{rnd.program.year}")
