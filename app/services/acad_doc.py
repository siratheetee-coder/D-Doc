# -*- coding: utf-8 -*-
"""
acad_doc.py - เอกสารงานวิชาการ
- ปพ.5 : แบบบันทึกผลการพัฒนาคุณภาพผู้เรียน (รายวิชา x ห้อง) - แนวนอน
- ปพ.6 : แบบรายงานผลการพัฒนาคุณภาพผู้เรียนรายบุคคล (สมุดพก) - รายคน / ทั้งห้อง

ความกว้างตารางต้องไม่เกินพื้นที่พิมพ์ A4: แนวตั้ง 16.0 / แนวนอน 26.7 ซม.
(บทเรียนจากรอบไล่แก้ A4 - python-docx ไม่บีบให้เอง)
"""
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from app.services.doc_page import set_a4
from app.services.office_doc import _float_signature

from app.database import get_data_dir
from app.thai_utils import thai_date, is_secondary, be_date_input
from app.services.academic import (term_label, CHAR_ITEMS, CHAR_FIELDS, READ_DOMAINS,
                                   TH_MONTHS, TH_MONTH_FULL, TERM_MONTHS, quality_of_avg, char_avg, read_avg,
                                   effective_eval, weighted_avg, activities_for,
                                   activity_summary, onet_for, is_exit_level, ONET_SUBJECTS,
                                   count_marks, parse_marks, parse_days_csv)

THAI_FONT = "TH Sarabun New"


def _safe(text: str) -> str:
    for ch in '<>:"/\\|?*':
        text = text.replace(ch, "_")
    return text.strip()


def _doc(landscape: bool = False):
    doc = Document(); set_a4(doc, landscape=landscape)
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.2)
    base = doc.styles["Normal"]; base.font.name = THAI_FONT; base.font.size = Pt(14)
    base._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    return doc


def _p(doc, text="", *, align="left", bold=False, size=14, after=2, page_break=False):
    p = doc.add_paragraph()
    # ขึ้นหน้าใหม่โดยติดที่ย่อหน้านี้เอง (ไม่ใช่ add_page_break ที่สร้างย่อหน้าเปล่า
    # ทับกับตารางที่เต็มหน้าพอดี -> เกิดหน้าว่างคั่น)
    if page_break:
        p.paragraph_format.page_break_before = True
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.name = THAI_FONT
    r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    return p


def _cell(cell, text, *, bold=False, align="center", size=13, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(size); r.font.name = THAI_FONT
    r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    if fill:
        tcpr = cell._tc.get_or_add_tcPr()
        tcpr.append(tcpr.makeelement(qn("w:shd"), {qn("w:val"): "clear",
                                                   qn("w:color"): "auto", qn("w:fill"): fill}))


def _widths(table, widths):
    for row in table.rows:
        for c, w in zip(row.cells, widths):
            c.width = w
    _center(table)


def _center(table):
    """จัดตารางให้อยู่กึ่งกลางหน้ากระดาษ (python-docx ชิดซ้ายเป็นค่าปริยาย)"""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False


def _tight_cells(table, *, lr=15):
    """ลดระยะขอบซ้าย/ขวาภายในเซลล์ทั้งตาราง (กันตัวเลขโดนบีบในคอลัมน์แคบ) หน่วย twips
    ค่าปริยายของ Word = 108 twips (~0.19 ซม.) ต่อข้าง · ตั้ง 15 = เกือบชิดขอบ"""
    from docx.oxml import OxmlElement
    tblpr = table._tbl.tblPr
    mar = tblpr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar"); tblpr.append(mar)
    for side in ("left", "right"):
        e = mar.find(qn("w:" + side))
        if e is None:
            e = OxmlElement("w:" + side); mar.append(e)
        e.set(qn("w:w"), str(lr)); e.set(qn("w:type"), "dxa")


def _new_section(doc, *, landscape: bool):
    """ขึ้น section ใหม่พร้อมตั้งแนวกระดาษเอง
    (ห้ามเรียก set_a4 ซ้ำ เพราะมันบังคับ *ทุก* section ให้เป็นแนวเดียวกัน
     - เอกสารนี้ต้องผสม ปกแนวตั้ง + เนื้อในแนวนอน)"""
    from docx.enum.section import WD_SECTION
    from app.services.doc_page import A4_W, A4_H
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    if landscape:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = A4_H, A4_W
    else:
        sec.orientation = WD_ORIENT.PORTRAIT
        sec.page_width, sec.page_height = A4_W, A4_H
    sec.left_margin = sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.2)
    return sec


def _class_label(c) -> str:
    return f"{c.level}/{c.room}" if (c.room or "").strip() else (c.level or "")


def _sign_block(doc, name, role, *, after=0, size=10):
    """บล็อกลงนาม + วางลายเซ็นลอยทับถ้าผู้ลงนามอัปโหลดไว้"""
    p = _p(doc, "(ลงชื่อ).............................................", align="center", after=0, size=size)
    if name:
        _float_signature(p, name)
    _p(doc, f"( {name or '.......................................'} )", align="center", after=0, size=size)
    _p(doc, role, align="center", after=after, size=size)


def _logo_header(doc, school, *, page_break=False, height_cm=2.0, after=2):
    """แทรกโลโก้/ตราโรงเรียนกึ่งกลางหัวเอกสาร (ถ้าตั้งค่าไว้) · คืน True ถ้าใส่สำเร็จ
    รับ page_break มาแปะที่ย่อหน้าโลโก้ เพื่อให้ตัวเรียกไม่ต้องขึ้นหน้าซ้ำ"""
    import io
    logo = getattr(school, "logo", None)
    if not logo:
        return False
    p = doc.add_paragraph()
    if page_break:
        p.paragraph_format.page_break_before = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    try:
        p.add_run().add_picture(io.BytesIO(logo), height=Cm(height_cm))
    except Exception:      # ไฟล์เสีย/ฟอร์แมตไม่รองรับ - ข้ามโลโก้ ไม่ให้ล้มทั้งเอกสาร
        p._element.getparent().remove(p._element)
        return False
    return True


# ============================ ปพ.5 ============================
def _assignments_of(subject, db, term):
    """ชิ้นงานเก็บคะแนน (เรียงตามลำดับ) + คะแนนรายคน {(sid, aid): score} ของภาคนั้น"""
    from app.models import AcadAssignment, AcadAssignmentScore
    assigns = (db.query(AcadAssignment).filter_by(subject_id=subject.id, term=term)
               .order_by(AcadAssignment.seq, AcadAssignment.id).all())
    ascore = {}
    if assigns:
        aids = [a.id for a in assigns]
        for r in (db.query(AcadAssignmentScore)
                  .filter(AcadAssignmentScore.assignment_id.in_(aids)).all()):
            ascore[(r.acad_student_id, r.assignment_id)] = r.score
    return assigns, ascore


def _pp5_score_secondary(doc, subject, db, students, mmax, fmax, term):
    """ตารางคะแนนมัธยม (แนวนอน) - ที่/ชื่อ/[งานเก็บ]/คะแนนเก็บ/สอบกลางภาค/ปลายภาค/รวม/ผล/หมายเหตุ
    (ไม่มีเลขประจำตัว · งานเก็บ = เท่าที่มีจริง · สอบกลางภาค แยกช่องต่างหาก ไม่รวมในคะแนนเก็บ)"""
    from app.models import AcadScore
    scores = {x.acad_student_id: x for x in
              db.query(AcadScore).filter_by(subject_id=subject.id, term=term).all()}
    assigns, ascore = _assignments_of(subject, db, term)
    work = [a for a in assigns if not a.is_midterm]
    mids = [a for a in assigns if a.is_midterm]
    nw, has_mid = len(work), bool(mids)
    km_full = sum(a.max_score or 0 for a in work)  # เต็มคะแนนเก็บ = ผลรวมเต็มงานเก็บจริง
    mm_full = sum(a.max_score or 0 for a in mids)   # เต็มสอบกลางภาค = ผลรวมเต็มงานกลางภาคจริง
    total_full = km_full + mm_full + fmax
    keep_col = 2 + nw
    mid_col = keep_col + 1 if has_mid else None
    final_col = keep_col + 1 + (1 if has_mid else 0)
    total_col, grade_col, note_col = final_col + 1, final_col + 2, final_col + 3
    ncol = note_col + 1
    t = doc.add_table(rows=2, cols=ncol); t.style = "Table Grid"
    r0, r1 = t.rows[0].cells, t.rows[1].cells

    def vm2(idx, text):
        c = r0[idx].merge(r1[idx]); _cell(c, text, bold=True, fill="EDE9FE", size=12)

    vm2(0, "ที่"); vm2(1, "ชื่อ-นามสกุล")
    if nw:
        c = r0[2].merge(r0[2 + nw - 1]); _cell(c, "การเก็บคะแนนระหว่างภาค", bold=True, fill="EDE9FE", size=12)
        for j, a in enumerate(work):
            _cell(r1[2 + j], f"{j + 1}\n({a.max_score:g})", bold=True, fill="F1F5F9", size=11)
    vm2(keep_col, f"คะแนนเก็บ\n(เต็ม {km_full:g})")
    if has_mid:
        vm2(mid_col, f"สอบกลางภาค\n(เต็ม {mm_full:g})")
    vm2(final_col, f"คะแนนปลายภาค\n(เต็ม {fmax:g})")
    vm2(total_col, f"รวม\n(เต็ม {total_full:g})")
    vm2(grade_col, "ผลการเรียน")
    vm2(note_col, "หมายเหตุ")
    for s in students:
        sc = scores.get(s.id)
        cells = t.add_row().cells
        _cell(cells[0], s.seq or "", size=12)
        _cell(cells[1], s.name, align="left", size=12)
        work_vals = [ascore.get((s.id, a.id)) for a in work]
        for j, v in enumerate(work_vals):
            _cell(cells[2 + j], f"{v:g}" if v is not None else "", size=12)
        mid_sum = sum(ascore[(s.id, a.id)] for a in mids if ascore.get((s.id, a.id)) is not None)
        has_work = any(v is not None for v in work_vals)
        work_sum = sum(v for v in work_vals if v is not None)
        keep = work_sum if has_work else ((sc.score_mid - mid_sum) if (sc and sc.score_mid is not None) else None)
        _cell(cells[keep_col], f"{keep:g}" if keep is not None else "", size=12)
        if has_mid:
            _cell(cells[mid_col], f"{mid_sum:g}" if any(ascore.get((s.id, a.id)) is not None for a in mids) else "", size=12)
        _cell(cells[final_col], f"{sc.score_final:g}" if sc and sc.score_final is not None else "", size=12)
        _cell(cells[total_col], f"{sc.score:g}" if sc and sc.score is not None else "", size=12, bold=True)
        _cell(cells[grade_col], sc.grade if sc else "", bold=True, size=12)
        _cell(cells[note_col], "", size=12)
    # แนวนอนเกือบเต็มหน้า ~26 ซม.
    # คอลัมน์สรุป: คะแนนเก็บ [สอบกลางภาค] ปลายภาค รวม ผล หมายเหตุ
    summary_w = [2.4] + ([2.4] if has_mid else []) + [2.4, 2.0, 2.0, 2.2]
    fixed = 1.0 + 5.2 + sum(summary_w)
    aw = min(1.6, max(0.7, (26.0 - fixed) / nw)) if nw else 0
    _widths(t, [Cm(1.0), Cm(5.2)] + [Cm(aw)] * nw + [Cm(w) for w in summary_w])
    _tight_cells(t)
    if work:
        legend = " | ".join(f"{j + 1}. {a.name or 'งานที่ ' + str(j + 1)} (เต็ม {a.max_score:g})"
                            for j, a in enumerate(work))
        _p(doc, "งานเก็บคะแนน: " + legend, size=11, after=0, align="left")


def _pp5_score_page(doc, school, klass, subject, db, *, page_break: bool = False):
    """หน้าใบคะแนนของ 1 รายวิชา (ใช้ทั้งแบบแผ่นเดี่ยวและในเล่มรวม)"""
    from app.models import AcadScore, AcadTeaching
    students = sorted(klass.students, key=lambda s: (s.seq or 999, s.name))
    term = subject.term if subject.term is not None else 0
    scores = {s.acad_student_id: s for s in
              db.query(AcadScore).filter_by(subject_id=subject.id, term=term).all()}
    teach = db.query(AcadTeaching).filter_by(class_id=klass.id, subject_id=subject.id).first()
    teacher = teach.teacher.name if (teach and teach.teacher) else ""

    _p(doc, "แบบบันทึกผลการเรียน", align="center", bold=True, size=16, after=0, page_break=page_break)
    _p(doc, school.name or "", align="center", bold=True, size=16, after=0)
    head = (f"รายวิชา {subject.code or ''} {subject.name}   ชั้น {_class_label(klass)}   "
            f"ปีการศึกษา {klass.year}   {term_label(term)}")
    _p(doc, head, align="center", size=12, after=2)
    meta = []
    if subject.learn_group:
        meta.append(f"กลุ่มสาระการเรียนรู้{subject.learn_group}")
    if subject.hours:
        meta.append(f"เวลาเรียน {subject.hours} ชั่วโมง")
    if subject.credit:
        meta.append(f"{subject.credit:g} หน่วยกิต")
    if meta:
        _p(doc, "  |  ".join(meta), align="center", size=12, after=8)

    mmax = subject.mid_max if (subject.mid_max or 0) > 0 else 70
    fmax = subject.final_max if (subject.final_max or 0) > 0 else 30
    if not is_secondary(klass.level):
        _pp5_score_2term(doc, subject, db, students)          # ประถม: 2 ภาคเรียน + เฉลี่ย + เกรดรายปี
    else:
        _pp5_score_secondary(doc, subject, db, students, mmax, fmax, term)

    _p(doc, "", after=10)
    _sign_block(doc, teacher, "ครูผู้สอน", size=14)


def _vcell(cell, text, *, bold=False, size=11, fill=None):
    """เซลล์ตัวหนังสือแนวตั้ง (อ่านจากล่างขึ้นบน) สำหรับหัวคอลัมน์แคบ"""
    _cell(cell, text, bold=bold, align="center", size=size, fill=fill)
    tcpr = cell._tc.get_or_add_tcPr()
    tcpr.append(tcpr.makeelement(qn("w:textDirection"), {qn("w:val"): "btLr"}))


def _pp5_score_2term(doc, subject, db, students):
    """ตารางคะแนนแบบประถม 2 ภาคเรียน - ก็อปตามชีต 'พิมพ์วิชา' ของไฟล์ Excel จริง (ป.6):
    หัว 5 แถว · มุมซ้ายบน 'วิชา/รหัส:ชื่อ' · ภาค 1/2 (ระหว่างภาค-ปลายภาค-รวม-ผล แนวตั้ง) +
    แถวคะแนนเต็ม 70/30/100 · คะแนนเฉลี่ย2ภาค/ผลตลอดปี (แนวตั้ง) · กล่องคะแนนเฉลี่ยชั้น + สรุปผ่าน/ไม่ผ่าน
    ฟอนต์ TH Sarabun New 14 ทั้งตาราง"""
    from app.models import AcadScore
    mmax = subject.mid_max if (subject.mid_max or 0) > 0 else 70
    fmax = subject.final_max if (subject.final_max or 0) > 0 else 30
    sids = [s.id for s in students]
    rows = {}
    if sids:
        for x in (db.query(AcadScore).filter(AcadScore.subject_id == subject.id,
                                             AcadScore.acad_student_id.in_(sids)).all()):
            rows[(x.acad_student_id, x.term)] = x
    # งานเก็บคะแนนต่อภาค (แยกสอบกลางภาคออกจากงานเก็บ · จำนวนช่อง = เท่าที่มีจริง)
    a1, as1 = _assignments_of(subject, db, 1)
    a2, as2 = _assignments_of(subject, db, 2)
    w1 = [a for a in a1 if not a.is_midterm]; m1 = [a for a in a1 if a.is_midterm]
    w2 = [a for a in a2 if not a.is_midterm]; m2 = [a for a in a2 if a.is_midterm]
    blk1 = len(w1) + 4 + (1 if m1 else 0)     # งาน + ระหว่างภาค[+กลางภาค]ปลายภาค/รวม/ผล
    blk2 = len(w2) + 4 + (1 if m2 else 0)
    b1 = 2                        # เริ่มบล็อกภาค 1
    b2 = b1 + blk1                # เริ่มบล็อกภาค 2
    avg_col = b2 + blk2           # คะแนนเฉลี่ย 2 ภาค
    tot_col = avg_col + 1         # ผลการเรียนตลอดปี
    box_l = tot_col + 1           # กล่องสรุป (ผ่าน)
    box_r = box_l + 1             # ไม่ผ่าน
    ncol = box_r + 1
    t = doc.add_table(rows=5, cols=ncol); t.style = "Table Grid"

    def C(r, c):
        return t.cell(r, c)

    def mg(r1, c1, r2, c2):
        return C(r1, c1).merge(C(r2, c2))

    # มุมซ้ายบน: วิชา + รหัส:ชื่อวิชา (แถว 0-1) · ที่ + ชื่อ (แถว 2-4)
    _cell(mg(0, 0, 1, 0), "วิชา", bold=True, fill="EDE9FE", size=14)
    _cell(mg(0, 1, 1, 1), f"{subject.code or ''} : {subject.name}".strip(" :"),
          bold=True, fill="EDE9FE", size=14, align="left")
    _cell(mg(2, 0, 4, 0), "ที่", bold=True, fill="EDE9FE", size=14)
    _cell(mg(2, 1, 4, 1), "ชื่อ - นามสกุล", bold=True, fill="EDE9FE", size=14)
    # ภาคเรียนที่ 1 / 2 : หัวรวม + งานเก็บ + ระหว่างภาค[+สอบกลางภาค]/ปลายภาค/รวม/ผล + แถวคะแนนเต็ม
    for base, lab, work, mids in [(b1, "ภาคเรียนที่ 1", w1, m1), (b2, "ภาคเรียนที่ 2", w2, m2)]:
        nw, hm = len(work), bool(mids)
        kmf = sum(a.max_score or 0 for a in work)   # เต็มระหว่างภาค = ผลรวมเต็มงานเก็บจริง
        mmf = sum(a.max_score or 0 for a in mids)    # เต็มสอบกลางภาค = ผลรวมเต็มงานกลางภาคจริง
        blk = nw + 4 + (1 if hm else 0)
        _cell(mg(0, base, 0, base + blk - 1), lab, bold=True, fill="EDE9FE", size=14)
        if nw:
            _cell(mg(1, base, 1, base + nw - 1), "การเก็บคะแนนระหว่างภาค", bold=True, fill="F1F5F9", size=12)
            for j, a in enumerate(work):
                _cell(mg(2, base + j, 3, base + j), str(j + 1), bold=True, fill="F1F5F9", size=12)
                _cell(C(4, base + j), f"{a.max_score:g}", bold=True, fill="F8FAFC", size=12)
        sb = base + nw
        _vcell(mg(1, sb, 3, sb), "ระหว่างภาค", bold=True, fill="F1F5F9", size=12)
        _cell(C(4, sb), f"{kmf:g}", bold=True, fill="F8FAFC", size=12)
        col = sb + 1
        if hm:
            _vcell(mg(1, col, 3, col), "สอบกลางภาค", bold=True, fill="F1F5F9", size=12)
            _cell(C(4, col), f"{mmf:g}", bold=True, fill="F8FAFC", size=12); col += 1
        _vcell(mg(1, col, 3, col), "ปลายภาค", bold=True, fill="F1F5F9", size=12)
        _cell(C(4, col), str(fmax), bold=True, fill="F8FAFC", size=12)
        _vcell(mg(1, col + 1, 3, col + 1), "รวม", bold=True, fill="F1F5F9", size=12)
        _cell(C(4, col + 1), f"{kmf + mmf + fmax:g}", bold=True, fill="F8FAFC", size=12)
        _vcell(mg(1, col + 2, 4, col + 2), "ผลการเรียน", bold=True, fill="F1F5F9", size=12)
    # คะแนนเฉลี่ย 2 ภาค / ผลการเรียนตลอดปี (แนวตั้ง ผสาน 5 แถว)
    _vcell(mg(0, avg_col, 4, avg_col), "คะแนนเฉลี่ย 2 ภาคเรียน", bold=True, fill="EDE9FE", size=13)
    _vcell(mg(0, tot_col, 4, tot_col), "ผลการเรียนตลอดปี", bold=True, fill="EDE9FE", size=13)
    # กล่องสรุป: คะแนนเฉลี่ยชั้น + สรุปผลการประเมิน (ผ่าน/ไม่ผ่าน)
    _cell(mg(0, box_l, 0, box_r), "คะแนนเฉลี่ย", bold=True, fill="EDE9FE", size=13)
    avals = [a.score for a in (rows.get((s.id, 0)) for s in students) if a and a.score is not None]
    cavg = sum(avals) / len(avals) if avals else None
    _cell(mg(1, box_l, 1, box_r), f"{cavg:.2f}" if cavg is not None else "", bold=True, fill="F8FAFC", size=13)
    _cell(mg(2, box_l, 3, box_r), "สรุปผลการประเมิน", bold=True, fill="EDE9FE", size=13)
    _cell(C(4, box_l), "ผ่าน", bold=True, fill="F1F5F9", size=13)
    _cell(C(4, box_r), "ไม่ผ่าน", bold=True, fill="F1F5F9", size=13)

    for s in students:
        cells = t.add_row().cells
        _cell(cells[0], s.seq or "", size=13)
        _cell(cells[1], s.name, align="left", size=13)
        for base, tno, work, mids, ascore in [(b1, 1, w1, m1, as1), (b2, 2, w2, m2, as2)]:
            nw, hm = len(work), bool(mids)
            work_vals = [ascore.get((s.id, a.id)) for a in work]
            for j, v in enumerate(work_vals):
                _cell(cells[base + j], f"{v:g}" if v is not None else "", size=13)
            r = rows.get((s.id, tno))
            mid_sum = sum(ascore[(s.id, a.id)] for a in mids if ascore.get((s.id, a.id)) is not None)
            has_work = any(v is not None for v in work_vals)
            work_sum = sum(v for v in work_vals if v is not None)
            keep = work_sum if has_work else ((r.score_mid - mid_sum) if (r and r.score_mid is not None) else None)
            sb = base + nw
            _cell(cells[sb], f"{keep:g}" if keep is not None else "", size=13)
            col = sb + 1
            if hm:
                has = any(ascore.get((s.id, a.id)) is not None for a in mids)
                _cell(cells[col], f"{mid_sum:g}" if has else "", size=13); col += 1
            _cell(cells[col], f"{r.score_final:g}" if r and r.score_final is not None else "", size=13)
            _cell(cells[col + 1], f"{r.score:g}" if r and r.score is not None else "", size=13)
            _cell(cells[col + 2], r.grade if r else "", size=13, bold=True)
        ann = rows.get((s.id, 0))
        _cell(cells[avg_col], f"{ann.score:g}" if ann and ann.score is not None else "", size=13)
        _cell(cells[tot_col], ann.grade if ann else "", size=13, bold=True)
        ok = bool(ann and (ann.grade or "").strip() not in ("", "0", "ร", "มส"))
        has = bool(ann and ann.grade)
        _cell(cells[box_l], "P" if (has and ok) else "", size=13, bold=True)
        _cell(cells[box_r], "P" if (has and not ok) else "", size=13, bold=True)
    # แนวนอน เกือบเต็มหน้า ~26 ซม.
    nsum = (4 + (1 if m1 else 0)) + (4 + (1 if m2 else 0))   # คอลัมน์สรุป 2 ภาค
    nwk = len(w1) + len(w2)
    fixed = 0.9 + 3.5 + 1.05 * nsum + 1.1 + 1.1 + 0.9 + 0.9
    aw = min(1.1, max(0.7, (26.0 - fixed) / nwk)) if nwk else 0
    col_w = [0.9, 3.5]
    for work, mids in [(w1, m1), (w2, m2)]:
        col_w += [aw] * len(work) + [1.05] * (4 + (1 if mids else 0))
    col_w += [1.1, 1.1, 0.9, 0.9]
    _widths(t, [Cm(w) for w in col_w])
    _tight_cells(t)
    if w1 or w2:
        def _leg(assigns):
            return " | ".join(f"{j + 1}. {a.name or 'งานที่ ' + str(j + 1)} (เต็ม {a.max_score:g})"
                              for j, a in enumerate(assigns))
        parts = []
        if w1:
            parts.append("ภาค 1 - " + _leg(w1))
        if w2:
            parts.append("ภาค 2 - " + _leg(w2))
        _p(doc, "งานเก็บคะแนน: " + "   ||   ".join(parts), size=11, after=0, align="left")


def _pp5_subject_cover(doc, school, klass, subject, db, students):
    """หน้าปก ปพ.5 รายวิชา (ตามชีต 'พิมพ์ปกปพ.5' ของไฟล์รายวิชาจริง) - แนวตั้ง ฟอนต์ 14"""
    from app.models import AcadScore, AcadTeaching, AcadCharEval, AcadReadEval
    from app.services.academic import char_avg, read_avg, quality_of_avg, QUALITY_LEVELS
    term = subject.term if subject.term is not None else 0
    sid_set = {s.id for s in students}
    scores = [x for x in db.query(AcadScore).filter_by(subject_id=subject.id, term=term).all()
              if x.acad_student_id in sid_set]
    teach = db.query(AcadTeaching).filter_by(class_id=klass.id, subject_id=subject.id).first()
    teacher = teach.teacher.name if (teach and teach.teacher) else ""

    _logo_header(doc, school, height_cm=1.6, after=2)
    _p(doc, "สมุดบันทึกการพัฒนาคุณภาพผู้เรียน (ปพ.5)", align="center", bold=True, size=18, after=4)
    snm = (school.name or "").strip()
    loc = [snm if snm.startswith("โรงเรียน") else f"โรงเรียน{snm}"]
    if (school.district or "").strip():
        loc.append(f"อำเภอ{school.district.strip()}")
    if (school.province or "").strip():
        loc.append(f"จังหวัด{school.province.strip()}")
    _p(doc, "   ".join(loc), align="center", size=14, after=0)
    if (school.area_office or "").strip():
        _ao = school.area_office.strip()
        _p(doc, _ao if _ao.startswith(("สำนักงาน", "สพ")) else f"สำนักงานเขตพื้นที่การศึกษา{_ao}",
           align="center", size=14, after=4)
    _p(doc, f"ปีการศึกษา {klass.year}          ชั้น {_class_label(klass)}", align="center", bold=True, size=14, after=0)
    if subject.learn_group:
        _p(doc, f"กลุ่มสาระการเรียนรู้{subject.learn_group}", align="center", size=14, after=0)
    _p(doc, f"รหัสวิชา {subject.code or ''}          รายวิชา {subject.name}", align="center", size=14, after=0)
    meta = []
    if subject.hours:
        meta.append(f"เวลาเรียน {subject.hours} ชั่วโมง")
    if subject.credit:
        meta.append(f"{subject.credit:g} หน่วยกิต")
    if meta:
        _p(doc, "          ".join(meta), align="center", size=14, after=0)
    _p(doc, f"ครูผู้สอน {teacher}", align="center", size=14, after=0)
    _p(doc, f"ครูประจำชั้น {klass.homeroom.name if klass.homeroom else ''}", align="center", size=14, after=6)

    _p(doc, "การอนุมัติผลการเรียน", align="center", bold=True, size=14, after=2)
    grade_cols = ["4", "3.5", "3", "2.5", "2", "1.5", "1", "0", "ร", "มส"]
    gmap = {x.acad_student_id: (x.grade or "") for x in scores}
    cnt = _grade_counts(gmap.get(s.id, "") for s in students)
    svals = [x.score for x in scores if x.score is not None]
    savg = sum(svals) / len(svals) if svals else None
    gt = doc.add_table(rows=2, cols=1 + len(grade_cols)); gt.style = "Table Grid"
    _cell(gt.rows[0].cells[0], "ระดับผลการเรียน", bold=True, fill="EDE9FE", size=14)
    for i, g in enumerate(grade_cols):
        _cell(gt.rows[0].cells[1 + i], g, bold=True, fill="EDE9FE", size=14)
    _cell(gt.rows[1].cells[0], "จำนวน (คน)", bold=True, size=14)
    for i, g in enumerate(grade_cols):
        _cell(gt.rows[1].cells[1 + i], cnt[g] or "", size=14)
    _widths(gt, [Cm(4.4)] + [Cm(1.36)] * len(grade_cols))
    _p(doc, f"นักเรียนทั้งหมด {len(students)} คน"
       + (f"          คะแนนเฉลี่ยรายวิชา {savg:.2f}" if savg is not None else ""),
       align="center", size=14, after=4)

    ce, rd = {}, {}      # รวมทุกภาคเรียน (ประถมประเมินแยกภาค) -> เฉลี่ยเป็นผลรวมทั้งปี
    for r in db.query(AcadCharEval).filter_by(subject_id=subject.id).all():
        ce.setdefault(r.acad_student_id, []).append(r)
    for r in db.query(AcadReadEval).filter_by(subject_id=subject.id).all():
        rd.setdefault(r.acad_student_id, []).append(r)
    ccnt = {q: 0 for q in QUALITY_LEVELS}
    rcnt = {q: 0 for q in QUALITY_LEVELS}

    def _agg(rows, avg_fn):
        if not rows:
            return None
        vals = [v for v in (avg_fn(r) for r in rows) if v is not None]
        return (sum(vals) / len(vals)) if vals else None

    for s in students:
        ca = _agg(ce.get(s.id), char_avg)
        if ca is not None:
            ccnt[quality_of_avg(ca)[1]] += 1
        ra = _agg(rd.get(s.id), read_avg)
        if ra is not None:
            rcnt[quality_of_avg(ra)[1]] += 1
    qt = doc.add_table(rows=3, cols=1 + len(QUALITY_LEVELS)); qt.style = "Table Grid"
    _cell(qt.rows[0].cells[0], "ผลการประเมิน", bold=True, fill="EDE9FE", size=14)
    for i, q in enumerate(QUALITY_LEVELS):
        _cell(qt.rows[0].cells[1 + i], q, bold=True, fill="EDE9FE", size=14)
    _cell(qt.rows[1].cells[0], "คุณลักษณะฯ (คน)", align="left", size=14)
    _cell(qt.rows[2].cells[0], "อ่าน คิดวิเคราะห์ เขียน (คน)", align="left", size=14)
    for i, q in enumerate(QUALITY_LEVELS):
        _cell(qt.rows[1].cells[1 + i], ccnt[q] or "", size=14)
        _cell(qt.rows[2].cells[1 + i], rcnt[q] or "", size=14)
    _widths(qt, [Cm(6.0)] + [Cm(3.0)] * len(QUALITY_LEVELS))
    _p(doc, "", after=4)

    st = doc.add_table(rows=1, cols=2)
    for cell, (nm, role) in zip(st.rows[0].cells,
                                [(teacher, "ครูผู้สอน"), ("", "หัวหน้ากลุ่มสาระการเรียนรู้")]):
        for i, txt in enumerate(["(ลงชื่อ)..........................................",
                                 f"( {nm or '.....................................'} )", role]):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(txt); r.font.size = Pt(14); r.font.name = THAI_FONT
            r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
            if i == 0 and nm:
                _float_signature(p, nm)
    _widths(st, [Cm(9.0), Cm(9.0)])
    _p(doc, "", after=2)
    _p(doc, "เรียนเสนอเพื่อพิจารณา", align="center", size=14, after=2)
    _sign_block(doc, (getattr(school, "academic_head_name", "") or "").strip(),
                "หัวหน้า/รองผู้อำนวยการ ฝ่ายวิชาการ", size=14, after=6)
    _p(doc, "[   ] อนุมัติ          [   ] ไม่อนุมัติ", align="center", size=14, after=4)
    director = (getattr(school, "director_name", "") or "").strip()
    dpos = ("ผู้อำนวยการ" + school.name) if (school.name or "").startswith("โรงเรียน") \
        else "ผู้อำนวยการโรงเรียน"
    _sign_block(doc, director, dpos, size=14)
    _p(doc, "วันที่ …........../…..................../…...........", align="center", size=14, after=0)


def _pp5_subject_summary(doc, school, klass, subject, db, students, *, page_break=True):
    """หน้าสรุปการประเมินผลการพัฒนาคุณภาพผู้เรียน (รายวิชา) - ตามชีต 'พิมพ์สรุปผลพัฒนาผู้เรียน'
    ที่/ชื่อ/คะแนนเก็บ/ปลายภาค/ประเมินผลรายปี(คะแนน+ผล)/ตัวชี้วัดที่ผ่าน/คุณลักษณะ/อ่านคิดเขียน · แนวตั้ง"""
    from app.models import AcadScore, AcadIndicatorResult, AcadCharEval, AcadReadEval
    from app.services.academic import char_avg, read_avg, quality_of_avg
    from app.services.curriculum import selected_indicators
    term = subject.term if subject.term is not None else 0
    mmax = subject.mid_max if (subject.mid_max or 0) > 0 else 70
    fmax = subject.final_max if (subject.final_max or 0) > 0 else 30
    sids = [s.id for s in students]
    scores = {x.acad_student_id: x for x in
              db.query(AcadScore).filter_by(subject_id=subject.id, term=term).all()}
    n_ind = len(selected_indicators(subject))
    ind_pass = {}
    if sids and n_ind:
        for r in (db.query(AcadIndicatorResult)
                  .filter(AcadIndicatorResult.subject_id == subject.id,
                          AcadIndicatorResult.acad_student_id.in_(sids)).all()):
            if r.passed:
                ind_pass[r.acad_student_id] = ind_pass.get(r.acad_student_id, 0) + 1
    ce, rd = {}, {}
    for r in db.query(AcadCharEval).filter_by(subject_id=subject.id).all():
        ce.setdefault(r.acad_student_id, []).append(r)
    for r in db.query(AcadReadEval).filter_by(subject_id=subject.id).all():
        rd.setdefault(r.acad_student_id, []).append(r)

    def qnum(rows, avg_fn):
        if not rows:
            return ""
        vals = [v for v in (avg_fn(r) for r in rows) if v is not None]
        return str(quality_of_avg(sum(vals) / len(vals))[0]) if vals else ""

    _p(doc, "สรุปการประเมินผลการพัฒนาคุณภาพผู้เรียน (รายวิชา)",
       align="center", bold=True, size=16, after=0, page_break=page_break)
    _p(doc, f"รายวิชา {subject.code or ''} {subject.name}   ชั้น {_class_label(klass)}   "
            f"ปีการศึกษา {klass.year}", align="center", size=12, after=6)
    t = doc.add_table(rows=2, cols=9); t.style = "Table Grid"
    r0, r1 = t.rows[0].cells, t.rows[1].cells

    def vm(idx, txt):
        c = r0[idx].merge(r1[idx]); _cell(c, txt, bold=True, fill="EDE9FE", size=12)

    vm(0, "ที่"); vm(1, "ชื่อ-นามสกุล")
    vm(2, f"คะแนนเก็บ\n(เต็ม {mmax:g})"); vm(3, f"คะแนนปลายภาค\n(เต็ม {fmax:g})")
    c = r0[4].merge(r0[5]); _cell(c, "ประเมินผลรายปี", bold=True, fill="EDE9FE", size=12)
    _cell(r1[4], f"คะแนน\n(เต็ม {mmax + fmax:g})", bold=True, fill="F1F5F9", size=11)
    _cell(r1[5], "ผลการเรียน", bold=True, fill="F1F5F9", size=11)
    vm(6, f"ตัวชี้วัด\nที่ผ่าน/{n_ind}" if n_ind else "ตัวชี้วัด")
    vm(7, "ผลประเมิน\nคุณลักษณะ"); vm(8, "ผลประเมิน\nอ่านคิดเขียน")
    for s in students:
        sc = scores.get(s.id)
        cells = t.add_row().cells
        _cell(cells[0], s.seq or "", size=14)
        _cell(cells[1], s.name, align="left", size=14)
        _cell(cells[2], f"{sc.score_mid:g}" if sc and sc.score_mid is not None else "", size=14)
        _cell(cells[3], f"{sc.score_final:g}" if sc and sc.score_final is not None else "", size=14)
        _cell(cells[4], f"{sc.score:g}" if sc and sc.score is not None else "", size=14)
        _cell(cells[5], sc.grade if sc else "", bold=True, size=14)
        _cell(cells[6], str(ind_pass.get(s.id, "")) if n_ind else "", size=14)
        _cell(cells[7], qnum(ce.get(s.id), char_avg), size=14)
        _cell(cells[8], qnum(rd.get(s.id), read_avg), size=14)
    _widths(t, [Cm(1.0), Cm(4.2), Cm(2.1), Cm(2.3), Cm(1.7), Cm(1.5), Cm(1.7), Cm(1.7), Cm(1.7)])
    _tight_cells(t)
    _p(doc, "ตัวชี้วัด = จำนวนที่ผ่าน · คุณลักษณะ/อ่านคิดเขียน: 3 ดีเยี่ยม · 2 ดี · 1 ผ่าน · 0 ไม่ผ่าน",
       size=11, after=0, align="center")


def render_pp5(school, klass, subject, db) -> str:
    """ปพ.5 รายวิชา (เล่มของครูรายวิชา): ปก -> รายชื่อ -> เวลาเรียน -> ตัวชี้วัด
    -> คะแนน -> คุณลักษณะ -> อ่านคิดเขียน -> สรุปผลพัฒนา -> เกณฑ์ (สลับแนวตามความกว้างตาราง)"""
    from app.services.curriculum import selected_indicators as _sel
    doc = _doc(landscape=False)          # ปกแนวตั้ง
    students = sorted(klass.students, key=lambda s: (s.seq or 999, s.name))
    _pp5_subject_cover(doc, school, klass, subject, db, students)
    _new_section(doc, landscape=True)
    _pp5_roster(doc, klass, students, db)
    _new_section(doc, landscape=False)
    att_months = TERM_MONTHS.get(subject.term) if is_secondary(klass.level) else None
    att_sub = subject.id if getattr(school, "attendance_by_subject", False) else None
    _pp5_attendance_daily(doc, klass, students, db, page_break=False, months_ok=att_months,
                          subject_id=att_sub)
    if _sel(subject):
        _new_section(doc, landscape=True)
        _pp5_indicator_page(doc, klass, subject, db, students, page_break=False)
    _new_section(doc, landscape=True)
    _pp5_score_page(doc, school, klass, subject, db, page_break=False)
    # คุณลักษณะ/อ่านคิดเขียน · ประถมประเมินแยกภาคเรียน (พิมพ์ทั้ง 2 ภาค) · มัธยม = ภาคเดียว
    eval_terms = (1, 2) if not is_secondary(klass.level) else (0,)
    for et in eval_terms:
        _pp5_char_page(doc, klass, subject, db, students, page_break=True, term=et)
    for et in eval_terms:
        _pp5_read_page(doc, klass, subject, db, students, page_break=True, term=et)
    # สรุปผลพัฒนาผู้เรียน (แนวตั้ง) -> เกณฑ์/ปกหลัง (แนวตั้ง)
    _new_section(doc, landscape=False)
    _pp5_subject_summary(doc, school, klass, subject, db, students, page_break=False)
    _pp5_criteria_page(doc, page_break=True)
    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    out = out_dir / (_safe(f"ปพ.5_{subject.name}_{_class_label(klass)}_{klass.year}") + ".docx")
    doc.save(str(out))
    return str(out)


# ---------------- ปพ.5 ทั้งเล่ม ----------------
def _grade_counts(grades):
    """นับจำนวนคนต่อระดับผลการเรียน -> dict ตามคอลัมน์ปก"""
    cols = ["4", "3.5", "3", "2.5", "2", "1.5", "1", "0", "ร", "มส"]
    n = {c: 0 for c in cols}
    for g in grades:
        g = (str(g) if g is not None else "").strip()
        if g in n:
            n[g] += 1
    return n


def _student_age(birth) -> str:
    """อายุเต็มปี ณ วันนี้ จากวันเกิด (คืน '' ถ้าไม่มี)"""
    if not birth:
        return ""
    from datetime import date
    t = date.today()
    y = t.year - birth.year - ((t.month, t.day) < (birth.month, birth.day))
    return str(y) if y >= 0 else ""


def _pp5_roster(doc, klass, students, db):
    """หน้ารายชื่อนักเรียน + ข้อมูลบิดา/มารดา/ผู้ปกครอง/ที่อยู่ (ดึงจากทะเบียนกลาง)"""
    from app.models import Student
    _p(doc, f"รายชื่อและข้อมูลนักเรียน ชั้น {_class_label(klass)} ปีการศึกษา {klass.year}",
       align="center", bold=True, size=16, after=6)
    ids = [s.student_id for s in students if s.student_id]
    cmap = {}
    if ids:
        for st in db.query(Student).filter(Student.id.in_(ids)).all():
            cmap[st.id] = st
    heads = ["เลขที่", "ชื่อ-นามสกุล", "วัน/เดือน/ปี เกิด", "อายุ", "ชื่อบิดา", "อาชีพ",
             "ชื่อมารดา", "อาชีพ", "ผู้ปกครอง", "เกี่ยวข้อง", "อาชีพ", "ที่อยู่ปัจจุบัน"]
    t = doc.add_table(rows=1, cols=len(heads)); t.style = "Table Grid"
    for i, h in enumerate(heads):
        _cell(t.rows[0].cells[i], h, bold=True, fill="EDE9FE", size=10)
    for s in students:
        st = cmap.get(s.student_id)

        def g(a):
            return (getattr(st, a, "") or "") if st else ""

        cells = t.add_row().cells
        _cell(cells[0], s.seq or "", size=10)
        _cell(cells[1], s.name, align="left", size=10)
        _cell(cells[2], thai_date(st.birthdate) if (st and st.birthdate) else "", size=10)
        _cell(cells[3], _student_age(st.birthdate) if st else "", size=10)
        _cell(cells[4], g("father_name"), align="left", size=10)
        _cell(cells[5], g("father_job"), size=10)
        _cell(cells[6], g("mother_name"), align="left", size=10)
        _cell(cells[7], g("mother_job"), size=10)
        _cell(cells[8], g("guardian_name"), align="left", size=10)
        _cell(cells[9], g("guardian_relation"), size=10)
        _cell(cells[10], g("guardian_job"), size=10)
        _cell(cells[11], _fmt_addr(st) if st else "", align="left", size=10)
    _widths(t, [Cm(1.0), Cm(3.5), Cm(2.9), Cm(1.0), Cm(3.0), Cm(1.8),
                Cm(3.0), Cm(1.8), Cm(2.6), Cm(1.5), Cm(1.6), Cm(3.0)])   # รวม 26.7


def _pp5_indicator_page(doc, klass, subject, db, students, *, page_break=True):
    """หน้าผลการประเมินตัวชี้วัดรายวิชา (สรุปผ่าน/ทั้งหมด รายมาตรฐาน) + รายการตัวชี้วัด
    ข้ามถ้ายังไม่มีชุดตัวชี้วัดของกลุ่มสาระ+ชั้นนั้น · คืน True ถ้าออกหน้าจริง (ใช้จัดหน้าแรกของ section)"""
    from app.services.curriculum import selected_indicators
    from app.models import AcadIndicatorResult
    inds = selected_indicators(subject)
    if not inds:
        return False
    # จัดกลุ่มตามมาตรฐาน (คงลำดับ)
    stds = []
    for it in inds:
        if not stds or stds[-1][0] != it["std"]:
            stds.append((it["std"], []))
        stds[-1][1].append(it)
    # คะแนน 0-3 ต่อตัวชี้วัด (เมทริกซ์ นักเรียน x ตัวชี้วัด ตามไฟล์จริง)
    scmap = {}
    sids = [s.id for s in students]
    if sids:
        for r in (db.query(AcadIndicatorResult)
                  .filter(AcadIndicatorResult.subject_id == subject.id,
                          AcadIndicatorResult.acad_student_id.in_(sids)).all()):
            sc = r.score if r.score is not None else (3 if r.passed else None)
            if sc is not None:
                scmap[(r.acad_student_id, r.code)] = sc

    n = len(inds)
    _p(doc, "ผลการประเมินตัวชี้วัดรายวิชา",
       align="center", bold=True, size=16, after=0, page_break=page_break)
    _p(doc, f"รายวิชา {subject.code or ''} {subject.name}   ชั้น {_class_label(klass)}   "
            f"ปีการศึกษา {klass.year}   ({n} ตัวชี้วัด)", align="center", size=12, after=4)

    # หัว 2 แถว: มาตรฐาน (ผสานตามกลุ่ม) + เลขข้อ (1..n)
    t = doc.add_table(rows=2, cols=2 + n + 2); t.style = "Table Grid"
    h0, h1 = t.rows[0].cells, t.rows[1].cells
    h0[0].merge(h1[0]); _cell(h0[0], "ที่", bold=True, fill="EDE9FE", size=14)
    h0[1].merge(h1[1]); _cell(h0[1], "ชื่อ-นามสกุล", bold=True, fill="EDE9FE", size=14)
    col = 2
    for st, items in stds:
        first = col
        for _ in items:
            col += 1
        merged = h0[first]
        for c in range(first + 1, col):
            merged = merged.merge(h0[c])
        _cell(merged, st, bold=True, fill="EDE9FE", size=14)
    for j, it in enumerate(inds):
        _cell(h1[2 + j], str(it["seq"]), bold=True, fill="F1F5F9", size=14)   # เลขข้อ = ลำดับภายในมาตรฐาน
    h0[2 + n].merge(h1[2 + n]); _cell(h0[2 + n], f"ผ่าน/{n}", bold=True, fill="EDE9FE", size=14)
    h0[3 + n].merge(h1[3 + n]); _cell(h0[3 + n], "ผล", bold=True, fill="EDE9FE", size=14)

    for s in students:
        cells = t.add_row().cells
        _cell(cells[0], s.seq or "", size=14)
        _cell(cells[1], s.name, align="left", size=14)
        npass, nfill = 0, 0
        for j, it in enumerate(inds):
            sc = scmap.get((s.id, it["code"]))
            _cell(cells[2 + j], sc if sc is not None else "", size=14)
            if sc is not None:
                nfill += 1
                if sc >= 1:
                    npass += 1
        _cell(cells[2 + n], str(npass), size=14, bold=True)
        _cell(cells[3 + n], ("ผ่าน" if npass == n else "ไม่ผ่าน") if nfill else "", size=14, bold=True)
    iw = min(0.62, 18.5 / max(1, n))
    _widths(t, [Cm(0.9), Cm(4.4)] + [Cm(iw)] * n + [Cm(1.4), Cm(1.5)])
    _tight_cells(t)      # ลดขอบเซลล์ กันเลขในคอลัมน์ตัวชี้วัดที่แคบโดนบีบ
    _p(doc, "คะแนน 0-3 ต่อตัวชี้วัด (3 ดีเยี่ยม · 2 ดี · 1 ผ่าน · 0 ไม่ผ่าน) | ผ่านตัวชี้วัด = ได้ตั้งแต่ 1 | "
            "เกณฑ์: ต้องผ่านครบทุกตัวชี้วัด", size=10, after=4, align="center")
    # รายการตัวชี้วัดอ้างอิง แยกตามมาตรฐาน (เลขข้อเริ่มใหม่ทุกมาตรฐาน)
    _p(doc, "ตัวชี้วัด:", bold=True, size=10, after=1)
    for st, items in stds:
        _p(doc, f"มาตรฐาน {st}", bold=True, size=10, after=0)
        for it in items:
            _p(doc, f"   {it['seq']}. {it['text']}", size=10, after=0)
    return True


def _pp5_attendance_daily(doc, klass, students, db, *, page_break=True, months_ok=None,
                          subject_id=None):
    """หน้าบันทึกเวลาเรียน (รายวัน) - ตารางเช็กชื่อต่อเดือน (เฉพาะเดือนที่มีปฏิทิน)
    months_ok = set เดือนที่อนุญาต (มัธยม = เฉพาะเดือนในภาคเรียนนั้น) · None = ทุกเดือน
    subject_id = ค่า = เวลาเรียนแยกรายวิชานั้น · None = รายห้อง (subject_id IS NULL)"""
    from app.models import AcadCalendar, AcadAttendance
    sids = [s.id for s in students]
    att = {}
    if sids:
        cond = (AcadAttendance.subject_id == subject_id) if subject_id \
            else AcadAttendance.subject_id.is_(None)
        for a in (db.query(AcadAttendance)
                  .filter(AcadAttendance.acad_student_id.in_(sids), cond).all()):
            att[(a.acad_student_id, a.month)] = parse_marks(a.marks)
    cals = {r.month: parse_days_csv(r.days_csv)
            for r in db.query(AcadCalendar).filter_by(year=klass.year).all()}
    months = [(m, nm) for m, nm in TH_MONTHS
              if cals.get(m) and (months_ok is None or m in months_ok)]
    if not months:
        return
    _p(doc, f"บันทึกเวลาเรียน (รายวัน) ชั้น {_class_label(klass)} ปีการศึกษา {klass.year}",
       align="center", bold=True, size=16, after=2, page_break=page_break)
    _p(doc, "/ = มา · ป = ป่วย · ล = ลากิจ · ข = ขาด", size=14, after=6, align="center")
    for m, nm in months:
        days = cals[m]
        _p(doc, f"เดือน{TH_MONTH_FULL.get(m, nm)}", bold=True, size=16, after=2)
        t = doc.add_table(rows=1, cols=2 + len(days) + 1); t.style = "Table Grid"
        _cell(t.rows[0].cells[0], "เลขที่", bold=True, fill="EDE9FE", size=14)
        _cell(t.rows[0].cells[1], "ชื่อ-นามสกุล", bold=True, fill="EDE9FE", size=14)
        # หัวเลขวันที่คุมขนาดตามความกว้างคอลัมน์ (กัน 2 หลักตกบรรทัด) แต่ไม่เกิน 14
        dayw = max(0.42, min(1.0, (18.0 - 1.44 - 5.56 - 1.5) / max(1, len(days))))
        dhead = 14 if dayw >= 0.72 else (13 if dayw >= 0.6 else 12)
        for j, d in enumerate(days):
            _cell(t.rows[0].cells[2 + j], str(d), bold=True, fill="EDE9FE", size=dhead)
        _cell(t.rows[0].cells[-1], "มา", bold=True, fill="EDE9FE", size=14)
        for s in students:
            marks = att.get((s.id, m), {})
            cells = t.add_row().cells
            _cell(cells[0], s.seq or "", size=14)
            _cell(cells[1], s.name, align="left", size=14)
            npres = 0
            for j, d in enumerate(days):
                ch = marks.get(d, "")
                _cell(cells[2 + j], ch, size=14)
                if ch == "/":
                    npres += 1
            _cell(cells[-1], npres, size=14, bold=True)
        # ตามไฟล์จริง: เลขที่ 1.44 · ชื่อ 5.56 · มา 1.5 · วันที่เหลือกระจายพอดี (แนวตั้ง 18 ซม.)
        _widths(t, [Cm(1.44), Cm(5.56)] + [Cm(dayw)] * len(days) + [Cm(1.5)])
        _tight_cells(t)      # ลดขอบเซลล์ กันเลขวันที่โดนบีบ (ตามที่ขอ -0.3)
        _p(doc, "", after=4)
    return True


def _pp5_char_page(doc, klass, subject, db, students, *, page_break=True, term=0):
    """หน้าประเมินคุณลักษณะอันพึงประสงค์ 8 ข้อ ของ 1 รายวิชา (ประถมแยกภาคเรียน)"""
    from app.models import AcadCharEval
    tsuf = f" (ภาคเรียนที่ {term})" if term in (1, 2) else ""
    _p(doc, "ผลการประเมินคุณลักษณะอันพึงประสงค์ (รายวิชา)" + tsuf,
       align="center", bold=True, size=16, after=0, page_break=page_break)
    _p(doc, f"รายวิชา {subject.code or ''} {subject.name}   ชั้น {_class_label(klass)}   "
            f"ปีการศึกษา {klass.year}", align="center", size=14, after=6)

    chars = {r.acad_student_id: r for r in
             db.query(AcadCharEval).filter_by(subject_id=subject.id, term=term).all()}
    heads = ["เลขที่", "ชื่อ-นามสกุล"] + [f"ข้อ {i}" for i in range(1, 9)] + ["เฉลี่ย", "ผล"]
    t = doc.add_table(rows=1, cols=len(heads)); t.style = "Table Grid"
    for i, h in enumerate(heads):
        _cell(t.rows[0].cells[i], h, bold=True, fill="EDE9FE", size=14)
    for s in students:
        r = chars.get(s.id)
        cells = t.add_row().cells
        _cell(cells[0], s.seq or "", size=14)
        _cell(cells[1], s.name, align="left", size=14)
        for j, f in enumerate(CHAR_FIELDS):
            v = getattr(r, f) if r else None
            _cell(cells[2 + j], v if v is not None else "", size=14)
        avg = char_avg(r) if r else None
        _cell(cells[10], f"{avg:.2f}" if avg is not None else "", size=14, bold=True)
        _cell(cells[11], quality_of_avg(avg)[1], size=14, bold=True)
    _widths(t, [Cm(1.2), Cm(5.6)] + [Cm(1.55)] * 8 + [Cm(1.6), Cm(2.4)])   # รวม 23.2
    _tight_cells(t)
    _p(doc, "คุณลักษณะฯ: " + " | ".join(f"ข้อ {i} {nm}" for i, nm in enumerate(CHAR_ITEMS, 1)),
       size=10, after=2, align="center")
    _p(doc, "คะแนน 0-3 ต่อข้อ | เฉลี่ย ≥2.5 ดีเยี่ยม | 1.5-2.49 ดี | 1-1.49 ผ่าน | ต่ำกว่า 1 ไม่ผ่าน "
            "| ช่องว่าง = ยังไม่ประเมิน", size=10, after=0, align="center")


def _pp5_read_page(doc, klass, subject, db, students, *, page_break=True, term=0):
    """หน้าประเมินการอ่าน คิดวิเคราะห์ และเขียน 3 ด้าน ของ 1 รายวิชา (ประถมแยกภาคเรียน)"""
    from app.models import AcadReadEval
    tsuf = f" (ภาคเรียนที่ {term})" if term in (1, 2) else ""
    _p(doc, "ผลการประเมินการอ่าน คิดวิเคราะห์ และเขียน (รายวิชา)" + tsuf,
       align="center", bold=True, size=16, after=0, page_break=page_break)
    _p(doc, f"รายวิชา {subject.code or ''} {subject.name}   ชั้น {_class_label(klass)}   "
            f"ปีการศึกษา {klass.year}", align="center", size=14, after=6)

    reads = {r.acad_student_id: r for r in
             db.query(AcadReadEval).filter_by(subject_id=subject.id, term=term).all()}
    heads2 = ["เลขที่", "ชื่อ-นามสกุล"] + [lb for _, lb in READ_DOMAINS] + ["เฉลี่ย", "ผล"]
    t2 = doc.add_table(rows=1, cols=len(heads2)); t2.style = "Table Grid"
    for i, h in enumerate(heads2):
        _cell(t2.rows[0].cells[i], h, bold=True, fill="EDE9FE", size=14)
    for s in students:
        r = reads.get(s.id)
        cells = t2.add_row().cells
        _cell(cells[0], s.seq or "", size=14)
        _cell(cells[1], s.name, align="left", size=14)
        for j, (f, _lb) in enumerate(READ_DOMAINS):
            v = getattr(r, f) if r else None
            _cell(cells[2 + j], v if v is not None else "", size=14)
        avg = read_avg(r) if r else None
        _cell(cells[5], f"{avg:.2f}" if avg is not None else "", size=14, bold=True)
        _cell(cells[6], quality_of_avg(avg)[1], size=14, bold=True)
    _widths(t2, [Cm(1.2), Cm(5.6), Cm(3.4), Cm(3.4), Cm(3.4), Cm(1.6), Cm(2.4)])   # รวม 21.0
    _tight_cells(t2)
    _p(doc, "การอ่าน คิดวิเคราะห์ และเขียน: " + " | ".join(lb for _, lb in READ_DOMAINS),
       size=10, after=2, align="center")
    _p(doc, "คะแนน 0-3 ต่อด้าน | เฉลี่ย ≥2.5 ดีเยี่ยม | 1.5-2.49 ดี | 1-1.49 ผ่าน | ต่ำกว่า 1 ไม่ผ่าน "
            "| ช่องว่าง = ยังไม่ประเมิน", size=10, after=0, align="center")


def _pp5_quality_summary(doc, klass, subjects, students, db, kind, title):
    """สรุปทั้งปี: นักเรียน x วิชา (ผลรายวิชาเป็นเลข 0-3) -> เฉลี่ย -> ผลสุดท้าย
    (ตามชีต 'พิมพ์สรุปคุณลักษณะทั้งปี' / 'พิมพ์สรุปอ่านคิดเขียนทั้งปี' ของไฟล์จริง)"""
    from app.models import AcadCharEval, AcadReadEval
    Model = AcadCharEval if kind == "char" else AcadReadEval
    avg_fn = char_avg if kind == "char" else read_avg
    _p(doc, f"{title} ชั้น {_class_label(klass)} ปีการศึกษา {klass.year}",
       align="center", bold=True, size=16, after=6, page_break=True)
    if not subjects:
        return
    rows_by = {}      # (student, subject, term) -> row
    for r in (db.query(Model)
              .filter(Model.subject_id.in_([x.id for x in subjects])).all()):
        rows_by[(r.acad_student_id, r.subject_id, r.term or 0)] = r
    nsub = len(subjects)
    two_term = not is_secondary(klass.level)      # ประถม = ประเมินแยก 2 ภาคเรียน

    if two_term:
        ncol = 2 + nsub * 2 + 2
        t = doc.add_table(rows=2, cols=ncol); t.style = "Table Grid"
        r0, r1 = t.rows[0].cells, t.rows[1].cells
        r0[0].merge(r1[0]); _cell(r0[0], "เลขที่", bold=True, fill="EDE9FE", size=11)
        r0[1].merge(r1[1]); _cell(r0[1], "ชื่อ-นามสกุล", bold=True, fill="EDE9FE", size=11)
        for i, sub in enumerate(subjects):
            b = 2 + i * 2
            r0[b].merge(r0[b + 1]); _cell(r0[b], sub.code or sub.name[:6], bold=True, fill="EDE9FE", size=10)
            _cell(r1[b], "ภ.1", bold=True, fill="F1F5F9", size=9)
            _cell(r1[b + 1], "ภ.2", bold=True, fill="F1F5F9", size=9)
        avgc, resc = 2 + nsub * 2, 2 + nsub * 2 + 1
        r0[avgc].merge(r1[avgc]); _cell(r0[avgc], "เฉลี่ย", bold=True, fill="EDE9FE", size=11)
        r0[resc].merge(r1[resc]); _cell(r0[resc], "ผล", bold=True, fill="EDE9FE", size=11)
        for s in students:
            cells = t.add_row().cells
            _cell(cells[0], s.seq or "", size=11)
            _cell(cells[1], s.name, align="left", size=11)
            nums = []
            for i, sub in enumerate(subjects):
                b = 2 + i * 2
                for k, tno in enumerate((1, 2)):
                    r = rows_by.get((s.id, sub.id, tno))
                    n = quality_of_avg(avg_fn(r))[0] if r else ""
                    _cell(cells[b + k], n if n != "" else "", size=11)
                    if n != "":
                        nums.append(n)
            avg = (sum(nums) / len(nums)) if nums else None
            _cell(cells[avgc], f"{avg:.2f}" if avg is not None else "", size=11, bold=True)
            _cell(cells[resc], quality_of_avg(avg)[1], size=11, bold=True)
        sw = min(1.5, max(0.6, (26.0 - 1.2 - 4.5 - 1.6 - 2.0) / max(1, nsub * 2)))
        _widths(t, [Cm(1.2), Cm(4.5)] + [Cm(sw)] * (nsub * 2) + [Cm(1.6), Cm(2.0)])
        _tight_cells(t)
        _p(doc, "ตัวเลข = ผลรายวิชาต่อภาค (3 ดีเยี่ยม | 2 ดี | 1 ผ่าน | 0 ไม่ผ่าน) · ภ.1/ภ.2 = ภาคเรียนที่ 1/2 | "
                "เฉลี่ยรวมทุกวิชาทุกภาค", size=11, after=0, align="center")
    else:
        sw = min(2.2, 15.5 / nsub)
        t = doc.add_table(rows=1, cols=2 + nsub + 2); t.style = "Table Grid"
        _cell(t.rows[0].cells[0], "เลขที่", bold=True, fill="EDE9FE", size=11)
        _cell(t.rows[0].cells[1], "ชื่อ-นามสกุล", bold=True, fill="EDE9FE", size=11)
        for i, sub in enumerate(subjects):
            _cell(t.rows[0].cells[2 + i], sub.code or sub.name[:6], bold=True, fill="EDE9FE", size=10)
        _cell(t.rows[0].cells[-2], "เฉลี่ย", bold=True, fill="EDE9FE", size=11)
        _cell(t.rows[0].cells[-1], "ผล", bold=True, fill="EDE9FE", size=11)
        for s in students:
            cells = t.add_row().cells
            _cell(cells[0], s.seq or "", size=11)
            _cell(cells[1], s.name, align="left", size=11)
            nums = []
            for j, sub in enumerate(subjects):
                r = rows_by.get((s.id, sub.id, 0))
                n = quality_of_avg(avg_fn(r))[0] if r else ""
                _cell(cells[2 + j], n if n != "" else "", size=11)
                if n != "":
                    nums.append(n)
            avg = (sum(nums) / len(nums)) if nums else None
            _cell(cells[-2], f"{avg:.2f}" if avg is not None else "", size=11, bold=True)
            _cell(cells[-1], quality_of_avg(avg)[1], size=11, bold=True)
        _widths(t, [Cm(1.2), Cm(6.0)] + [Cm(sw)] * nsub + [Cm(1.6), Cm(2.2)])
        _p(doc, "ตัวเลขในตาราง = ผลรายวิชา (3 ดีเยี่ยม | 2 ดี | 1 ผ่าน | 0 ไม่ผ่าน) | "
                "ผลสุดท้ายมาจากเฉลี่ยข้ามวิชาด้วยเกณฑ์เดียวกัน", size=11, after=0, align="center")


def render_pp5_book(school, klass, db, term: int | None = None) -> str:
    """ปพ.5 ทั้งเล่ม: ปก (แนวตั้ง · ลายเซ็นครบทุกคนอยู่หน้าแรก)
    -> รายชื่อ -> สรุปเวลาเรียน -> คะแนนรายวิชา (วิชาละหน้า)
    -> สรุปผลทุกวิชา -> สรุปผลการประเมินทั้งปี  (ส่วนนี้แนวนอน)
    มัธยม: เล่มรายภาค (term 1/2) · ประถม: ทั้งปี (term 0)"""
    from app.models import AcadScore, AcadSubject
    from app.services.academic import weighted_avg, QUALITY_LEVELS

    sec = is_secondary(klass.level)
    t = (term if term in (1, 2) else 1) if sec else 0
    doc = _doc()                      # ปกเป็นแนวตั้ง แล้วค่อยสลับเป็นแนวนอนหลังปก
    students = sorted(klass.students, key=lambda s: (s.seq or 999, s.name))
    # ค่าที่ใช้จริงต่อคน (คำนวณจากรายวิชา/รายเดือนถ้ามี · ไม่มีก็ค่า manual) - จุดตัดสินใจเดียว
    effs = {s.id: effective_eval(s, db) for s in students}
    subjects = (db.query(AcadSubject).filter_by(year=klass.year, level=klass.level, term=t)
                .order_by(AcadSubject.seq, AcadSubject.code).all())
    sub_ids = [x.id for x in subjects]
    # เกรดทุกคน x ทุกวิชาของภาคนี้ (คิวรีเดียว)
    sc_map = {}
    if sub_ids:
        for row in db.query(AcadScore).filter(AcadScore.subject_id.in_(sub_ids),
                                              AcadScore.term == t).all():
            sc_map[(row.acad_student_id, row.subject_id)] = row
    term_txt = f"ภาคเรียนที่ {t}" if sec else "ตลอดปีการศึกษา"

    # ---------- หน้า 1: ปก (แนวตั้ง · พื้นที่พิมพ์ 18.0 ซม.) ----------
    _logo_header(doc, school, height_cm=2.2, after=4)      # ตราโรงเรียนบนปก (ถ้าตั้งค่าไว้)
    _p(doc, "สมุดบันทึกผลการพัฒนาคุณภาพผู้เรียน (ปพ.5)", align="center", bold=True, size=16, after=2)
    _p(doc, f"ชั้น {_class_label(klass)}   ปีการศึกษา {klass.year}" + (f"   {term_txt}" if sec else ""),
       align="center", bold=True, size=14, after=0)
    loc = [school.name or ""]
    if (school.district or "").strip():
        loc.append(f"อำเภอ{school.district.strip()}")
    if (school.province or "").strip():
        loc.append(f"จังหวัด{school.province.strip()}")
    _p(doc, "  ".join(loc), align="center", size=14, after=0)
    if (school.area_office or "").strip():
        _ao = school.area_office.strip()
        _p(doc, _ao if _ao.startswith(("สำนักงาน", "สพ")) else f"สำนักงานเขตพื้นที่การศึกษา{_ao}",
           align="center", size=14, after=0)
    boys = sum(1 for s in students if s.sex == "M")
    girls = sum(1 for s in students if s.sex == "F")
    _p(doc, f"นักเรียนทั้งหมด {len(students)} คน  (ชาย {boys} | หญิง {girls})",
       align="center", size=14, after=6)

    # ตารางแจกแจงระดับผลการเรียนรายวิชา
    grade_cols = ["4", "3.5", "3", "2.5", "2", "1.5", "1", "0", "ร", "มส"]
    gt = doc.add_table(rows=2, cols=2 + len(grade_cols)); gt.style = "Table Grid"
    gt.rows[0].cells[0].merge(gt.rows[1].cells[0])
    gt.rows[0].cells[1].merge(gt.rows[1].cells[1])
    _cell(gt.rows[0].cells[0], "ที่", bold=True, fill="EDE9FE")
    _cell(gt.rows[0].cells[1], "รายวิชา", bold=True, fill="EDE9FE")
    top = gt.rows[0].cells[2]
    for c in gt.rows[0].cells[3:]:
        top = top.merge(c)
    _cell(top, "จำนวนนักเรียนแยกตามระดับผลการเรียน (คน)", bold=True, fill="EDE9FE", size=14)
    for i, g in enumerate(grade_cols):
        _cell(gt.rows[1].cells[2 + i], g, bold=True, fill="EDE9FE", size=14)
    for i, sub in enumerate(subjects, start=1):
        cells = gt.add_row().cells
        _cell(cells[0], i, size=14)
        _cell(cells[1], f"{sub.code or ''} {sub.name}".strip(), align="left", size=14)
        cnt = _grade_counts((sc_map.get((s.id, sub.id)).grade
                             if sc_map.get((s.id, sub.id)) else "") for s in students)
        for j, g in enumerate(grade_cols):
            _cell(cells[2 + j], cnt[g] or "", size=14)
    # แนวตั้ง: 1.0 + 6.0 + 10x1.1 = 18.0 ซม. พอดีพื้นที่พิมพ์
    _widths(gt, [Cm(1.0), Cm(6.0)] + [Cm(1.1)] * len(grade_cols))

    # ตารางเล็ก: คุณลักษณะฯ + อ่านคิดฯ (ค่าที่ใช้จริง - คำนวณจากรายวิชาถ้ามี)
    _p(doc, "", after=4)
    qt = doc.add_table(rows=1, cols=1 + len(QUALITY_LEVELS)); qt.style = "Table Grid"
    _cell(qt.rows[0].cells[0], "ผลการประเมิน (คน)", bold=True, fill="EDE9FE", size=14)
    for i, q in enumerate(QUALITY_LEVELS):
        _cell(qt.rows[0].cells[1 + i], q, bold=True, fill="EDE9FE", size=14)
    for lab, key in [("คุณลักษณะอันพึงประสงค์", "desired_char"),
                     ("การอ่าน คิดวิเคราะห์ และเขียน", "read_think")]:
        cells = qt.add_row().cells
        _cell(cells[0], lab, align="left", size=14)
        for i, q in enumerate(QUALITY_LEVELS):
            n = sum(1 for s in students if effs[s.id][key] == q)
            _cell(cells[1 + i], n or "", size=14)
    _widths(qt, [Cm(7.2)] + [Cm(2.7)] * len(QUALITY_LEVELS))     # รวม 18.0

    # ---------- ลายเซ็นทุกคนต้องอยู่หน้าปก ----------
    # จัดเป็นตาราง 3 คอลัมน์ ครูประจำชั้น(1-2) + หัวหน้าฝ่ายวิชาการ ในแถวเดียว
    # แล้ว ผอ. อยู่ตรงกลางด้านล่าง - กินพื้นที่แนวตั้งน้อยกว่าเรียงลงมาทีละบล็อก
    _p(doc, "", after=8)
    head = (getattr(school, "academic_head_name", "") or "").strip()
    signers = [(p.name, "ครูประจำชั้น") for p in (klass.homeroom, klass.co_homeroom) if p]
    signers.append((head, "หัวหน้าฝ่ายวิชาการ"))
    st = doc.add_table(rows=1, cols=len(signers))
    for cell, (nm, role) in zip(st.rows[0].cells, signers):
        for i, txt in enumerate(["(ลงชื่อ)......................................",
                                 f"( {nm or '.....................................'} )", role]):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(txt); r.font.size = Pt(14); r.font.name = THAI_FONT
            r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
            if i == 0 and nm:
                _float_signature(p, nm)
    _widths(st, [Cm(18.0 / len(signers))] * len(signers))

    _p(doc, "", after=8)
    _p(doc, "ผลการตรวจสอบ    [   ] อนุมัติ        [   ] ไม่อนุมัติ", align="center", size=14, after=6)
    director = (getattr(school, "director_name", "") or "").strip()
    dpos = ("ผู้อำนวยการ" + school.name) if (school.name or "").startswith("โรงเรียน") \
        else "ผู้อำนวยการโรงเรียน"
    _sign_block(doc, director, dpos, size=14)
    _p(doc, "วันที่ ........ เดือน ......................... พ.ศ. ..........", align="center", size=14, after=0)

    # ---------- รายชื่อนักเรียน (แนวนอน) ----------
    _new_section(doc, landscape=True)
    _pp5_roster(doc, klass, students, db)

    # ---------- เวลาเรียนรายวัน + สรุปเวลาเรียน (แนวตั้ง) ----------
    # มัธยม = เอาเฉพาะเดือนในภาคเรียนนั้น · ประถม = ทั้งปี
    att_months = TERM_MONTHS.get(t) if sec else None
    _new_section(doc, landscape=False)
    _pp5_attendance_daily(doc, klass, students, db, page_break=False, months_ok=att_months)
    # สรุปเวลาเรียน = แนวนอน เต็มหน้า (ตามที่เจ้าของขอ)
    _new_section(doc, landscape=True)
    _p(doc, f"สรุปเวลาเรียน ชั้น {_class_label(klass)} ปีการศึกษา {klass.year}",
       align="center", bold=True, size=16, after=6, page_break=False)
    monthly = any(effs[s.id]["months"] for s in students)
    if monthly:
        # แบบรายเดือน: เดือน + รวม + ป่วย/ลา/ขาด + ร้อยละ (มัธยม = เฉพาะเดือนในภาคเรียน)
        from app.models import AcadClassMonth, AcadAttendance
        opens = {m.month: m.days_open for m in
                 db.query(AcadClassMonth).filter_by(class_id=klass.id).all()}
        month_list = [(m, nm) for m, nm in TH_MONTHS if (att_months is None or m in att_months)]
        nmon = len(month_list)
        sids = [s.id for s in students]
        marks_by = {}
        if sids:
            for a in db.query(AcadAttendance).filter(
                    AcadAttendance.acad_student_id.in_(sids),
                    AcadAttendance.subject_id.is_(None)).all():   # เล่มรวมใช้เวลารายห้อง (โฮมรูม)
                if att_months is None or a.month in att_months:
                    marks_by[(a.acad_student_id, a.month)] = count_marks(a.marks)

        def _term_att(s):
            """คืน (present ต่อเดือน, มา, ป่วย, ลา, ขาด) ของเทอมนี้ - คำนวณจาก marks ถ้ามี"""
            ef = effs[s.id]
            has_marks = any((s.id, m) in marks_by for m, _ in month_list)
            if att_months is None and not has_marks:      # ทั้งปี + ไม่มี marks รายวัน -> ยอดรวมเดิม
                return ({m: ef["months"].get(m) for m, _ in month_list},
                        ef["days_present"], ef["days_sick"], ef["days_leave"], ef["days_absent"])
            pm, pres, sick, leave, absent = {}, 0, 0, 0, 0
            for m, _ in month_list:
                cm = marks_by.get((s.id, m))
                if cm:
                    pm[m] = cm.get("/", 0)
                    pres += cm.get("/", 0); sick += cm.get("ป", 0)
                    leave += cm.get("ล", 0); absent += cm.get("ข", 0)
                else:
                    pm[m] = ef["months"].get(m)
                    pres += ef["months"].get(m) or 0
            return (pm, pres, sick if has_marks else None,
                    leave if has_marks else None, absent if has_marks else None)

        nm_first, nm_last = 2, 2 + nmon - 1
        s_first = nm_last + 1                       # คอลัมน์ มา
        pct_col, res_col = s_first + 4, s_first + 5
        ncol = s_first + 6
        at = doc.add_table(rows=3, cols=ncol); at.style = "Table Grid"
        r0, r1, r2 = at.rows[0].cells, at.rows[1].cells, at.rows[2].cells
        c = r0[0].merge(r1[0]).merge(r2[0]); _cell(c, "เลขที่", bold=True, fill="EDE9FE", size=14)
        c = r0[1].merge(r1[1]).merge(r2[1]); _cell(c, "ชื่อ-นามสกุล", bold=True, fill="EDE9FE", size=14)
        c = r0[nm_first].merge(r1[nm_last]); _cell(c, "เดือน", bold=True, fill="EDE9FE", size=14)
        for j, (m, nm) in enumerate(month_list):
            _cell(r2[nm_first + j], nm, bold=True, fill="EDE9FE", size=14)
        tot_hdr = sum(v for m, v in opens.items() if v and (att_months is None or m in att_months))
        c = r0[s_first].merge(r0[s_first + 3]); _cell(c, "รวมเวลาเรียน", bold=True, fill="EDE9FE", size=14)
        c = r1[s_first].merge(r1[s_first + 3])
        _cell(c, f"{tot_hdr} วัน" if tot_hdr else "จำนวนวัน", bold=True, fill="EDE9FE", size=14)
        for j, lab in enumerate(["มา", "ป่วย", "ลา", "ขาด"]):
            _cell(r2[s_first + j], lab, bold=True, fill="EDE9FE", size=14)
        c = r0[pct_col].merge(r1[pct_col]).merge(r2[pct_col])
        _cell(c, "มาเรียน\nร้อยละ", bold=True, fill="EDE9FE", size=14)
        c = r0[res_col].merge(r1[res_col]).merge(r2[res_col])
        _cell(c, "ผลการประเมิน", bold=True, fill="EDE9FE", size=14)
        cells = at.add_row().cells      # แถววันเปิดเรียนของห้อง (ตัวหารร้อยละ)
        _cell(cells[0], "", size=14)
        _cell(cells[1], "วันเปิดเรียน", align="left", bold=True, size=14)
        for j, (m, nm) in enumerate(month_list):
            v = opens.get(m)
            _cell(cells[nm_first + j], v if v is not None else "", bold=True, size=14)
        _cell(cells[s_first], tot_hdr or "", bold=True, size=14)
        for k in range(s_first + 1, ncol):
            _cell(cells[k], "", size=14)
        open_term = sum(opens.get(m) or 0 for m, _ in month_list)
        for s in students:
            pm, pres, sick, leave, absent = _term_att(s)
            cells = at.add_row().cells
            _cell(cells[0], s.seq or "", size=14)
            _cell(cells[1], s.name, align="left", size=14)
            for j, (m, nm) in enumerate(month_list):
                v = pm.get(m)
                _cell(cells[nm_first + j], v if v is not None else "", size=14)
            _cell(cells[s_first], pres if pres is not None else "", bold=True, size=14)
            _cell(cells[s_first + 1], "" if sick is None else sick, size=14)
            _cell(cells[s_first + 2], "" if leave is None else leave, size=14)
            _cell(cells[s_first + 3], "" if absent is None else absent, size=14)
            pct = (pres * 100.0 / open_term) if (open_term > 0 and pres is not None) else None
            _cell(cells[pct_col], f"{pct:.1f}" if pct is not None else "", size=14)
            _cell(cells[res_col], ("ผ่าน" if pct >= 80 else "ไม่ผ่าน") if pct is not None else "",
                  bold=True, size=14)
        # แนวนอน: กระจายเดือนให้เต็ม (เดือนน้อย=กว้างขึ้น) รวม ≤26.7
        fixed = 1.4 + 4.4 + (1.3 + 1.05 + 1.05 + 1.05 + 1.7 + 2.6)
        mw = min(1.5, max(0.7, (26.0 - fixed) / max(1, nmon)))
        _widths(at, [Cm(1.4), Cm(4.4)] + [Cm(mw)] * nmon +
                [Cm(1.3), Cm(1.05), Cm(1.05), Cm(1.05), Cm(1.7), Cm(2.6)])
        _tight_cells(at)   # กันตัวเลขเดือนโดนบีบ
    else:
        # แบบยอดรวมทั้งปี (โรงเรียนที่ไม่กรอกรายเดือน)
        at = doc.add_table(rows=1, cols=10); at.style = "Table Grid"
        for i, h in enumerate(["เลขที่", "เลขประจำตัว", "ชื่อ-นามสกุล", "วันเปิดเรียน", "มาเรียน",
                               "ป่วย", "ลา", "ขาด", "ร้อยละ", "ผล"]):
            _cell(at.rows[0].cells[i], h, bold=True, fill="EDE9FE", size=14)
        for s in students:
            ef = effs[s.id]
            cells = at.add_row().cells
            _cell(cells[0], s.seq or "", size=14)
            _cell(cells[1], s.student_no or "", size=14)
            _cell(cells[2], s.name, align="left", size=14)
            vals = [ef["days_open"], ef["days_present"],
                    ef["days_sick"], ef["days_leave"], ef["days_absent"]]
            for j, v in enumerate(vals):
                _cell(cells[3 + j], v if v is not None else "", size=14)
            pct = None
            if (ef["days_open"] or 0) > 0 and ef["days_present"] is not None:
                pct = ef["days_present"] * 100.0 / ef["days_open"]
            _cell(cells[8], f"{pct:.1f}" if pct is not None else "", size=14)
            _cell(cells[9], ("ผ่าน" if pct >= 80 else "ไม่ผ่าน") if pct is not None else "", bold=True, size=14)
        # แนวนอน เต็มหน้า ~26 ซม.
        _widths(at, [Cm(1.5), Cm(2.7), Cm(6.8), Cm(2.5), Cm(2.4),
                     Cm(1.9), Cm(1.9), Cm(1.9), Cm(2.3), Cm(2.6)])   # รวม ~26.5
    _p(doc, "เกณฑ์การผ่าน: มีเวลาเรียนไม่น้อยกว่าร้อยละ 80 ของเวลาเรียนทั้งหมด", size=12, after=0)

    # ---------- ตัวชี้วัดรายวิชา (ทุกวิชารวมเป็นชุด, แนวนอน) - มาก่อนคะแนน ตามไฟล์จริง ----------
    from app.services.curriculum import selected_indicators as _sel_ind
    if any(_sel_ind(sub) for sub in subjects):
        _new_section(doc, landscape=True)
        ind_done = False
        for sub in subjects:
            if _pp5_indicator_page(doc, klass, sub, db, students, page_break=ind_done):
                ind_done = True
    # ---------- คะแนนรายวิชา / แบบบันทึกผลการเรียน (ทุกวิชา, แนวนอน) ----------
    # (คุณลักษณะ/อ่านคิดเขียน ไม่พิมพ์รายวิชา - พิมพ์เฉพาะสรุปทั้งปีด้านล่าง ตามไฟล์จริง)
    _new_section(doc, landscape=True)
    for i, sub in enumerate(subjects):
        _pp5_score_page(doc, school, klass, sub, db, page_break=(i > 0))

    # ---------- สรุปผลการเรียนทุกวิชา + คุณลักษณะ/อ่านเขียน/ผลประเมิน (แนวนอน) ----------
    _new_section(doc, landscape=True)
    _p(doc, f"สรุปผลการเรียนทุกรายวิชา ชั้น {_class_label(klass)} ปีการศึกษา {klass.year}"
       + (f" {term_txt}" if sec else ""), align="center", bold=True, size=16, after=6, page_break=False)
    if subjects:
        sw = min(2.6, 17.5 / len(subjects))
        mt = doc.add_table(rows=1, cols=2 + len(subjects) + 1); mt.style = "Table Grid"
        _cell(mt.rows[0].cells[0], "เลขที่", bold=True, fill="EDE9FE", size=11)
        _cell(mt.rows[0].cells[1], "ชื่อ-นามสกุล", bold=True, fill="EDE9FE", size=11)
        for i, sub in enumerate(subjects):
            _cell(mt.rows[0].cells[2 + i], sub.code or sub.name[:6], bold=True, fill="EDE9FE", size=10)
        _cell(mt.rows[0].cells[-1], "เฉลี่ย", bold=True, fill="EDE9FE", size=11)
        for s in students:
            cells = mt.add_row().cells
            _cell(cells[0], s.seq or "", size=11)
            _cell(cells[1], s.name, align="left", size=11)
            pairs = []
            for i, sub in enumerate(subjects):
                row = sc_map.get((s.id, sub.id))
                g = row.grade if row else ""
                _cell(cells[2 + i], g or "", size=11, bold=True)
                pairs.append((g, sub.credit if sec else sub.hours))
            avg = weighted_avg(pairs)
            _cell(cells[-1], f"{avg:.2f}" if avg is not None else "", bold=True, size=11)
        _widths(mt, [Cm(1.2), Cm(6.0)] + [Cm(sw)] * len(subjects) + [Cm(1.8)])
        _p(doc, "เฉลี่ยถ่วงน้ำหนักด้วย" + ("หน่วยกิต" if sec else "เวลาเรียน")
           + " | ร/มส/ผ/มผ ไม่นำมาคิดเฉลี่ย", size=11, after=0, align="center")

    # ---------- สรุปคุณลักษณะฯ / อ่านคิดเขียน ทุกวิชา ----------
    _pp5_quality_summary(doc, klass, subjects, students, db, "char",
                         "สรุปผลการประเมินคุณลักษณะอันพึงประสงค์ทุกรายวิชา")
    _pp5_quality_summary(doc, klass, subjects, students, db, "read",
                         "สรุปผลการประเมินการอ่าน คิดวิเคราะห์ และเขียนทุกรายวิชา")

    # ---------- สรุปการประเมินผลการเรียน (กริดรวม ก็อปตามชีต 'พิมพ์สรุปผลการประเมินทั้งปี') ----------
    _p(doc, f"สรุปการประเมินผลการเรียน ชั้น{_class_label(klass)} โรงเรียน {school.name or ''} "
       f"ปีการศึกษา {klass.year}" + (f" {term_txt}" if sec else ""),
       align="center", bold=True, size=16, after=6, page_break=True)
    from app.models import AcadIndicatorResult
    from app.services.curriculum import selected_indicators as _sel_ind
    from app.services.academic import activities_for, activity_summary
    acts = activities_for(klass.year, klass.level, db)      # ใช้ตัดสิน ผ/มผ ภายใน (ไม่โชว์คอลัมน์)
    ind_total = sum(len(_sel_ind(sub)) for sub in subjects)
    ind_pass = {}
    if subjects and students:
        for r in db.query(AcadIndicatorResult).filter(
                AcadIndicatorResult.subject_id.in_([sub.id for sub in subjects]),
                AcadIndicatorResult.acad_student_id.in_([s.id for s in students])).all():
            if r.passed:
                ind_pass[r.acad_student_id] = ind_pass.get(r.acad_student_id, 0) + 1
    _qnum = {"ดีเยี่ยม": "3", "ดี": "2", "ผ่าน": "1", "ไม่ผ่าน": "0"}
    nsub = len(subjects)
    ncol = 1 + nsub + 5
    ft = doc.add_table(rows=4, cols=ncol); ft.style = "Table Grid"

    def _fc(r, c):
        return ft.cell(r, c)

    def _fmg(r1, c1, r2, c2):
        return _fc(r1, c1).merge(_fc(r2, c2))
    # คอลัมน์แรก = ป้ายหัว 4 แถว (วิชาที่/ประเภท/วิชา/รหัสวิชา) แล้วเป็น เลขที่ ในแถวข้อมูล
    for r, lab in enumerate(["วิชาที่", "ประเภท", "วิชา", "รหัสวิชา"]):
        _cell(_fc(r, 0), lab, bold=True, fill="EDE9FE", size=14)
    # คอลัมน์รายวิชา: เลขวิชา / ประเภท / ชื่อวิชา(แนวตั้ง) / รหัสวิชา
    for i, sub in enumerate(subjects, start=1):
        _cell(_fc(0, i), str(i), bold=True, fill="EDE9FE", size=14)
        _cell(_fc(1, i), sub.kind or "พื้นฐาน", bold=True, fill="F1F5F9", size=14)
        _vcell(_fc(2, i), sub.name, bold=True, fill="F8FAFC", size=14)
        _cell(_fc(3, i), sub.code or "", bold=True, fill="F1F5F9", size=14)
    # คอลัมน์ประเมิน (หัวแนวตั้ง) - ตัวชี้วัดเว้นแถวล่างใส่ยอดรวม (ตามชีต R2:R4 + R5=จำนวนรวม)
    a0 = 1 + nsub
    _vcell(_fmg(0, a0, 3, a0), "ผลการเรียนเฉลี่ย", bold=True, fill="EDE9FE", size=14)
    _vcell(_fmg(0, a0 + 1, 2, a0 + 1), "ผลประเมินตัวชี้วัด", bold=True, fill="EDE9FE", size=14)
    _cell(_fc(3, a0 + 1), str(ind_total) if ind_total else "", bold=True, fill="F8FAFC", size=14)
    _vcell(_fmg(0, a0 + 2, 3, a0 + 2), "ผลประเมินคุณลักษณะฯ", bold=True, fill="EDE9FE", size=14)
    _vcell(_fmg(0, a0 + 3, 3, a0 + 3), "ผลประเมินอ่านคิดเขียน", bold=True, fill="EDE9FE", size=14)
    _vcell(_fmg(0, a0 + 4, 3, a0 + 4), "ผลการเรียนตลอดปี", bold=True, fill="EDE9FE", size=14)
    for s in students:
        cells = ft.add_row().cells
        _cell(cells[0], s.seq or "", size=14)
        pairs, grades = [], []
        for i, sub in enumerate(subjects, start=1):
            row = sc_map.get((s.id, sub.id))
            g = row.grade if row else ""
            grades.append(g)
            pairs.append((g, sub.credit if sec else sub.hours))
            _cell(cells[i], g or "", size=14, bold=True)
        avg = weighted_avg(pairs)
        ef = effs[s.id]
        _cell(cells[a0], f"{avg:.2f}" if avg is not None else "", bold=True, size=14)
        _cell(cells[a0 + 1], str(ind_pass.get(s.id, "")) if ind_total else "", size=14)
        _cell(cells[a0 + 2], _qnum.get(ef["desired_char"], ""), size=14)
        _cell(cells[a0 + 3], _qnum.get(ef["read_think"], ""), size=14)
        overall = ""
        asum = activity_summary(s, db) if acts else "ผ"    # ไม่มีกิจกรรม = ไม่กันด้วยกิจกรรม
        if subjects and all((g or "").strip() for g in grades):
            bad_grade = any((g or "").strip() in ("0", "ร", "มส") for g in grades)
            bad_qual = "ไม่ผ่าน" in (ef["desired_char"], ef["read_think"])
            no_qual = not ef["desired_char"].strip() or not ef["read_think"].strip()
            if not no_qual and asum != "":
                overall = "มผ" if (bad_grade or bad_qual or asum == "มผ") else "ผ"
        _cell(cells[a0 + 4], overall, bold=True, size=14)
    # ความกว้าง (แนวนอน): เลขที่ + วิชา (กว้างพอให้รหัสไม่ตกบรรทัด) + 5 คอลัมน์ประเมิน
    assess = [1.5, 1.5, 1.5, 1.5, 1.5]
    sw = min(1.6, (26.0 - 1.4 - sum(assess)) / max(1, nsub))
    _widths(ft, [Cm(1.4)] + [Cm(sw)] * nsub + [Cm(x) for x in assess])
    _tight_cells(ft)
    # แถวชื่อวิชา (แนวตั้ง) สูงพอให้ชื่อยาวไม่โดนตัด แต่ไม่สูงเกิน (ย่อหัวตาราง)
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    maxlen = max((len(sub.name or "") for sub in subjects), default=8)
    ft.rows[2].height = Cm(min(7.5, maxlen * 0.28 + 0.4))
    ft.rows[2].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    _p(doc, f"ตัวชี้วัด = จำนวนที่ผ่านจากทั้งหมด {ind_total} ตัว | คุณลักษณะ·อ่านคิดเขียน = 3 ดีเยี่ยม 2 ดี 1 ผ่าน 0 ไม่ผ่าน | "
            "ผล ผ = ผ่านครบทุกวิชา คุณลักษณะ อ่านเขียน และกิจกรรม (ข้อมูลไม่ครบ = เว้นว่าง)",
       size=12, after=0, align="center")

    # ---------- หน้าสุดท้าย: เกณฑ์การประเมิน (แนวตั้ง) ----------
    _new_section(doc, landscape=False)
    _pp5_criteria_page(doc, page_break=False)

    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    suffix = f"_ภาค{t}" if sec else ""
    out = out_dir / (_safe(f"ปพ.5_ทั้งเล่ม_{_class_label(klass)}_{klass.year}{suffix}") + ".docx")
    doc.save(str(out))
    return str(out)


def _pp5_criteria_page(doc, *, page_break=True):
    """หน้าเกณฑ์การประเมิน (ปิดท้ายเล่ม ปพ.5) ตามไฟล์จริง"""
    _p(doc, "เกณฑ์การประเมินผลการเรียน", align="center", bold=True, size=16, after=6, page_break=page_break)
    # ระดับผลการเรียน
    t = doc.add_table(rows=1, cols=3); t.style = "Table Grid"
    for i, h in enumerate(["ระดับผลการเรียน", "ความหมาย", "ช่วงคะแนน"]):
        _cell(t.rows[0].cells[i], h, bold=True, fill="EDE9FE")
    for lv, mean, rng in [("4", "ผลการเรียนเยี่ยม", "80 - 100"), ("3.5", "ผลการเรียนดีมาก", "75 - 79"),
                          ("3", "ผลการเรียนดี", "70 - 74"), ("2.5", "ผลการเรียนค่อนข้างดี", "65 - 69"),
                          ("2", "ผลการเรียนน่าพอใจ", "60 - 64"), ("1.5", "ผลการเรียนพอใช้", "55 - 59"),
                          ("1", "ผลการเรียนผ่านเกณฑ์", "50 - 54"), ("0", "ผลการเรียนต่ำกว่าเกณฑ์", "0 - 49")]:
        c = t.add_row().cells
        _cell(c[0], lv); _cell(c[1], mean, align="left"); _cell(c[2], rng)
    _widths(t, [Cm(4.5), Cm(7.0), Cm(4.5)])

    _p(doc, "", after=4)
    _p(doc, "บันทึกเวลาเรียน:  / หมายถึง มาเรียน · ป หมายถึง ป่วย · ล หมายถึง ลากิจ · ข หมายถึง ขาด",
       size=12, after=2)
    _p(doc, "กิจกรรมพัฒนาผู้เรียน:  มี 2 เกณฑ์ คือ ผ่าน (ผ) และ ไม่ผ่าน (มผ)", size=12, after=6)

    _p(doc, "เกณฑ์การตัดสินคุณลักษณะอันพึงประสงค์ และการอ่าน คิดวิเคราะห์ และเขียนสื่อความ",
       bold=True, size=16, after=4)
    t2 = doc.add_table(rows=1, cols=3); t2.style = "Table Grid"
    for i, h in enumerate(["คะแนน", "ความหมาย", "ช่วงคะแนนเฉลี่ย"]):
        _cell(t2.rows[0].cells[i], h, bold=True, fill="EDE9FE")
    for lv, mean, rng in [("3", "ดีเยี่ยม", "2.5 - 3"), ("2", "ดี", "1.5 - 2.49"),
                          ("1", "ผ่าน", "1 - 1.49"), ("0", "ไม่ผ่าน", "0 - 0.99")]:
        c = t2.add_row().cells
        _cell(c[0], lv); _cell(c[1], mean); _cell(c[2], rng)
    _widths(t2, [Cm(4.5), Cm(7.0), Cm(4.5)])

    _p(doc, "", after=4)
    _p(doc, "เกณฑ์การตัดสินตัวชี้วัดรายวิชา:  ต้องผ่านทุกตัวชี้วัด (ร้อยละ 100)", bold=True, size=12, after=6)
    _p(doc, "สูตรร้อยละของเวลาเรียน = (เวลามาเรียน × 100) ÷ เวลาเรียนทั้งหมด", size=12, after=2)
    _p(doc, "สูตรร้อยละของตัวชี้วัด = (จำนวนตัวชี้วัดที่ผ่าน × 100) ÷ จำนวนตัวชี้วัดทั้งหมด", size=12, after=2)
    _p(doc, "สูตรค่าเฉลี่ยของแบบประเมิน = ผลรวมของคะแนนประเมินทั้งหมด ÷ จำนวนข้อของแบบประเมิน",
       size=12, after=0)


# ============================ ปพ.6 ============================
_DOTS = "................................................................................................"


def _pp6_central(s, db):
    """คืนทะเบียนกลาง (Student) ของ AcadStudent - None ถ้าเพิ่มมือ (ไม่ได้ดึงจากทะเบียน)"""
    from app.models import Student
    return db.get(Student, s.student_id) if s.student_id else None


def _pp6_cover(doc, school, s, db, *, page_break):
    """หน้า 1 : ปกสมุดรายงานประจำตัวนักเรียน"""
    klass = s.klass
    added = _logo_header(doc, school, page_break=page_break, height_cm=3.0, after=6)
    _p(doc, "", after=0, page_break=(page_break and not added))
    if not added:
        _p(doc, "", after=0)
    _p(doc, "สมุดรายงานประจำตัวนักเรียน", align="center", bold=True, size=26, after=6)
    _p(doc, "แบบรายงานผลการพัฒนาคุณภาพผู้เรียนรายบุคคล (ปพ.6)", align="center", bold=True, size=16, after=2)
    _p(doc, f"ปีการศึกษา {klass.year}", align="center", size=15, after=18)
    _p(doc, school.name or "", align="center", bold=True, size=17, after=2)
    loc = []
    if getattr(school, "district", ""):
        loc.append(f"อำเภอ{school.district}")
    if getattr(school, "province", ""):
        loc.append(f"จังหวัด{school.province}")
    if loc:
        _p(doc, "  ".join(loc), align="center", size=14, after=2)
    if getattr(school, "area_office", ""):
        _p(doc, school.area_office, align="center", size=14, after=24)
    else:
        _p(doc, "", after=24)
    # รูปนักเรียนบนปก (ถ้าผูกทะเบียนกลาง + อัปโหลดรูปไว้)
    st = _pp6_central(s, db)
    if st and getattr(st, "photo", None):
        import io as _io
        pic = _p(doc, "", align="center", after=6)
        try:
            pic.add_run().add_picture(_io.BytesIO(st.photo), height=Cm(4.0))
        except Exception:
            pic._element.getparent().remove(pic._element)
    _p(doc, f"เลขที่ {s.seq or '.....'}", align="center", size=15, after=4)
    _p(doc, f"ชื่อ  {s.name}", align="center", bold=True, size=18, after=4)
    _p(doc, f"เลขประจำตัวนักเรียน  {s.student_no or '................'}", align="center", size=14, after=4)
    _p(doc, f"ชั้น  {_class_label(klass)}", align="center", size=15, after=24)
    homerooms = [p.name for p in (klass.homeroom, klass.co_homeroom) if p]
    _p(doc, "ครูประจำชั้น", align="center", size=14, after=2)
    if homerooms:
        for nm in homerooms:
            _p(doc, nm, align="center", bold=True, size=15, after=2)
    else:
        _p(doc, ".............................................", align="center", size=14, after=2)


def _fmt_addr(st) -> str:
    """เรียงที่อยู่จากฟิลด์แยกเป็นบรรทัดเดียว"""
    parts = []
    if st.addr_no:
        parts.append(f"บ้านเลขที่ {st.addr_no}")
    if st.addr_moo:
        parts.append(f"หมู่ {st.addr_moo}")
    if st.addr_soi:
        parts.append(f"ซอย{st.addr_soi}")
    if st.addr_road:
        parts.append(f"ถนน{st.addr_road}")
    if st.addr_tambon:
        parts.append(f"ตำบล{st.addr_tambon}")
    if st.addr_amphoe:
        parts.append(f"อำเภอ{st.addr_amphoe}")
    if st.addr_province:
        parts.append(f"จังหวัด{st.addr_province}")
    if st.addr_zip:
        parts.append(st.addr_zip)
    return " ".join(parts)


def _pp6_personal(doc, school, s, db):
    """หน้า 2 : ข้อมูลส่วนตัว (จากทะเบียนนักเรียนกลาง)"""
    st = _pp6_central(s, db)
    _p(doc, "ข้อมูลส่วนตัว", align="center", bold=True, size=17, after=6, page_break=True)

    def g(attr):
        return (getattr(st, attr, "") or "") if st else ""

    rows = [
        ("ชื่อ-นามสกุล", s.name),
        ("เลขประจำตัวนักเรียน", s.student_no or ""),
        ("เลขประจำตัวประชาชน", g("id_card")),
        ("วันเกิด", thai_date(st.birthdate) if (st and st.birthdate) else ""),
        ("เชื้อชาติ / สัญชาติ / ศาสนา",
         " / ".join(x for x in [g("race"), g("nationality"), g("religion")] if x)),
        ("หมู่เลือด", g("blood_group")),
        ("โรคประจำตัว", g("congenital_disease")),
        ("ที่อยู่", _fmt_addr(st) if st else ""),
        ("โทรศัพท์", g("phone")),
        ("ชื่อบิดา", g("father_name")),
        ("ชื่อมารดา", g("mother_name")),
        ("โรงเรียนเดิม", g("prev_school")),
        ("วันเข้าเรียน", thai_date(st.enroll_date) if (st and st.enroll_date) else ""),
    ]
    t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
    for lab, val in rows:
        cells = t.add_row().cells
        _cell(cells[0], lab, bold=True, align="left", fill="F1F5F9")
        _cell(cells[1], val or "", align="left")
    _widths(t, [Cm(5.0), Cm(11.0)])

    # รูปนักเรียน: ฝังรูปจริงถ้ามี ไม่งั้นเว้นกล่องกรอบให้ติดรูป
    _p(doc, "", after=6)
    photo = doc.add_table(rows=1, cols=1); photo.style = "Table Grid"
    photo.rows[0].height = Cm(4.0)
    if st and getattr(st, "photo", None):
        import io as _io
        cell = photo.rows[0].cells[0]
        cell.text = ""
        pp = cell.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            pp.add_run().add_picture(_io.BytesIO(st.photo), height=Cm(3.8))
        except Exception:
            _cell(cell, "รูปนักเรียน", align="center", size=12)
    else:
        _cell(photo.rows[0].cells[0], "รูปนักเรียน\n(ขนาด 1-2 นิ้ว)", align="center", size=12)
    _widths(photo, [Cm(3.5)])
    if not st:
        _p(doc, "", after=2)
        _p(doc, "(นักเรียนคนนี้เพิ่มด้วยมือ ยังไม่ได้ผูกทะเบียนนักเรียนกลาง จึงยังไม่มีข้อมูลส่วนตัว)",
           size=12, after=0)


def _pp6_attendance_growth(doc, school, s, db, ef):
    """หน้า 3 : เวลาเรียนรายเดือน + น้ำหนัก/ส่วนสูง (ภาวะโภชนาการ)"""
    from app.models import AcadAttendance
    from app.services import growth
    klass = s.klass
    _p(doc, "เวลาเรียน", align="center", bold=True, size=17, after=6, page_break=True)

    att = {a.month: a for a in db.query(AcadAttendance)
           .filter(AcadAttendance.acad_student_id == s.id,
                   AcadAttendance.subject_id.is_(None)).all()}   # ปพ.6 ใช้เวลารายห้อง (โฮมรูม)
    heads = ["เดือน", "มา", "ป่วย", "ลา", "ขาด"]
    t = doc.add_table(rows=1, cols=5); t.style = "Table Grid"
    for i, h in enumerate(heads):
        _cell(t.rows[0].cells[i], h, bold=True, fill="EDE9FE")
    tot = {"/": 0, "ป": 0, "ล": 0, "ข": 0}
    for m, abbr in TH_MONTHS:
        a = att.get(m)
        cm = count_marks(a.marks) if (a and a.marks) else None
        cells = t.add_row().cells
        _cell(cells[0], abbr, align="left")
        if cm:
            _cell(cells[1], cm["/"] or "")
            _cell(cells[2], cm["ป"] or "")
            _cell(cells[3], cm["ล"] or "")
            _cell(cells[4], cm["ข"] or "")
            for k in tot:
                tot[k] += cm[k]
        elif a and a.present is not None:
            _cell(cells[1], a.present)
            for c in cells[2:]:
                _cell(c, "")
            tot["/"] += a.present
        else:
            for c in cells[1:]:
                _cell(c, "")
    rc = t.add_row().cells
    _cell(rc[0], "รวม", bold=True, fill="F1F5F9")
    _cell(rc[1], tot["/"] or "", bold=True); _cell(rc[2], tot["ป"] or "", bold=True)
    _cell(rc[3], tot["ล"] or "", bold=True); _cell(rc[4], tot["ข"] or "", bold=True)
    _widths(t, [Cm(4.0), Cm(3.0), Cm(3.0), Cm(3.0), Cm(3.0)])

    days_open = ef.get("days_open")
    present = ef.get("days_present")
    if days_open and present is not None:
        pct = 100.0 * present / days_open if days_open else 0
        _p(doc, f"วันเปิดเรียนทั้งปี {days_open} วัน  มาเรียน {present} วัน  คิดเป็นร้อยละ {pct:.1f}",
           size=13, after=6)
    else:
        _p(doc, "", after=4)

    # ---- น้ำหนัก/ส่วนสูง (ภาวะโภชนาการ) ----
    _p(doc, "น้ำหนัก / ส่วนสูง", bold=True, size=15, after=2)
    st = _pp6_central(s, db)
    who = st or s     # ใช้ทะเบียนกลาง (มีวันเกิด) ถ้ามี ไม่งั้น AcadStudent
    ms = growth.measures_for(db, s.student_id, klass.year) if s.student_id else {}
    gh = ["ภาคเรียน", "ครั้งที่", "วันที่", "น้ำหนัก (กก.)", "ส่วนสูง (ซม.)", "ภาวะโภชนาการ"]
    gt = doc.add_table(rows=1, cols=6); gt.style = "Table Grid"
    for i, h in enumerate(gh):
        _cell(gt.rows[0].cells[i], h, bold=True, fill="EDE9FE")
    for term in (1, 2):
        m = ms.get(term)
        cells = gt.add_row().cells
        _cell(cells[0], term)
        _cell(cells[1], 1)
        _cell(cells[2], be_date_input(m.date) if (m and m.date) else "")
        _cell(cells[3], f"{m.weight:g}" if (m and m.weight) else "")
        _cell(cells[4], f"{m.height:g}" if (m and m.height) else "")
        res = growth.measure_result(who, m) if m else None
        _cell(cells[5], res["wh"] if (res and res.get("wh")) else "")
    _widths(gt, [Cm(2.8), Cm(2.2), Cm(3.0), Cm(3.0), Cm(3.0), Cm(3.6)])
    _p(doc, "เกณฑ์อ้างอิง: กราฟการเจริญเติบโตของกรมอนามัย", align="center", size=12, after=0)


def _pp6_grades(doc, school, s, db, ef):
    """หน้า 4 : ผลการเรียนรายวิชา (เว้นคอลัมน์ตัวชี้วัด) + GPA"""
    from app.models import AcadScore, AcadSubject
    klass = s.klass
    sec = is_secondary(klass.level)
    _p(doc, "ผลการเรียน", align="center", bold=True, size=17, after=6, page_break=True)

    subs = (db.query(AcadSubject).filter_by(year=klass.year, level=klass.level)
            .order_by(AcadSubject.seq, AcadSubject.code).all())
    my = {(x.subject_id, x.term): x for x in
          db.query(AcadScore).filter_by(acad_student_id=s.id).all()}
    gpa_pairs = []

    if sec:
        heads = ["ที่", "รหัสวิชา", "รายวิชา", "หน่วยกิต", "ตัวชี้วัด", "ภาค 1", "ภาค 2"]
        ws = [Cm(1.0), Cm(2.2), Cm(6.0), Cm(1.8), Cm(2.0), Cm(2.0), Cm(2.0)]
    else:
        heads = ["ที่", "รหัสวิชา", "รายวิชา", "เวลาเรียน", "ตัวชี้วัด", "คะแนน", "ผลการเรียน"]
        ws = [Cm(1.0), Cm(2.4), Cm(6.2), Cm(2.0), Cm(2.0), Cm(2.0), Cm(2.4)]
    t = doc.add_table(rows=1, cols=len(heads)); t.style = "Table Grid"
    for i, h in enumerate(heads):
        _cell(t.rows[0].cells[i], h, bold=True, fill="EDE9FE")

    seen, no = set(), 0
    for sub in subs:
        key = (sub.code or "", sub.name)
        if sec and key in seen:
            continue
        seen.add(key)
        no += 1
        cells = t.add_row().cells
        _cell(cells[0], no)
        _cell(cells[1], sub.code or "")
        _cell(cells[2], sub.name, align="left")
        if sec:
            _cell(cells[3], f"{sub.credit:g}" if sub.credit else "")
            _cell(cells[4], "")   # ตัวชี้วัด: เว้นว่าง (ยังไม่ติดตามรายตัวชี้วัด)
            for i, tm in enumerate((1, 2)):
                ids = [x.id for x in subs if (x.code or "", x.name) == key and x.term == tm]
                g = ""
                for sid2 in ids:
                    row = my.get((sid2, tm))
                    if row and row.grade:
                        g = row.grade
                _cell(cells[5 + i], g, bold=True)
                gpa_pairs.append((g, sub.credit or 1))
        else:
            _cell(cells[3], sub.hours or "")
            _cell(cells[4], "")   # ตัวชี้วัด: เว้นว่าง
            row = my.get((sub.id, 0))
            _cell(cells[5], f"{row.score:g}" if (row and row.score is not None) else "")
            _cell(cells[6], row.grade if row else "", bold=True)
            gpa_pairs.append(((row.grade if row else ""), sub.credit or 1))
    _widths(t, ws)

    gpa = weighted_avg(gpa_pairs)
    _p(doc, "", after=2)
    _p(doc, f"ผลการเรียนเฉลี่ย (GPA): {gpa:.2f}" if gpa is not None else "ผลการเรียนเฉลี่ย (GPA): -",
       bold=True, size=14, after=0)


def _pp6_assess_table(doc, s, db, title, model, avg_fn, summary_val):
    """หน้าประเมินรายวิชา (คุณลักษณะ / อ่านคิดเขียน) : รายวิชา -> ผลการประเมิน + สรุป
    เรียงรายวิชาตามลำดับเดียวกับหน้าผลการเรียน (seq, รหัส) ให้ทุกหน้าตรงกัน"""
    from app.models import AcadSubject
    klass = s.klass
    _p(doc, title, align="center", bold=True, size=17, after=6, page_break=True)
    subs = (db.query(AcadSubject).filter_by(year=klass.year, level=klass.level)
            .order_by(AcadSubject.seq, AcadSubject.code).all())
    by_subj = {r.subject_id: r for r in db.query(model).filter_by(acad_student_id=s.id).all()}
    heads = ["รหัสวิชา", "รายวิชา", "ผลการประเมิน"]
    t = doc.add_table(rows=1, cols=3); t.style = "Table Grid"
    for i, h in enumerate(heads):
        _cell(t.rows[0].cells[i], h, bold=True, fill="EDE9FE")
    any_row = False
    seen = set()
    for sub in subs:
        key = (sub.code or "", sub.name)
        if key in seen:
            continue
        seen.add(key)
        r = by_subj.get(sub.id)
        label = quality_of_avg(avg_fn(r))[1] if r else ""
        if not label:
            continue
        any_row = True
        cells = t.add_row().cells
        _cell(cells[0], sub.code or "")
        _cell(cells[1], sub.name or "", align="left")
        _cell(cells[2], label, bold=True)
    if not any_row:
        cells = t.add_row().cells
        _cell(cells[0], "")
        _cell(cells[1], "(ยังไม่ได้ประเมินรายวิชา)", align="left")
        _cell(cells[2], "")
    _widths(t, [Cm(2.6), Cm(9.4), Cm(4.0)])
    _p(doc, "", after=2)
    _p(doc, f"สรุปผลการประเมิน: {summary_val or '-'}", bold=True, size=14, after=0)


def _pp6_activities_onet(doc, school, s, db):
    """หน้า 7 : กิจกรรมพัฒนาผู้เรียน + ผลการทดสอบระดับชาติ (O-NET)"""
    from app.models import AcadActivityResult
    klass = s.klass
    _p(doc, "กิจกรรมพัฒนาผู้เรียน", align="center", bold=True, size=17, after=6, page_break=True)
    acts = activities_for(klass.year, klass.level, db)
    my_act = {r.activity_id: (r.result or "").strip() for r in
              db.query(AcadActivityResult).filter_by(acad_student_id=s.id).all()}
    t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
    _cell(t.rows[0].cells[0], "กิจกรรม", bold=True, fill="EDE9FE", align="left")
    _cell(t.rows[0].cells[1], "ผลการประเมิน", bold=True, fill="EDE9FE")
    if acts:
        for a in acts:
            cells = t.add_row().cells
            _cell(cells[0], a.name, align="left")
            _cell(cells[1], my_act.get(a.id, "") or "")
    else:
        cells = t.add_row().cells
        _cell(cells[0], "(ยังไม่ได้ตั้งกิจกรรมพัฒนาผู้เรียน)", align="left")
        _cell(cells[1], "")
    _widths(t, [Cm(12.0), Cm(4.0)])
    _p(doc, "", after=2)
    _p(doc, f"สรุปผลกิจกรรมพัฒนาผู้เรียน: {activity_summary(s, db) or '-'}", bold=True, size=14, after=6)

    # ---- O-NET เฉพาะชั้นปลายทาง ----
    if is_exit_level(klass.level):
        _p(doc, "ผลการทดสอบทางการศึกษาระดับชาติขั้นพื้นฐาน (O-NET)", bold=True, size=14, after=2)
        o = onet_for(s, db)
        oh = ["วิชา", "คะแนนเต็ม", "คะแนนที่ได้"]
        ot = doc.add_table(rows=1, cols=3); ot.style = "Table Grid"
        for i, h in enumerate(oh):
            _cell(ot.rows[0].cells[i], h, bold=True, fill="EDE9FE")
        for subj in ONET_SUBJECTS:
            row = o.get(subj)
            cells = ot.add_row().cells
            _cell(cells[0], subj, align="left")
            _cell(cells[1], f"{row.full_score:g}" if (row and row.full_score is not None) else "")
            _cell(cells[2], f"{row.score:g}" if (row and row.score is not None) else "")
        _widths(ot, [Cm(6.0), Cm(4.0), Cm(4.0)])


_CDOTS = "..............................................................................................."


def _comment_table(doc, title, *, lines_per_term=4):
    """ตารางกรอบความคิดเห็น 1 ชุด: หัวรวม + [ภาคเรียน | ความคิดเห็น...] + แถวภาค 1/2 (เว้นเขียนมือ)"""
    t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
    # หัวรวม (ผสาน 2 ช่อง)
    hr = t.add_row().cells
    hr[0].merge(hr[1])
    _cell(hr[0], title, bold=True, size=15, fill="F1F5F9")
    # หัวคอลัมน์
    sh = t.add_row().cells
    _cell(sh[0], "ภาคเรียน", bold=True, size=13, fill="F8FAFC")
    _cell(sh[1], title, bold=True, size=13, fill="F8FAFC")
    for term in (1, 2):
        cells = t.add_row().cells
        _cell(cells[0], f"ภาคเรียนที่ {term}", size=13)
        cells[1].text = ""
        first = cells[1].paragraphs[0]
        for i in range(lines_per_term):
            p = first if i == 0 else cells[1].add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(_CDOTS); r.font.size = Pt(13); r.font.name = THAI_FONT
            r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    _widths(t, [Cm(3.0), Cm(13.0)])
    return t


def _pp6_comments(doc, school, s, db):
    """หน้า 8 : ความคิดเห็นครูประจำชั้น + ผู้ปกครอง (ตารางกรอบ แยกภาคเรียน 1/2)"""
    _p(doc, "ความคิดเห็นของครูประจำชั้นและผู้ปกครอง", align="center", bold=True, size=17, after=8,
       page_break=True)
    _comment_table(doc, "ความคิดเห็นของครูประจำชั้น")
    _p(doc, "", after=8)
    _comment_table(doc, "ความคิดเห็นของผู้ปกครอง")


def _pp6_summary(doc, school, s, db, ef):
    """หน้า 9 : แบบสรุปผลการประเมินการเรียน (ตาม ปพ.6) + เกณฑ์ตัดสิน + ลงนาม"""
    from app.models import AcadScore, AcadSubject
    klass = s.klass
    sec = is_secondary(klass.level)
    _p(doc, "ผลการประเมินการเรียนระดับชั้น", align="center", bold=True, size=17, after=8,
       page_break=True)

    # เลขที่ / ชื่อ-นามสกุล
    _p(doc, f"เลขที่  {s.seq or '.......'}     ชื่อ-นามสกุล  {s.name}", size=14, after=4)
    # ย้ายเข้าระหว่างปีการศึกษา (เว้นให้ติ๊ก/กรอกวันที่เอง)
    _p(doc, "ย้ายเข้าระหว่างปีการศึกษา   ☐ ใช่   ☐ ไม่ใช่      วันที่ย้ายเข้า .............................",
       size=14, after=6)

    subs = (db.query(AcadSubject).filter_by(year=klass.year, level=klass.level).all())
    my = {(x.subject_id, x.term): x for x in
          db.query(AcadScore).filter_by(acad_student_id=s.id).all()}
    pairs, grades = [], []
    for sub in subs:
        row = my.get((sub.id, sub.term if sec else 0))
        if row:
            pairs.append((row.grade, sub.credit or 1))
            grades.append((str(row.grade) if row.grade is not None else "").strip())
    gpa = weighted_avg(pairs)

    present = ef.get("days_present")
    days_open = ef.get("days_open")
    att_pct = (100.0 * present / days_open) if (days_open and present is not None) else None
    att_txt = f"{att_pct:.2f}" if att_pct is not None else ""

    char = (ef.get("desired_char") or "").strip()
    read = (ef.get("read_think") or "").strip()
    act = (activity_summary(s, db) or "").strip()
    has_zero = any(g == "0" for g in grades)

    # ตัวชี้วัดรายวิชา: ร้อยละที่ผ่าน (เฉพาะตัวชี้วัดที่ครูเลือกใช้)
    from app.services.curriculum import selected_indicators
    from app.models import AcadIndicatorResult
    inds_by_sub, total_ind = {}, 0
    for sub in subs:
        its = selected_indicators(sub)
        if its:
            inds_by_sub[sub.id] = its
            total_ind += len(its)
    ind_pct = None
    if inds_by_sub:
        passset = {(r.subject_id, r.code) for r in
                   db.query(AcadIndicatorResult)
                   .filter(AcadIndicatorResult.acad_student_id == s.id,
                           AcadIndicatorResult.subject_id.in_(inds_by_sub.keys())).all()
                   if r.passed}
        passed_ind = sum(1 for sid_, its in inds_by_sub.items()
                         for it in its if (sid_, it["code"]) in passset)
        ind_pct = (100.0 * passed_ind / total_ind) if total_ind else None

    def _pass(ok):
        return "ผ่าน" if ok else "ไม่ผ่าน"

    # ที่ | การประเมิน | จำนวน/คะแนน/ร้อยละ | ผลการประเมิน
    rows = [
        ("มีเวลาเรียนตลอดปีการศึกษาร้อยละ 80 ขึ้นไป", att_txt,
         _pass(att_pct is None or att_pct >= 80)),
        ("ผ่านการประเมินตัวชี้วัดรายวิชา ร้อยละ 100",
         f"{ind_pct:.2f}" if ind_pct is not None else "100.00",
         _pass(ind_pct is None or ind_pct >= 100)),
        ("ผ่านการประเมินกิจกรรมพัฒนาผู้เรียนทุกกิจกรรม", act,
         _pass(act != "มผ")),
        ("การตัดสินคุณลักษณะอันพึงประสงค์", char or "-",
         _pass(char not in ("", "ไม่ผ่าน"))),
        ("การตัดสินการอ่าน คิดวิเคราะห์ เขียนสื่อความ", read or "-",
         _pass(read not in ("", "ไม่ผ่าน"))),
        ("ไม่มีรายวิชาที่มีผลการเรียน 0",
         "มีรายวิชาได้ 0" if has_zero else "ผ่านทุกรายวิชา", _pass(not has_zero)),
    ]
    t = doc.add_table(rows=1, cols=4); t.style = "Table Grid"
    for i, h in enumerate(["ที่", "การประเมิน", "จำนวน/คะแนน/ร้อยละ", "ผลการ\nประเมิน"]):
        _cell(t.rows[0].cells[i], h.replace("\n", " "), bold=True, fill="EDE9FE")
    for i, (lab, num, res) in enumerate(rows, 1):
        cells = t.add_row().cells
        _cell(cells[0], i)
        _cell(cells[1], lab, align="left")
        _cell(cells[2], num)
        _cell(cells[3], res, bold=True)
    # แถวสรุป GPA
    fr = t.add_row().cells
    fr[0].merge(fr[1])
    _cell(fr[0], "ผลการเรียนเฉลี่ยตลอดหนึ่งปีการศึกษา", align="left", bold=True, fill="F1F5F9")
    _cell(fr[2], f"{gpa:.2f}" if gpa is not None else "-", bold=True)
    _cell(fr[3], f"ได้  {len(pairs)}", bold=True)
    _widths(t, [Cm(1.2), Cm(8.3), Cm(3.7), Cm(2.8)])

    all_pass = (att_pct is None or att_pct >= 80) and not has_zero and \
        char not in ("", "ไม่ผ่าน") and read not in ("", "ไม่ผ่าน") and \
        (ind_pct is None or ind_pct >= 100)
    grad = is_exit_level(klass.level)
    verb = "จบการศึกษา" if grad else "เลื่อนชั้น"
    _p(doc, "", after=2)
    _p(doc, f"หมายเหตุ   นักเรียน{'ผ่าน' if all_pass else 'ยังไม่ผ่าน'}การประเมินทุกรายการ "
            f"{'สามารถ ' + verb + 'ได้' if all_pass else 'ต้องปรับปรุงก่อน' + verb}",
       size=13, after=6)
    tick_pass = "☑" if all_pass else "☐"
    tick_repeat = "☐" if all_pass else "☑"
    _p(doc, f"สรุปผลการประเมิน   {tick_pass} ได้{verb}      {tick_repeat} ซ้ำชั้น",
       size=14, after=4)
    _p(doc, "วันที่อนุมัติการเรียน  ...............................................", size=14, after=16)

    homerooms = [p.name for p in (klass.homeroom, klass.co_homeroom) if p]
    # ลงนามครูประจำชั้น (สูงสุด 2 คน) วางเคียงกันในตารางไร้เส้น
    st = doc.add_table(rows=1, cols=2)
    _no_border_signs(st.rows[0].cells[0], homerooms[0] if len(homerooms) > 0 else "", "ครูประจำชั้น")
    _no_border_signs(st.rows[0].cells[1], homerooms[1] if len(homerooms) > 1 else "", "ครูประจำชั้น")
    _widths(st, [Cm(8.0), Cm(8.0)])

    _p(doc, "", after=10)
    director = (getattr(school, "director_name", "") or "").strip()
    dpos = ("ผู้อำนวยการ" + school.name) if (school.name or "").startswith("โรงเรียน") \
        else "ผู้อำนวยการโรงเรียน"
    _sign_block(doc, director, dpos)


def _no_border_signs(cell, name, role):
    """ช่องลงนามไร้เส้นในตาราง (ลายเซ็นลอยถ้ามี)"""
    from docx.oxml.ns import qn as _qn
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.makeelement(_qn("w:tcBorders"), {})
    for edge in ("top", "left", "bottom", "right"):
        borders.append(borders.makeelement(_qn(f"w:{edge}"), {_qn("w:val"): "nil"}))
    tcpr.append(borders)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("(ลงชื่อ)........................................."); r.font.size = Pt(14); r.font.name = THAI_FONT
    r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    if name:
        _float_signature(p, name)
    for txt in (f"( {name or '.......................................'} )", role):
        pp = cell.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_after = Pt(0)
        rr = pp.add_run(txt); rr.font.size = Pt(14); rr.font.name = THAI_FONT
        rr._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)


def _pp6_body(doc, school, s, db, *, page_break: bool = False):
    """สมุดพก ปพ.6 ของนักเรียน 1 คน (9 หน้า) - แนวตั้ง"""
    from app.models import AcadCharEval, AcadReadEval
    ef = effective_eval(s, db)
    _pp6_cover(doc, school, s, db, page_break=page_break)               # 1
    _pp6_personal(doc, school, s, db)                                  # 2
    _pp6_attendance_growth(doc, school, s, db, ef)                     # 3
    _pp6_grades(doc, school, s, db, ef)                                # 4
    _pp6_assess_table(doc, s, db, "คุณลักษณะอันพึงประสงค์",           # 5
                      AcadCharEval, char_avg, ef.get("desired_char"))
    _pp6_assess_table(doc, s, db, "การอ่าน คิดวิเคราะห์ และเขียน",     # 6
                      AcadReadEval, read_avg, ef.get("read_think"))
    _pp6_activities_onet(doc, school, s, db)                           # 7
    _pp6_comments(doc, school, s, db)                                  # 8
    _pp6_summary(doc, school, s, db, ef)                               # 9


def render_pp6(school, s, db) -> str:
    doc = _doc()
    _pp6_body(doc, school, s, db)
    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    out = out_dir / (_safe(f"ปพ.6_{s.name}_{s.klass.year}") + ".docx")
    doc.save(str(out))
    return str(out)


def render_pp6_class(school, klass, db) -> str:
    """สมุดพกทั้งห้อง รวมเป็นไฟล์เดียว (คนละหน้า)"""
    doc = _doc()
    students = sorted(klass.students, key=lambda s: (s.seq or 999, s.name))
    for i, s in enumerate(students):
        _pp6_body(doc, school, s, db, page_break=(i > 0))
    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    out = out_dir / (_safe(f"ปพ.6_ทั้งห้อง_{_class_label(klass)}_{klass.year}") + ".docx")
    doc.save(str(out))
    return str(out)


# ============================ เวลาเรียนรายวัน (พิมพ์จากหน้าเช็กชื่อ) ============================
_WD_ABBR = {"จันทร์": "จ", "อังคาร": "อ", "พุธ": "พ", "พฤหัสบดี": "พฤ",
            "ศุกร์": "ศ", "เสาร์": "ส", "อาทิตย์": "อา"}


def _attendance_month_section(doc, school, klass, db, month, subject=None, page_break=False):
    """หนึ่งหน้าของแบบบันทึกเวลาเรียน (หัว + ตาราง + สรุป + ลงนาม) ของเดือนที่ระบุ
    ใช้ร่วมกันทั้งพิมพ์เดือนเดียวและพิมพ์ทั้งภาคเรียน (หลายเดือน คนละหน้า)"""
    from app.models import AcadAttendance, AcadCalendar, AcadTeaching, Person
    from app.services.academic import month_weekdays

    y = klass.year
    students = sorted(klass.students, key=lambda s: (s.seq or 999, s.name))
    cal = db.query(AcadCalendar).filter_by(year=y, month=month).first()
    open_days = parse_days_csv(cal.days_csv if cal else "")
    wd = month_weekdays(y, month)
    sid = subject.id if subject else None
    subj_cond = (AcadAttendance.subject_id == sid) if sid else AcadAttendance.subject_id.is_(None)
    rows = {}
    if students:
        rows = {a.acad_student_id: a for a in db.query(AcadAttendance)
                .filter(AcadAttendance.acad_student_id.in_([s.id for s in students]),
                        AcadAttendance.month == month, subj_cond).all()}
    # ครูผู้บันทึก: รายวิชา = ครูผู้สอน · โฮมรูม = ครูประจำชั้น
    teacher = ""
    if subject:
        tt = db.query(AcadTeaching).filter_by(class_id=klass.id, subject_id=sid).first()
        teacher = tt.teacher.name if (tt and tt.teacher) else ""
    else:
        teacher = klass.homeroom.name if klass.homeroom else ""

    has_logo = _logo_header(doc, school, height_cm=1.5, page_break=page_break)
    _p(doc, school.name or "", align="center", bold=True, size=15, after=0,
       page_break=(page_break and not has_logo))
    _p(doc, "แบบบันทึกเวลาเรียน (รายวัน)", align="center", bold=True, size=13, after=1)
    # เน้นชื่อเดือนให้เด่น ครูดูออกทันทีว่าเดือนอะไร
    _p(doc, f"ประจำเดือน{TH_MONTH_FULL[month]}  ปีการศึกษา {y}", align="center", bold=True, size=16, after=1)
    sub = f"ชั้น {_class_label(klass)}"
    if subject:
        sub += f"    วิชา {subject.code or ''} {subject.name}"
    if teacher:
        sub += f"    ครูผู้สอน/ประจำชั้น {teacher}"
    _p(doc, sub, align="center", size=12, after=3)

    nday = len(open_days)
    ncol = 2 + nday + 5           # ที่ + ชื่อ + วัน + (มา/ป่วย/ลา/ขาด/ร้อยละ)
    t = doc.add_table(rows=2 + len(students) + 1, cols=ncol)
    t.style = "Table Grid"
    hdr = t.rows[0].cells

    def vm(idx, text, size=10):
        c = t.rows[0].cells[idx].merge(t.rows[1].cells[idx]); _cell(c, text, bold=True, fill="EDE9FE", size=size)

    vm(0, "ที่"); vm(1, "ชื่อ-นามสกุล")
    if nday:
        c = hdr[2].merge(hdr[2 + nday - 1]); _cell(c, "วันที่ (เฉพาะวันเปิดเรียน)", bold=True, fill="EDE9FE", size=10)
        for j, d in enumerate(open_days):
            _cell(t.rows[1].cells[2 + j], f"{_WD_ABBR.get(wd.get(d, ''), '')}\n{d}", bold=True, fill="F1F5F9", size=8)
    labels = ["มา", "ป่วย", "ลา", "ขาด", "ร้อยละ"]
    for k, lb in enumerate(labels):
        vm(2 + nday + k, lb, size=10)

    # แถวนักเรียน
    day_present = [0] * nday          # นับ "มา" รายวัน (แถวสรุปท้าย)
    tot = {"/": 0, "ป": 0, "ล": 0, "ข": 0}
    for i, s in enumerate(students):
        cells = t.rows[2 + i].cells
        _cell(cells[0], s.seq or (i + 1), size=9)
        _cell(cells[1], s.name, align="left", size=9)
        a = rows.get(s.id)
        mk = parse_marks(a.marks) if a else {}
        cnt = count_marks(a.marks) if a else {"/": 0, "ป": 0, "ล": 0, "ข": 0}
        for j, d in enumerate(open_days):
            ch = mk.get(d, "")
            _cell(cells[2 + j], ch, size=9)
            if ch == "/":
                day_present[j] += 1
        for k in ("/", "ป", "ล", "ข"):
            tot[k] += cnt.get(k, 0)
        pct = round(cnt.get("/", 0) * 100.0 / nday, 1) if nday else ""
        vals = [cnt.get("/", 0), cnt.get("ป", 0), cnt.get("ล", 0), cnt.get("ข", 0), pct]
        for k, v in enumerate(vals):
            _cell(cells[2 + nday + k], v if v != 0 else ("0" if k < 4 else v), size=9)

    # แถวสรุปรวมท้ายตาราง: จำนวนมาเรียนรายวัน + ยอดรวมทั้งห้อง
    last = t.rows[2 + len(students)].cells
    sc = last[0].merge(last[1]); _cell(sc, "รวมมาเรียน (คน/วัน)", bold=True, align="right", fill="F1F5F9", size=9)
    for j in range(nday):
        _cell(last[2 + j], day_present[j], bold=True, fill="F1F5F9", size=9)
    ndc = len(students) * nday
    class_pct = round(tot["/"] * 100.0 / ndc, 1) if ndc else ""
    for k, v in enumerate([tot["/"], tot["ป"], tot["ล"], tot["ข"], class_pct]):
        _cell(last[2 + nday + k], v, bold=True, fill="F1F5F9", size=9)

    # ความกว้างคอลัมน์ให้พอดีแนวนอน (พื้นที่พิมพ์ ~26.7 ซม.)
    fixed = 0.9 + 4.2 + (0.9 * 4) + 1.3      # ที่+ชื่อ+4ช่องนับ+ร้อยละ
    dayw = min(0.85, (26.7 - fixed) / nday) if nday else 0.85
    widths = [Cm(0.9), Cm(4.2)] + [Cm(dayw)] * nday + [Cm(0.9)] * 4 + [Cm(1.3)]
    _widths(t, widths)
    _tight_cells(t)
    # บีบความสูงแถว (line spacing แบบตายตัว) ให้ตารางเตี้ยลง -> พอดีหน้าเดียว
    for row in t.rows:
        for cell in row.cells:
            for pp in cell.paragraphs:
                pf = pp.paragraph_format
                pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = Pt(12)

    # สรุปข้อความท้ายเอกสาร (ตัดบรรทัดหมายเหตุ 'รหัส' ออกตามตัวอย่าง)
    total_pres = tot["/"]
    _p(doc, f"สรุป: วันเปิดเรียนทั้งเดือน {nday} วัน · นักเรียน {len(students)} คน · "
            f"รวมวันมาเรียนทั้งห้อง {total_pres} คน-วัน · เฉลี่ยการมาเรียน {class_pct}%  "
            f"(ป่วย {tot['ป']} · ลา {tot['ล']} · ขาด {tot['ข']})",
       align="left", size=11, after=4)
    # ลงนามแบบบรรทัดเดียว (ประหยัดพื้นที่ให้แต่ละเดือนพอดีหน้าเดียว)
    sp = _p(doc, "ลงชื่อ .......................................... ครูผู้บันทึกเวลาเรียน   "
                 f"( {teacher or '..........................................'} )",
            align="right", size=11, after=0)
    if teacher:
        _float_signature(sp, teacher)


def _attendance_doc(landscape=True):
    doc = _doc(landscape=landscape)
    sec0 = doc.sections[0]
    sec0.top_margin = Cm(0.8); sec0.bottom_margin = Cm(0.6)   # บีบให้พอดีหน้าเดียว
    return doc


def render_attendance_month(school, klass, db, month, subject=None) -> str:
    """แบบบันทึกเวลาเรียนรายวันของเดือนหนึ่ง (แนวนอน หน้าเดียว)"""
    doc = _attendance_doc()
    _attendance_month_section(doc, school, klass, db, month, subject=subject)
    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    tag = f"_{subject.code or subject.name}" if subject else ""
    out = out_dir / (_safe(f"เวลาเรียน_{_class_label(klass)}_{TH_MONTH_FULL[month]}{tag}_{klass.year}") + ".docx")
    doc.save(str(out))
    return str(out)


def render_attendance_term(school, klass, db, term, subject=None) -> str:
    """แบบบันทึกเวลาเรียนทั้งภาคเรียน - รวมเดือนของภาคนั้นที่ตั้งปฏิทินแล้ว (คนละหน้า)"""
    from app.models import AcadCalendar
    term_months = [m for m, _ in TH_MONTHS if m in TERM_MONTHS.get(term, set())]
    have = {r.month for r in db.query(AcadCalendar).filter_by(year=klass.year).all()
            if parse_days_csv(r.days_csv)}
    months = [m for m in term_months if m in have] or term_months
    doc = _attendance_doc()
    if not months:
        _p(doc, "ยังไม่มีเดือนในภาคเรียนนี้", align="center", size=14)
    for i, m in enumerate(months):
        _attendance_month_section(doc, school, klass, db, m, subject=subject, page_break=(i > 0))
    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    tag = f"_{subject.code or subject.name}" if subject else ""
    out = out_dir / (_safe(f"เวลาเรียน_ภาค{term}_{_class_label(klass)}{tag}_{klass.year}") + ".docx")
    doc.save(str(out))
    return str(out)


# ============================ ตารางเรียน ============================
_TT_DAYS = [(1, "จันทร์"), (2, "อังคาร"), (3, "พุธ"), (4, "พฤหัสบดี"), (5, "ศุกร์")]

# สีอ่อนแยกวิชา (ต้องตรงกับ JS ในหน้าจัดตาราง _TT_PALETTE)
_TT_PALETTE = ["E0F2FE", "DCFCE7", "FEF3C7", "FCE7F3", "EDE9FE", "FFE4E6",
               "D1FAE5", "FEF9C3", "E0E7FF", "FFEDD5", "CCFBF1", "F3E8FF"]


def _tt_fill(key: str):
    """เลือกสีอ่อนคงที่ต่อวิชา จาก key (เช่น 's5' / 'n:โฮมรูม') - hash ตรงกับฝั่ง JS"""
    if not key:
        return None
    return _TT_PALETTE[sum(ord(ch) for ch in key) % len(_TT_PALETTE)]


def _tt_cell_lines(cell, lines, *, fill=None):
    """เขียนหลายบรรทัดในเซลล์เดียว - แต่ละบรรทัด (ข้อความ, ขนาด, หนา) กำหนดเองได้
    ใช้ทำช่องตารางเรียน: บรรทัดแรก=ชื่อวิชา (ใหญ่) บรรทัดถัดไป=ชื่อครู (เล็กลง)"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    for idx, (text, size, bold) in enumerate(lines):
        r = p.add_run(str(text)); r.font.size = Pt(size); r.font.name = THAI_FONT; r.bold = bold
        r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
        if idx < len(lines) - 1:
            r.add_break()
    if fill:
        tcpr = cell._tc.get_or_add_tcPr()
        tcpr.append(tcpr.makeelement(qn("w:shd"), {qn("w:val"): "clear",
                                                   qn("w:color"): "auto", qn("w:fill"): fill}))


def _tt_grid_doc(doc, periods, day_cells):
    """วาดตารางเรียน (แถว=วัน · คอลัมน์=คาบ) ให้ใหญ่เต็มหน้า · day_cells -> list[(ข้อความ,ขนาด,หนา)]"""
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
    ncol = 1 + len(periods)
    t = doc.add_table(rows=1 + len(_TT_DAYS), cols=ncol); t.style = "Table Grid"
    hdr = t.rows[0].cells
    _tt_cell_lines(hdr[0], [("วัน / คาบ", 14, True)], fill="EDE9FE")
    for j, p in enumerate(periods):
        lines = [(p.name, 14, True)]
        if p.time_label:
            lines.append((p.time_label, 12, False))
        _tt_cell_lines(hdr[1 + j], lines, fill=("F1F5F9" if p.is_break else "EDE9FE"))
    t.rows[0].height = Cm(1.2); t.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for i, (dnum, dname) in enumerate(_TT_DAYS):
        cs = t.rows[1 + i].cells
        _tt_cell_lines(cs[0], [(dname, 14, True)], fill="FAF8FF")
        for j, p in enumerate(periods):
            if p.is_break:
                _tt_cell_lines(cs[1 + j], [("พัก", 13, False)], fill="F8FAFC")
            else:
                res = day_cells(dnum, p)
                lines, fill = res if (isinstance(res, tuple) and len(res) == 2) else (res, None)
                _tt_cell_lines(cs[1 + j], lines, fill=fill)
        t.rows[1 + i].height = Cm(1.5); t.rows[1 + i].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    fixed = 2.2
    pw = (26.7 - fixed) / len(periods) if periods else 2.9
    _widths(t, [Cm(2.2)] + [Cm(pw)] * len(periods))
    _tight_cells(t)
    for row in t.rows:                     # จัดกึ่งกลางแนวตั้งให้ดูเป็นระเบียบ
        for c in row.cells:
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return t


def render_timetable_class(school, klass, db) -> str:
    """ตารางเรียนของห้อง (แนวนอน) - วิชา + ครูผู้สอนในแต่ละคาบ"""
    from app.models import AcadPeriod, AcadTimetable, AcadTeaching
    y = klass.year
    periods = (db.query(AcadPeriod).filter_by(year=y)
               .order_by(AcadPeriod.seq, AcadPeriod.id).all())
    cells = {(r.day, r.period_id): r for r in
             db.query(AcadTimetable).filter_by(class_id=klass.id).all()}

    def tname(sid):
        if not sid:
            return ""
        tt = db.query(AcadTeaching).filter_by(class_id=klass.id, subject_id=sid).first()
        return tt.teacher.name if (tt and tt.teacher) else ""

    def cell(dnum, p):
        r = cells.get((dnum, p.id))
        if not r:
            return []
        if r.subject_id:
            nm = r.subject.name if r.subject else ""
            tn = tname(r.subject_id)
            lines = [(nm, 14, False)]
            if tn:
                lines.append((tn, 12, False))     # ชื่อครูผู้สอนใต้ชื่อวิชา
            return lines, _tt_fill(f"s{r.subject_id}")
        note = r.note or ""
        return [(note, 14, False)], (_tt_fill(f"n:{note}") if note else None)

    doc = _doc(landscape=True)
    _logo_header(doc, school)
    _p(doc, "ตารางเรียน", align="center", bold=True, size=18, after=0)
    _p(doc, school.name or "", align="center", bold=True, size=18, after=1)
    hm = f"   ครูประจำชั้น {klass.homeroom.name}" if klass.homeroom else ""
    _p(doc, f"ชั้น {_class_label(klass)}   ปีการศึกษา {y}{hm}", align="center", size=16, after=4)
    _tt_grid_doc(doc, periods, cell)
    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    out = out_dir / (_safe(f"ตารางเรียน_{_class_label(klass)}_{y}") + ".docx")
    doc.save(str(out))
    return str(out)


def render_timetable_teacher(school, person, db, year) -> str:
    """ตารางสอนของครู 1 คน (แนวนอน) - วิชา + ห้องในแต่ละคาบ (ทุกห้องที่สอน)"""
    from app.models import AcadPeriod, AcadTimetable, AcadTeaching, AcadClass
    periods = (db.query(AcadPeriod).filter_by(year=year)
               .order_by(AcadPeriod.seq, AcadPeriod.id).all())
    pairs = {(t.class_id, t.subject_id) for t in
             db.query(AcadTeaching).filter_by(teacher_id=person.id).all()}
    class_ids = {c.id for c in db.query(AcadClass).filter_by(year=year).all()}
    grid = {}
    for r in db.query(AcadTimetable).all():
        if r.class_id in class_ids and (r.class_id, r.subject_id) in pairs:
            c = db.get(AcadClass, r.class_id)
            grid.setdefault((r.day, r.period_id), []).append(
                (r.subject.name if r.subject else "", _class_label(c), r.subject_id))

    def cell(dnum, p):
        lst = grid.get((dnum, p.id), [])
        lines = []
        for nm, cl, _sid in lst:
            lines.append((nm, 14, False))
            lines.append((f"({cl})", 12, False))     # ห้องที่สอนใต้ชื่อวิชา
        fill = _tt_fill(f"s{lst[0][2]}") if lst else None    # แยกสีตามวิชาแรกในคาบ
        return lines, fill

    doc = _doc(landscape=True)
    _logo_header(doc, school)
    _p(doc, "ตารางสอน", align="center", bold=True, size=18, after=0)
    _p(doc, school.name or "", align="center", bold=True, size=18, after=1)
    _p(doc, f"ครู {person.name}   ปีการศึกษา {year}", align="center", size=16, after=4)
    _tt_grid_doc(doc, periods, cell)
    out_dir = get_data_dir() / "documents"; out_dir.mkdir(exist_ok=True)
    out = out_dir / (_safe(f"ตารางสอน_{person.name}_{year}") + ".docx")
    doc.save(str(out))
    return str(out)
