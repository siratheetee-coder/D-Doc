# -*- coding: utf-8 -*-
"""
sale_doc.py — ออกใบเสนอราคา / ใบเสร็จรับเงิน (Word) จากข้อมูลผู้ขาย (seller_config) + ลูกค้า (lead)
ฟอนต์ TH Sarabun + โลโก้ + รวมเป็นตัวอักษร (บาทถ้วน)
"""
from pathlib import Path
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from app.database import get_data_dir
from app.thai_utils import thai_date, bahttext

THAI_FONT = "TH Sarabun New"
_STATIC = Path(__file__).resolve().parent.parent / "static"
BLUE = RGBColor(0x1d, 0x4e, 0xd8)


def _safe(text: str) -> str:
    for ch in '<>:"/\\|?*':
        text = text.replace(ch, "_")
    return text.strip()


def _fmt(v) -> str:
    return "{:,.2f}".format(v or 0)


def _run(p, text, *, bold=False, size=14, color=None):
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = THAI_FONT
    if color is not None:
        r.font.color.rgb = color
    r._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    r._element.rPr.rFonts.set(qn("w:ascii"), THAI_FONT)
    r._element.rPr.rFonts.set(qn("w:hAnsi"), THAI_FONT)
    return r


def _p(doc, text="", *, align="left", bold=False, size=14, after=2, color=None):
    p = doc.add_paragraph()
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        _run(p, text, bold=bold, size=size, color=color)
    return p


def _set_cell(cell, text, *, bold=False, align="left", size=13.5, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(0)
    _run(p, text, bold=bold, size=size)
    if fill is not None:
        shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {
            qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): fill})
        cell._tc.get_or_add_tcPr().append(shd)


def _header(doc, seller, title, doc_no, doc_date):
    """หัวเอกสาร: โลโก้+ผู้ขาย (ซ้าย) / ชื่อเอกสาร+เลขที่+วันที่ (ขวา)"""
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Cm(11.0)
    t.columns[1].width = Cm(6.5)
    left, right = t.rows[0].cells
    left.width = Cm(11.0); right.width = Cm(6.5)

    # ซ้าย: โลโก้ + ข้อมูลผู้ขาย
    logo = _STATIC / "logo.png"
    lp = left.paragraphs[0]; lp.paragraph_format.space_after = Pt(2)
    if logo.exists():
        try:
            lp.add_run().add_picture(str(logo), height=Cm(1.3))
        except Exception:
            pass
    _run(lp, "  D-Doc", bold=True, size=20, color=BLUE)
    _set_cell_lines(left, [
        (seller.get("name", ""), True, 15),
        (seller.get("address", ""), False, 13.5),
        (f"เลขประจำตัวผู้เสียภาษี {seller.get('tax_id','')}", False, 13.5),
        (f"โทร. {seller.get('phone','')}" + (f"  อีเมล {seller['email']}" if seller.get("email") else ""), False, 13.5),
    ])

    # ขวา: ชื่อเอกสาร + เลขที่ + วันที่
    rp = right.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(4)
    _run(rp, title, bold=True, size=22, color=BLUE)
    _set_cell_lines(right, [
        (f"เลขที่  {doc_no}", False, 14),
        (f"วันที่  {thai_date(doc_date)}", False, 14),
    ], align="right")
    _no_borders(t)


def _set_cell_lines(cell, lines, align="left"):
    a = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
         "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    for text, bold, size in lines:
        p = cell.add_paragraph()
        p.alignment = a
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        _run(p, text, bold=bold, size=size)


def _no_borders(table):
    tblPr = table._element.tblPr
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(tblPr.makeelement(qn("w:" + edge), {qn("w:val"): "none"}))
    tblPr.append(borders)


def _customer_box(doc, label, lead):
    _p(doc, "", after=4)
    box = doc.add_table(rows=1, cols=1)
    box.style = "Table Grid"
    c = box.rows[0].cells[0]
    _set_cell(c, "", size=6)
    lines = [(f"{label}  {lead.get('school_name') or '-'}", True, 14.5)]
    if lead.get("address"):
        lines.append((f"ที่อยู่  {lead['address']}", False, 13.5))
    if lead.get("tax_id"):
        lines.append((f"เลขประจำตัวผู้เสียภาษี  {lead['tax_id']}", False, 13.5))
    contact = []
    if lead.get("contact_name"):
        contact.append("ผู้ติดต่อ " + lead["contact_name"])
    if lead.get("phone"):
        contact.append("โทร. " + lead["phone"])
    if contact:
        lines.append(("  ·  ".join(contact), False, 13.5))
    _set_cell_lines(c, lines)


def _items_table(doc, lead):
    amount = float(lead.get("amount") or 0)
    desc = f"ค่าบริการระบบบริหารงานเอกสารโรงเรียน D-Doc — {lead.get('packages') or 'ครบทุกงาน'} (สมาชิกรายปี)"
    headers = ["ลำดับ", "รายการ", "จำนวน", "ราคาต่อหน่วย", "จำนวนเงิน (บาท)"]
    widths = [Cm(1.4), Cm(8.6), Cm(1.9), Cm(2.7), Cm(2.9)]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (h, w) in enumerate(zip(headers, widths)):
        _set_cell(table.rows[0].cells[i], h, bold=True, align="center", fill="DCE6F7")
        table.rows[0].cells[i].width = w
    r = table.add_row().cells
    _set_cell(r[0], "1", align="center")
    _set_cell(r[1], desc)
    _set_cell(r[2], "1 ปี", align="center")
    _set_cell(r[3], _fmt(amount), align="right")
    _set_cell(r[4], _fmt(amount), align="right")
    for i, w in enumerate(widths):
        r[i].width = w
    # เว้นแถวว่างให้ดูเป็นตารางจริง
    for _ in range(2):
        e = table.add_row().cells
        for i, w in enumerate(widths):
            _set_cell(e[i], ""); e[i].width = w
    # แถวรวม
    tot = table.add_row().cells
    tot[0].merge(tot[3])
    _set_cell(tot[0], "รวมเป็นเงินทั้งสิ้น", bold=True, align="right")
    _set_cell(tot[4], _fmt(amount), bold=True, align="right")
    for i, w in enumerate(widths):
        try:
            tot[i].width = w
        except Exception:
            pass
    _p(doc, f"({bahttext(amount)})", align="center", bold=True, size=14, after=6)
    return amount


def _sign2(doc, seller, left_role, right_role, left_name=""):
    """ช่องลงนาม 2 คอลัมน์: ซ้าย=ผู้สั่งซื้อ(ลูกค้า) / ขวา=ผู้เสนอราคาหรือผู้รับเงิน(ผู้ขาย)"""
    _p(doc, "", after=16)
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Cm(8.6); t.columns[1].width = Cm(8.6)
    blank_name = "(...........................................)"
    cols = [(t.rows[0].cells[0], left_role, (f"( {left_name} )" if left_name else blank_name)),
            (t.rows[0].cells[1], right_role, f"( {seller.get('signer','')} )")]
    for cell, role, name in cols:
        cell.width = Cm(8.6)
        _set_cell_lines(cell, [
            ("ลงชื่อ ...........................................", False, 14),
            (name, False, 14),
            (role, False, 14),
        ], align="center")
    _no_borders(t)


def _finish(doc, fname):
    out_dir = get_data_dir() / "documents"
    out_dir.mkdir(parents=True, exist_ok=True)
    try:
        path = out_dir / fname
        doc.save(str(path))
        return str(path)
    except OSError:
        import tempfile
        path = Path(tempfile.gettempdir()) / fname
        doc.save(str(path))
        return str(path)


def _base_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(1.8)
    sec.top_margin = Cm(1.3)
    sec.bottom_margin = Cm(1.3)
    base = doc.styles["Normal"]
    base.font.name = THAI_FONT
    base.font.size = Pt(14)
    base._element.rPr.rFonts.set(qn("w:cs"), THAI_FONT)
    return doc


# ---------------- ใบเสนอราคา ----------------
def render_quotation(lead, seller, doc_no, doc_date=None) -> str:
    doc_date = doc_date or datetime.now()
    doc = _base_doc()
    _header(doc, seller, "ใบเสนอราคา", doc_no, doc_date)
    _p(doc, "", after=2)
    _customer_box(doc, "เรียน", lead)
    _p(doc, "", after=4)
    _p(doc, "บริษัท/ผู้ประกอบการ มีความยินดีเสนอราคาสินค้า/บริการ ตามรายการดังต่อไปนี้",
       align="left", after=6)
    _items_table(doc, lead)
    valid = int(seller.get("quote_valid_days", 30) or 30)
    _p(doc, f"เงื่อนไข: ยืนราคา {valid} วัน นับจากวันที่ในใบเสนอราคา",
       align="left", size=13.5, after=2)
    _sign2(doc, seller, "ผู้สั่งซื้อ", "ผู้เสนอราคา", left_name=lead.get("contact_name", ""))
    return _finish(doc, _safe(f"ใบเสนอราคา_{doc_no}_{lead.get('school_name','')}") + ".docx")


# ---------------- ใบเสร็จรับเงิน ----------------
def render_receipt(lead, seller, doc_no, doc_date=None) -> str:
    doc_date = doc_date or datetime.now()
    doc = _base_doc()
    _header(doc, seller, "ใบเสร็จรับเงิน", doc_no, doc_date)
    _p(doc, "", after=2)
    _customer_box(doc, "ได้รับเงินจาก", lead)
    _p(doc, "", after=4)
    _p(doc, "ได้รับชำระเงินค่าสินค้า/บริการ ตามรายการดังต่อไปนี้", align="left", after=6)
    amount = _items_table(doc, lead)
    _p(doc, "การชำระเงิน: โอนเงินผ่านพร้อมเพย์/ธนาคาร ได้รับเงินไว้เป็นการถูกต้องเรียบร้อยแล้ว",
       align="left", size=13.5, after=2)
    _p(doc, "หมายเหตุ: ผู้ขายไม่ได้จดทะเบียนภาษีมูลค่าเพิ่ม จำนวนเงินข้างต้นไม่รวมภาษีมูลค่าเพิ่ม",
       align="left", size=12.5, after=2)
    _sign2(doc, seller, "ผู้สั่งซื้อ", "ผู้รับเงิน", left_name=lead.get("contact_name", ""))
    return _finish(doc, _safe(f"ใบเสร็จรับเงิน_{doc_no}_{lead.get('school_name','')}") + ".docx")
