"""พิมพ์ 'แบบการจัดตารางสอนแทนครูที่ไม่มาปฏิบัติราชการ' (ตามต้นฉบับโรงเรียน)

สร้างเอกสาร .docx เป็นตารางที่มีการ merge หัวตารางและคอลัมน์วันที่/ชื่อครู/สาเหตุ
ให้เหมือนแบบฟอร์มจริง แล้วเติมข้อมูลจาก SubstituteAssignment ที่ฝ่ายวิชาการจัดไว้
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH as _AL
from docx.enum.table import WD_ALIGN_VERTICAL as _VA
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.database import get_data_dir
from app.services.gov_forms import _full_date, _safe

_FONT = "TH Sarabun New"
_TH_MONTHS_TERM = None  # (ไม่ใช้)


def _term_of(d):
    """คาดคะเนภาคเรียนจากเดือน (พ.ค.-ต.ค. = 1, อื่น ๆ = 2)"""
    if not d:
        return ""
    return "1" if 5 <= d.month <= 10 else "2"


def _class_label(klass):
    if not klass:
        return ""
    lvl = (klass.level or "").strip()
    room = (klass.room or "").strip()
    return lvl + ("/" + room if room and room not in ("", "0") else "")


def _period_no(per):
    if not per:
        return ""
    if per.seq:
        return str(per.seq)
    import re
    m = re.search(r"\d+", per.name or "")
    return m.group(0) if m else (per.name or "")


def _set_cell(cell, text, *, bold=False, size=15, align=_AL.CENTER, valign=_VA.CENTER):
    cell.vertical_alignment = valign
    # ยุบให้เหลือย่อหน้าเดียว (เซลล์ที่ merge มาจะมีหลายย่อหน้า -> ข้อความซ้ำ)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    p = cell.paragraphs[0]
    p.alignment = align
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = _FONT
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:cs"), _FONT); rf.set(qn("w:ascii"), _FONT); rf.set(qn("w:hAnsi"), _FONT)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)


def _reason_for(db, a):
    """สาเหตุที่ไม่มาปฏิบัติราชการ จาก source ของ assignment"""
    src = (a.source or "").strip()
    if src == "travel":
        return "ไปราชการ"
    if src == "leave" and a.source_id and db is not None:
        from app.models import LeaveRequest
        lr = db.get(LeaveRequest, a.source_id)
        if lr:
            return (lr.reason or "").strip() or (lr.leave_type or "").strip() or "ลา"
        return "ลา"
    return ""


def render_substitute_schedule(school, db, assignments, *, term="", year="",
                               director_name="", director_position="",
                               print_date=None) -> str:
    """คืน path .docx ของแบบการจัดตารางสอนแทน
    assignments: list[SubstituteAssignment] (เรียงแล้วหรือไม่ก็ได้)
    """
    from app.models import Person, AcadClass, AcadSubject, AcadPeriod

    _p = {p.id: p for p in db.query(AcadPeriod).all()} if db is not None else {}

    def _person(pid):
        return db.get(Person, pid) if (pid and db is not None) else None

    # ---- สร้างแถวข้อมูล (คลี่รายวิชา) ----
    rows = []
    for a in assignments:
        klass = db.get(AcadClass, a.class_id) if (a.class_id and db is not None) else None
        subj = db.get(AcadSubject, a.subject_id) if (a.subject_id and db is not None) else None
        absent = _person(a.absent_person_id)
        sub = _person(a.substitute_person_id)
        per = _p.get(a.period_id)
        rows.append({
            "date": a.date, "absent_id": a.absent_person_id,
            "absent": absent.name if absent else "",
            "reason": _reason_for(db, a),
            "subject": subj.name if subj else "",
            "period": _period_no(per),
            "period_seq": (per.seq if per else 0),
            "klass": _class_label(klass),
            "sub": sub.name if sub else "",
        })
    rows.sort(key=lambda r: ((r["date"].toordinal() if r["date"] else 0),
                             r["absent"], r["period_seq"]))

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.5))

    def _center(text, size, bold=False, after=0):
        p = doc.add_paragraph()
        p.alignment = _AL.CENTER
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(text)
        run.bold = bold; run.font.size = Pt(size); run.font.name = _FONT
        rpr = run._element.get_or_add_rPr()
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:cs"), _FONT); rf.set(qn("w:ascii"), _FONT); rf.set(qn("w:hAnsi"), _FONT)
        rpr.append(rf)
        return p

    _center(school.name or "", 18, bold=True, after=0)
    _center("แบบการจัดตารางสอนแทนครูที่ไม่มาปฏิบัติราชการ", 18, bold=True, after=0)
    _center(f"ภาคเรียนที่ {term or '.....'}  ปีการศึกษา {year or '.........'}", 16, after=6)

    # ---- ตาราง 8 คอลัมน์ ----
    ncol = 8
    table = doc.add_table(rows=2, cols=ncol)
    table.style = "Table Grid"
    table.alignment = 1  # center

    widths = [2.6, 3.6, 2.8, 3.4, 1.3, 1.3, 3.8, 2.4]  # ซม.
    for row in table.rows:
        for c, w in zip(row.cells, widths):
            c.width = Cm(w)

    h0, h1 = table.rows[0].cells, table.rows[1].cells
    _set_cell(h0[0], "วันที่สอนแทน", bold=True)
    _set_cell(h1[0], "วันเดือนปี", bold=True)
    a = table.cell(0, 1).merge(table.cell(1, 1)); _set_cell(a, "รายชื่อครูที่ไม่มาปฏิบัติราชการ", bold=True)
    a = table.cell(0, 2).merge(table.cell(1, 2)); _set_cell(a, "สาเหตุที่ไม่มาปฏิบัติราชการ", bold=True)
    a = table.cell(0, 3).merge(table.cell(0, 6)); _set_cell(a, "การสอนแทน", bold=True)
    _set_cell(h1[3], "รายวิชา", bold=True)
    _set_cell(h1[4], "คาบที่", bold=True)
    _set_cell(h1[5], "ชั้น", bold=True)
    _set_cell(h1[6], "ครูที่ปฏิบัติราชการแทน", bold=True)
    a = table.cell(0, 7).merge(table.cell(1, 7)); _set_cell(a, "ลายมือชื่อผู้สอนแทน", bold=True)

    # ---- แถวข้อมูล ----
    HDR = 2
    if not rows:
        rc = table.add_row().cells
        m = rc[0]
        for k in range(1, ncol):
            m = m.merge(rc[k])
        _set_cell(m, "— ยังไม่มีการจัดครูสอนแทน —")
    for r in rows:
        rc = table.add_row().cells
        _set_cell(rc[0], "")            # เติมภายหลังหลัง merge (กันข้อความซ้ำ)
        _set_cell(rc[1], "", align=_AL.LEFT)
        _set_cell(rc[2], "")
        _set_cell(rc[3], r["subject"], align=_AL.LEFT)
        _set_cell(rc[4], r["period"])
        _set_cell(rc[5], r["klass"])
        _set_cell(rc[6], r["sub"], align=_AL.LEFT)
        _set_cell(rc[7], "")

    # ---- merge แนวตั้ง: วันเดือนปี (ตามวันที่) และ ชื่อครู/สาเหตุ (ตามครูในวันเดียวกัน) ----
    def _merge_col(col, start, end, keep_text):
        top = table.cell(start, col)
        for k in range(start + 1, end + 1):
            top = top.merge(table.cell(k, col))
        _set_cell(top, keep_text, align=(_AL.LEFT if col == 1 else _AL.CENTER))

    n = len(rows)
    i = 0
    while i < n:
        j = i
        while j + 1 < n and rows[j + 1]["date"] == rows[i]["date"]:
            j += 1
        _merge_col(0, HDR + i, HDR + j, _full_date(rows[i]["date"]))
        # ภายในบล็อกวันที่ merge ตามครู
        k = i
        while k <= j:
            m = k
            while m + 1 <= j and rows[m + 1]["absent_id"] == rows[k]["absent_id"]:
                m += 1
            _merge_col(1, HDR + k, HDR + m, rows[k]["absent"])
            _merge_col(2, HDR + k, HDR + m, rows[k]["reason"])
            k = m + 1
        i = j + 1

    # ---- ลงชื่อ ผอ. ----
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    pd = print_date
    _center("ลงชื่อ...............................................ผู้อำนวยการโรงเรียน", 16, after=0)
    _center("(" + (director_name or "").strip() + ")", 16, after=0)
    if pd:
        _center(_full_date(pd), 16, after=0)

    out = get_data_dir() / "documents"; out.mkdir(exist_ok=True)
    path = out / (_safe("แบบจัดตารางสอนแทน") + ".docx")
    doc.save(str(path))
    return str(path)
